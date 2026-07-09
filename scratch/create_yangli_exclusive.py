# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_yangli_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM YANGLI TIÊU CHUẨN")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp các hệ cửa nhôm Yangli vát cạnh vát 3D hiện đại\nGiải pháp tối ưu hiệu quả kinh tế bền vững bảo hành sơn tĩnh điện lên đến 15 năm")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM YANGLI", level=1)
    add_paragraph_with_spacing(doc, 
        "Yangli là một trong những dòng nhôm định hình tầm trung - cận cao cấp rất phổ biến và được ưa chuộng "
        "tại thị trường Việt Nam. Sản phẩm được sản xuất từ phôi nhôm sạch mác 6063-T5 đảm bảo độ cứng chắc vật lý "
        "và khả năng chống biến dạng co ngót cao. Điểm nhấn thẩm mỹ lớn nhất của Yangli là thiết kế 'vát cạnh chéo' (vát góc 35 độ) "
        "tạo hiệu ứng thị giác 3D mềm mại, giúp dễ lau chùi và hạn chế đọng nước tối đa. Hệ nhôm Yangli sở hữu bề mặt sơn tĩnh điện "
        "chất lượng cao được nhà sản xuất cam kết bảo hành lên đến 15 năm, đem lại giải pháp cửa nhôm kính có chất lượng tốt, "
        "thẩm mỹ đẹp cùng hiệu quả kinh tế rất cao cho các công trình nhà phố, văn phòng, và chung cư trung cấp."
    )
    
    add_paragraph_with_spacing(doc, "Yangli cung cấp các hệ nhôm đa dạng phù hợp cho nhiều mục đích sử dụng:")
    
    # Table of 8 systems overview in Word
    table = doc.add_table(rows=9, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.0), Inches(2.0), Inches(1.2), Inches(2.8)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "Hệ 55 Mở Quay", "Yangli hệ 55 vát cạnh", "1.2mm - 1.4mm", "Cửa sổ/đi mở quay vát cạnh nghiêng 35 độ thẩm mỹ, thoát nước tốt. Dòng bán chạy nhất."),
        ("2", "Hệ 55 Trượt Lùa", "Yangli lùa sổ/đi 55", "1.2mm - 1.4mm", "Cửa lùa trượt tiết kiệm không gian mở cánh cho cửa sổ và phòng ngủ."),
        ("3", "Hệ Slim", "Yangli Slim nội thất", "1.2mm", "Vách ngăn kính, cửa đi lùa treo ray trên khung siêu mỏng tối giản."),
        ("4", "Hệ 60", "Yangli mở quay bản trung", "1.4mm - 1.6mm", "Cửa đi mở quay phòng ngủ, cửa ban công cứng vững."),
        ("5", "Hệ 65", "Yangli mở quay bản to", "1.6mm - 2.0mm", "Cửa đi chính đại sảnh mở quay chịu lực gió va đập tốt."),
        ("6", "Hệ 93", "Yangli trượt lùa 2-3 ray", "1.4mm - 1.8mm", "Cửa lùa 4 cánh, lùa 6 cánh bản cánh hộp chạy ray inox chống xệ."),
        ("7", "Hệ Mặt Dựng", "Vách mặt dựng Yangli", "2.0mm - 2.5mm", "Vách kính ngoài trời đố đứng lấy sáng cho showroom và văn phòng tầm trung."),
        ("8", "Hệ Nội Thất", "Tủ cánh kính Yangli", "1.0mm - 1.2mm", "Khung cánh tủ quần áo, tủ trưng bày cánh kính slim sang trọng.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM YANGLI", level=1)
    
    yangli_details = [
        ("1. Hệ 55 Mở Quay - Thiết kế vát cạnh 3D (Bán chạy nhất)",
         "• Độ dày nhôm: 1.2mm cho cửa sổ, 1.4mm cho cửa đi. Tại vị trí chịu lực bản lề dày tới 1.6mm - 1.8mm.\n"
         "• Bản khuôn bao: 55mm.\n"
         "• Công năng: Cửa sổ mở quay, cửa sổ mở hất ngoài ban công, cửa đi thông phòng ngủ, cửa toilet.\n"
         "• Điểm nhấn vát cạnh: Góc vát chéo nghiêng 35 độ giúp cửa có gờ mềm mại kiểu dáng hiện đại, cản gió giật mạnh và trôi nước nhanh."),
         
        ("2. Hệ 55 Trượt Lùa - Tiết kiệm diện tích",
         "• Độ dày nhôm: 1.2mm - 1.4mm.\n"
         "• Bản khuôn bao: 55mm.\n"
         "• Công năng: Cửa sổ trượt lùa 2 cánh, cửa lùa ngăn phòng hẹp.\n"
         "• Ưu điểm: Phẳng mặt, lướt êm ái trên ray trượt nhôm định hình."),
         
        ("3. Hệ Slim nội thất - Phong cách tối giản",
         "• Độ dày nhôm: 1.2mm bản nhôm rộng siêu mảnh chỉ 16mm.\n"
         "• Công năng: Cửa đi lùa treo trong nhà, vách kính phân vùng bếp - phòng khách.\n"
         "• Phụ kiện: Hệ bánh xe trượt treo giảm chấn đóng mở khít kín nhẹ nhàng không cần ray dưới sàn nhà."),
         
        ("4. Hệ 60 - Bản quay trung cấp cách âm tốt",
         "• Độ dày nhôm: 1.4mm - 1.6mm.\n"
         "• Công năng: Cửa đi mở quay ban công, cửa mở quay thông phòng lớn.\n"
         "• Ưu điểm: Bản cánh rộng dập góc thủy lực chắc chắn chịu lực vặn xoắn tốt."),
         
        ("5. Hệ 65 - Cửa mở quay bản to chịu lực gió",
         "• Độ dày nhôm: 1.6mm - 2.0mm.\n"
         "• Bản rộng khuôn bao: 65mm.\n"
         "• Công năng: Phù hợp làm cửa ra vào chính 2 cánh, 4 cánh mở quay mặt tiền cho nhà phố, văn phòng thương mại.\n"
         "• Điểm nhấn: Kết cấu gân gia cường dày chịu gió bão lớn tốt."),
         
        ("6. Hệ 93 - Cửa trượt lùa ray inox",
         "• Độ dày nhôm: 1.4mm - 1.8mm.\n"
         "• Công năng: Cửa đi trượt lùa phòng khách ra ban công sân vườn lớn (2 ray 4 cánh hoặc 3 ray 6 cánh).\n"
         "• Ưu điểm: Vận hành rất đầm tay, ray inox chịu mài mòn giúp bánh xe chạy êm, cách âm giảm tiếng ồn ngoài đường hiệu quả."),
         
        ("7. Hệ Vách Mặt Dựng kính (Curtain Wall)",
         "• Độ dày nhôm: 2.0mm - 2.5mm đố đứng rỗng.\n"
         "• Công năng: Vách ngăn kính lấy sáng sảnh showroom, cửa hàng kinh doanh, tòa nhà văn phòng tầm trung.\n"
         "• Ưu điểm: Thi công lắp dựng nhanh tại công trường, chống thấm nước hoàn hảo nhờ hệ keo kết cấu và gioăng EPDM."),
         
        ("8. Hệ Nội Thất tủ cánh kính slim",
         "• Độ dày nhôm: 1.0mm - 1.2mm siêu nhẹ.\n"
         "• Công năng: Gia công làm khung cửa kính tủ quần áo sang trọng, tủ trưng bày rượu, tủ bếp.\n"
         "• Thẩm mỹ: Sơn phủ bóng màu đen bóng, titan tạo hiệu ứng ánh kim vô cùng cao cấp.")
    ]
    
    for title, desc in yangli_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Warranty and Gasket Specs ---
    add_heading_with_spacing(doc, "QUY TRÌNH SƠN TĨNH ĐIỆN & CHẾ ĐỘ BẢO HÀNH 15 NĂM", level=1)
    add_paragraph_with_spacing(doc, 
        "Nhôm Yangli nổi tiếng với bề mặt sơn tĩnh điện đồng đều và bám dính cực tốt nhờ quy trình xử lý phốt-phát "
        "hóa bề mặt và nung sấy sơn ở nhiệt độ cao theo tiêu chuẩn ISO. Nhà cung cấp áp dụng chế độ bảo hành bề mặt sơn tĩnh điện "
        "lên tới 15 năm chống bong tróc, bay màu hay rỉ sét trong điều kiện thời tiết nhiệt đới khắc nghiệt tại Việt Nam.\n\n"
        "Vật tư phụ đi kèm khuyên dùng:\n"
        "• Hệ gioăng cao su EPDM 2 lớp độ đàn hồi dẻo dai cao, chống lão hóa nhiệt độ.\n"
        "• Keo silicon trung tính Dow Corning chịu nước, keo bọt nở chống thấm khe tường PU.\n"
        "• Hệ phụ kiện đồng bộ Kinlong, PMA, Draho, Bogo chính hãng mang lại tuổi thọ hoạt động trên 100.000 lần mở cửa."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ YANGLI]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Yangli.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Yangli Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_yangli_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Yangli"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for premium
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM KÍNH YANGLI TIÊU CHUẨN"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        'Phụ Kiện', 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Dòng Cạnh Vát (Bán chạy)", "Hệ 55 Mở Quay", "Yangli hệ 55 vát cạnh chéo", 
         "Cửa sổ mở quay, cửa sổ mở hất góc 45 độ chữ A định vị, cửa đi 1 cánh phòng ngủ, cửa toilet.", 
         "1.2mm - 1.4mm", "Phụ kiện Kinlong chính hãng, PMA, Draho", "Gioăng EPDM dẻo dai, keo silicone trung tính", "Góc vát nghiêng 35 độ mềm mại hiện đại, bảo hành sơn 15 năm."),
         
        ("2", "Dòng Tiết Kiệm (Economy)", "Hệ 55 Trượt lùa", "Yangli lùa sổ/đi hệ 55", 
         "Cửa sổ mở trượt ngang 2 cánh, vách ngăn kính lùa trượt, tối ưu cho phòng hẹp.", 
         "1.2mm - 1.4mm", "Bánh xe đơn & chốt sập Kinlong, Draho", "Gioăng lông cản bụi, keo bọt nở PU", "Giải pháp cửa sổ lùa tiết kiệm chi phí thi công."),
         
        ("3", "Dòng Tối Giản (Minimalist)", "Hệ Slim nội thất", "Yangli Slim lùa treo trong nhà", 
         "Cửa đi lùa treo ray trên không ray dưới sàn tránh vấp ngã, vách kính slim phân vùng nội thất.", 
         "1.2mm", "Bộ trượt lùa treo giảm chấn Yangli", "Gioăng cao su đệm cánh, keo silicon", "Đố nhôm bản siêu mỏng rộng 16mm tối giản sang trọng."),
         
        ("4", "Dòng Trung Cấp", "Hệ 60 mở quay", "Yangli mở quay bản cánh trung", 
         "Cửa đi mở quay 1 cánh, 2 cánh ra ban công căn hộ, biệt thự liền kề vừa.", 
         "1.4mm - 1.6mm", "Khóa đa điểm & bản lề 3D Draho, Kinlong", "Gioăng EPDM kép khít kín tốt", "Bản nhôm dày dặn hơn hệ 55, cách âm trung bình."),
         
        ("5", "Dòng Trung Cấp", "Hệ 65 mở quay", "Yangli mở quay bản to chịu lực", 
         "Cửa đi mặt tiền mở quay chính 2 cánh, 4 cánh chịu gió bão va đập rung lắc.", 
         "1.6mm - 2.0mm", "Phụ kiện khóa tay nắm đa điểm Kinlong, Draho", "Gioăng EPDM chống nước mặn và nhiệt", "Cấu trúc khoang rỗng gân gia cường dày chịu gió lớn."),
         
        ("6", "Dòng Tiết Kiệm (Economy)", "Hệ 93 trượt lùa", "Yangli lùa trượt 2-3 ray", 
         "Cửa đi lùa ban công lớn 4 cánh trượt ngang, lồng cánh dồn góc gọn gàng.", 
         "1.4mm - 1.8mm", "Bánh xe kép chịu lực & chốt âm Draho", "Gioăng lông, keo kết cấu Dow Corning", "Ray inox chịu lực hạn chế mài mòn bánh xe, chống xệ tốt."),
         
        ("7", "Dòng Dự Án (B2B)", "Hệ Mặt Dựng", "Vách mặt dựng kính Yangli", 
         "Mặt dựng kính ngoài trời lấy sáng thông tầng văn phòng, sảnh showroom kinh doanh.", 
         "2.0mm - 2.5mm", "Ke định vị mặt dựng & keo silicon kết cấu", "Gioăng EPDM nhiều khoang, keo dán chuyên dụng", "Khả năng chịu áp lực gió cấp bão lớn cho khối đế."),
         
        ("8", "Dòng Nội Thất", "Hệ Tủ Cánh Kính", "Tủ quần áo cánh kính slim", 
         "Khung cánh bao slim siêu mỏng nhẹ bọc kính cường lực màu xám/trà làm tủ rượu, tủ áo.", 
         "1.0mm - 1.2mm", "Bản lề hơi giảm chấn & tay nắm mạ Anode", "Gioăng cao su chèn mép kính", "Sơn màu vàng titan, đen bóng tạo hiệu ứng lộng lẫy hiện đại.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Tiết Kiệm" in seg:
            seg_fill = fill_economy
        elif "Vát Cạnh" in seg or "Tối Giản" in seg:
            seg_fill = fill_premium
        else:
            seg_fill = fill_standard
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Yangli.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Yangli Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_yangli_docx()
    generate_yangli_xlsx()
