# -*- coding: utf-8 -*-
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Sets light borders for the table."""
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def main():
    doc = Document()
    
    # 1. Page Setup & Margins (1 inch / 2.54 cm all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add page numbering setup later or headers/footers
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Cơ Khí Sao Vàng (CKSV) | Báo cáo Định hướng & Chiến lược Phát triển Thương hiệu"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = 'Arial'
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Công ty Cổ phần Sản xuất Cơ khí Sao Vàng - Tài liệu lưu hành nội bộ"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.name = 'Arial'
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # 2. Styling Helpers
    PRIMARY_COLOR = RGBColor(11, 60, 93)      # Dark Navy #0B3C5D
    SECONDARY_COLOR = RGBColor(176, 126, 40)  # Muted Gold #B07E28
    TEXT_COLOR = RGBColor(38, 38, 38)          # Off-Black #262626
    
    def add_heading_1(text, space_before=18, space_after=6):
        h = doc.add_heading('', level=1)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text, space_before=12, space_after=4):
        h = doc.add_heading('', level=2)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return h

    def add_heading_3(text, space_before=8, space_after=2):
        h = doc.add_heading('', level=3)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        return h

    def add_para(text, bold_prefix="", space_after=6, is_bullet=False, indent_level=0):
        style = 'List Bullet' if is_bullet else 'Normal'
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.2
        
        if indent_level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = TEXT_COLOR
            
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = TEXT_COLOR
        return p

    def add_callout(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(10.5)
            r_bold.font.bold = True
            r_bold.font.italic = True
            r_bold.font.color.rgb = PRIMARY_COLOR
            
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(10.5)
        r_text.font.italic = True
        r_text.font.color.rgb = RGBColor(80, 80, 80)
        return p

    # 3. Cover Page / Title Block
    p_comp = doc.add_paragraph()
    p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_comp.paragraph_format.space_before = Pt(36)
    p_comp.paragraph_format.space_after = Pt(18)
    run_comp = p_comp.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_comp.font.name = 'Arial'
    run_comp.font.size = Pt(12)
    run_comp.font.bold = True
    run_comp.font.color.rgb = PRIMARY_COLOR

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("BÁO CÁO NỘI DUNG & ĐỊNH HƯỚNG CHIẾN LƯỢC PHÁT TRIỂN")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_before = Pt(6)
    p_subtitle.paragraph_format.space_after = Pt(72)
    run_sub = p_subtitle.add_run("PHÂN KHÚC: CƠ KHÍ SAO VÀNG (CKSV)\nTái Định Vị thành Đơn Vị Cung Cấp Giải Pháp Cơ Khí Tổng Thể (EPCM)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.left_indent = Inches(1.5)
    p_meta.paragraph_format.space_after = Pt(4)
    r_meta_rec = p_meta.add_run("Kính gửi: ")
    r_meta_rec.font.bold = True
    r_meta_rec.font.name = 'Arial'
    r_meta_rec_val = p_meta.add_run("Ban Giám đốc Công ty Cổ phần Sản xuất Cơ khí Sao Vàng")
    r_meta_rec_val.font.name = 'Arial'

    p_meta2 = doc.add_paragraph()
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta2.paragraph_format.left_indent = Inches(1.5)
    p_meta2.paragraph_format.space_after = Pt(4)
    r_meta_pre = p_meta2.add_run("Người trình bày: ")
    r_meta_pre.font.bold = True
    r_meta_pre.font.name = 'Arial'
    r_meta_pre_val = p_meta2.add_run("Bộ phận Marketing & Kỹ thuật")
    r_meta_pre_val.font.name = 'Arial'

    p_meta3 = doc.add_paragraph()
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta3.paragraph_format.left_indent = Inches(1.5)
    p_meta3.paragraph_format.space_after = Pt(12)
    r_meta_date = p_meta3.add_run("Thời gian: ")
    r_meta_date.font.bold = True
    r_meta_date.font.name = 'Arial'
    r_meta_date_val = p_meta3.add_run("Tháng 07 năm 2026")
    r_meta_date_val.font.name = 'Arial'

    doc.add_page_break()

    # 4. Content Sections
    
    # --- LỜI MỞ ĐẦU ---
    add_heading_1("LỜI MỞ ĐẦU: LÝ DO TÁI ĐỊNH VỊ THƯƠNG HIỆU CƠ KHÍ SAO VÀNG")
    
    add_para("Kính gửi Sếp và Ban Giám đốc,", space_after=10)
    
    add_para("Hiện nay, hơn 95% doanh nghiệp cơ khí quy mô vừa và nhỏ tại Việt Nam đang giới thiệu năng lực theo cách truyền thống: liệt kê danh sách máy móc (máy cắt laser, máy chấn, máy CNC) và các dịch vụ gia công đơn lẻ. Cách tiếp cận này vô tình đưa doanh nghiệp vào cuộc chiến cạnh tranh khốc liệt về giá (Commodity Trap) và làm mờ nhạt đi năng lực thực sự cùng chất xám của đội ngũ kỹ sư.")
    
    add_para("Khách hàng hiện đại — đặc biệt là các Chủ đầu tư lớn, nhà thầu chính, và các doanh nghiệp FDI — không đơn thuần mua 'thời gian chạy máy CNC' hay 'tấn thép thô'. Họ mua một kết quả hoàn chỉnh, một tiến độ được cam kết chặt chẽ, và một đầu mối duy nhất chịu trách nhiệm toàn bộ từ A đến Z.")
    
    add_para("Vì vậy, chúng tôi đề xuất tái định vị Cơ Khí Sao Vàng (CKSV) từ một 'Xưởng gia công cơ khí' thông thường thành ")
    
    p_last = doc.paragraphs[-1]
    r_em = p_last.add_run("Đơn vị Cung cấp Giải pháp Cơ khí Tổng thể (Integrated Mechanical Solutions Provider)")
    r_em.font.bold = True
    r_em.font.name = 'Arial'
    r_last = p_last.add_run(" theo mô hình chuỗi giá trị khép kín EPCM (Engineering – Procurement – Construction – Maintenance).")
    r_last.font.name = 'Arial'
    
    add_para("Tài liệu dưới đây trình bày toàn bộ nội dung chi tiết, định hướng chiến lược thương hiệu và cấu trúc thông tin của Cơ Khí Sao Vàng để Sếp duyệt trước khi triển khai số hóa lên hệ thống website mới và tài liệu bán hàng (Catalogue/Hồ sơ năng lực).")

    # --- PHẦN A: TẦM NHÌN, SỨ MỆNH, GIÁ TRỊ CỐT LÕI ---
    add_heading_1("PHẦN A: ĐỊNH HƯỚNG TẦM NHÌN, SỨ MỆNH & GIÁ TRỊ CỐT LÕI")
    
    add_heading_2("1. TẦM NHÌN (Vision)")
    add_para("Trở thành thương hiệu hàng đầu tại Việt Nam và khu vực trong lĩnh vực cung cấp giải pháp cơ khí tổng thể cho các công trình công nghiệp quy mô lớn và kiến trúc dân dụng cao cấp.")
    add_para("Tiên phong chuẩn hóa và vận hành chuỗi giá trị khép kín từ Thiết kế - Chế tạo - Lắp dựng đến Bảo trì (EPCM). CKSV hướng tới vị thế một thương hiệu cơ khí chất lượng cao, mang lại giá trị bền vững và là đối tác chiến lược tin cậy hàng đầu của các tập đoàn, doanh nghiệp FDI.")

    add_heading_2("2. SỬ MỆNH (Mission)")
    add_para("Đồng hành từ ý tưởng sơ khởi đến khi dự án đi vào vận hành an toàn và bền bỉ. Chúng tôi tối ưu hóa giá trị đầu tư cho khách hàng bằng giải pháp trọn gói kỹ thuật cao, giúp loại bỏ mọi rủi ro trung gian.", is_bullet=True)
    add_para("Thiết lập môi trường làm việc an toàn, chuyên nghiệp, truyền cảm hứng và không ngừng phát triển cho đội ngũ kỹ sư và thợ kỹ nghệ lành nghề.", is_bullet=True)
    add_para("Góp phần nâng tầm thương hiệu cơ khí Việt Nam thông qua việc kiểm soát chất lượng chuẩn mực, ứng dụng công nghệ hiện đại và thiết kế thông minh (DFM).", is_bullet=True)

    add_heading_2("3. GIÁ TRỊ CỐT LÕI (Core Values)")
    add_para("Cam kết trung thực tuyệt đối về nguồn gốc vật tư (100% CO/CQ), minh bạch thông số kỹ thuật và bảo mật hoàn hảo tài sản trí tuệ của đối tác.", bold_prefix="Chính trực (Integrity): ", is_bullet=True)
    add_para("Quy trình khép kín, kiểm soát chất lượng nghiêm ngặt, tuân thủ đúng tiến độ và các tiêu chuẩn kỹ thuật quốc tế (AWS, ASME).", bold_prefix="Chuyên nghiệp (Professionalism): ", is_bullet=True)
    add_para("Đặt thực tế sản xuất và thi công làm trọng tâm trong khâu thiết kế (Triết lý DFM). Không ngừng cải tiến kỹ thuật để mang lại phương án chịu tải và kết cấu tối ưu nhất.", bold_prefix="Tư duy Thực chiến (DFM Mindset): ", is_bullet=True)
    add_para("Kiến tạo công trình bền bỉ qua thời gian (bảo hành kết cấu lên đến 5 năm), bảo đảm an toàn lao động tuyệt đối và phát triển mối quan hệ bền vững cùng đối tác.", bold_prefix="Bền vững (Sustainability): ", is_bullet=True)

    # --- PHẦN B: CKSV GIẢI QUYẾT NHỮNG BÀI TOÁN GÌ? ---
    add_heading_1("PHẦN B: CKSV GIẢI QUYẾT NHỮNG BÀI TOÁN GÌ CHO KHÁCH HÀNG?")
    add_para("Khách hàng doanh nghiệp (chủ đầu tư, tổng thầu, doanh nghiệp FDI) thường gặp rất nhiều khó khăn khi làm việc với các cơ sở gia công truyền thống. CKSV tập trung giải quyết triệt để 5 bài toán cốt lõi sau:")

    add_heading_2("Bài toán 1: Rủi ro đứt gãy thông tin & gánh nặng quản lý nhiều nhà thầu phụ")
    add_callout("Khách hàng phải tự kết nối đơn vị thiết kế -> xưởng cắt chấn -> đội thợ lắp dựng hiện trường -> đơn vị bảo trì. Khi xảy ra sai sót hoặc lệch kích thước, các bên đùn đẩy trách nhiệm, dẫn tới dự án bị trì trệ.", "Nỗi đau thực tế: ")
    add_para("Giải quyết bằng chuỗi dịch vụ khép kín EPCM. CKSV là đầu mối chịu trách nhiệm duy nhất (Single Point of Contact) từ khâu đo đạc hiện trạng, thiết kế DFM, sản xuất tại xưởng cho đến thi công lắp đặt và bảo trì. Khách hàng tiết kiệm được 15-20% chi phí quản lý và loại bỏ hoàn toàn rủi ro đứt gãy thông tin.", "Giải pháp của CKSV: ")

    add_heading_2("Bài toán 2: Bản vẽ lý thuyết xa rời thực tế sản xuất và thi công")
    add_callout("Bản vẽ thiết kế kỹ thuật từ các đơn vị tư vấn thiết kế thuần túy thường rất đẹp về mặt lý thuyết nhưng không khả thi khi đưa vào sản xuất (sai dung sai phôi) hoặc không thể lắp dựng ngoài công trường do thiếu không gian thao tác cẩu nâng hoặc mối hàn ở góc khuất bất khả thi.", "Nỗi đau thực tế: ")
    add_para("Ứng dụng triết lý Thiết kế để Sản xuất & Thi công (DFM). Đội ngũ kỹ sư CKSV am hiểu sâu sắc biện pháp thi công thực địa và máy móc thiết bị nhà xưởng. Chúng tôi chủ động kiểm soát và hiệu chỉnh kết cấu ngay từ bản vẽ thiết kế 3D, giúp triệt tiêu 98% sai lệch và hao phí phôi khi lắp dựng thực tế.", "Giải pháp của CKSV: ")

    add_heading_2("Bài toán 3: Vật liệu đầu vào kém chất lượng, trôi nổi không rõ nguồn gốc")
    add_callout("Các xưởng nhỏ lẻ thường pha trộn thép thứ phẩm, inox kém chất lượng không rõ nguồn gốc để hạ giá thành. Hậu quả là công trình nhanh rỉ sét, nứt gãy mối hàn, gây nguy hiểm nghiêm trọng trong quá trình vận hành.", "Nỗi đau thực tế: ")
    add_para("Cam kết 100% vật tư đầu vào có đầy đủ chứng chỉ xuất xứ và chất lượng (CO/CQ) từ các thương hiệu uy tín toàn cầu và nội địa (Hòa Phát, Posco, Inox Outokumpu...). CKSV sẵn sàng cung cấp đầy đủ hồ sơ kiểm định chất lượng trước khi khởi công.", "Giải pháp của CKSV: ")

    add_heading_2("Bài toán 4: Chậm tiến độ công trình & Phát sinh chi phí ngoài tầm kiểm soát")
    add_callout("Sai hỏng phải làm lại, tiến độ thi công chậm chạp của thợ ngoài hiện trường khiến dự án bị kéo dài, gây trễ hạn đưa công trình vào khai thác và phát sinh hàng loạt chi phí nhân công, quản lý.", "Nỗi đau thực tế: ")
    add_para("Nhờ sự đồng bộ khép kín từ thiết kế - sản xuất đến thi công, CKSV giúp đẩy nhanh tiến độ dự án hơn 20% so với thị trường. Đồng thời, chúng tôi cam kết tiến độ bằng các điều khoản phạt rõ ràng và đảm bảo không phát sinh bất kỳ chi phí ngoài hợp đồng nào.", "Giải pháp của CKSV: ")

    add_heading_2("Bài toán 5: Rò rỉ thông tin công nghệ & Bản vẽ kỹ thuật độc quyền")
    add_callout("Các đối tác FDI và doanh nghiệp chế tạo máy luôn lo sợ bản vẽ chi tiết máy, công nghệ dây chuyền độc quyền bị xưởng gia công rò rỉ cho bên thứ ba hoặc bị sao chép thiết kế.", "Nỗi đau thực tế: ")
    add_para("Xây dựng quy trình bảo mật thông tin nghiêm ngặt và sẵn sàng ký kết văn bản thỏa thuận bảo mật (NDA - Non-Disclosure Agreement) trước khi tiếp nhận tài liệu kỹ thuật của đối tác. Toàn bộ thông tin được mã hóa và chỉ phân quyền cho nhân sự trực tiếp thực hiện dự án.", "Giải pháp của CKSV: ")

    # --- PHẦN I: HỆ THỐNG NHẬN DIỆN THƯƠNG HIỆU & THÔNG ĐIỆP ĐỊNH VỊ ---
    add_heading_1("PHẦN I: HỆ THỐNG NHẬN DIỆN THƯƠNG HIỆU & THÔNG ĐIỆP ĐỊNH VỊ")
    
    add_heading_2("1. Tuyên ngôn định vị thương hiệu (Brand Positioning)")
    add_para('"Cơ Khí Sao Vàng không chỉ gia công cơ khí đơn thuần — Chúng tôi cung cấp giải pháp cơ khí trọn gói, đồng hành cùng khách hàng từ ý tưởng, thiết kế, sản xuất đến thi công lắp đặt và vận hành công trình."')
    
    add_heading_2("2. Thông điệp truyền thông cốt lõi (Slogan & Key Message)")
    add_para("GIẢI PHÁP CƠ KHÍ TOÀN DIỆN — TỪ Ý TƯỞNG ĐẾN CÔNG TRÌNH", bold_prefix="Slogan chính: ")
    add_para("Một Đối Tác – Một Quy Trình – Một Giải Pháp Hoàn Chỉnh.", bold_prefix="Triết lý vận hành: ")
    add_para("Khách hàng chỉ làm việc với một đầu mối duy nhất, giảm thiểu tối đa chi phí quản lý trung gian và thời gian điều phối liên lạc.", bold_prefix="Một Đối Tác: ", indent_level=1)
    add_para("Sự liền mạch tuyệt đối từ khâu khảo sát hiện trạng, lên phương án thiết kế đến khi chế tạo xong và bàn giao dự án nghiệm thu thực tế.", bold_prefix="Một Quy Trình: ", indent_level=1)
    add_para("Kết quả cuối cùng đạt hiệu quả tối ưu nhất về cả tính kỹ thuật, độ bền vững, tính thẩm mỹ và hiệu quả kinh tế.", bold_prefix="Một Giải Pháp Hoàn Chỉnh: ", indent_level=1)

    # --- PHẦN II: 5 LỢI THẾ CẠNH TRANH CỐT LÕI (USPs) ---
    add_heading_1("PHẦN II: 5 LỢI THẾ CẠNH TRANH CỐT LÕI (USPs)")
    add_para("Để thuyết phục các đối tác lớn, tổng thầu và doanh nghiệp FDI ngay từ cái nhìn đầu tiên, CKSV tự tin khẳng định 5 thế mạnh vượt trội:")
    
    add_para("Chúng tôi đảm nhận toàn bộ vòng đời của sản phẩm/dự án cơ khí: Khảo sát hiện trạng → Tư vấn giải pháp kỹ thuật → Thiết kế DFM (2D/3D) → Chế tạo tại xưởng → Thi công lắp đặt → Bảo trì định kỳ. Khách hàng không cần tốn thời gian tìm kiếm và kết nối nhiều nhà thầu phụ cho cùng một hạng mục.", bold_prefix="1. Chuỗi dịch vụ khép kín hoàn chỉnh (End-to-End Capability): ", is_bullet=True)
    
    add_para("Năng lực thiết kế của CKSV luôn gắn liền với thực tế gia công tại xưởng và điều kiện lắp dựng trên công trường. Đội ngũ kỹ sư của chúng tôi hiểu sâu sắc các biện pháp thi công, giúp tối ưu hóa kết cấu ngay từ bản vẽ, giảm thiểu hao phí phôi và triệt tiêu 98% các sai lệch lắp ráp ngoài công trường.", bold_prefix="2. Triết lý Thiết kế để Sản xuất & Thi công (DFM): ", is_bullet=True)
    
    add_para("Chúng tôi không bán các sản phẩm có sẵn để ép khách hàng thích nghi. CKSV chuyên sâu thiết kế và gia công tùy biến theo đúng nhu cầu đặc thù và bản vẽ kỹ thuật chi tiết của từng đối tác, đảm bảo tính độc bản và sự tương thích tuyệt đối.", bold_prefix="3. Năng lực gia công tùy biến cao (Tailored Customization): ", is_bullet=True)
    
    add_para("Kiểm soát chặt chẽ quy trình chất lượng gồm ba yếu tố cốt lõi:", bold_prefix="4. Cam kết tiêu chuẩn chất lượng B2B chuẩn mực: ", is_bullet=True)
    add_para("100% phôi thép, inox, nhôm định hình có đầy đủ chứng chỉ xuất xứ và chất lượng (CO/CQ), nói không với nguyên liệu trôi nổi.", bold_prefix="Vật liệu đầu vào: ", indent_level=1)
    add_para("Đội ngũ thợ hàn tay nghề cao, được đào tạo và kiểm tra định kỳ theo tiêu chuẩn hàn quốc tế (AWS D1.1, ASME).", bold_prefix="Tiêu chuẩn hàn: ", indent_level=1)
    add_para("Đảm bảo nghiêm ngặt dung sai kỹ thuật theo đúng thỏa thuận hợp đồng thông qua các thiết bị đo kiểm hiện đại.", bold_prefix="Độ chính xác kích thước: ", indent_level=1)
    
    add_para("CKSV mang lại chi phí tối ưu nhất cho khách hàng thông qua thiết kế hợp lý tránh dư thừa chịu tải không cần thiết, quy trình sản xuất đồng bộ giảm thiểu sai hỏng phải làm lại và tiến độ nhanh chóng giúp dự án sớm đưa vào khai thác thương mại.", bold_prefix="5. Giải pháp tối ưu chi phí tổng thể (Total Cost Optimization): ", is_bullet=True)

    # --- PHẦN III: 4 TRỤ CỘT NĂNG LỰC DỊCH VỤ (BUSINESS UNITS) ---
    add_heading_1("PHẦN III: 4 TRỤ CỘT NĂNG LỰC DỊCH VỤ (BUSINESS UNITS)")
    add_para("Cơ cấu năng lực của CKSV được tổ chức chặt chẽ theo 4 mảng chiến lược, đại diện cho chuỗi giá trị đầy đủ của một dự án cơ khí hiện đại:")

    # 4 Pillars Summary Table
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table, color="B0C4DE")
    
    # Headers
    headers = [
        "01. ENGINEERING\n(Kỹ Thuật & Thiết Kế)",
        "02. MANUFACTURING\n(Gia Công & Chế Tạo)",
        "03. INSTALLATION\n(Thi Công & Lắp Đặt)",
        "04. MAINTENANCE\n(Bảo Trì & Vận Hành)"
    ]
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_background(cell, "0B3C5D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Sub-headers
    sub_headers = [
        "Biến ý tưởng thành thiết kế sản xuất khả thi",
        "Sản xuất chính xác bằng máy móc CNC & tay nghề cao",
        "Lắp dựng hiện trường nhanh chóng, an toàn tuyệt đối",
        "Đồng hành trọn đời, xử lý sự cố nhanh trong 24 giờ"
    ]
    for col_idx, sub_text in enumerate(sub_headers):
        cell = table.cell(1, col_idx)
        set_cell_background(cell, "F2F6F9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(sub_text)
        run.font.name = 'Arial'
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = SECONDARY_COLOR

    # Brief description
    descriptions = [
        "Khảo sát đo đạc hiện trạng, thiết kế 2D/3D SolidWorks, Shop Drawing DFM, phân tích ứng suất FEA, lập bảng kê vật tư BOM.",
        "Phay CNC, Tiện CNC, Cắt laser fiber kim loại tấm, Chấn gấp CNC, chế tạo kết cấu thép, gia công Inox chuyên dụng (bồn bể, ống công nghệ), gia công nhôm kính, OEM/ODM.",
        "Dựng khung nhà kết cấu thép bằng thiết bị chuyên dụng, lắp đặt mặt dựng kính vách kính biệt thự/cao ốc, căn chỉnh băng tải bồn bể, nghiệm thu NDT.",
        "Bảo trì kết cấu định kỳ, sửa chữa khẩn cấp sự cố mối hàn rò rỉ hoặc hư hỏng cửa kính trong 24h, tính toán gia cường cải tạo nâng cấp tải trọng kết cấu."
    ]
    for col_idx, desc_text in enumerate(descriptions):
        cell = table.cell(2, col_idx)
        set_cell_background(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(desc_text)
        run.font.name = 'Arial'
        run.font.size = Pt(8.5)
        run.font.color.rgb = TEXT_COLOR

    add_para("", space_after=12) # Spacer

    add_heading_2("TRỤ CỘT 01: ENGINEERING SERVICES (Kỹ Thuật & Thiết Kế)")
    add_para("Đây là giai đoạn khởi đầu quyết định sự thành bại và tính kinh tế của toàn bộ dự án.")
    add_callout('"Chúng tôi biến mọi ý tưởng sơ khởi thành giải pháp kỹ thuật có tính khả thi cao và sẵn sàng đưa vào sản xuất."', "Thông điệp chủ đạo: ")
    add_para("Trực tiếp cử kỹ sư đến công trường đo đạc, khảo sát địa hình, không gian lắp đặt để thu thập dữ liệu gốc chính xác nhất.", bold_prefix="Khảo sát & Đo đạc hiện trạng: ", is_bullet=True)
    add_para("Phân tích nhu cầu để đề xuất phương án kết cấu, mác vật liệu phù hợp nhằm đảm bảo khả năng chịu lực, độ bền và tối ưu chi phí.", bold_prefix="Tư vấn kỹ thuật & Tối ưu hóa giải pháp: ", is_bullet=True)
    add_para("Mô phỏng hình học chi tiết sản phẩm trên các phần mềm chuyên nghiệp như AutoCAD, SolidWorks, Inventor.", bold_prefix="Thiết kế kỹ thuật 2D/3D: ", is_bullet=True)
    add_para("Thiết kế chi tiết phục vụ gia công tại xưởng và biện pháp lắp dựng ngoài hiện trường.", bold_prefix="Triển khai bản vẽ Shop Drawing: ", is_bullet=True)
    add_para("Phân tích chịu tải, ứng suất uốn, kéo, nén của các khung dầm, bồn áp lực trước khi chế tạo để đảm bảo an toàn tuyệt đối.", bold_prefix="Tính toán kết cấu & Mô phỏng ứng suất (FEA): ", is_bullet=True)
    add_para("Lập bảng kê vật tư chi tiết, chính xác để phục vụ mua sắm vật liệu hiệu quả, tránh dư thừa gây lãng phí.", bold_prefix="Bóc tách khối lượng (BOM - Bill of Materials): ", is_bullet=True)

    add_heading_2("TRỤ CỘT 02: MANUFACTURING SERVICES (Gia Công & Chế Tạo)")
    add_para("Trái tim của Cơ Khí Sao Vàng nằm ở năng lực sản xuất đồng bộ, kết hợp giữa công nghệ CNC chính xác và kỹ nghệ thủ công tinh xảo.")
    add_callout('"Sản xuất chính xác từng micromet, đáp ứng từ chi tiết máy đơn lẻ đến các kết cấu phức tạp quy mô lớn."', "Thông điệp chủ đạo: ")
    add_para("Phay CNC, Tiện CNC trên các trung tâm gia công hiện đại. Chuyên gia công các chi tiết máy đòi hỏi độ chính xác cao, đồ gá (Jig), khuôn mẫu.", bold_prefix="Gia công cơ khí chính xác: ", is_bullet=True)
    add_para("Cắt laser fiber tốc độ cao trên thép, inox, nhôm định hình; Chấn gấp CNC thủy lực đảm bảo độ vuông góc và bán kính uốn chuẩn xác; Đột dập CNC tạo hình lỗ nghệ thuật hoặc kỹ thuật.", bold_prefix="Gia công kim loại tấm CNC: ", is_bullet=True)
    add_para("Hàn tổ hợp khung nhà thép tiền chế, dầm chịu lực, kết cấu bao che, cầu thang thoát hiểm, giàn không gian.", bold_prefix="Chế tạo kết cấu chịu lực: ", is_bullet=True)
    add_para("Chuyên chế tạo bồn chứa hóa chất, đường ống công nghệ thực phẩm, lan can inox kính, và các phụ kiện chịu ăn mòn muối biển cho du thuyền cao cấp (Marine Grade).", bold_prefix="Gia công Inox chuyên dụng (Inox 304, Inox 316L): ", is_bullet=True)
    add_para("Gia công cắt chấn các profile nhôm định hình, hệ mặt dựng kính khung nhôm, cửa đi, cửa sổ nhôm kính cao cấp.", bold_prefix="Gia công Nhôm & Cửa nhôm kính: ", is_bullet=True)
    add_para("Chế tạo linh kiện cơ khí, vỏ hộp máy, khung bệ máy theo thiết kế riêng biệt của khách hàng để phục vụ lắp ráp hàng loạt.", bold_prefix="Sản xuất OEM/ODM & Hợp đồng sản xuất hàng loạt: ", is_bullet=True)

    add_heading_2("TRỤ CỘT 03: INSTALLATION SERVICES (Thi Công & Lắp Đặt)")
    add_para("Đưa sản phẩm từ nhà xưởng lắp dựng an toàn và hoàn thiện thẩm mỹ tại công trường.")
    add_callout('"Đội ngũ kỹ thuật trực tiếp triển khai lắp dựng, đảm bảo an toàn lao động, tiến độ cam kết và chất lượng thẩm mỹ cao nhất."', "Thông điệp chủ đạo: ")
    add_para("Thi công dựng khung dầm, cột thép nhà xưởng, nhà tiền chế ngoài hiện trường bằng thiết bị cẩu nâng chuyên dụng, căn chỉnh độ thẳng đứng và lực xiết bu-lông liên kết đạt chuẩn.", bold_prefix="Lắp dựng kết cấu thép: ", is_bullet=True)
    add_para("Thi công mặt dựng kính (curtain wall), vách kính, mái đón kính cường lực, cầu thang xoắn thép bản nghệ thuật, lan can, ban công biệt thự.", bold_prefix="Lắp đặt hệ thống nhôm kính & kiến trúc dân dụng: ", is_bullet=True)
    add_para("Định vị, cân chỉnh đồng tâm và lắp đặt các hệ thống băng tải, bồn bể công nghiệp, giá kệ kho thông minh, hệ thống đường ống công nghệ trong các nhà máy.", bold_prefix="Lắp đặt hệ thống thiết bị công nghiệp: ", is_bullet=True)
    add_para("Tiến hành đo đạc độ võng, kiểm tra mối hàn hiện trường (NDT nếu yêu cầu), chạy thử không tải/có tải đối với thiết bị và bàn giao mặt bằng sạch sẽ cho chủ đầu tư.", bold_prefix="Công tác nghiệm thu & Bàn giao: ", is_bullet=True)

    add_heading_2("TRỤ CỘT 04: MAINTENANCE SERVICES (Bảo Trì & Vận Hành)")
    add_para("Sự đồng hành của CKSV không dừng lại sau khi bàn giao dự án. Chúng tôi bảo vệ giá trị đầu tư của khách hàng qua thời gian.")
    add_callout('"Đồng hành dài hạn cùng doanh nghiệp, bảo dưỡng định kỳ và khắc phục nhanh chóng mọi sự cố kỹ thuật."', "Thông điệp chủ đạo: ")
    add_para("Kiểm tra độ võng kết cấu thép, kiểm tra độ kính khít vách kính/cửa nhôm, bôi trơn hệ thống băng tải con lăn, đo kiểm an toàn áp lực bồn bể định kỳ.", bold_prefix="Bảo trì định kỳ: ", is_bullet=True)
    add_para("Cử đội phản ứng nhanh xử lý các sự cố gãy hỏng kết cấu, rò rỉ đường ống công nghệ hoặc hư hỏng cửa nhôm kính tại công trình trong vòng 24 giờ.", bold_prefix="Dịch vụ sửa chữa khẩn cấp: ", is_bullet=True)
    add_para("Tính toán và thực hiện gia cường kết cấu thép để tăng tải trọng sàn, kéo dài băng tải, nâng cấp công nghệ hệ thống đường ống để mở rộng quy mô sản xuất cho nhà máy.", bold_prefix="Cải tạo & Nâng cấp: ", is_bullet=True)

    # --- PHẦN IV: CHIẾN LƯỢC PHÂN TÁCH ĐỐI TƯỢNG KHÁCH HÀNG ---
    add_heading_1("PHẦN IV: CHIẾN LƯỢC PHÂN TÁCH ĐỐI TƯỢNG KHÁCH HÀNG")
    add_para("Để nội dung truyền thông đạt hiệu quả cao, chúng tôi phân chia rõ hai phân khúc khách hàng cốt lõi của CKSV trên các kênh thông tin:")
    
    add_heading_2("1. Phân khúc Khách hàng Dự án & Công nghiệp (B2B)")
    add_para("Các Tổng thầu xây dựng (Coteccons, Hòa Bình, Ricons...), Chủ đầu tư dự án công nghiệp, nhà máy FDI (Nhật Bản, Hàn Quốc, Đài Loan...), các đơn vị vận hành cảng biển/du thuyền.", bold_prefix="Đối tượng: ")
    add_para("Sự an toàn, năng lực pháp lý (năng lực hồ sơ thầu), chứng chỉ vật liệu (CO/CQ), tiêu chuẩn kỹ thuật mối hàn (AWS), tiến độ giao hàng và năng lực tài chính.", bold_prefix="Tiêu chí quyết định mua hàng: ")
    add_para("Sử dụng văn phong kỹ thuật, chuyên nghiệp, đưa các số liệu thực tế về nhà xưởng, năng lực thiết bị và các dự án công nghiệp quy mô lớn đã hoàn thành.", bold_prefix="Cách trình bày nội dung: ")

    add_heading_2("2. Phân khúc Khách hàng Kiến trúc & Dân dụng (B2C & B2B2C)")
    add_para("Chủ biệt thự lâu đài, kiến trúc sư thiết kế nội/ngoại thất, nhà thầu xây dựng nhà dân dụng cao cấp.", bold_prefix="Đối tượng: ")
    add_para("Tính thẩm mỹ tinh xảo, độ hoàn thiện chi tiết bề mặt (đánh bóng gương, sơn tĩnh điện ngoài trời chống rỉ sét), phong cách kiến trúc (tân cổ điển, hiện đại) và sự uy tín của đơn vị thi công.", bold_prefix="Tiêu chí quyết định mua hàng: ")
    add_para("Tập trung vào hình ảnh sản phẩm sắc nét, các đoạn video quay cận cảnh chi tiết mối hàn vô hình, độ uốn cong mượt mà của cầu thang xoắn hay vẻ sang trọng của cửa nhôm kính mạ anodized.", bold_prefix="Cách trình bày nội dung: ")

    # --- PHẦN V: CHÍNH SÁCH BẢO MẬT & KIỂM SOÁT CHẤT LƯỢNG ---
    add_heading_1("PHẦN V: CHÍNH SÁCH BẢO MẬT & KIỂM SOÁT CHẤT LƯỢNG (TIÊU CHUẨN TRÌNH BÀY WEB)")
    add_para("Để tăng độ uy tín (Authority) trên website, chúng tôi đề xuất đưa vào 2 chính sách cam kết pháp lý mạnh mẽ:")
    
    add_heading_2("1. Cam kết bảo mật thông tin bản vẽ kỹ thuật (NDA Commitment)")
    add_callout('"Cơ Khí Sao Vàng cam kết bảo vệ tuyệt đối tài sản trí tuệ và thông tin công nghệ của khách hàng. Mọi bản vẽ thiết kế 2D, 3D, thông số kỹ thuật được gửi tới chúng tôi đều được bảo mật theo quy trình nghiêm ngặt. CKSV sẵn sàng ký kết Thỏa thuận bảo mật thông tin (NDA - Non-Disclosure Agreement) trước khi tiếp nhận tài liệu chi tiết từ phía đối tác."')
    add_para("Xóa bỏ rào cản e ngại của các đối tác FDI hoặc đơn vị chế tạo máy khi muốn tải bản vẽ lên website để yêu cầu báo giá nhanh.", bold_prefix="Mục đích: ")

    add_heading_2("2. Chính sách quản lý chất lượng (Quality Assurance Policy)")
    add_para("Chỉ nhập khẩu và sử dụng thép, inox từ các thương hiệu uy tín toàn cầu hoặc trong nước (như Hòa Phát, Posco, Inox Outokumpu) với đầy đủ giấy tờ CO/CQ chứng minh chất lượng.", bold_prefix="Cam kết nguồn gốc phôi: ", is_bullet=True)
    add_para("Mỗi chi tiết máy sau khi gia công CNC đều trải qua quy trình đo kiểm nghiêm ngặt bằng các thiết bị đo cơ học chính xác hoặc máy đo tọa độ (CMM) trước khi đóng gói xuất xưởng.", bold_prefix="Quy trình kiểm soát dung sai: ", is_bullet=True)
    add_para("Cam kết bảo hành lên tới 5 năm cho các hạng mục kết cấu thép chịu lực và 2 năm cho các thiết bị phụ trợ khác.", bold_prefix="Bảo hành kết cấu dài hạn: ", is_bullet=True)

    # --- PHẦN VI: KẾ HOẠCH BỐ CỤC NỘI DUNG TRÊN WEBSITE MỚI ---
    add_heading_1("PHẦN VI: KẾ HOẠCH BỐ CỤC NỘI DUNG TRÊN WEBSITE MỚI")
    add_para("Bố cục được tối ưu hóa để người dùng trải nghiệm mượt mà, dễ dàng hành động:")
    
    add_para("Giới thiệu thông điệp tổng thể 'Giải pháp cơ khí trọn gói', 5 thế mạnh vượt trội (USPs), tóm tắt 4 trụ cột năng lực dưới dạng các Tab tương tác nhanh, slide dự án tiêu biểu và Form đăng ký gọi lại.", bold_prefix="Trang chủ Cơ Khí (Homepage): ", is_bullet=True)
    add_para("Diễn giải sâu về 4 trụ cột (Engineering, Manufacturing, Construction, Maintenance) kèm theo hình ảnh thực tế xưởng sản xuất và công trường.", bold_prefix="Trang chi tiết Lĩnh vực hoạt động (linh-vuc-co-khi.html): ", is_bullet=True)
    add_para("Thư viện hình ảnh thực tế của các dự án đã hoàn thành (Ví dụ: Dự án Cầu thang thép bản Penthouse, Dự án Lan can du thuyền inox 316L, Khung kết cấu thép nhà xưởng).", bold_prefix="Trang Dự án (du-an.html & du-an-chi-tiet.html): ", is_bullet=True)
    add_para("Catalog hiển thị các sản phẩm cơ khí chi tiết như Jig gá, băng tải con lăn, bồn chứa inox, chi tiết máy chính xác để bộ phận mua hàng dễ tham chiếu.", bold_prefix="Trang Sản phẩm (san-pham.html): ", is_bullet=True)
    add_para("Tích hợp nút upload file bản vẽ kỹ thuật (.pdf, .dwg, .dxf, .step, .zip) cùng dòng cam kết NDA bảo mật để thu hút thông tin dự án trực tiếp.", bold_prefix="Form Báo giá thông minh (Get a Quote Form): ", is_bullet=True)

    # --- PHẦN VII: TỔNG KẾT CHIẾN LƯỢC ---
    add_heading_1("PHẦN VII: TỔNG KẾT CHIẾN LƯỢC PHÁT TRIỂN")
    
    add_heading_2("1. Cơ Khí Sao Vàng giỏi nhất ở đâu? (Core Expertise)")
    add_para("Cơ Khí Sao Vàng làm tốt nhất và có thế mạnh tuyệt đối ở hai khía cạnh sau:")
    add_para("CKSV sở hữu đội ngũ kỹ sư có tư duy thực chiến tại nhà xưởng và hiện trường lắp dựng. Chúng tôi không chỉ nhận bản vẽ và gia công rập khuôn; chúng tôi có khả năng phát hiện lỗi kỹ thuật trên bản vẽ của đối tác, tư vấn tối ưu hóa thiết kế, và chuyển đổi chúng thành bản vẽ Shop Drawing hoàn hảo để đưa vào chế tạo cực nhanh, đạt hiệu quả tối ưu về chịu tải và tiết kiệm phôi vật tư.", bold_prefix="Năng lực thiết kế DFM thực chiến vượt trội: ", is_bullet=True)
    add_para("Rất ít đơn vị cơ khí vừa và nhỏ có thể đồng thời gia công tốt và thi công kết hợp các hạng mục đa vật liệu đòi hỏi tiêu chuẩn khắt khe (như thép kết cấu nặng chịu lực, inox 316L chống ăn mòn muối biển đánh bóng gương, và hệ mặt dựng nhôm kính kiến trúc cao cấp). CKSV làm chủ kỹ thuật hàn chất lượng cao (AWS, ASME) và kiểm soát dung sai chặt chẽ, tạo nên các sản phẩm tích hợp có tính thẩm mỹ và kỹ thuật vượt trội.", bold_prefix="Gia công & Tích hợp kết cấu đa vật liệu phức tạp: ", is_bullet=True)

    add_heading_2("2. Tại sao khách hàng phải chọn Cơ Khí Sao Vàng? (Competitive Advantages)")
    add_para("Giữa hàng ngàn doanh nghiệp cơ khí trên thị trường, khách hàng lựa chọn CKSV vì những giá trị khác biệt duy nhất sau:")
    add_para("CKSV chịu trách nhiệm toàn diện từ A - Z (từ đo đạc, thiết kế, sản xuất đến thi công và bảo trì). Khách hàng không cần phải quản lý và phối hợp nhiều nhà thầu phụ độc lập, loại bỏ hoàn toàn tình trạng đùn đẩy trách nhiệm khi xảy ra lỗi lắp ráp.", bold_prefix="Giải pháp trọn gói 1 đầu mối khép kín (Single-point Responsibility): ", is_bullet=True)
    add_para("Thông qua thiết kế DFM tối ưu lực chịu tải và quy trình chế tạo đồng bộ hóa, CKSV giúp khách hàng giảm thiểu hao phí vật tư, hạn chế 98% sai sót sửa chữa ngoài hiện trường, rút ngắn tiến độ đến 20%, giúp dự án nhanh chóng đi vào khai thác sinh lời.", bold_prefix="Tối ưu hóa tổng chi phí thực tế (Value Engineering): ", is_bullet=True)
    add_para("Chúng tôi công khai xuất xứ vật liệu (100% CO/CQ), chất lượng thợ hàn đạt chứng chỉ quốc tế và quy trình kiểm tra dung sai nghiêm ngặt. Sự minh bạch này giúp khách hàng hoàn toàn an tâm trong công tác nghiệm thu.", bold_prefix="Chất lượng B2B chuẩn mực và minh bạch tối đa: ", is_bullet=True)
    add_para("Chúng tôi cam kết bảo hành kết cấu thép lên tới 5 năm - một cam kết mạnh mẽ chứng minh sự tự tin tuyệt đối vào chất lượng bền vững của công trình.", bold_prefix="Chính sách bảo hành kết cấu vượt trội (Lên đến 5 năm): ", is_bullet=True)
    add_para("Quy trình bảo mật NDA nghiêm ngặt giúp các tập đoàn chế tạo máy và doanh nghiệp FDI hoàn toàn an tâm khi bàn giao bản vẽ công nghệ độc quyền để gia công.", bold_prefix="Cam kết bảo mật tuyệt đối (NDA): ", is_bullet=True)

    # Save document
    output_filename = "BaoCao_ChienLuoc_TaiDinhVi_CKSV.docx"
    doc.save(output_filename)
    print(f"Document successfully created and saved as: {output_filename}")

if __name__ == "__main__":
    main()
