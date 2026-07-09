# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng PMA (Đã tích hợp đầy đủ thông tin từ catalogue Đại Phúc):
1. Catalogue_HeCuaNhom_PMA.xlsx
2. ThuVien_HeCuaNhom_PMA.docx
Phiên bản 3.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

PMA_SYSTEMS = [
    ('1', 'Phổ thông Vát cạnh', 'PMA 55', 'Cửa sổ & đi vát cạnh PMA 55', 
     'Cửa sổ, cửa đi mở quay thiết kế vát góc thẩm mỹ cao, nẹp liền sập dễ sản xuất, tối ưu kinh tế cho công trình dân dụng.', 
     '1.2mm - 1.4mm', 'Phụ kiện đồng bộ PMA / Kinlong rãnh 22', 'Gioăng EPDM đơn / Silicon Apollo A500', 
     'Giá thành rẻ, dễ gia công, linh kiện phổ thông, tuy nhiên khả năng cách âm trung bình.'),
     
    ('2', 'Cao cấp Platinum', 'PMA Platinum', 'Hệ cửa nhôm phào kép PMA Platinum', 
     'Cửa đi, cửa sổ chất lượng vượt trội tích hợp phào kép nghệ thuật ôm tường, sơn AkzoNobel 15 năm.', 
     '1.4mm - 2.0mm', 'Phụ kiện CMECH / BOGO cao cấp rãnh C Châu Âu', 'Gioăng EPDM đa khoang / Silicon Dow', 
     'Thiết kế cánh phẳng sang trọng, độ kín nước cao. Chi phí cao hơn hệ thường.'),
     
    ('3', 'Rãnh C Classic', 'PMA Classic 58', 'Cửa đi & sổ rãnh C Classic 58', 
     'Cửa sổ, cửa đi chuẩn Châu Âu rãnh C tích hợp đồng bộ phụ kiện Châu Âu, chống dột nước tốt.', 
     '1.3mm - 1.4mm', 'Phụ kiện rãnh C Châu Âu (Sigico, Hopo, CMECH)', 'Gioăng cao su chèn nỉ / Silicon Sika', 
     'Dễ lắp phụ kiện cao cấp, cách âm vượt trội trong phân khúc tầm trung.'),
     
    ('4', 'Rãnh C Classic', 'PMA 65', 'Cửa mở quay bản lớn PMA 65', 
     'Cửa đi, cửa sổ mở quay bản dày 65mm cực chắc khỏe, hỗ trợ lắp kính hộp độ dày lớn.', 
     '2.0mm - 3.0mm', 'Phụ kiện CMECH / Roto rãnh C Châu Âu', 'Gioăng EPDM đa khoang / Silicon Dow', 
     'Độ dày đắp ứng tốt cho cả các ứng dụng công nghiệp chịu lực lớn.'),
     
    ('5', 'Rãnh C Classic', 'PMA 75', 'Cửa đi mở quay siêu bền PMA 75', 
     'Cửa đi bản khung 75mm dày dặn chuyên biệt cho các vị trí cửa chịu áp lực gió lớn ngoại thất.', 
     '2.5mm - 3.0mm', 'Bản lề chịu tải CMECH & khóa đa điểm Hopo', 'Gioăng EPDM đúc chịu nhiệt / Sika', 
     'Thích hợp cho các ứng dụng khắt khe như hệ thống đóng tự động và khu vận chuyển công nghiệp.'),
     
    ('6', 'Trượt lùa Trung cấp', 'PMA 93', 'Cửa lùa trượt PMA 93', 
     'Cửa đi trượt lùa 2 ray hoặc 3 ray bản rộng 93mm tiết kiệm không gian đóng mở tối đa.', 
     '1.4mm - 2.0mm', 'Bánh xe chịu lực & khóa bán nguyệt PMA', 'Gioăng EPDM chống nước chèn nỉ', 
     'Giá bình dân, tối ưu diện tích. Độ kín khít cách âm ở mức trung bình.'),
     
    ('7', 'Trượt lùa Trung cấp', 'PMA 95', 'Cửa lùa dày dặn PMA 95', 
     'Cửa đi trượt lùa bản 95mm cải tiến độ dày thanh profile tăng tính vững chãi khi trượt.', 
     '1.6mm - 2.0mm', 'Bánh xe trượt lùa Hopo & khóa tay gạt', 'Gioăng EPDM kép / Apollo A500', 
     'Cứng cáp hơn hệ lùa 93, chống rung tốt. Trọng lượng cánh nặng hơn.'),
     
    ('8', 'Trượt lùa Cao cấp', 'PMA 115', 'Cửa lùa 3 ray PMA 115', 
     'Hệ cửa đi trượt lùa 3 ray tích hợp sẵn dải ray chạy lưới chống côn trùng đồng bộ.', 
     '1.8mm - 2.0mm', 'Phụ kiện Hopo / CMECH đồng bộ', 'Gioăng EPDM đa lớp / Dow Corning', 
     'Tích hợp lưới chống muỗi tiện lợi, sang trọng biệt thự vườn. Chiếm diện tích khuôn tường.'),
     
    ('9', 'Lift & Slide Cao cấp', 'PMA Lift & Slide', 'Cửa trượt nâng PMA Lift & Slide', 
     'Cửa đi trượt nâng cao cấp, cơ cấu gạt tay nắm nâng cánh trượt nhẹ và hạ cánh ép gioăng phẳng kín khít.', 
     '2.0mm', 'Phụ kiện trượt nâng Sobinco (Bỉ) / CMECH chịu tải 300kg', 'Gioăng nén EPDM phẳng sàn', 
     'Kín khít tuyệt đối cách âm tốt, chống bão ven biển. Giá thành phụ kiện rất đắt.'),
     
    ('10', 'Slim Ngoại thất Cao cấp', 'PMA Slim', 'Cửa lùa mỏng tối giản PMA Slim', 
     'Cửa lùa Slim bản cánh siêu nhỏ tối giản tinh tế mở rộng tối đa 95% diện tích mặt kính panorama.', 
     '1.6mm - 2.0mm', 'Ray inox SUS316L & bánh xe chịu tải treo giảm chấn', 'Gioăng chèn kính chìm chống co', 
     'Vẻ đẹp hiện đại tối giản cực sang trọng. Khả năng chống bão ngoài trời hạn chế, chủ yếu làm nội thất.'),
     
    ('11', 'Cầu cách nhiệt Cao cấp', 'PMA Thermal 65', 'Cửa cách nhiệt cầu PAG Thermal 65', 
     'Cửa mở quay cách âm cách nhiệt tốt nhờ cầu cách nhiệt Polyamide dải sợi thủy tinh ở giữa.', 
     '1.4mm - 1.8mm', 'Phụ kiện rãnh C Sobinco / CMECH cách nhiệt', 'Gioăng EPDM đa khoang / Silicon Dow', 
     'Giảm hóa đơn tiền điện điều hòa 30%, cách âm 38dB. Giá thành cao, gia công ép góc CNC phức tạp.'),
     
    ('12', 'Cầu cách nhiệt Siêu cao cấp', 'PMA Thermal 75', 'Cửa cách nhiệt bản lớn Thermal 75', 
     'Hệ cửa cầu cách nhiệt bản lớn cực dày, cách âm đạt trên 40dB chuyên trị phòng ngủ biệt thự hướng Tây.', 
     '2.0mm', 'Phụ kiện an ninh chống cạy RC2 Sobinco', 'Gioăng đúc xốp EPDM cầu trung tâm', 
     'Hiệu năng cản nhiệt cách âm cao nhất thị trường. Giá thành siêu đắt đỏ.'),
     
    ('13', 'Curtain Wall Stick', 'PMA Curtain Wall', 'Hệ vách mặt dựng Stick PMA', 
     'Vách kính mặt dựng bao quanh tòa nhà cao tầng khung xương nhôm chìm hoặc nổi thẩm mỹ hiện đại.', 
     '2.0mm - 2.5mm', 'Ke kết cấu chịu lực chuyên dụng', 'Silicon kết cấu chống thấm Dow Corning 895', 
     'Khả năng chịu gió rung gió bão tốt, lấy sáng tự nhiên tối đa. Đòi hỏi kỹ thuật thi công leo giàn giáo phức tạp.'),
     
    ('14', 'Phổ thông', 'PMA Railing', 'Hệ lan can hợp kim nhôm PMA', 
     'Lan can nhôm định hình lắp ghép modul không hàn, bề mặt sơn tĩnh điện chống rỉ sét muối biển.', 
     '1.5mm - 2.0mm', 'Phụ kiện lắp ghép ốc vít inox SUS304', 'Không dùng keo / Gioăng chặn chân nhôm', 
     'Đồng bộ thẩm mỹ biệt thự, không han gỉ như sắt/inox. Kiểu dáng khuôn cố định ít tùy biến hoa văn.'),
     
    ('15', 'Phổ thông', 'PMA Partition', 'Vách ngăn văn phòng PMA', 
     'Vách ngăn phân chia không gian văn phòng, nhà xưởng gọn nhẹ, lắp đặt di động tháo dỡ nhanh chóng.', 
     '1.0mm - 1.2mm', 'Ke nhảy liên kết góc đơn giản', 'Gioăng chèn nỉ / Silicon Apollo', 
     'Lắp đặt cực nhanh, chi phí siêu rẻ. Khả năng cách âm thấp, kết cấu chịu tải kém.'),
     
    ('16', 'Phổ thông', 'PMA 50', 'Hệ nhôm PMA 50 đa năng', 
     'Cửa sổ mở quay, cửa thông phòng tắm và tường ngăn văn phòng khối lượng nhẹ cực kỳ linh hoạt.', 
     '1.2mm - 2.0mm', 'Phụ kiện đồng bộ PMA / Kinlong nhẹ', 'Gioăng EPDM đơn / Silicon Apollo', 
     'Ứng dụng cho các kết cấu tải trọng nhẹ, độ thẩm mỹ tốt, chi phí cực kỳ tối ưu.'),
     
    ('17', 'Phổ thông', 'PMA 63', 'Cửa đi & cửa sổ lớn PMA 63', 
     'Hệ nhôm chuyên dụng cho các khuôn viên cửa sổ lớn, cửa cổng sân vườn chịu tác động môi trường khắc nghiệt.', 
     '1.5mm - 2.5mm', 'Phụ kiện chịu tải đồng bộ PMA', 'Gioăng EPDM kép cách nước', 
     'Thanh profile thiết kế tối ưu tăng cứng lực chịu lực vặn, chống sệ cánh cửa lớn.'),
     
    ('18', 'Phổ thông', 'PMA 72', 'Cửa đi kích thước lớn PMA 72', 
     'Hệ nhôm đặc biệt cho cửa sổ lớn và cửa cổng mở rộng khẩu độ lớn vượt tiêu chuẩn thông thường.', 
     '2.0mm - 3.0mm', 'Bản lề cối chịu lực lớn & khóa đa điểm', 'Gioăng kép đúc chịu lực chèn chìm', 
     'Chống rung chấn khi đóng mở cửa lớn, thích hợp biệt thự ven biển hoặc căn hộ penhouse rộng.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_PMA.xlsx
# ─────────────────────────────────────────────
def create_pma_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue PMA"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC ĐẦY ĐỦ 18 HỆ THỐNG CỬA NHÔM ĐỊNH HÌNH PMA (VIỆT NAM)"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="0B57D0")
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="EDF7ED")
    for r_idx, row_vals in enumerate(PMA_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_PMA.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel PMA saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel PMA to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_PMA.docx
# ─────────────────────────────────────────────
def create_pma_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x0D,0x22,0x40))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM ĐỊNH HÌNH PMA (MỞ RỘNG 18 HỆ)", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp toàn diện 18 hệ thống cửa nhôm định hình hãng PMA\nTích hợp đầy đủ thông tin kỹ thuật từ catalogue chính thức Đại Phúc và PMA Vina", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM PMA (VIỆT NAM)", bold=True, size=14, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    add_p(
        "Nhôm định hình PMA là thương hiệu uy tín hàng đầu do Nhà máy Nhôm định hình PMA chế tạo đùn ép trực tiếp tại Việt Nam. "
        "Với dây chuyền đùn nhôm hiện đại, sử dụng phôi hợp kim nhôm tiêu chuẩn quốc tế AL6063-T5, sản phẩm có độ cứng vững cao và chống rung chấn xuất sắc. "
        "Ngoài các hệ vát cạnh 55 truyền thống vô cùng kinh tế, PMA hiện nay đã mở rộng và đa dạng hóa dải sản phẩm kỹ thuật rất mạnh mẽ: "
        "Từ hệ cửa nhôm 50 nhẹ cho toilet và phòng ngủ, hệ cửa nhôm 63 cho cửa sổ lớn chịu tác động thời tiết khắc nghiệt, "
        "hệ nhôm rãnh C Classic (58, 65, 75) hỗ trợ cách âm cách nhiệt tốt, hệ trượt nâng và trượt lùa bản lớn (93, 95, 115) tích hợp ray lưới chống côn trùng, "
        "cho đến các dòng vách mặt dựng Stick, lan can modul lắp ghép thông minh và các hệ nhôm cầu cách nhiệt Thermal 65/75 đạt chuẩn Châu Âu cao cấp."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D 18 HỆ NHÔM PMA", bold=True, size=14, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    
    # 1. PMA 55
    add_p("1. PMA 55 Vát cạnh (Phổ thông)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.4mm.")
    add_bullet("Ưu điểm R&D: Giá thành hợp lý, dễ gia công sản xuất nhanh nhờ nẹp liền sập, phụ kiện Kinlong/PMA phổ biến trên thị trường, rất phù hợp lắp nhà phố dân dụng bình dân.")
    add_bullet("Nhược điểm R&D: Thiết kế truyền thống rãnh 22mm không tương thích phụ kiện rãnh C Châu Âu cao cấp, cách âm và cách nhiệt thấp hơn các hệ nhôm dày khác, không nên làm cánh khẩu độ quá lớn.")
    
    # 2. PMA Platinum
    add_p("2. PMA Platinum (Cao cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 2.0mm cứng vững.")
    add_bullet("Ưu điểm R&D: Cánh phẳng trơn nhẵn mang phong cách sang trọng hiện đại, phụ kiện CMECH/Bogo đồng bộ rãnh C Châu Âu cực bền bỉ, độ kín nước cao vượt trội nhờ thiết kế rãnh cách nước kín khít.")
    add_bullet("Nhược điểm R&D: Chi phí nhôm thanh và phụ kiện đi kèm khá cao, yêu cầu người thợ gia công cắt ép góc CNC chính xác cao để bảo đảm độ liền lạc của lớp phào kép nghệ thuật.")

    # 3. PMA Classic 58
    add_p("3. PMA Classic 58 (Rãnh C Châu Âu)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.3mm - 1.4mm.")
    add_bullet("Ưu điểm R&D: Tiêu chuẩn rãnh C chuẩn Châu Âu tích hợp đồng bộ với phụ kiện rãnh C cao cấp, lớp sơn phủ AkzoNobel bảo hành 10-15 năm.")
    add_bullet("Nhược điểm R&D: Phụ kiện rãnh C có giá thành cao hơn phụ kiện rãnh 22 thông thường, cần dưỡng khoét khóa đặc chủng.")

    # 4. PMA 65
    add_p("4. PMA 65 (Rãnh C Classic)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm - 3.0mm (Đáp ứng chuẩn chịu lực công nghiệp nặng).")
    add_bullet("Ưu điểm R&D: Bản khung 65mm dày dặn cực kỳ chắc khỏe, chuyên dùng cho các kết cấu chịu lực gió lớn hoặc ứng dụng kính hộp cách âm dày chân không.")
    add_bullet("Nhược điểm R&D: Cánh nặng, giá thành nhôm thanh cao do trọng lượng nặng.")

    # 5. PMA 75
    add_p("5. PMA 75 (Rãnh C Classic)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.5mm - 3.0mm.")
    add_bullet("Ưu điểm R&D: Bản khung bao 75mm dày dặn vững chắc chịu lực cực tốt ngoại thất, phù hợp các ứng dụng khắt khe như hệ thống đóng tự động, khu vực vận chuyển công nghiệp nặng.")
    add_bullet("Nhược điểm R&D: Chi phí đầu tư cao, lắp đặt nặng.")

    # 6. PMA 93
    add_p("6. PMA 93 (Trượt lùa Trung cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Lùa trượt bản 93 phổ thông, tiết kiệm 100% diện tích không gian đóng mở, giá nhôm thanh rất bình dân dễ bán.")
    add_bullet("Nhược điểm R&D: Độ kín khít chắn mưa bụi kém hơn cửa đi mở quay, dễ mòn ray nhôm nếu tần suất trượt cao.")

    # 7. PMA 95
    add_p("7. PMA 95 (Trượt lùa Trung cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.6mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Khung dày dặn hơn hệ lùa 93 truyền thống, chống rung chịu lực đẩy của gió tốt.")
    add_bullet("Nhược điểm R&D: Bánh xe và phụ kiện đi kèm đắt tiền hơn, trọng lượng vận hành đầm tay hơn.")

    # 8. PMA 115
    add_p("8. PMA 115 (Trượt lùa Cao cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.8mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Thiết kế 3 ray có ray chạy lưới chống côn trùng riêng biệt rất tiện lợi, lấy gió thoáng cực tốt biệt thự sân vườn.")
    add_bullet("Nhược điểm R&D: Chiếm nhiều diện tích độ dày khuôn tường khi lắp dựng.")

    # 9. PMA Lift & Slide
    add_p("9. PMA Lift & Slide (Trượt nâng cao cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm.")
    add_bullet("Ưu điểm R&D: Cơ cấu trượt nâng hạ nén gioăng kín khít tuyệt đối cách âm tốt ngang cửa quay, chịu tải cánh lớn lên đến 300kg trượt lướt êm ái.")
    add_bullet("Nhược điểm R&D: Giá bộ phụ kiện nâng hạ nhập khẩu rất đắt đỏ, đòi hỏi độ phẳng mặt nền sàn thi công cao.")

    # 10. PMA Slim
    add_p("10. PMA Slim (Slim Ngoại thất)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.6mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Khung cánh siêu mỏng cho tầm nhìn panorama 95%, ray inox SUS316L trượt nhẹ nhàng tinh tế.")
    add_bullet("Nhược điểm R&D: Bản cánh mỏng chịu gió giật bão lớn kém, hạn chế chống ngập nước, ưu tiên dùng khu vực trong nhà hoặc ban công có mái che.")

    # 11. PMA Thermal 65
    add_p("11. PMA Thermal 65 (Cầu cách nhiệt)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 1.8mm.")
    add_bullet("Ưu điểm R&D: Cầu cách nhiệt Polyamide giúp giảm truyền nhiệt phòng điều hòa tới 30%, cách âm tốt đạt 38dB.")
    add_bullet("Nhược điểm R&D: Chi phí cao hơn nhôm không cầu cách nhiệt 40-50%, kỹ thuật ép góc bơm keo cấu trúc phức tạp.")

    # 12. PMA Thermal 75
    add_p("12. PMA Thermal 75 (Cầu cách nhiệt)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm bản lớn.")
    add_bullet("Ưu điểm R&D: Cách nhiệt cách âm tối đa đạt trên 40dB, kết cấu đa khoang bản lớn cực kỳ vững chãi.")
    add_bullet("Nhược điểm R&D: Giá thành siêu cao phân khúc Luxury.")

    # 13. PMA Curtain Wall
    add_p("13. PMA Curtain Wall (Mặt dựng Stick)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm - 2.5mm.")
    add_bullet("Ưu điểm R&D: Xương nhôm mặt dựng Stick chịu gió bão tốt, lấy sáng tự nhiên tối đa cho các tòa nhà văn phòng hiện đại.")
    add_bullet("Nhược điểm R&D: Thi công giàn giáo cao tầng bên ngoài chịu ảnh hưởng thời tiết phức tạp.")

    # 14. PMA Railing
    add_p("14. PMA Railing (Lan can nhôm định hình)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.5mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Lan can hợp kim nhôm đúc định hình lắp ghép ốc vít modul đồng bộ, không rỉ sét han gỉ như sắt/inox ngoài trời.")
    add_bullet("Nhược điểm R&D: Kiểu dáng định hình sẵn ít tùy biến hoa văn cổ điển.")

    # 15. PMA Partition
    add_p("15. PMA Partition (Vách ngăn văn phòng)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.0mm - 1.2mm.")
    add_bullet("Ưu điểm R&D: Vách ngăn phân chia không gian văn phòng di động gọn nhẹ, lắp đặt tháo dỡ nhanh, chi phí cực rẻ.")
    add_bullet("Nhược điểm R&D: Cách âm kém, kết cấu mỏng chịu tải kém.")

    # 16. PMA 50
    add_p("16. PMA 50 (Hệ cửa và vách đa năng tải nhẹ)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Thiết kế gọn nhẹ đa năng thích ứng tối đa cho cửa sổ mở quay, cửa phòng tắm, tường ngăn phân chia phòng.")
    add_bullet("Nhược điểm R&D: Khả năng chịu sức gió bão ngoài trời hạn chế, kết cấu tải trọng nhẹ không lắp được kính quá dày.")

    # 17. PMA 63
    add_p("17. PMA 63 (Hệ cửa sổ lớn & cửa cổng chịu tải)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.5mm - 2.5mm.")
    add_bullet("Ưu điểm R&D: Dùng hiệu quả cho các khuôn viên cửa sổ kính lớn, cửa cổng ra vào ngoài trời cần độ cứng cáp chịu tải gió giật.")
    add_bullet("Nhược điểm R&D: Yêu cầu thợ gia công cắt dập đúng ke tăng cứng chống sệ góc cánh lớn.")

    # 18. PMA 72
    add_p("18. PMA 72 (Hệ cửa mở rộng kích thước cực lớn)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm - 3.0mm.")
    add_bullet("Ưu điểm R&D: Độ dày thanh profile vượt trội chuyên trị cửa sổ lớn, cửa cổng mở rộng khẩu độ lớn không bị võng rung lắc.")
    add_bullet("Nhược điểm R&D: Giá thành cao, cánh nặng đòi hỏi bản lề cối và phụ kiện kim khí loại đặc chủng chịu lực tốt.")

    # ── PHẦN SO SÁNH MA TRẬN PHÂN KHÚC HỆ PMA 55 ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN R&D CẠNH TRANH HỆ 55 THỊ TRƯỜNG", bold=True, size=14, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    add_p(
        "Bảng so sánh chéo đặc tính kỹ thuật, phụ kiện và phân khúc giá giữa hệ cửa nhôm PMA 55 "
        "và các thương hiệu cạnh tranh trực tiếp trên thị trường giúp tư vấn bán hàng nhanh:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "PMA 55"
    hdr_c[2].text = "Topal XFAD"
    hdr_c[3].text = "Viralwindow VRA55"
    hdr_c[4].text = "Maxpro R55"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ('Độ Dày (mm)', "1.2mm - 1.4mm", "1.4mm", "1.4mm", "1.4mm"),
        ("Hệ phụ kiện đi kèm", "Kinlong / PMA đồng bộ", "Bogo / CMECH rãnh C", "Viralwindow đồng bộ", "Maxpro / CMECH rãnh C"),
        ("Chuẩn thiết kế rãnh", "Rãnh 22mm truyền thống", "Rãnh C Châu Âu", "Rãnh 22mm truyền thống", "Rãnh C Châu Âu"),
        ("Công nghệ sơn phủ", "Sơn tĩnh điện thường", "Sơn tĩnh điện AkzoNobel", "Sơn tĩnh điện thường/Akzo", "Mạ điện di Anodise ED"),
        ("Phân khúc giá bán", "$$ (Bình dân kinh tế)", "$$$ (Khá - Cao cấp)", "$$ (Tầm trung)", "$$$ (Cận cao cấp)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_PMA.docx")
    try:
        doc.save(path)
        print(f"  >> Docx PMA saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx PMA to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING expanded PMA docx & xlsx with 18 systems in BASE directory...")
    print("=" * 60)
    create_pma_xlsx()
    create_pma_docx()
    print("=" * 60)
