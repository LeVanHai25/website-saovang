# -*- coding: utf-8 -*-
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Sets light borders for the table."""
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def main():
    doc = Document()
    
    # 1. Page Setup & Margins (1 inch / 2.54 cm all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Nhôm Kính Sao Vàng (NKSV) | Báo cáo Định hướng & Chiến lược Phát triển Thương hiệu"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = 'Arial'
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Công ty Cổ phần Sản xuất Cơ khí Sao Vàng - Phân khúc Nhôm Kính - Lưu hành nội bộ"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.name = 'Arial'
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Styling Helpers
    PRIMARY_COLOR = RGBColor(0, 102, 102)      # Deep Teal #006666
    SECONDARY_COLOR = RGBColor(176, 126, 40)  # Muted Gold #B07E28
    TEXT_COLOR = RGBColor(38, 38, 38)          # Off-Black #262626
    
    def add_heading_1(text, space_before=18, space_after=6):
        h = doc.add_heading('', level=1)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text, space_before=12, space_after=4):
        h = doc.add_heading('', level=2)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return h

    def add_heading_3(text, space_before=8, space_after=2):
        h = doc.add_heading('', level=3)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        return h

    def add_para(text, bold_prefix="", space_after=6, is_bullet=False, indent_level=0):
        style = 'List Bullet' if is_bullet else 'Normal'
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.2
        
        if indent_level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = TEXT_COLOR
            
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = TEXT_COLOR
        return p

    def add_callout(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(10.5)
            r_bold.font.bold = True
            r_bold.font.italic = True
            r_bold.font.color.rgb = PRIMARY_COLOR
            
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(10.5)
        r_text.font.italic = True
        r_text.font.color.rgb = RGBColor(80, 80, 80)
        return p

    # --- Title Page ---
    p_comp = doc.add_paragraph()
    p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_comp.paragraph_format.space_before = Pt(36)
    p_comp.paragraph_format.space_after = Pt(18)
    run_comp = p_comp.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_comp.font.name = 'Arial'
    run_comp.font.size = Pt(12)
    run_comp.font.bold = True
    run_comp.font.color.rgb = PRIMARY_COLOR

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(48)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("BÁO CÁO CHI TIẾT & CHIẾN LƯỢC TÁI ĐỊNH VỊ THƯƠNG HIỆU")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_before = Pt(6)
    p_subtitle.paragraph_format.space_after = Pt(48)
    run_sub = p_subtitle.add_run("PHÂN KHÚC: NHÔM KÍNH SAO VÀNG (NKSV)\nNhà Cung Cấp Giải Phái Cửa & Mặt Dựng Nhôm Kính Kiến Trúc Tổng Thể\nTích hợp 19 Hệ Nhôm, Quy trình EPCM & Thư viện kỹ thuật số")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    # Visual Numbers Block
    p_nums = doc.add_paragraph()
    p_nums.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nums.paragraph_format.space_before = Pt(12)
    p_nums.paragraph_format.space_after = Pt(24)
    run_nums = p_nums.add_run("NĂNG LỰC SỐ HÓAẤN TƯỢNG (NĂM 2026)\n12+ Năm Kinh Nghiệm | 800+ Công Trình Bàn Giao | 19 Hệ Nhôm Chính Hãng\n25+ Kỹ Sư Thực Chiến | 150,000+ m² Lắp Dựng | 30+ Tỉnh Thành Phủ Sóng")
    run_nums.font.name = 'Arial'
    run_nums.font.size = Pt(10.5)
    run_nums.font.bold = True
    run_nums.font.color.rgb = PRIMARY_COLOR

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.left_indent = Inches(1.5)
    p_meta.paragraph_format.space_after = Pt(4)
    r_meta_rec = p_meta.add_run("Kính gửi: ")
    r_meta_rec.font.bold = True
    r_meta_rec.font.name = 'Arial'
    r_meta_rec_val = p_meta.add_run("Ban Giám đốc Công ty Cổ phần Sản xuất Cơ khí Sao Vàng")
    r_meta_rec_val.font.name = 'Arial'

    p_meta2 = doc.add_paragraph()
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta2.paragraph_format.left_indent = Inches(1.5)
    p_meta2.paragraph_format.space_after = Pt(4)
    r_meta_pre = p_meta2.add_run("Người trình bày: ")
    r_meta_pre.font.bold = True
    r_meta_pre.font.name = 'Arial'
    r_meta_pre_val = p_meta2.add_run("Bộ phận Marketing & Kỹ thuật")
    r_meta_pre_val.font.name = 'Arial'

    p_meta3 = doc.add_paragraph()
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta3.paragraph_format.left_indent = Inches(1.5)
    p_meta3.paragraph_format.space_after = Pt(12)
    r_meta_date = p_meta3.add_run("Thời gian: ")
    r_meta_date.font.bold = True
    r_meta_date.font.name = 'Arial'
    r_meta_date_val = p_meta3.add_run("Tháng 07 năm 2026")
    r_meta_date_val.font.name = 'Arial'

    doc.add_page_break()

    # --- LỜI MỞ ĐẦU ---
    add_heading_1("LỜI MỞ ĐẦU: LÝ DO TÁI ĐỊNH VỊ THƯƠNG HIỆU NHÔM KÍNH SAO VÀNG")
    add_para("Kính gửi Sếp và Ban Giám đốc,", space_after=10)
    add_para("Thị trường thi công nhôm kính hiện nay tại Việt Nam đang bị bão hòa nghiêm trọng trong phân khúc phổ thông. Hơn 95% đơn vị thi công trên thị trường là các xưởng sản xuất nhỏ lẻ, hoạt động theo mô hình truyền thống: báo giá m2 thô mập mờ, gia công cơ học thiếu máy móc chính xác và sử dụng vật tư trôi nổi không rõ nguồn gốc (nhôm Xingfa nhái, phụ kiện giả mạo). Điều này kéo doanh nghiệp vào cuộc chiến cạnh tranh khốc liệt về giá và làm suy giảm chất lượng công trình của khách hàng.")
    add_para("Khách hàng hiện đại ngày càng thông thái. Họ không chỉ mua các ô cửa nhôm thông thường; họ cần những giải pháp kiến trúc hoàn hảo cho siêu biệt thự, resort ven biển hoặc cao ốc văn phòng. Họ đòi hỏi sự chống chịu thời tiết muối biển khắc nghiệt, khả năng cản gió bão lớn, tính cách âm cách nhiệt tốt và tính thẩm mỹ nghệ thuật (như uốn vòm, cửa lùa slim cánh mỏng chịu lực, vách mặt dựng kính lớn).")
    add_para("Để đón đầu làn sóng số hóa và phục vụ đối tác lớn, chúng tôi đề xuất tái định vị phân khúc Nhôm Kính Sao Vàng (NKSV) từ một xưởng cửa nhôm thông thường trở thành ")
    p_last = doc.paragraphs[-1]
    r_em = p_last.add_run("Nhà Cung Cấp Giải Pháp Cửa & Mặt Dựng Nhôm Kính Kiến Trúc Tổng Thể (Architectural Aluminum & Glass Solutions Provider)")
    r_em.font.bold = True
    r_em.font.name = 'Arial'
    r_last = p_last.add_run(" theo mô hình chuỗi giá trị khép kín EPCM (Engineering – Procurement – Construction – Maintenance).")
    r_last.font.name = 'Arial'
    add_para("Tài liệu dưới đây trình bày toàn bộ chiến lược định hướng thương hiệu, phân khúc hệ nhôm, quy trình vận hành và thư viện kỹ thuật chi tiết để Sếp xem xét và phê duyệt trước khi số hóa lên hệ thống website mới.")

    # --- PHẦN A: ĐỊNH HƯỚNG TẦM NHÌN, SỨ MỆNH & GIÁ TRỊ CỐT LÕI ---
    add_heading_1("PHẦN A: ĐỊNH HƯỚNG TẦM NHÌN, SỨ MỆNH & GIÁ TRỊ CỐT LÕI (NKSV)")
    
    add_heading_2("1. TẦM NHÌN (Vision)")
    add_para("Trở thành thương hiệu hàng đầu Việt Nam trong việc cung cấp giải pháp cửa và mặt dựng nhôm kính kiến trúc phân khúc cao cấp và luxury, đi đầu trong việc ứng dụng công nghệ xử lý bề mặt xi mạ điện di Anodized ED đạt độ bền màu trên 40 năm và các hệ nhôm rãnh C Châu Âu cách âm cách nhiệt hoàn hảo.")

    add_heading_2("2. SỨ MỆNH (Mission)")
    add_para("Kiến tạo những khoảng mở hoàn mỹ, kết nối không gian thiên nhiên và con người bằng các hệ cửa nhôm kính đạt độ thẩm mỹ tinh xảo nhất và hiệu năng kỹ thuật an toàn cao nhất.", is_bullet=True)
    add_para("Đồng hành cùng các kiến trúc sư, nhà thầu chính để hiện thực hóa những thiết kế khó thi công (như uốn vòm cong, vách kính Spider nhịp lớn, cửa xếp trượt góc vuông không cột đỡ).", is_bullet=True)
    add_para("Cam kết loại bỏ vấn nạn hàng nhái, hàng giả linh phụ kiện trên thị trường bằng sự trung thực tuyệt đối về nguồn gốc vật tư xuất xứ chính hãng.", is_bullet=True)

    add_heading_2("3. GIÁ TRỊ CỐT LÕI (Core Values)")
    add_para("Khe ghép góc kín khít như sợi chỉ, silicon hoàn thiện mịn màng, cửa vận hành êm ái đầm chắc.", bold_prefix="Thẩm Mỹ Tinh Xảo (Aesthetics): ", is_bullet=True)
    add_para("Tính toán kỹ lưỡng tải trọng gió, hệ số cản nhiệt U-value, khả năng cách âm để tư vấn giải pháp nhôm kính đúng kỹ thuật cho từng công trình.", bold_prefix="Hiệu Năng Kỹ Thuật (Performance): ", is_bullet=True)
    add_para("100% thanh nhôm, phụ kiện, gioăng, keo đều chính hãng rõ ràng nguồn gốc chứng chỉ (CO/CQ). Nói không với hàng giả nhái.", bold_prefix="Minh Bạch Tuyệt Đối (Integrity): ", is_bullet=True)
    add_para("Bảo hành màu sắc bề mặt lên đến 40 năm. Đồng hành chăm sóc bảo dưỡng cửa định kỳ trọn vòng đời công trình.", bold_prefix="Cam Kết Dài Hạn (Sustainability): ", is_bullet=True)

    # --- PHẦN B: NKSV GIẢI QUYẾT NHỮNG BÀI TOÁN GÌ CHO KHÁCH HÀNG? ---
    add_heading_1("PHẦN B: NKSV GIẢI QUYẾT NHỮNG BÀI TOÁN GÌ CHO KHÁCH HÀNG?")
    add_para("Khách hàng B2B và B2C cao cấp thường đối mặt với các vấn đề nhức nhối khi thi công nhôm kính. NKSV tập trung giải quyết triệt để 5 bài toán lớn sau:")

    add_heading_2("Bài toán 1: Hàng nhái, hàng giả linh phụ kiện và thanh nhôm tràn lan")
    add_callout("Thanh nhôm Xingfa bị làm giả tem nhãn, đặc biệt là phụ kiện khóa bản lề bị nhái loại 2, loại 3 tinh vi. Công trình sử dụng vài tháng đã rỉ sét bản lề, kẹt khóa, xệ cánh.", "Nỗi đau thực tế: ")
    add_para("NKSV thiết lập quy trình kiểm soát chuỗi cung ứng nghiêm ngặt. Chúng tôi là đối tác trực tiếp của các hãng lớn (Civro Đức, Maxpro Nhật Bản, CMECH Mỹ...). Mọi lô hàng nhập xưởng đều đi kèm hồ sơ CO/CQ đầy đủ. Chúng tôi cam kết đền bù 200% nếu phát hiện hàng giả, hàng nhái.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 2: Ngấm nước mưa qua khe cửa & tiếng ồn gió rít cách âm kém")
    add_callout("Sau các trận bão lớn, nước mưa tràn qua các góc ép cánh cửa, thấm đẫm vào sàn gỗ đắt tiền của biệt thự. Hệ gioăng lỏng lẻo khiến tiếng ồn còi xe, tiếng gió rít lọt vào nhà gây mất ngủ.", "Nỗi đau thực tế: ")
    add_para("Ứng dụng biện pháp ép góc thủy lực kết hợp bơm keo chuyên dụng 2 thành phần (Xylotex/Dow Corning) tại xưởng. Sử dụng hệ gioăng EPDM 3 lớp khép kín, kết hợp phụ kiện rãnh C tiêu chuẩn Châu Âu đồng bộ ép chặt cánh cửa vào khung bao, giải quyết dứt điểm hiện tượng thấm dột và đạt hiệu số cách âm lên đến 40dB.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 3: Cửa bị xệ cánh và kẹt cọ nền nhà sau một thời gian sử dụng")
    add_callout("Cửa đi mở quay hoặc xếp lùa có kích thước lớn chịu tải trọng kính hộp rất nặng. Nếu xưởng sử dụng nhôm mỏng, liên kết góc lỏng lẻo và phụ kiện chịu lực kém, cửa sẽ nhanh chóng bị xệ cánh.", "Nỗi đau thực tế: ")
    add_para("Sử dụng các thanh nhôm profile có độ dày tiêu chuẩn từ 1.8mm - 2.5mm cho các hệ cửa lớn, liên kết bằng máy ép góc thủy lực lực ép 15 tấn kết hợp ke góc dày 4mm. NKSV phối hợp các thương hiệu phụ kiện chịu tải cao cấp như CMECH, Sobinco, Roto để đảm bảo độ bền đóng mở trên 100,000 lần không bị xệ cạ.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 4: Mất thẩm mỹ ở các góc liên kết và bong tróc sơn màu nhôm")
    add_callout("Các khe ghép góc 45 độ bị hở lớn, silicon bắn nham nhở, các đường uốn vòm tròn bị móp méo gãy khúc, bề mặt nhôm dễ bong tróc sơn do thời tiết nắng nóng mặn biển khắc nghiệt.", "Nỗi đau thực tế: ")
    add_para("Sản xuất trên dây chuyền máy cắt CNC hai đầu tự động, máy uốn vòm CNC kỹ thuật số chính xác. Thợ bắn silicone của NKSV có tay nghề cao đảm bảo đường keo phẳng mịn, đồng đều. Chúng tôi cung cấp các dòng nhôm công nghệ xi mạ Anodized ED chống phai màu do tia UV và chống muối biển tuyệt đối.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 5: Rủi ro đứt gãy kết nối giữa khung thép chịu lực và vách kính lớn")
    add_callout("Khi làm vách kính lớn hoặc mái đón kính trung tâm, phần kết cấu thép đỡ do một bên cơ khí làm, phần nhôm kính do bên cửa làm. Khi hai bên không khớp bản vẽ, kính lắp vào bị nứt vỡ hoặc không thể lắp vừa, hai nhà thầu đổ lỗi lẫn nhau.", "Nỗi đau thực tế: ")
    add_para("NKSV kế thừa nền tảng cơ khí chịu lực nặng từ CKSV. Chúng tôi có khả năng tự tính toán kết cấu thép đỡ kết hợp hệ nhôm kính mặt dựng Spider/Curtain Wall một cách đồng bộ. Thiết kế 3D tích hợp giúp loại bỏ hoàn toàn các sai lệch giao diện cơ khí - nhôm kính, chịu trách nhiệm trọn gói trước chủ đầu tư.", "Giải pháp của NKSV: ")

    # --- PHẦN C: 10 LÝ DO KHÁCH HÀNG LỰA CHỌN NKSV ---
    add_heading_1("PHẦN C: 10 LÝ DO KHÁCH HÀNG TIN TƯỞNG LỰA CHỌN NKSV")
    add_para("Để khẳng định ưu thế vượt trội và tạo sự an tâm tuyệt đối cho khách hàng, NKSV xây dựng 10 giá trị cam kết thực tế:")
    
    add_para("NKSV chịu trách nhiệm toàn diện từ khảo sát hiện trạng, thiết kế bản vẽ Shop Drawing, sản xuất tại xưởng đến thi công ngoài công trường. Quy trình khép kín giúp loại bỏ hoàn toàn các sai lệch kích thước giữa bản vẽ lý thuyết và thực tế lắp ráp ngoài công trường.", bold_prefix="01. Thiết kế - Chế tạo - Thi công đồng bộ khép kín: ", is_bullet=True)
    add_para("Làm chủ dữ liệu và phân phối 19 hệ nhôm lớn từ Luxury đến Phổ thông. Đội ngũ kỹ sư tư vấn khách quan dựa trên ngân sách thực tế và công năng cần thiết của chủ nhà, tuyệt đối không chèo kéo, không ép bán một thương hiệu cố định.", bold_prefix="02. Hệ sinh thái 19 dòng nhôm chính hãng phong phú: ", is_bullet=True)
    add_para("Kế thừa nhà xưởng cơ khí hiện đại giúp NKSV xử lý xuất sắc các hạng mục siêu khó mà xưởng nhôm kính thông thường từ chối: mái đón kính nhịp lớn kết hợp giàn thép đỡ, vách kính mặt dựng nhôm kính lớn, cửa đi cong uốn vòm, cửa xếp trượt siêu trường.", bold_prefix="03. Nền tảng kết cấu cơ khí chịu lực vững chắc: ", is_bullet=True)
    add_para("Mọi dự án đều được kỹ sư thiết kế trực tiếp đo đạc laser hiện trạng thu thập thông số gốc trước khi báo giá và lên bản vẽ. Nói không với việc báo giá online qua loa, ước lượng mập mờ.", bold_prefix="04. Đội ngũ kỹ sư trực tiếp khảo sát thực địa: ", is_bullet=True)
    add_para("Bảo hành màu sắc sơn tĩnh điện/xi mạ màu từ 10 - 40 năm. Bảo hành chống ngấm thấm nước mưa qua khe cửa lên đến 3 năm và hỗ trợ kỹ thuật trọn đời.", bold_prefix="05. Chính sách bảo hành minh bạch và dài hạn: ", is_bullet=True)
    add_para("Chỉ nhập khẩu thanh nhôm và phụ kiện có đầy đủ chứng chỉ xuất xứ CO/CQ. Nói không với hàng nhái cỏ, hàng kém chất lượng trôi nổi. Sẵn sàng hoàn tiền gấp đôi nếu phát hiện sai cam kết.", bold_prefix="06. Nguồn gốc vật tư minh bạch và rõ ràng: ", is_bullet=True)
    add_para("Đội thợ thi công lành nghề của NKSV tuân thủ quy chuẩn lắp đặt nghiêm ngặt: cân chỉnh khung bằng máy laser, liên kết bằng vít nở inox chịu lực chống rỉ sét, bơm keo bọt nở cách âm và đi keo silicone láng mịn thẩm mỹ.", bold_prefix="07. Quy chuẩn thi công lắp đặt chuẩn kỹ thuật B2B: ", is_bullet=True)
    add_para("Định kỳ 1 năm/lần cử kỹ sư đến kiểm tra độ mỏi bản lề, độ đàn hồi gioăng cao su, đường keo silicone và bảo dưỡng tra dầu mỡ phụ kiện cửa miễn phí trong thời hạn bảo hành.", bold_prefix="08. Chính sách bảo trì định kỳ hàng năm độc nhất: ", is_bullet=True)
    add_para("Có đủ năng lực thi công và quy trình quản lý phù hợp cho mọi công trình từ nhà dân dụng phố nhỏ, siêu biệt thự uốn vòm lâu đài sang trọng đến dự án resort mặt biển quy mô lớn, nhà máy FDI yêu cầu tiêu chuẩn kỹ thuật nghiêm ngặt.", bold_prefix="09. Khả năng đáp ứng đa dạng quy mô công trình: ", is_bullet=True)
    add_para("Khách hàng chỉ làm việc với một đầu mối chịu trách nhiệm duy nhất (Single Point of Contact) cho toàn bộ gói cửa vách kính, loại bỏ hoàn toàn tình trạng đùn đẩy trách nhiệm khi xảy ra sự cố ngấm dột hay kẹt xệ.", bold_prefix="10. Một đầu mối chịu trách nhiệm toàn diện: ", is_bullet=True)

    # --- PHẦN D: TƯ VẤN GIẢI PHÁP THEO LOẠI HÌNH CÔNG TRÌNH ---
    add_heading_1("PHẦN D: GIẢI PHÁP NHÔM KÍNH TỰ CẬP THEO TỪNG LOẠI CÔNG TRÌNH")
    add_para("Kiến trúc sư và chủ đầu tư dễ dàng tìm kiếm phương án phù hợp theo loại hình công trình của mình:")

    add_heading_2("1. Giải pháp cho Siêu Biệt Thự & Lâu Đài")
    add_para("Các hệ nhôm xi mạ điện di Anodized cao cấp (Maxpro.JP, Hondalex, Soco), hệ nhôm cầu cách nhiệt (Civro, Eurowindow) hoặc hệ nhôm lùa Slim siêu mỏng tối giản.", bold_prefix="Khuyến nghị hệ nhôm: ", is_bullet=True)
    add_para("Kính hộp 2 lớp cách âm cách nhiệt tốt tích hợp rèm trong kính (kính hộp Low-E), giúp tiết kiệm điện năng điều hòa và tạo không gian sống cực kỳ yên tĩnh.", bold_prefix="Khuyến nghị hệ kính: ", is_bullet=True)
    add_para("Hệ phụ kiện rãnh C Châu Âu CMECH mạ vàng/bạc cao cấp, khóa vân tay thông minh tích hợp nhà thông minh smarthome, uốn vòm đỉnh cong tân cổ điển uốn lượn nghệ thuật, lan can kính cong dán an toàn và mái kính vòm sảnh đón.", bold_prefix="Phụ kiện & Hạng mục đi kèm: ", is_bullet=True)

    add_heading_2("2. Giải pháp cho Resort & Khách Sạn Ven Biển")
    add_para("Các hệ nhôm sử dụng công nghệ mạ điện di Anodise ED tiêu chuẩn Nhật Bản (Maxpro.JP, Hondalex, Soco) có độ dày từ 1.8mm - 2.5mm chịu tải tốt.", bold_prefix="Khuyến nghị hệ nhôm: ", is_bullet=True)
    add_para("Thanh nhôm mạ ED bảo vệ bề mặt chống ăn mòn của axit và muối mặn biển tuyệt đối, cản tia UV, thiết kế bản cánh lớn chịu sức cản gió bão cấp 12. Kính dán an toàn 2 lớp hoặc kính cường lực dày.", bold_prefix="Đặc tính kỹ thuật chủ lực: ", is_bullet=True)
    add_para("Hệ cửa trượt lùa Panorama khẩu độ lớn tối đa hóa góc view ngắm biển, hệ mặt dựng nhôm kính lớn Spider chân nhện sảnh đón và mái kính chịu gió bão.", bold_prefix="Hạng mục đi kèm: ", is_bullet=True)

    add_heading_2("3. Giải pháp cho Nhà Phố & Chung Cư")
    add_para("Nhôm Xingfa nhập khẩu Quảng Đông tem đỏ, Topal Prima, PMA Platinum, Owin, Viralwindow.", bold_prefix="Khuyến nghị hệ nhôm: ", is_bullet=True)
    add_para("Kinh tế, bền bỉ, tính năng sử dụng ổn định lâu dài, dễ sửa chữa và thay thế phụ kiện. Sử dụng sơn tĩnh điện ngoài trời chất lượng tốt.", bold_prefix="Ưu thế nổi bật: ", is_bullet=True)
    add_para("Cửa sổ mở hất an toàn chống gió đập, cửa đi mở quay hệ 55 ép góc đúc thủy lực, vách kính ngăn phòng tắm, cửa nhôm mở lùa tiết kiệm diện tích hành lang.", bold_prefix="Hạng mục đi kèm: ", is_bullet=True)

    add_heading_2("4. Giải pháp cho Văn Phòng & Cao Ốc")
    add_para("Nhôm Slim trượt siêu mỏng, Eurowindow EA55, Xingfa hệ 65 vách mặt dựng lớn, PMI.", bold_prefix="Khuyến nghị hệ nhôm: ", is_bullet=True)
    add_para("Hệ vách kính mặt dựng lớn Curtain Wall (Unitized hoặc Semi-Unitized) bao che toàn bộ bề mặt tòa nhà, vách ngăn phòng họp kính hộp cách âm tốt, cửa kính thủy lực bản lề sàn lực đẩy nhẹ nhàng, vách kính trang trí sảnh.", bold_prefix="Hạng mục đi kèm: ", is_bullet=True)

    add_heading_2("5. Giải pháp cho Nhà Máy & Nhà Xưởng Công Nghiệp")
    add_para("Nhôm Xingfa Việt Nam, EuroVN, VietPhap, PMA.", bold_prefix="Khuyến nghị hệ nhôm: ", is_bullet=True)
    add_para("Cửa đi công nghiệp bản lớn chịu lực tần suất đóng mở liên tục, hệ cửa đi cửa sổ sổ hất thoát hiểm, vách ngăn kính nhà xưởng, và đặc biệt là hệ cửa chống cháy kiểm soát khói lửa đạt chuẩn PCCC.", bold_prefix="Hạng mục đi kèm: ", is_bullet=True)

    # --- PHẦN E: TƯ VẤN HỆ NHÔM THEO NGÂN SÁCH ---
    add_heading_1("PHẦN E: HƯỚNG DẪN LỰA CHỌN HỆ NHÔM THEO NGÂN SÁCH ĐẦU TƯ")
    add_para("NKSV phân rõ danh mục sản phẩm theo 4 phân khúc ngân sách cụ thể giúp khách hàng dễ dàng định hướng tài chính:")
    
    add_para("Phù hợp cho công trình nhà dân dụng bình dân, nhà cấp 4, phòng trọ cao cấp, nhà xưởng công nghiệp. Sử dụng các hệ nhôm nội địa tiết kiệm chi phí nhưng vẫn đảm bảo độ bền cơ học cơ bản.", bold_prefix="1. Ngân sách dưới 3.000.000 VNĐ/m² (Phân khúc Phổ thông): ", is_bullet=True)
    add_para("VietPhap, PMA vát cạnh, Yangli, VietY, EuroVN.", bold_prefix="Khuyên dùng các hãng: ", indent_level=1)
    
    add_para("Phù hợp cho nhà phố hiện đại, biệt thự mini, nhà vườn dân dụng chất lượng cao, các dự án tòa nhà chung cư tầm trung.", bold_prefix="2. Ngân sách từ 3.000.000 - 5.000.000 VNĐ/m² (Phân khúc Trung cấp): ", is_bullet=True)
    add_para("Xingfa Quảng Đông nhập khẩu tem đỏ, Topal Prima, Owin, Viralwindow.", bold_prefix="Khuyên dùng các hãng: ", indent_level=1)
    
    add_para("Phù hợp cho biệt thự phố sang trọng, nhà phố cao cấp khu đô thị, resort khách sạn nghỉ dưỡng quy mô vừa.", bold_prefix="3. Ngân sách từ 5.000.000 - 8.000.000 VNĐ/m² (Phân khúc Cao cấp): ", is_bullet=True)
    add_para("Maxpro.JP (Nhật Bản mạ ED), Hondalex (Nhật Bản mạ màu), Soco, Kogen, PMI (Malaysia), PAG.", bold_prefix="Khuyên dùng các hãng: ", indent_level=1)
    
    add_para("Phù hợp cho siêu biệt thự lâu đài cổ điển uốn vòm, dinh thự luxury, resort mặt biển cao cấp tiêu chuẩn 5 sao, penthouse sang trọng bậc nhất.", bold_prefix="4. Ngân sách trên 8.000.000 VNĐ/m² (Phân khúc Siêu Cao Cấp / Luxury): ", is_bullet=True)
    add_para("Civro (Đức cầu cách nhiệt), Maxpro.JP bản Luxury phụ kiện CMECH, Eurowindow cầu cách nhiệt kính hộp Low-E, Nhôm Slim ngoại thất bản siêu lớn nhập khẩu.", bold_prefix="Khuyên dùng các hãng: ", indent_level=1)

    # --- PHẦN F: BẢNG SO SÁNH TRỰC QUAN HỆ NHÔM TIÊU BIỂU ---
    add_heading_1("PHẦN F: BẢNG SO SÁNH TRỰC QUAN CÁC HỆ NHÔM TIÊU BIỂU")
    add_para("Bảng đối chiếu kỹ thuật giúp khách hàng và KTS nhanh chóng nhận diện điểm khác biệt giữa 4 hệ nhôm đại diện cho các phân khúc tại NKSV:")

    # Comparison Table
    table_comp = doc.add_table(rows=10, cols=5)
    set_table_borders(table_comp, color="CCCCCC")
    
    widths_comp = [Inches(1.5), Inches(1.3), Inches(1.3), Inches(1.3), Inches(1.3)]
    
    headers_comp = ["Tiêu chí so sánh", "CIVRO (Đức)", "MAXPRO.JP (Nhật)", "XINGFA (Quảng Đông)", "PMA (Việt Nam)"]
    for col_idx, h_text in enumerate(headers_comp):
        cell = table_comp.cell(0, col_idx)
        set_cell_background(cell, "006666")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    comp_data = [
        ("Phân khúc giá thô", "$$$$$ (Trên 8tr/m²)", "$$$$ (5 - 8tr/m²)", "$$$ (3 - 5tr/m²)", "$$ (Dưới 3tr/m²)"),
        ("Độ bền màu sắc màu nhôm", "★★★★★ (Trên 35 năm)", "★★★★★ (40 năm bảo hành)", "★★★★ (10 - 15 năm)", "★★★ (5 - 10 năm)"),
        ("Kháng muối biển, ăn mòn", "★★★★★ (Tuyệt hảo)", "★★★★★ (Tuyệt hảo mạ ED)", "★★★ (Khá)", "★★ (Trung bình)"),
        ("Khả năng cách âm", "★★★★★ (Đỉnh cao cầu nhôm)", "★★★★★ (Rất tốt gioăng rãnh C)", "★★★★ (Khá tốt)", "★★★ (Bình thường)"),
        ("Phụ kiện tương thích", "Sobinco (Bỉ), CMECH (Mỹ)", "CMECH (Mỹ), Sigico, Bogo", "Kinlong chính hãng, Bogo", "KMA, PMA đồng bộ, Kinlong"),
        ("Uốn vòm mỹ thuật", "✓ (Rất đẹp, nhôm dày)", "✓ (Bề mặt mạ ED bóng mịn)", "✓ (Rất phổ biến)", "o (Bản mỏng khó uốn đẹp)"),
        ("Hệ cửa đi trượt lùa Slim", "✓ (Thiết kế rãnh C cực sang)", "✓ (Rất chuộng biệt thự biển)", "o (Bản cánh dày, thô hơn)", "- (Không có hệ Slim chuyên dụng)"),
        ("Phù hợp công trình Biệt thự", "✓ (Khuyên dùng số 1)", "✓ (Khuyên dùng số 1)", "✓ (Phổ thông)", "- (Không phù hợp)"),
        ("Phù hợp công trình Nhà phố", "- (Chi phí quá cao)", "✓ (Tốt nếu tài chính dư dả)", "✓ (Khuyên dùng số 1)", "✓ (Kinh tế)")
    ]

    for row_idx, row_data in enumerate(comp_data, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table_comp.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.1
            
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            
            if col_idx == 0:
                run.font.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Zebra striping
            if row_idx % 2 == 0:
                set_cell_background(cell, "F9FBFB")
            else:
                set_cell_background(cell, "FFFFFF")

    for row in table_comp.rows:
        for idx, width in enumerate(widths_comp):
            row.cells[idx].width = width

    add_para("", space_after=12) # Spacer

    # --- PHẦN G: 8 BƯỚC QUY TRÌNH LÀM VIỆC TẠI NKSV ---
    add_heading_1("PHẦN G: QUY TRÌNH 8 BƯỚC VẬN HÀNH KHÉP KÍN EPCM")
    add_para("Quy trình 8 bước chuẩn mực giúp chủ đầu tư và nhà thầu kiểm soát tiến độ và chất lượng dự án hoàn hảo:")

    add_para("NKSV ghi nhận yêu cầu sơ khởi từ khách hàng qua hotline, website hoặc văn phòng đại diện. Phân loại phân khúc công trình (Biệt thự, Dự án, Nhà xưởng) để chuyển giao cho kỹ sư chuyên trách.", bold_prefix="Bước 1: Tiếp nhận nhu cầu khách hàng: ", is_bullet=True)
    add_para("Cử kỹ sư kết cấu trực tiếp đến công trình đo đạc laser kích thước mở cửa thực tế, khảo sát địa hình vận chuyển, hướng gió bão, không gian thi công, độ phẳng của tường xây dựng.", bold_prefix="Bước 2: Khảo sát hiện trạng thực địa: ", is_bullet=True)
    add_para("Từ dữ liệu khảo sát, tiến hành tư vấn hệ nhôm trong 19 dòng nhôm, loại kính, hãng phụ kiện phù hợp nhất với điều kiện khí hậu và ngân sách mong muốn của chủ nhà.", bold_prefix="Bước 3: Tư vấn kỹ thuật và Vật liệu tối ưu: ", is_bullet=True)
    add_para("Thiết kế bản vẽ kết cấu 2D/3D, mô phỏng mặt đứng công trình. Triển khai bản vẽ Shop Drawing chi tiết uốn vòm, góc lắp đặt, mặt cắt profile nhôm và liên kết silicone.", bold_prefix="Bước 4: Thiết kế kỹ thuật & Shop Drawing (DFM): ", is_bullet=True)
    add_para("Bóc tách khối lượng vật tư thanh nhôm, kính, phụ kiện chính xác. Lập báo giá chi tiết từng hạng mục, đính kèm cam kết chất lượng, tiến độ và dự thảo hợp đồng.", bold_prefix="Bước 5: Báo giá chi tiết & Ký kết hợp đồng (Kèm cam kết NDA): ", is_bullet=True)
    add_para("Gia công thanh nhôm trên dây chuyền máy cắt CNC, máy phay khóa tự động, ép góc thủy lực lực lớn bơm keo 2 thành phần đúc ép góc khít chặt tại nhà xưởng.", bold_prefix="Bước 6: Sản xuất và Chế tạo tại nhà xưởng CNC: ", is_bullet=True)
    add_para("Vận chuyển thành phẩm bao bọc màng bảo vệ an toàn đến hiện trường. Cân chỉnh laser khung bao chèn keo bọt nở cách âm, lắp kính hộp, thợ chuyên nghiệp bắn đường silicone hoàn thiện láng mịn.", bold_prefix="Bước 7: Thi công và Lắp dựng hoàn thiện tại công trường: ", is_bullet=True)
    add_para("Nghiệm thu đóng mở nhẹ nhàng êm ái, test nước thử chống ngấm thấm. Bàn giao chìa khóa và kích hoạt chế độ bảo hành bề mặt lên đến 40 năm, chính sách bảo trì định kỳ hàng năm.", bold_prefix="Bước 8: Nghiệm thu bàn giao & Bảo trì bảo hành định kỳ: ", is_bullet=True)

    # --- PHẦN H: 6 CAM KẾT CHẤT LƯỢNG VÀNG ---
    add_heading_1("PHẦN H: 6 CAM KẾT CHẤT LƯỢNG VÀNG TẠI NKSV")
    add_para("Sự tin cậy của khách hàng là tài sản lớn nhất. NKSV cam kết bằng văn bản pháp lý 6 điều khoản vàng sau:")
    
    add_para("Nói không với thanh nhôm và linh phụ kiện nhái. Cam kết đền bù gấp đôi giá trị hợp đồng nếu phát hiện bất kỳ vật tư giả mạo nào trên công trình.", bold_prefix="✓ CAM KẾT 1: Tuyệt đối không sử dụng hàng giả, hàng nhái: ", is_bullet=True)
    add_para("Quy trình xuất nhập kho và nghiệm thu vật tư tại công trường minh bạch. Không tráo đổi nhôm dày sang nhôm mỏng, không thay phụ kiện loại 1 sang loại 2.", bold_prefix="✓ CAM KẾT 2: Không tráo đổi nguyên vật tư đã thỏa thuận: ", is_bullet=True)
    add_para("Bảng dự toán chi tiết rõ ràng từng loại mác nhôm, độ dày ly, mác kính, hãng phụ kiện. Tuyệt đối không ghi chung chung gây hiểu nhầm để trục lợi.", bold_prefix="✓ CAM KẾT 3: Báo đúng vật liệu và đúng đơn giá: ", is_bullet=True)
    add_para("Mọi liên kết cơ học, ép góc bơm keo, số lượng vít nở định vị khung bao đều tuân thủ nghiêm ngặt bản vẽ thiết kế Shop Drawing DFM đã duyệt.", bold_prefix="✓ CAM KẾT 4: Thi công lắp đặt đúng chuẩn kỹ thuật bản vẽ: ", is_bullet=True)
    add_para("NKSV chịu trách nhiệm trọn gói theo đơn giá đã cam kết. Không tự ý vẽ ra các chi phí phát sinh vô lý ngoài hợp đồng trong suốt quá trình triển khai.", bold_prefix="✓ CAM KẾT 5: Tuyệt đối không phát sinh chi phí vô lý: ", is_bullet=True)
    add_para("Thực hiện đầy đủ nghĩa vụ bảo trì định kỳ 1 năm/lần. Cử đội phản ứng nhanh xử lý sự cố cửa trong vòng 24 giờ. Bảo hành bề mặt nhôm từ 10 - 40 năm.", bold_prefix="✓ CAM KẾT 6: Thực hiện nghĩa vụ bảo hành đúng hẹn: ", is_bullet=True)

    # --- PHẦN I: THƯ VIỆN KIẾN THỨC CỬA NHÔM KÍNH (SEO DATABASE) ---
    add_heading_1("PHẦN I: THƯ VIỆN KIẾN THỨC KỸ THUẬT (PHỤC VỤ TRUYỀN THÔNG & SEO)")
    add_para("Để xây dựng vị thế chuyên gia (Authority) trên website, NKSV xây dựng cổng thư viện kiến thức đồ sộ gồm các chủ đề chuyên sâu sau:")
    
    add_para("Độ dày nhôm tối thiểu cho cửa đi chính là bao nhiêu? Hướng dẫn so sánh độ bền cơ học giữa nhôm hệ vát cạnh mỏng và nhôm hệ Châu Âu rãnh C dày dặn.", bold_prefix="• Hướng dẫn chọn cửa nhôm phù hợp với kết cấu nhà: ", is_bullet=True)
    add_para("Nhận diện tem dán Xingfa chính gốc Quảng Đông. Cách kiểm tra mã QR code, kiểm tra ánh đồng bên trong lòng thanh nhôm và phân biệt phụ kiện Kinlong thật - giả qua lò xo, độ sắc nét logo.", bold_prefix="• Bí quyết phân biệt Xingfa nhập khẩu thật và giả trên thị trường: ", is_bullet=True)
    add_para("Tìm hiểu kết cấu cánh cửa và khung bao siêu mảnh (chỉ rộng 16mm - 30mm) giúp tối đa tầm nhìn. So sánh cửa Slim lùa treo giảm chấn nội thất và cửa lùa Slim ngoại thất chịu lực gió bão.", bold_prefix="• Cửa nhôm hệ Slim là gì? Khi nào nên ứng dụng cửa Slim?: ", is_bullet=True)
    add_para("Phân tích cơ chế mạ điện di Anodised ED (Anodizing + Electrodeposition Coating) tiêu chuẩn Nhật Bản JIS H8602 Class A1. Tại sao mạ ED chống ăn mòn muối biển tốt hơn sơn tĩnh điện thông thường?", bold_prefix="• Công nghệ xử lý bề mặt nhôm Anodised ED bền màu 40 năm: ", is_bullet=True)
    add_para("Thanh nhôm cấu tạo 2 nửa độc lập liên kết bởi dải cầu polyamit cách nhiệt ở giữa giúp ngăn cản dòng nhiệt truyền qua nhôm như thế nào? Phân tích hiệu quả tiết kiệm điện điều hòa.", bold_prefix="• Nhôm cầu cách nhiệt (Thermal Break) hoạt động ra sao?: ", is_bullet=True)
    add_para("Kính hộp 2 hoặc 3 lớp hút chân không bơm khí trơ Argon cách âm cách nhiệt tốt. So sánh hiệu quả cản nhiệt giữa kính dán an toàn thường và kính phủ màng Low-E ngăn tia hồng ngoại.", bold_prefix="• Kính hộp cách âm và Kính Low-E cản nhiệt thế hệ mới: ", is_bullet=True)
    add_para("Tại sao gioăng cao su tổng hợp EPDM có tính đàn hồi cao, chịu nhiệt và chống lão hóa dưới nắng mưa tốt hơn gioăng PVC giá rẻ? Tầm quan trọng của hệ gioăng 3 lớp khép kín.", bold_prefix="• Gioăng cao su EPDM - Trái tim quyết định độ kín khít của cửa: ", is_bullet=True)
    add_para("Khái niệm hệ mặt dựng kính lớn bao quanh tòa nhà cao tầng. So sánh 3 hệ mặt dựng: Unitized (sản xuất module tại xưởng lắp ráp nhanh), Semi-Unitized (lắp bán lắp ghép), và Stick (thi công thô tại công trường).", bold_prefix="• Phân biệt hệ mặt dựng kính Curtain Wall: Unitized, Semi-Unitized, Stick: ", is_bullet=True)
    add_para("Cấu tạo hệ liên kết kính chân nhện Spider bằng thép không rỉ SUS316 cường lực lớn. Ứng dụng thi công các mái sảnh đón kính cong nghệ thuật nhịp lớn văn phòng và trung tâm thương mại.", bold_prefix="• Hệ liên kết vách kính chân nhện Spider nhịp lớn sảnh đón: ", is_bullet=True)

    # --- PHẦN J: CASE STUDY - CÂU CHUYỆN DỰ ÁN THỰC TẾ ---
    add_heading_1("PHẦN J: CASE STUDY (CÂU CHUYỆN DỰ ÁN THỰC TẾ TIÊU BIỂU)")
    add_para("Thay vì liệt kê danh sách dự án đơn điệu, NKSV truyền thông bằng câu chuyện giải quyết vấn đề (Case Study) để tạo sự đồng cảm sâu sắc với khách hàng:")

    add_heading_2("Case Study 1: Resort ven biển Mũi Né - Bình Thuận (Chống chịu muối mặn)")
    add_para("Dự án khu nghỉ dưỡng sát biển gặp hiện tượng các bộ cửa nhôm kính của đơn vị cũ thi công bị hoen ố, rỉ sét phụ kiện chỉ sau 18 tháng vận hành. Bề mặt sơn tĩnh điện nhôm bị phồng rộp do hơi muối mặn ăn mòn sâu, cửa lùa bị kẹt cát muối không thể mở nhẹ nhàng.", bold_prefix="Vấn đề của khách hàng: ", is_bullet=True)
    add_para("Khảo sát hiện trạng hướng gió biển thổi trực tiếp chứa nồng độ muối lớn. Kỹ sư NKSV tư vấn thay thế toàn bộ bằng hệ nhôm Maxpro.JP công nghệ mạ điện di Anodise ED Nhật Bản (đạt tiêu chuẩn JIS H8602). Kết hợp hệ phụ kiện rãnh C cao cấp CMECH (Mỹ) mạ màu đặc chủng chống muối biển, gioăng EPDM chịu thời tiết cực đoan.", bold_prefix="Giải pháp của NKSV: ", is_bullet=True)
    add_para("Sau 3 năm đi vào hoạt động thực tế dưới nắng gió biển khắc nghiệt, toàn bộ hệ cửa vẫn sáng bóng như mới, hoàn toàn không bị oxy hóa hay hoen rỉ. Hệ phụ kiện vận hành êm ái, trơn tru. Chủ đầu tư hoàn toàn an tâm và tiếp tục giao thầu giai đoạn 2 dự án cho NKSV.", bold_prefix="Kết quả thực tế: ", is_bullet=True)

    add_heading_2("Case Study 2: Siêu Biệt Thự Lâu Đài tại Phú Mỹ Hưng - Quận 7 (Uốn vòm nghệ thuật)")
    add_para("Công trình lâu đài tân cổ điển yêu cầu hệ cửa đi chính uốn cong hình elip nghệ thuật bản cánh dày chịu lực kính hộp Low-E nặng 120kg. Nhiều xưởng nhôm kính thông thường từ chối thực hiện vì nhôm Xingfa bản dày rất khó uốn vòm cong tròn khít mà không làm gãy rạn bề mặt.", bold_prefix="Vấn đề của khách hàng: ", is_bullet=True)
    add_para("Kỹ sư NKSV uốn vòm trên máy uốn CNC thủy lực ba trục kỹ thuật số chính xác. Chúng tôi sử dụng thanh nhôm Maxpro uốn cong bán kính elip chuẩn xác. Liên kết góc ép đúc thủy lực bơm keo đặc chủng đảm bảo độ vững chắc cho cánh cửa chịu tải trọng kính hộp nặng.", bold_prefix="Giải pháp của NKSV: ", is_bullet=True)
    add_para("Bàn giao bộ cửa vòm cong elip nghệ thuật phẳng mịn tinh xảo. Các mép nhôm khít kín không có khe hở, cửa đóng mở đầm tay, êm ru. Bộ cửa trở thành điểm nhấn kiến trúc sang trọng bậc nhất của toàn lâu đài.", bold_prefix="Kết quả thực tế: ", is_bullet=True)

    # --- PHẦN K: 40 CÂU HỎI THƯỜNG GẶP (FAQ DATABASE) ---
    add_heading_1("PHẦN K: 40 CÂU HỎI THƯỜNG GẶP (FAQ DATABASE)")
    add_para("Bộ dữ liệu câu hỏi thường gặp giúp hỗ trợ khách hàng tự động trên website và cải thiện điểm SEO cực mạnh trên Google:")

    # Detailed FAQ definition (40 items)
    faqs = [
        ("Q1: Thời gian từ khi ký hợp đồng đến khi lắp đặt hoàn thiện cửa nhôm kính là bao lâu?",
         "A1: Thông thường tiến độ dao động từ 25 - 40 ngày. Trong đó: 15-20 ngày dùng để nhập vật tư thanh nhôm, kính và sản xuất CNC tại xưởng; 10-20 ngày thi công lắp ráp trực tiếp tại công trường tùy thuộc quy mô dự án."),
        
        ("Q2: Nhôm Kính Sao Vàng có nhận thi công các công trình ở tỉnh xa không?",
         "A2: Có. NKSV nhận thi công trọn gói trên toàn quốc. Chúng tôi đã hoàn thành nhiều dự án tại các tỉnh miền Tây, miền Trung và Tây Nguyên với quy trình vận chuyển bọc lót an toàn tuyệt đối."),
        
        ("Q3: Tại sao cửa nhôm kính của NKSV có giá cao hơn các xưởng nhôm kính nhỏ lẻ?",
         "A3: NKSV không cạnh tranh bằng nhôm mỏng giá rẻ. Giá của chúng tôi phản ánh chất lượng: 100% nhôm và phụ kiện chính hãng nhập khẩu, công nghệ ép góc thủy lực bơm keo 2 thành phần, gioăng EPDM 3 lớp, keo silicone Dow Corning chính hãng và chế độ bảo hành màu sắc bề mặt lên tới 40 năm kèm bảo trì hàng năm."),
        
        ("Q4: Cửa nhôm hệ rãnh C Châu Âu là gì và có điểm gì vượt trội?",
         "A4: Là thanh nhôm thiết kế rãnh phụ kiện dạng chữ C tiêu chuẩn quốc tế. Giúp lắp đặt phụ kiện (như CMECH, Roto) trực tiếp vào rãnh mà không cần khoan khoét lỗ phá vỡ kết cấu nhôm, tăng độ chịu lực, độ kín khít và tính thẩm mỹ cao hơn hệ nhôm truyền thống."),
        
        ("Q5: NKSV uốn được những dạng cửa cong như thế nào?",
         "A5: Chúng tôi uốn được vòm cong tròn, vòm elip, vòm bán nguyệt đối với hầu hết các hệ nhôm dày (Xingfa, Maxpro, Civro) trên máy uốn CNC thủy lực ba trục chuyên dụng."),
        
        ("Q6: Cửa lùa Slim của NKSV có dùng được cho cửa đi chính ngoài trời không?",
         "A6: Cửa Slim thông thường chỉ dùng cho nội thất. Tuy nhiên, NKSV có hệ nhôm Slim ngoại thất chuyên dụng với thanh cánh chịu tải lớn, gioăng kín nước và phụ kiện khóa đa điểm chịu được sức cản gió bão lớn ngoài trời."),
        
        ("Q7: NKSV có bảo hành nước mưa thấm qua khe cửa không?",
         "A7: Có. Chúng tôi cam kết bảo hành chống ngấm thấm nước mưa qua khe góc cửa và đường keo silicone hoàn thiện lên tới 3 năm bằng văn bản hợp đồng rõ ràng."),
        
        ("Q8: Tôi muốn xem trực tiếp các mẫu nhôm và phụ kiện thì xem ở đâu?",
         "A8: Sếp và quý khách hàng có thể ghé thăm trực tiếp Showroom Nhôm Kính Sao Vàng tại địa chỉ nhà xưởng xưởng hoặc liên hệ kỹ sư của chúng tôi mang các mẫu nhôm xi mạ Anodized, catalogue mẫu kính hộp thực tế đến tận công trình tư vấn."),
        
        ("Q9: NKSV có cung cấp bản vẽ CAD mặt cắt nhôm và shop drawing cho kiến trúc sư không?",
         "A9: Có. Cổng thư viện của chúng tôi sẵn sàng chia sẻ đầy đủ file CAD mặt cắt nhôm của 19 hệ nhôm phục vụ kiến trúc sư lên phương án thiết kế bản vẽ kỹ thuật."),
        
        ("Q10: NKSV có sẵn sàng ký kết thỏa thuận bảo mật NDA cho bản vẽ dự án B2B không?",
         "A10: Có. Chúng tôi cam kết bảo mật tuyệt đối bản vẽ và thông số kỹ thuật dự án của đối tác, ký kết NDA pháp lý trước khi tiếp nhận file thiết kế chi tiết."),
        
        ("Q11: Cửa nhôm xi mạ điện di Anodise ED khác gì cửa sơn tĩnh điện thông thường?",
         "A11: Sơn tĩnh điện bám màu bằng bột sơn phun nhiệt, dễ bong tróc phai màu sau 5-10 năm. Xi mạ điện di Anodise ED là công nghệ nhúng bể mạ điện hóa tạo lớp oxit nhôm siêu cứng tích hợp màng bảo vệ ED mịn màng chống ăn mòn muối biển, bảo hành màu sắc 40 năm."),
        
        ("Q12: Kính hộp Low-E cản nhiệt hoạt động như thế nào?",
         "A12: Kính hộp Low-E gồm 2 lớp kính cách nhau bằng khoảng không bơm khí trơ Argon. Lớp kính phía trong được phủ hợp chất cản nhiệt Low-E giúp phản xạ năng lượng mặt trời ngược lại, cản tia cực tím UV và cản nhiệt lượng truyền vào nhà."),
        
        ("Q13: NKSV sử dụng loại keo silicone nào để thi công lắp đặt ngoài trời?",
         "A13: Chúng tôi chỉ sử dụng các dòng keo silicone đặc chủng chịu thời tiết ngoài trời hàng đầu thế giới như Dow Corning (Dowsil) hoặc Xylotex của tập đoàn uy tín để đảm bảo độ bám dính cao và không nứt nẻ lão hóa dưới trời nắng gắt."),
        
        ("Q14: Bản lề cửa nhôm kính của NKSV chịu lực tối đa bao nhiêu kg?",
         "A14: Tùy thuộc dòng phụ kiện. Bản lề 3D Xingfa thông thường chịu lực 80-100kg/cánh. Bản lề chịu lực cao cấp CMECH (Mỹ) hoặc Sobinco (Bỉ) dùng cho cửa kính hộp bản lớn có thể chịu lực từ 120kg đến 200kg/cánh."),
        
        ("Q15: Tôi muốn lắp đặt cửa đi mở xếp trượt góc vuông không có cột ở giữa được không?",
         "A15: Được. NKSV thi công hệ cửa xếp trượt góc vuông 90 độ (Corner folding door). Khi mở ra hoàn toàn, hai hệ cửa xếp lùa về hai bên để lại khoảng mở góc vuông thoáng đãng không vướng cột giữa."),
        
        ("Q16: Cửa nhôm cầu cách nhiệt có thực sự giúp tiết kiệm tiền điện điều hòa không?",
         "A16: Có. Nghiên cứu thực tế cho thấy hệ cửa nhôm cầu cách nhiệt kết hợp kính hộp cách nhiệt giúp giảm đến 35% lượng nhiệt thất thoát qua cửa, giảm tải cho hệ thống điều hòa không khí và tiết kiệm chi phí điện năng hàng tháng."),
        
        ("Q17: NKSV có nhận sửa chữa, thay thế cửa nhôm kính cũ của đơn vị khác làm bị hỏng không?",
         "A17: Có. Chúng tôi có gói dịch vụ cải tạo và nâng cấp hệ thống cửa cũ bị xệ kẹt, thay thế kính hộp cách âm tốt hơn hoặc gia cố kết cấu nhôm kính cho tòa nhà."),
        
        ("Q18: Nhôm Maxpro.JP do nước nào sản xuất?",
         "A18: Nhôm Maxpro.JP được sản xuất bởi công nghệ mạ điện di Anodise ED Nhật Bản hiện đại bậc nhất, sản xuất theo tiêu chuẩn công nghiệp Nhật Bản JIS H8602 và được nhập khẩu phân phối uy tín tại Việt Nam."),
        
        ("Q19: NKSV có cung cấp cửa nhôm chống cháy đạt chuẩn kiểm định PCCC không?",
         "A19: Có. Chúng tôi cung cấp các hệ cửa nhôm kính chống cháy chuyên dụng (độ dày kính chống cháy và nhôm tăng cường bông khoáng cách nhiệt) có đầy đủ chứng nhận kiểm định chống cháy EI30, EI60, EI90 của Cục PCCC."),
        
        ("Q20: Vách kính Spider chân nhện khác gì vách kính Curtain Wall thông thường?",
         "A20: Vách kính Spider liên kết các tấm kính bằng các phụ kiện chân nhện inox SUS316 cường lực lớn bắt trực tiếp vào dầm thép đỡ, không có thanh đố nhôm chia ô, tạo tầm nhìn kính liền mạch tối đa sảnh đón tòa nhà."),
        
        ("Q21: Tại sao hệ gioăng EPDM lại quan trọng đối với cửa nhôm?",
         "A21: Gioăng EPDM là gioăng cao su tổng hợp chịu đàn hồi cực tốt, không bị mục nát hay nứt gãy dưới nắng mưa. Gioăng EPDM tạo độ kín khít tuyệt đối giữa cánh cửa và khung bao khi đóng cửa, giúp cửa cách âm, chống nước tràn vào."),
        
        ("Q22: Cửa sổ mở hất có ưu điểm gì so với cửa sổ mở quay?",
         "A22: Cửa sổ mở hất đẩy ra ngoài với góc hất 30 độ. Khi trời mưa nhỏ vẫn có thể mở hé cửa hất để lấy gió thoáng mát mà nước mưa không bị hắt vào trong nhà. Đồng thời chống gió bão giật cánh đập vỡ kính."),
        
        ("Q23: NKSV có cung cấp khóa cửa thông minh điều khiển qua app điện thoại không?",
         "A23: Có. Chúng tôi tích hợp khóa điện tử thông minh vân tay, thẻ từ, mã số của các thương hiệu hàng đầu (Kaadas, Yale, Bosch) tương thích đồng bộ với hệ cửa nhôm rãnh C cao cấp."),
        
        ("Q24: Mức dung sai kích thước cho phép khi lắp đặt cửa nhôm kính của NKSV là bao nhiêu?",
         "A24: Quy trình QA kiểm soát dung sai kích thước khung bao dưới 1mm. Khe hở giữa khung nhôm và tường chèn keo bọt nở rộng từ 5mm-8mm để đảm bảo giãn nở nhiệt học an toàn."),
        
        ("Q25: Tôi ở ven biển, lắp nhôm nào thì bền nhất không rỉ sét bề mặt?",
         "A25: NKSV khuyên dùng các dòng nhôm xi mạ điện di Anodized ED Nhật Bản (như Maxpro.JP, Hondalex) hoặc dòng cao cấp như Civro (Đức) vì công nghệ xử lý bề mặt này triệt tiêu hoàn toàn sự ăn mòn hóa học của muối biển."),
        
        ("Q26: Kính dán an toàn 2 lớp khác kính cường lực như thế nào?",
         "A26: Kính cường lực khi vỡ sẽ vụn thành những hạt nhỏ vô hại. Kính dán an toàn gồm 2 lớp kính liên kết bởi màng phim PVB ở giữa, khi vỡ các mảnh kính vẫn bám chặt vào màng phim PVB không bị rơi xuống, đảm bảo an toàn chống đột nhập."),
        
        ("Q27: NKSV có nhận uốn vòm kính cong dán keo an toàn không?",
         "A27: Có. Chúng tôi nhận gia công uốn vòm nhôm kết hợp kính cong (kính gia nhiệt cong hoặc kính dán cong an toàn) đúng biên dạng uốn vòm của khung cửa nhôm."),
        
        ("Q28: Chi phí vận chuyển cửa đi các tỉnh có bị tính thêm vào báo giá không?",
         "A28: Bảng báo giá chi tiết của NKSV luôn ghi rõ hạng mục vận chuyển bốc xếp trọn gói đến công trình để đảm bảo không phát sinh chi phí phát sinh bất ngờ cho chủ đầu tư."),
        
        ("Q29: Phụ kiện CMECH của nước nào sản xuất?",
         "A29: CMECH là thương hiệu phụ kiện cửa nhôm cao cấp hàng đầu thế giới có trụ sở tại Mỹ. Phụ kiện CMECH nổi tiếng với thiết kế thẩm mỹ sang trọng, mạ màu bền bỉ chống muối biển và đạt kiểm định đóng mở trên 100,000 lần."),
        
        ("Q30: NKSV có nhận sản xuất OEM/ODM cửa nhôm kính cho các đơn vị khác không?",
         "A30: Có. Với quy mô nhà xưởng máy móc CNC hiện đại, chúng tôi nhận gia công cơ khí nhôm, uốn vòm CNC và lắp ghép cửa thành phẩm theo yêu cầu đặt hàng của các nhà thầu thi công nhôm kính khác."),
        
        ("Q31: Tôi lắp đặt kính hộp thì có bị đọng hơi nước bên trong lòng kính không?",
         "A31: Không. Kính hộp chất lượng cao của NKSV được đi keo butyl liên kết xung quanh thanh nhôm đệm có chứa các hạt hút ẩm cự tốt. Đảm bảo hút sạch hơi ẩm bên trong hộp kính và không bao giờ bị đọng sương lòng kính."),
        
        ("Q32: Cửa mở xếp trượt có lắp được nhiều cánh lớn không?",
         "A32: Có. Hệ cửa đi xếp trượt có thể thiết kế tối đa từ 3 cánh đến 10 cánh xếp gọn về một hoặc hai bên tường, mở rộng không gian đi lại tối đa cho sảnh tiệc hay lối ra sân vườn."),
        
        ("Q33: Nhôm cỏ hay nhôm hệ vát cạnh có ưu điểm gì?",
         "A33: Nhôm cỏ (nhôm Việt Pháp đời đầu) có giá thành rất rẻ, thi công nhanh. Nhôm hệ vát cạnh (PMA, Yangli) thiết kế mặt cắt vát xiên nghệ thuật giúp tăng độ cứng chịu lực tốt hơn nhôm cỏ mà vẫn rất tiết kiệm phôi nhôm."),
        
        ("Q34: Keo silicone bọt nở (Foam keo) dùng để làm gì trong lắp đặt cửa?",
         "A34: Keo bọt nở được bơm vào khoảng hở giữa khung bao nhôm và tường gạch bê tông. Keo tự nở đầy khe hở giúp tiêu âm cản ồn, chống nước mưa ngấm vào tường và tạo độ liên kết đàn hồi chịu rung chấn tốt."),
        
        ("Q35: Bản lề sàn dùng cho loại cửa nào?",
         "A35: Bản lề sàn (bản lề thủy lực âm nền) chuyên dùng cho cửa đi mở xoay 180 độ hai chiều bằng kính cường lực tấm lớn (cửa thủy lực sảnh văn phòng, cửa đi chính nhà phố mặt tiền)."),
        
        ("Q36: NKSV có chính sách ưu đãi chiết khấu gì cho các nhà thầu xây dựng dài hạn không?",
         "A36: Có. Chúng tôi xây dựng chính sách đối tác lâu dài với mức chiết khấu thầu ưu đãi và hỗ trợ tư vấn biện pháp kỹ thuật, vẽ shop drawing miễn phí cho các Tổng thầu, KTS đối tác dài hạn."),
        
        ("Q37: Mặt dựng kính hệ Unitized khác Stick như thế nào?",
         "A37: Mặt dựng Stick được thi công thô trực tiếp tại hiện trường từng thanh nhôm đố dọc đố ngang và lắp kính. Mặt dựng Unitized được sản xuất gia công thành các tấm module hoàn chỉnh (nhôm + kính + gioăng keo) tại nhà xưởng rồi cẩu nâng lắp ghép cực nhanh ngoài hiện trường."),
        
        ("Q38: Tôi muốn làm lan can kính cường lực ngoài trời không dùng trụ có an toàn không?",
         "A38: An toàn. NKSV sử dụng hệ pát inox SUS316 định vị chân vách kính hoặc âm sàn, kết hợp kính dán cường lực an toàn dày từ 12mm - 19mm đảm bảo khả năng chịu lực tác động đập mạnh mà vẫn cực sang trọng hiện đại."),
        
        ("Q39: NKSV có hỗ trợ kiểm định chất lượng thợ lắp đặt định kỳ không?",
         "A39: Có. Đội ngũ thợ kỹ nghệ thi công của NKSV được huấn luyện đào tạo an toàn lao động thường xuyên và kiểm tra chất lượng tay nghề đi silicone, ép góc định kỳ hàng quý."),
        
        ("Q40: Làm sao để liên hệ bảo hành sửa chữa nhanh nhất khi cửa gặp sự cố?",
         "A40: Sếp và quý khách hàng chỉ cần liên hệ trực tiếp hotline bảo hành hiển thị trên website. Đội phản ứng nhanh NKSV tại khu vực gần nhất sẽ có mặt hỗ trợ xử lý sự cố trong vòng 24 giờ kể từ khi tiếp nhận thông tin.")
    ]

    for q, a in faqs:
        add_para(q, bold_prefix="", space_after=2)
        p_ans = add_para(a, bold_prefix="", space_after=8)
        p_ans.runs[0].font.italic = True
        p_ans.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    # --- PHẦN L: KẾ HOẠCH BỐ CỤC NỘI DUNG TRÊN WEBSITE MỚI ---
    add_heading_1("PHẦN L: KẾ HOẠCH BỐ CỤC NỘI DUNG TRÊN WEBSITE MỚI")
    add_para("Bố cục được tối ưu hóa để người dùng trải nghiệm mượt mà, tăng tính chuyển đổi B2B/B2C:")
    
    add_para("Hiển thị nổi bật thông điệp 'Giải pháp cửa & mặt dựng nhôm kính kiến trúc tổng thể', slide hình ảnh công trình thực tế, tóm tắt 10 lý do lựa chọn NKSV, bảng so sánh hệ nhôm trực quan, và form nhận tư vấn nhanh.", bold_prefix="Trang chủ Nhôm Kính: ", is_bullet=True)
    add_para("Diễn giải chi tiết về quy trình EPCM khép kín kèm theo video thực tế sản xuất tại nhà xưởng CNC và thợ bắn silicone ngoài công trường.", bold_prefix="Trang Lĩnh vực hoạt động (linh-vuc-hoat-dong.html): ", is_bullet=True)
    add_para("Nơi số hóa toàn bộ dữ liệu 19 dòng nhôm. Mỗi hệ nhôm sẽ có 1 trang con giới thiệu chi tiết về thông số kỹ thuật, catalogue PDF chính hãng, và nút tải File AutoCAD mặt cắt phục vụ KTS vẽ thiết kế.", bold_prefix="Trang Thư viện hệ nhôm (thu-vien-he-nhom.html): ", is_bullet=True)
    add_para("Thư viện ảnh thực tế phân chia theo loại hình công trình: Siêu biệt thự, lâu đài uốn vòm, resort ven biển chống mặn, mặt dựng kính văn phòng, nhà xưởng công nghiệp.", bold_prefix="Trang Dự án tiêu biểu (du-an.html): ", is_bullet=True)
    add_para("Tích hợp công cụ tính toán chi phí theo diện tích và phân khúc nhôm, kết hợp nút upload bản vẽ thiết kế nhận báo giá nhanh, đính kèm dòng cam kết bảo mật NDA hiển thị rõ bên cạnh nút gửi.", bold_prefix="Form Báo giá thông minh (Get a Quote Form): ", is_bullet=True)

    # --- PHẦN M: TỔNG KẾT CHIẾN LƯỢC PHÁT TRIỂN ---
    add_heading_1("PHẦN M: TỔNG KẾT CHIẾN LƯỢC PHÁT TRIỂN (NKSV GIỎI NHẤT Ở ĐÂU? TẠI SAO CHỌN NKSV?)")
    
    add_heading_2("1. Nhôm Kính Sao Vàng giỏi nhất ở đâu? (Core Expertise)")
    add_para("Năng lực sản xuất uốn vòm elip, vòm tròn bản lớn và hệ cửa đi cong nghệ thuật mà không bị biến dạng profile nhôm hay rạn nứt bề mặt xi mạ sơn tĩnh điện.", bold_prefix="• Kỹ nghệ uốn vòm nhôm kính nghệ thuật độ khó cao: ", is_bullet=True)
    add_para("Thiết kế và gia công các dòng cửa lùa Slim ngoại thất chịu tải lớn cánh siêu mỏng, cửa đi xếp trượt góc vuông 90 độ Corner Folding không cột góc nâng tầm không gian sống.", bold_prefix="• Vận hành hệ nhôm xi mạ Anodized ED & Cửa Slim thế hệ mới: ", is_bullet=True)
    add_para("Kế thừa nhà xưởng kết cấu cơ khí giúp NKSV dễ dàng liên kết đồng bộ các hệ mái kính sảnh đón nhịp lớn có khung thép chịu lực đỡ phức tạp hoặc mặt dựng vách kính Spider nhịp lớn.", bold_prefix="• Tích hợp đồng bộ khung chịu lực thép và vách kính lớn: ", is_bullet=True)

    add_heading_2("2. Tại sao khách hàng phải chọn Nhôm Kính Sao Vàng?")
    add_para("Với 10 lý do lựa chọn NKSV, 6 cam kết vàng bằng văn bản, hệ sinh thái 19 dòng nhôm đa dạng tư vấn khách quan theo ngân sách, và dịch vụ bảo trì định kỳ trọn đời, NKSV là đối tác duy nhất đáp ứng trọn vẹn mọi yêu cầu khắt khe của KTS và Chủ đầu tư.", bold_prefix="Tóm tắt giá trị cốt lõi: ")
    add_para("NKSV cam kết mang lại sự an tâm tuyệt đối, tính thẩm mỹ tinh xảo từng góc cạnh cửa và hiệu năng kỹ thuật cách âm chống thấm tối ưu cho mọi công trình Việt.")

    # Save document
    output_filename = "BaoCao_ChienLuoc_TaiDinhVi_NKSV_Premium.docx"
    doc.save(output_filename)
    print(f"Document successfully created and saved as: {output_filename}")

if __name__ == "__main__":
    main()
