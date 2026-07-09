# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng KOSO (Cửa nhựa ABS KOS - K Kumovina Hàn Quốc):
1. Catalogue_HeCuaNhom_KOSO.xlsx
2. ThuVien_HeCuaNhom_KOSO.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

KOSO_SYSTEMS = [
    ('1', 'Nhựa ABS Hàn Quốc', 'KOS Flat', 'Cửa nhựa ABS KOS cánh phẳng', 
     'Cửa thông phòng, cửa phòng ngủ cánh phẳng phủ lớp vân gỗ Deco Sheet nghệ thuật.', 
     '35mm - 38mm (Cánh)', 'Khóa tay gạt KOS / khóa tròn trơn & bản lề lá inox', 'Gioăng giảm chấn chèn khung bao / Silicon Apollo', 
     'Kháng nước 100%, không mối mọt cong vênh, rất nhẹ nhàng êm ái đóng mở.'),
     
    ('2', 'Nhựa ABS Hàn Quốc', 'KOS Panel', 'Cửa nhựa ABS KOS dập pano', 
     'Cửa phòng ngủ phong cách tân cổ điển dập pano họa tiết chìm nổi nghệ thuật giống cửa gỗ tự nhiên.', 
     '35mm - 38mm (Cánh)', 'Khóa tay gạt cao cấp & bản lề inox SUS304', 'Gioăng giảm chấn chèn sẵn trên khung', 
     'Vẻ đẹp truyền thống ấm cúng, cách âm tốt nhờ lớp lõi giấy tổ ong Honeycomb bên trong.'),
     
    ('3', 'Nhựa ABS Hàn Quốc', 'KOS Glass', 'Cửa nhựa ABS KOS kết hợp ô kính', 
     'Cửa toilet, cửa phòng tắm tích hợp thêm ô kính mờ hoặc kính cường lực lấy sáng nhẹ.', 
     '35mm - 38mm (Cánh)', 'Kính cường lực 5mm mạ mờ & khóa tròn KOS', 'Gioăng đệm nẹp kính / Silicon nội thất', 
     'Chống nước ẩm ướt tuyệt đối, lấy sáng tự nhiên nhẹ, dễ lau chùi vệ sinh.'),
     
    ('4', 'Nhựa ABS Hàn Quốc', 'KOS Line', 'Cửa nhựa ABS KOS chạy chỉ nhôm', 
     'Cửa thông phòng căn hộ chạy các đường chỉ nhôm trang trí bạc/vàng hiện đại trẻ trung.', 
     '35mm - 38mm (Cánh)', 'Khóa phân thể hiện đại / khóa điện tử vân tay', 'Gioăng giảm chấn / Silicon nội thất', 
     'Tạo điểm nhấn thẩm mỹ hiện đại đột phá. Giá thành cao hơn dòng phẳng thường một chút.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_KOSO.xlsx
# ─────────────────────────────────────────────
def create_koso_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue KOSO"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC CỬA NHỰA ABS KOS HÀN QUỐC (K KUMOVINA)"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="E65100") # Màu cam gỗ đặc trưng KOS
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="FFF3E0")
    for r_idx, row_vals in enumerate(KOSO_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_KOSO.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel KOSO saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel KOSO to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_KOSO.docx
# ─────────────────────────────────────────────
def create_koso_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xE6,0x51,0x00))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHỰA THÔNG PHÒNG ABS KOS HÀN QUỐC", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp hệ cửa nhựa ABS thương hiệu KOS (K Kumovina) Hàn Quốc\nTích hợp phân tích ưu nhược điểm R&D và ma trận so sánh phân khúc nội thất", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU CỬA NHỰA ABS KOS (K KUMOVINA)", bold=True, size=14, color=RGBColor(0xE6,0x51,0x00), space_before=12)
    add_p(
        "Cửa nhựa ABS thương hiệu KOS do Công ty K Kumovina (Hàn Quốc) đầu tư nhà máy sản xuất trực tiếp tại Việt Nam. "
        "Với nguyên liệu hạt nhựa ABS sinh học an toàn nhập khẩu 100% từ Hàn Quốc, "
        "thanh profile được ép chân không tấm Deco Sheet vân gỗ nghệ thuật chống trầy xước, "
        "kết hợp lớp lõi giấy tổ ong Honeycomb tăng cường cách âm và lớp gỗ PVC chạy quanh tăng cứng gia cố khóa. "
        "Cửa nhựa ABS KOS vô cùng nổi tiếng trong phân khúc cửa nội thất thông phòng và cửa vệ sinh tại Việt Nam "
        "nhờ đặc tính kháng nước tuyệt đối 100%, không bị cong vênh co ngót do thời tiết và hoàn toàn chống mối mọt."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D 4 DÒNG CỬA NHỰA KOS", bold=True, size=14, color=RGBColor(0xE6,0x51,0x00), space_before=12)
    
    # 1. KOS Flat
    add_p("1. KOS Flat (Cửa nhựa ABS cánh phẳng giả vân gỗ)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày cánh: 35mm - 38mm phủ màng Deco Sheet giả vân gỗ Hàn Quốc.")
    add_bullet("Ưu điểm R&D: Thiết kế phẳng mượt tối giản hiện đại phù hợp căn hộ chung cư cao cấp, kháng nước tuyệt đối, trọng lượng nhẹ giúp đóng mở êm ái không sệ cánh.")
    add_bullet("Nhược điểm R&D: Chỉ dùng làm cửa thông phòng trong nhà, tuyệt đối không lắp ngoài trời chịu ánh nắng trực tiếp hoặc mưa bão.")
    
    # 2. KOS Panel
    add_p("2. KOS Panel (Cửa nhựa ABS dập pano cổ điển)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày cánh: 35mm - 38mm dập nén pano chìm nổi.")
    add_bullet("Ưu điểm R&D: Vân gỗ pano nổi khối nghệ thuật giống hệt cửa gỗ tự nhiên ấm cúng cổ điển, cách âm tốt nhờ lõi Honeycomb dày dặn.")
    add_bullet("Nhược điểm R&D: Các góc cạnh pano dập sâu yêu cầu công nghệ hút chân không nhiệt cao để màng vân gỗ bám dính chắc chắn.")

    # 3. KOS Glass
    add_p("3. KOS Glass (Cửa nhựa ABS tích hợp ô kính lấy sáng)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày cánh: 35mm - 38mm có khung ô kính mờ.")
    add_bullet("Ưu điểm R&D: Thích hợp tuyệt đối cho cửa toilet vệ sinh hoặc phòng làm việc cần lấy sáng nhẹ và quan sát bên trong, chống ngập ẩm tuyệt vời.")
    add_bullet("Nhược điểm R&D: Giảm hiệu suất cách âm do các ô kính ghép nối.")

    # 4. KOS Line
    add_p("4. KOS Line (Cửa nhựa ABS chạy trang trí chỉ nhôm)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày cánh: 35mm - 38mm có nẹp chỉ nhôm chìm.")
    add_bullet("Ưu điểm R&D: Chỉ nhôm bạc hoặc vàng chạy ngang/dọc tạo điểm nhấn hiện đại sang trọng, kết hợp tốt với khóa vân tay điện tử.")
    add_bullet("Nhược điểm R&D: Chi phí hoàn thiện nẹp nhôm trang trí cao hơn dòng phẳng thường.")

    # ── PHẦN SO SÁNH MA TRẬN ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN PHÂN KHÚC CỬA NHỰA ABS KOS VỚI ĐỐI THỦ CẠNH TRANH", bold=True, size=14, color=RGBColor(0xE6,0x51,0x00), space_before=12)
    add_p(
        "Bảng ma trận so sánh chéo đặc tính kỹ thuật cửa nhựa ABS KOS Hàn Quốc "
        "với các dòng cửa gỗ công nghiệp và cửa nhựa giả gỗ khác trên thị trường nội thất Việt Nam:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "Cửa nhựa ABS KOS"
    hdr_c[2].text = "Cửa nhựa Composite"
    hdr_c[3].text = "Cửa gỗ MDF Melamine"
    hdr_c[4].text = "Cửa nhôm kính Xingfa"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ("Khả năng chịu nước", "Kháng nước 100% tuyệt đối", "Kháng nước 100% tuyệt đối", "Dễ nở mục khi ẩm ướt", "Kháng nước 100% tuyệt đối"),
        ("Khả năng chống cháy", "Chậm cháy lan (Nhựa ABS)", "Chậm cháy lan (Nhựa PVC)", "Dễ bắt lửa cháy bùng", "Không bắt lửa (Nhôm đặc)"),
        ("Khả năng chống mối mọt", "Chống mối mọt 100%", "Chống mối mọt 100%", "Dễ bị mối mọt đục phá", "Chống mối mọt 100%"),
        ("Độ nặng cánh cửa", "Cực kỳ nhẹ (giảm sệ bản lề)", "Đầm nặng chắc chắn", "Nặng trung bình", "Nặng chắc chắn"),
        ("Phân khúc ứng dụng", "Nội thất thông phòng, WC", "Nội thất thông phòng, WC", "Nội thất khô phòng ngủ", "Ngoại thất + Nội thất"),
        ("Phân khúc giá bán", "$$ (Bình dân kinh tế)", "$$$ (Khá - Tầm trung)", "$$$ (Khá - Tầm trung)", "$$$ (Khá - Tầm trung)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_KOSO.docx")
    try:
        doc.save(path)
        print(f"  >> Docx KOSO saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx KOSO to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone KOSO docx & xlsx in BASE directory...")
    print("=" * 60)
    create_koso_xlsx()
    create_koso_docx()
    print("=" * 60)
