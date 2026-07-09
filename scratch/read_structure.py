# -*- coding: utf-8 -*-
"""
Script đọc nhanh cấu trúc file Civro mẫu
"""
import os, sys
try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass
from openpyxl import load_workbook
from docx import Document

base_dir = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

# 1. Đọc Excel
excel_path = os.path.join(base_dir, "Catalogue_HeCuaNhom_Civro.xlsx")
if os.path.exists(excel_path):
    wb = load_workbook(excel_path)
    ws = wb.active
    print("EXCEL ROWS 1 to 10:")
    for r in range(1, 11):
        row_vals = [cell.value for cell in ws[r]]
        print(f"Row {r}: {row_vals}")
else:
    print("Excel Civro not found")

# 2. Đọc Word
docx_path = os.path.join(base_dir, "ThuVien_HeCuaNhom_Civro.docx")
if os.path.exists(docx_path):
    doc = Document(docx_path)
    print("\nWORD PARAGRAPHS (First 15):")
    for i, p in enumerate(doc.paragraphs[:15]):
        if p.text.strip():
            print(f"[{i}]: {p.text}")
    print("\nWORD TABLES:")
    print(f"Total tables: {len(doc.tables)}")
    if len(doc.tables) > 0:
        tbl = doc.tables[0]
        print("Table 0 Headers:")
        hdr = [cell.text for cell in tbl.rows[0].cells]
        print(hdr)
else:
    print("Docx Civro not found")
