# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_maxpro_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM MAXPRO JP NHẬT BẢN")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp 10 hệ cửa và cấu kiện nhôm Maxpro.JP cao cấp\nĐược xử lý bề mặt bằng công nghệ Anodise ED độc quyền bảo hành 25 năm")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM MAXPRO JP", level=1)
    add_paragraph_with_spacing(doc, 
        "Maxpro.JP là dòng nhôm cao cấp được hợp tác và sản xuất theo tiêu chuẩn công nghệ Nhật Bản "
        "(liên kết giữa công ty Trung Chính và tập đoàn LIXIL danh tiếng). Đặc điểm vượt trội nhất làm nên thương hiệu "
        "Maxpro JP là ứng dụng công nghệ xử lý bề mặt Electrodeposition Anodise ED (phủ cứng bề mặt qua nhiều lớp). "
        "Thanh nhôm Maxpro đạt chứng nhận quốc tế QUALANOD và tiêu chuẩn công nghiệp Nhật Bản JIS H8602 - Class A1, "
        "giúp bảo vệ sản phẩm khỏi ăn mòn axit, muối biển, chống trầy xước và bám bụi hoàn hảo với độ bền màu trên 40 năm. "
        "Sản phẩm được bảo hành bề mặt chính hãng lên đến 25 năm, cực kỳ phù hợp với các công trình ven biển, resort, biệt thự và căn hộ hạng sang."
    )
    
    add_paragraph_with_spacing(doc, "Maxpro JP cung cấp 10 hệ nhôm và giải pháp cấu kiện tiêu biểu được phân chia rõ ràng:")
    
    # Table of 10 systems overview in Word
    table = doc.add_table(rows=11, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.0), Inches(2.0), Inches(1.2), Inches(2.8)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "Hệ 55", "Cửa mở quay Classic/Luxury", "1.4mm - 2.0mm", "Cửa sổ/Cửa đi mở quay tiêu chuẩn. Dòng tiết kiệm và thông dụng nhất."),
        ("2", "Hệ 65", "Cửa mở quay cao cấp", "1.6mm - 2.0mm", "Cửa đi mở quay chính rãnh C Châu Âu, thiết kế vuông cạnh hiện đại."),
        ("3", "Hệ 70", "Cửa sổ/đi vát cạnh", "1.4mm - 1.8mm", "Cửa mở quay vát cạnh nghiêng thoát nước mưa và tăng cách âm."),
        ("4", "Hệ 83", "Cửa mở quay tân cổ điển", "1.6mm - 2.0mm", "Khung dày 83mm, bản cánh bo viền cổ kính tinh tế, kính hộp lên đến 32mm."),
        ("5", "Hệ 93", "Cửa lùa trượt tiết kiệm", "1.4mm - 1.6mm", "Cửa đi và cửa sổ mở trượt lùa lồng cánh kinh tế."),
        ("6", "Hệ 115", "Cửa trượt lùa nâng siêu chịu lực", "2.0mm - 3.0mm", "Lift & Slide chịu áp lực bão gió 2500 Pa ngoài ban công, chống bão ven biển."),
        ("7", "Hệ 80", "Cửa xếp gấp trượt xếp lùa", "1.8mm - 2.2mm", "Cửa đi xếp trượt nhiều cánh xếp gọn mở rộng tối đa lối ra sân vườn."),
        ("8", "Hệ MD65", "Vách mặt dựng kính cường lực", "2.0mm - 3.0mm", "Mặt dựng đố đứng chịu lực cho các tòa nhà showroom, văn phòng."),
        ("9", "Hệ Thủy Lực", "Cửa thủy lực bản lớn phào nổi", "2.0mm", "Cửa đi bản lề sàn đại sảnh phào nổi thẩm mỹ phong cách tân cổ điển."),
        ("10", "Hệ Lan Can", "Lan can nhôm không chôn", "2.5mm - 3.0mm", "Hệ lan can liên kết cơ khí chân đế chịu tải lớn, dễ bảo trì.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM MAXPRO JP", level=1)
    
    maxpro_details = [
        ("1. Hệ 55 - Cửa mở quay tiêu chuẩn (Classic / Luxury)",
         "• Độ dày nhôm: 1.4mm cho cửa sổ, 2.0mm cho cửa đi.\n"
         "• Bản khuôn bao: 55mm.\n"
         "• Công năng: Ứng dụng rộng rãi làm cửa sổ mở hất ngoài trời, cửa sổ mở quay, cửa đi 1 cánh phòng ngủ, cửa toilet.\n"
         "• Ưu điểm: Phổ biến, dễ gia công sản xuất, giá thành kinh tế nhất trong catalogue Maxpro."),
         
        ("2. Hệ 65 - Cửa mở quay rãnh C Châu Âu",
         "• Độ dày nhôm: 1.6mm - 2.0mm.\n"
         "• Thiết kế rãnh C: Cho phép đồng bộ hoàn toàn với các thương hiệu phụ kiện cao cấp hàng đầu thế giới như CMECH, Roto.\n"
         "• Công năng: Làm cửa đi ban công, cửa ra vào chính vuông cạnh vuông vức khỏe khoắn phong cách hiện đại."),
         
        ("3. Hệ 70 - Cửa vát cạnh nghiêng chuyên biệt",
         "• Độ dày nhôm: 1.4mm - 1.8mm.\n"
         "• Điểm nhấn thiết kế: Khung cánh có độ dốc vát nghiêng 45 độ hướng ra ngoài giúp thoát nước mưa cực nhanh, hạn chế đọng bụi bẩn mép gờ kính.\n"
         "• Công năng: Cửa sổ mở hất, cửa đi căn hộ phong cách tối giản."),
         
        ("4. Hệ 83 - Cửa đi/sổ Tân cổ điển Premium",
         "• Độ dày nhôm: 1.6mm - 2.0mm. Bản dày khuôn bao lên tới 83mm cứng vững.\n"
         "• Thiết kế phào bo tròn: Tạo các đường gờ bo tròn nghệ thuật mô phỏng chỉ gỗ tân cổ điển độc đáo, rất phù hợp biệt thự kiểu Pháp.\n"
         "• Kính tương thích: Khoang nhôm rộng cho phép lắp kính hộp Low-E cản nhiệt dày tới 32mm để cách âm cách nhiệt tốt nhất."),
         
        ("5. Hệ 93 - Cửa trượt lùa kinh tế",
         "• Độ dày nhôm: 1.4mm - 1.6mm.\n"
         "• Công năng: Làm cửa sổ trượt ngang 2 ray, cửa đi lùa phòng khách ban công nhà dân dụng vừa và nhỏ.\n"
         "• Điểm nhấn: Trượt bánh xe êm nhẹ trên ray inox tròn đúc."),
         
        ("6. Hệ 115 - Cửa trượt lùa nâng chống bão (Lift & Slide)",
         "• Độ dày nhôm: 2.0mm - 3.0mm siêu dày chịu lực gió bão biển.\n"
         "• Đặc tính cơ khí: Cơ cấu Lift & Slide gạt nâng cánh giúp cánh kính trượt nhẹ như bông, khi hạ cánh sẽ nén ép gioăng EPDM nằm chèn phẳng sàn cách âm tuyệt đối.\n"
         "• Khả năng chịu bão: Chịu áp lực bão gió cấp 14 (tương đương 2500 Pa), chuyên dùng cho cửa đi lùa biệt thự ven biển hoặc căn hộ penthouse siêu cao tầng."),
         
        ("7. Hệ 80 - Cửa đi xếp trượt cánh lớn",
         "• Độ dày nhôm: 1.8mm - 2.2mm.\n"
         "• Công năng: Xếp gấp gọn gàng 3 cánh, 4 cánh, 5 cánh, 6 cánh dồn về một góc mở rộng không gian 90%.\n"
         "• Điểm nhấn: Hệ ray treo trên chịu lực cực khỏe chống rung lắc giật cánh cửa khi vận hành."),
         
        ("8. Hệ Mặt Dựng MD65 (Curtain Wall)",
         "• Độ dày nhôm: 2.0mm - 3.0mm tùy thuộc vào tính toán chịu lực gió của mặt dựng.\n"
         "• Bản nhôm: 65mm đố đứng dọc lớn chịu mô-men xoắn cao.\n"
         "• Công năng: Vách dựng kính khổ lớn cho showroom xe hơi, mặt tiền khối đế trung tâm thương mại hoặc các tòa nhà văn phòng hiện đại."),
         
        ("9. Hệ Cửa Thủy Lực bản cánh lớn đại sảnh",
         "• Độ dày nhôm: 2.0mm cứng cáp.\n"
         "• Bản cánh nhôm: Siêu rộng 120mm - 150mm phào chỉ nổi nghệ thuật tân cổ điển kết hợp mạ màu Anodize Champagne/Bronze lộng lẫy.\n"
         "• Ứng dụng: Cánh cửa đi chính đại sảnh biệt thự lắp bản lề sàn thủy lực nặng chịu lực đóng mở liên tục."),
         
        ("10. Hệ Lan Can Nhôm Maxpro công nghệ không chôn",
         "• Độ dày nhôm: 2.5mm - 3.0mm chịu uốn va đập mạnh.\n"
         "• Thiết kế thông minh: Liên kết chân đế bằng hệ bu-lông hóa chất đặc chủng âm bề mặt sàn, không cần chôn sâu chân sắt truyền thống gỉ sét gây nứt vỡ bê tông.\n"
         "• Ưu điểm: Đảm bảo thẩm mỹ liền mạch mặt sàn gỗ/đá, thoát nước tốt và cực kỳ thuận tiện cho việc thay thế bảo trì kính hoặc nhôm sau này.")
    ]
    
    for title, desc in maxpro_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Đặc trưng bề mặt ED ---
    add_heading_with_spacing(doc, "CÔNG NGHỆ BỀ MẶT ANODISE ED ĐỘC QUYỀN", level=1)
    add_paragraph_with_spacing(doc, 
        "Khác biệt lớn nhất của nhôm Maxpro JP so với các hệ nhôm sơn tĩnh điện thông thường là quy trình xử lý "
        "bề mặt Electrodeposition Anodise ED qua 4 bước công nghệ Nhật Bản:\n\n"
        "1. Xử lý bề mặt sạch cơ học và hóa học.\n"
        "2. Anodize hóa tạo lớp màng Oxit nhôm tự nhiên mỏng chống ăn mòn.\n"
        "3. Phủ màu điện di (Anode màu) tạo màu sắc sâu thẳm Champagne, Nâu thu.\n"
        "4. Lớp phủ Electrodeposition (ED) ngoài cùng bảo vệ bề mặt siêu cứng chống bám bẩn vân tay và chống trầy xước cấp độ cao.\n\n"
        "Bề mặt ED của Maxpro JP đã vượt qua bài kiểm tra phun sương muối mặn trong phòng thí nghiệm trên 2000 giờ liên tục, "
        "đảm bảo không rỉ sét và bong tróc, là dòng nhôm bền bỉ nhất cho môi trường ven biển Việt Nam."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ MAXPRO JP]", before=200, italic=True)
    
    # Ensure directory exists and save
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Maxpro.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Maxpro Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_maxpro_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Maxpro JP"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM CÔNG NGHỆ ANODISE ED MAXPRO JP"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        "Phụ Kiện Đồng Bộ", 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Dòng Tiết Kiệm (Economy)", "Hệ 55", "Cửa mở quay Classic/Luxury", 
         "Cửa sổ và cửa đi mở quay 1 cánh, 2 cánh, cửa sổ mở hất định vị góc, vách kính tĩnh.", 
         "1.4mm - 2.0mm", "Phụ kiện Maxpro đồng bộ, Draho", "Gioăng EPDM kép, keo Dow Corning", "Sơn Anodise ED bảo hành 25 năm. Thích hợp nhà dân dụng."),
         
        ("2", "Dòng Cao Cấp (Premium)", "Hệ 65", "Cửa mở quay cao cấp Châu Âu", 
         "Cửa đi mở quay chính, cửa sổ mở quay hất rãnh C Châu Âu liên kết ke góc ép thủy lực lực nén lớn.", 
         "1.6mm - 2.0mm", "Phụ kiện Cmech, Roto rãnh C", "Gioăng EPDM 3 tầng kín khít, silicone", "Thiết kế vuông cạnh hiện đại thanh lịch."),
         
        ("3", "Dòng Chuyên Biệt", "Hệ 70", "Cửa sổ/đi vát cạnh nghiêng", 
         "Cửa đi mở quay, cửa sổ mở quay lật có cạnh vát dốc nghiêng thoát nước nhanh.", 
         "1.4mm - 1.8mm", "Tay nắm & Khóa Sigico, Cmech", "Gioăng EPDM, keo silicone trung tính", "Hạn chế đọng bụi bẩn mép kính, tăng cách âm."),
         
        ("4", "Dòng Tân Cổ Điển", "Hệ 83", "Cửa mở quay phong cách tân cổ điển", 
         "Cửa đi/sổ mở quay bản cánh hộp to bo viền phào tròn thẩm mỹ nghệ thuật kiểu Pháp.", 
         "1.6mm - 2.0mm", "Tay nắm gạt & bản lề CMECH, Roto", "Gioăng EPDM cao cấp, keo bọt nở PU", "Khung bao rộng 83mm, lắp được kính hộp tới 32mm."),
         
        ("5", "Dòng Tiết Kiệm (Economy)", "Hệ 93", "Cửa trượt lùa phổ thông", 
         "Cửa sổ trượt lùa 2 cánh, 4 cánh trượt ngang ray inox đúc tròn.", 
         "1.4mm - 1.6mm", "Bánh xe & khóa sập Maxpro, Draho", "Gioăng lông cản bụi chống ồn", "Tiết kiệm diện tích mở cánh, vận hành êm."),
         
        ("6", "Dòng Siêu Chịu Lực (Luxury)", "Hệ 115", "Cửa lùa trượt nâng chống bão", 
         "Cửa đi mở trượt nâng Lift & Slide siêu chịu tải trọng gió penthouse cao tầng, biệt thự biển.", 
         "2.0mm - 3.0mm", "Phụ kiện nâng hạ CMECH, Roto cao cấp", "Gioăng nén EPDM chặt khi khóa cánh", "Chống bão bão gió áp lực cực cao tới 2500 Pa."),
         
        ("7", "Dòng Cao Cấp (Premium)", "Hệ 80", "Cửa xếp trượt folding", 
         "Cửa đi xếp trượt lùa gấp xếp gọn nhiều cánh dồn góc mở rộng lối sân vườn đại sảnh.", 
         "1.8mm - 2.2mm", "Bánh xe chịu lực treo trên & khóa Cmech", "Gioăng EPDM đệm đàn hồi cao", "Ray treo trên chịu tải lớn cản rung xệ cánh."),
         
        ("8", "Dòng Dự Án (B2B)", "Hệ MD65", "Vách mặt dựng kính chịu lực", 
         "Mặt dựng đố đứng chìm hoặc nổi liên kết keo cấu trúc chịu mô-men xoắn lớn.", 
         "2.0mm - 3.0mm", "Ke liên kết chịu lực chuyên dụng", "Keo Dow Corning 791 / 995 kết cấu", "Lắp đặt mặt tiền kính showroom văn phòng lớn."),
         
        ("9", "Dòng Chuyên Dụng (Luxury)", "Hệ Cửa Thủy Lực", "Cửa thủy lực bản lớn phào nổi đại sảnh", 
         "Cửa đi chính đại sảnh lắp bản lề sàn thủy lực đóng mở 2 chiều bản cánh phào nổi nghệ thuật.", 
         "2.0mm", "Bản lề sàn Adler/Hafele, tay nắm kéo lớn", "Keo kết cấu chuyên dụng, gioăng kép", "Bản nhôm siêu rộng 120-150mm phủ ED Champagne."),
         
        ("10", "Dòng Cấu Kiện Phụ", "Hệ Lan Can", "Lan can nhôm không chôn kính", 
         "Lan can liên kết chân đế bu-lông hóa chất trên mặt sàn gỗ/đá, dễ thay thế kính.", 
         "2.5mm - 3.0mm", "Phụ kiện lan can cơ khí Maxpro", "Gioăng cao su đệm kẹp giữ kính", "Không chôn chân kim loại gây gỉ sét nứt bê tông.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Tiết Kiệm" in seg:
            seg_fill = fill_economy
        elif "Siêu Chịu Lực" in seg or "Luxury" in seg or "Thủy Lực" in seg:
            seg_fill = fill_premium
        else:
            seg_fill = fill_standard
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Maxpro.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Maxpro Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_maxpro_docx()
    generate_maxpro_xlsx()
