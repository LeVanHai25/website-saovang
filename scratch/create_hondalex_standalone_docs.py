# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng Hondalex (Long Vân Group / Honda Metal Industries Japan):
1. Catalogue_HeCuaNhom_Hondalex.xlsx
2. ThuVien_HeCuaNhom_Hondalex.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

HONDALEX_SYSTEMS = [
    ('1', 'Anodized Cao cấp', 'Hondalex LV 60', 'Cửa đi mở quay LV 60', 
     'Cửa đi mở quay cao cấp bản dày 60mm cực kỳ chắc khỏe, bề mặt xử lý mạ điện di Anodize ED cao cấp chống ăn mòn muối biển.', 
     '1.8mm', 'Phụ kiện Hopo / CMECH / Long Vân đồng bộ rãnh C', 'Gioăng EPDM Nhật Bản / Silicon Dow Corning', 
     'Độ bền màu sơn mạ anodize ED lên đến trên 40 năm, đạt chuẩn công nghiệp JIS H4100 Nhật Bản.'),
     
    ('2', 'Anodized Trung cấp', 'Hondalex LV 34', 'Cửa đi mở quay LV 34', 
     'Cửa đi thông phòng, cửa đi căn hộ bản nhỏ gọn nhẹ, dễ lắp ráp sản xuất.', 
     '1.2mm - 1.3mm', 'Phụ kiện đồng bộ Long Vân rãnh 22', 'Gioăng cao su chèn sập / Silicon Apollo', 
     'Thẩm mỹ gọn nhẹ tinh tế phong cách tối giản Nhật Bản, giá thành hợp lý.'),
     
    ('3', 'Anodized Trung cấp', 'Hondalex LV 70', 'Cửa đi mở quay LV 70', 
     'Hệ thống cửa đi mở quay kinh tế bản 70mm, trọng lượng thanh nhẹ giúp dễ vận chuyển lắp dựng.', 
     '1.2mm - 1.3mm', 'Phụ kiện chốt gạt Long Vân / Kinlong', 'Gioăng EPDM / Silicon Apollo A500', 
     'Trọng lượng tối ưu giúp đóng mở rất nhẹ nhàng, bền bỉ mạ điện di Anodize.'),
     
    ('4', 'Anodized Cao cấp', 'Hondalex LV 60 Window', 'Cửa sổ bật hất LV 60', 
     'Cửa sổ mở hất hoặc mở quay bản 60mm chắc chắn, cách âm và chống bụi xâm nhập tốt.', 
     '1.5mm', 'Bản lề chữ A & tay gạt đa điểm rãnh C', 'Gioăng EPDM đa lớp / Silicon Dow', 
     'Độ kín khít chống gió giật cực cao, phù hợp lắp đặt tòa nhà cao tầng.'),
     
    ('5', 'Anodized Phổ thông', 'Hondalex LV 38 Window', 'Cửa sổ bật hất 38', 
     'Cửa sổ mở hất gọn nhẹ bản 38mm cho các vị trí cửa phụ nhỏ hoặc ô lấy sáng.', 
     '0.9mm - 1.2mm', 'Khớp chốt bật & phụ kiện nhẹ Long Vân', 'Gioăng cao su chèn sập / Apollo', 
     'Chi phí siêu tiết kiệm cho công trình nhà dân phổ thông.'),
     
    ('6', 'Anodized Trung cấp', 'Hondalex LV 76', 'Cửa lùa trượt LV 76', 
     'Cửa đi trượt lùa hoặc cửa sổ trượt lùa bản 76mm 2 ray đóng mở êm ái, tiết kiệm không gian.', 
     '1.2mm', 'Bánh xe trượt Long Vân & khóa bán nguyệt', 'Gioăng chèn nỉ mịn chống rung nước', 
     'Trượt lướt êm dịu phong cách Nhật Bản, chống ngập nước dột mưa tốt.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_Hondalex.xlsx
# ─────────────────────────────────────────────
def create_hondalex_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue Hondalex"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM MẠ ANODIZE ED HONDALEX JAPAN (LONG VÂN)"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="0D47A1") # Xanh dương đậm Nhật Bản
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="E3F2FD")
    for r_idx, row_vals in enumerate(HONDALEX_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_Hondalex.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel Hondalex saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel Hondalex to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_Hondalex.docx
# ─────────────────────────────────────────────
def create_hondalex_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x0D,0x47,0x41))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM MẠ ANODIZE ED HONDALEX JAPAN", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp hệ cửa nhôm liên doanh đùn ép Honda Metal Industries Nhật Bản\nTích hợp phân tích ưu nhược điểm R&D và ma trận so sánh phân khúc thị trường", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM HONDALEX (LONG VÂN)", bold=True, size=14, color=RGBColor(0x0D,0x47,0x41), space_before=12)
    add_p(
        "Nhôm Hondalex là sản phẩm hợp tác chiến lược giữa Công ty TNHH Long Vân NTV Việt Nam và Tập đoàn nhôm danh tiếng Honda Metal Industries (Nhật Bản). "
        "Với nhà máy LOVAL hiện đại tại Bình Dương và nhà máy Honda Metal Industries khánh thành tại KCN VSIP, "
        "Hondalex đi tiên phong cung cấp công nghệ xử lý bề mặt nhôm đùn ép mạ điện di Anodize ED hoàn toàn tự động đạt công suất 1000 tấn/tháng. "
        "Sản phẩm tuân thủ nghiêm ngặt tiêu chuẩn quốc tế ISO 9001, JIS H4100 Nhật Bản, JIS H8601/H8602 cho xử lý anodize, và tiêu chuẩn ASTM B211M của Mỹ. "
        "Đặc trưng lớn nhất của nhôm Hondalex là trọng lượng cực nhẹ nhưng kết cấu bền chắc dẻo dai, nước mạ bóng mịn chống ăn mòn muối biển mặn tuyệt đối "
        "đã được tin dùng trong các siêu công trình như cao ốc Long Hải, hầm Thủ Thiêm, và tòa nhà Charm Residence."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D 6 HỆ NHÔM HONDALEX", bold=True, size=14, color=RGBColor(0x0D,0x47,0x41), space_before=12)
    
    # 1. LV 60
    add_p("1. Hondalex LV 60 (Cửa đi mở quay Anodize ED cao cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.8mm bản dày.")
    add_bullet("Ưu điểm R&D: Độ cứng vững chịu lực gió bão hoàn hảo, công nghệ xi mạ điện di ED bóng mịn tự nhiên vô cùng sang trọng bền màu trên 40 năm, đạt chuẩn JIS Nhật Bản cao cấp nhất.")
    add_bullet("Nhược điểm R&D: Giá nhôm thanh xi mạ ED và phụ kiện đi kèm khá cao, kén thợ gia công.")
    
    # 2. LV 34
    add_p("2. Hondalex LV 34 (Cửa đi mở quay tối giản)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.3mm.")
    add_bullet("Ưu điểm R&D: Thiết kế tối giản tinh tế phong cách Nhật Bản, lắp ráp đơn giản, chi phí nguyên liệu tối ưu cho nhà phố dân dụng tầm trung.")
    add_bullet("Nhược điểm R&D: Bản nhôm mỏng không lắp được kính hộp quá dày cách âm.")

    # 3. LV 70
    add_p("3. Hondalex LV 70 (Cửa đi mở quay kinh tế)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.3mm.")
    add_bullet("Ưu điểm R&D: Trọng lượng thanh nhôm tối ưu giúp vận hành cửa đóng mở cực kỳ nhẹ tay êm ái, xi mạ Anodize ED chống xước xát rất tốt.")
    add_bullet("Nhược điểm R&D: Độ chịu gió giật ngoại thất kém hơn bản LV 60 dày.")

    # 4. LV 60 Window
    add_p("4. Hondalex LV 60 Window (Cửa sổ bật hất)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.5mm.")
    add_bullet("Ưu điểm R&D: Khung cánh thiết kế chuyên biệt chống dột nước mưa chèn gioăng EPDM đúc Nhật Bản kín khít tuyệt đối, chống gió bão giật cao tầng tốt.")
    add_bullet("Nhược điểm R&D: Trọng lượng nặng hơn các dòng cửa sổ nhôm cỏ.")

    # 5. LV 38 Window
    add_p("5. Hondalex LV 38 Window (Cửa sổ hất phổ thông)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 0.9mm - 1.2mm.")
    add_bullet("Ưu điểm R&D: Bản nhôm siêu gọn mỏng nhẹ, giá thành rẻ nhất phân khúc, thi công gia công dập ke góc cực nhanh.")
    add_bullet("Nhược điểm R&D: Độ bền chịu tải kém, không lắp được ở các vị trí chịu bão lớn.")

    # 6. LV 76
    add_p("6. Hondalex LV 76 (Cửa đi & sổ trượt lùa)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm.")
    add_bullet("Ưu điểm R&D: Lùa trượt 2 ray tiết kiệm không gian, hệ thống gioăng chèn nỉ cao cấp giúp lướt êm dịu không tiếng ồn chuẩn Nhật.")
    add_bullet("Nhược điểm R&D: Kín khít chắn nước bão ở mức khá, cần vệ sinh ray trượt định kỳ.")

    # ── PHẦN SO SÁNH MA TRẬN ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN PHÂN KHÚC HỆ HONDALEX LV 60 VỚI ĐỐI THỦ CẠNH TRANH", bold=True, size=14, color=RGBColor(0x0D,0x47,0x41), space_before=12)
    add_p(
        "Bảng ma trận so sánh chéo đặc tính kỹ thuật hệ nhôm Hondalex LV 60 xi mạ điện di ED Nhật Bản "
        "với các dòng nhôm cùng phân khúc xi mạ cao cấp trên thị trường:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "Hondalex LV 60"
    hdr_c[2].text = "Maxpro R55"
    hdr_c[3].text = "Civro AW55"
    hdr_c[4].text = "Xingfa Class A"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ("Nguồn gốc xuất xứ", "Liên doanh Nhật - Việt", "Nhật Bản - đùn ép TQ/VN", "Đức/Trung Quốc (Luxury)", "Quảng Đông (Trung Quốc)"),
        ("Công nghệ bề mặt", "Mạ điện di Anodise ED", "Mạ điện di Anodise ED", "Sơn phủ Matt / Anodize", "Mạ Anodized cải tiến"),
        ("Độ dày cánh cửa (mm)", "1.8mm", "1.4mm - 2.0mm", "1.6mm - 1.8mm", "2.0mm"),
        ("Tiêu chuẩn kỹ thuật", "JIS H4100 (Nhật Bản)", "JIS H4100 (Nhật Bản)", "DIN EN (Châu Âu)", "Qualicoat / TCVN"),
        ("Độ bền màu bảo hành", "Bảo hành trên 40 năm", "Bảo hành 25 năm", "Bảo hành 20 năm", "Bảo hành 20-25 năm"),
        ("Phân khúc giá bán", "$$$$ (Cao cấp)", "$$$ (Cận cao cấp)", "$$$$$ (Siêu cao cấp)", "$$$$ (Cao cấp)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_Hondalex.docx")
    try:
        doc.save(path)
        print(f"  >> Docx Hondalex saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx Hondalex to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone Hondalex docx & xlsx in BASE directory...")
    print("=" * 60)
    create_hondalex_xlsx()
    create_hondalex_docx()
    print("=" * 60)
