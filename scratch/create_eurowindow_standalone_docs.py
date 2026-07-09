# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng Eurowindow (Cửa nhựa uPVC & Nhôm cao cấp):
1. Catalogue_HeCuaNhom_Eurowindow.xlsx
2. ThuVien_HeCuaNhom_Eurowindow.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

EUROWINDOW_SYSTEMS = [
    ('1', 'Nhựa lõi thép uPVC', 'EW uPVC Asia60', 'Cửa đi & sổ nhựa uPVC Koemmerling', 
     'Hệ thống cửa nhựa lõi thép uPVC cao cấp chống ồn, chống bụi vượt trội nhờ profile Koemmerling Đức.', 
     '2.2mm - 2.5mm (vỏ uPVC) + lõi thép 1.5mm', 'Phụ kiện Roto / G-U / Eurowindow đồng bộ', 'Gioăng TPE đa lớp ép đúc trực tiếp', 
     'Khả năng cách âm đứng đầu các loại cửa, kết cấu nhựa rỗng bền bỉ không bị ăn mòn rỉ sét.'),
     
    ('2', 'Nhôm Cao cấp', 'EW EA55', 'Cửa nhôm mở quay EA55', 
     'Cửa đi và cửa sổ mở quay nhôm Eurowindow thế hệ mới rãnh C Châu Âu thiết kế bo tròn tinh tế.', 
     '1.4mm - 2.0mm', 'Phụ kiện Roto / CMECH / Eurowindow rãnh C Châu Âu', 'Gioăng EPDM đúc 3 lớp / Silicon Dow Corning', 
     'Thẩm mỹ cao cấp, sơn tĩnh điện Eurowindow bảo hành lên tới 10-20 năm chống tia cực tím.'),
     
    ('3', 'Nhôm Cao cấp', 'EW EA70', 'Cửa nhôm lùa trượt EA70', 
     'Cửa đi lùa trượt hoặc cửa sổ trượt lùa bản mỏng thế hệ mới giúp tối ưu hóa không gian ban công phòng ngủ.', 
     '1.6mm - 1.8mm', 'Bánh xe chịu lực và khóa đa điểm Eurowindow', 'Gioăng chèn nỉ mịn kép chống rít chấn nước', 
     'Trượt lướt êm ái, thích hợp cho chung cư trung và cao cấp.'),
     
    ('4', 'Nhôm Cao cấp', 'EW EA90', 'Vách mặt dựng Stick EA90', 
     'Hệ thống vách kính mặt dựng Stick chịu lực gió lớn bao che bên ngoài tòa nhà lớn và showroom.', 
     '2.2mm - 3.0mm', 'Bulong móng nở kết cấu chịu lực chuyên dụng', 'Silicon kết cấu chống thấm Dow Corning 791/895', 
     'Khả năng chịu bão siêu việt, tạo nên diện mạo kính liền khối hiện đại cho công trình.'),
     
    ('5', 'Cầu cách nhiệt', 'EW Thermal', 'Cửa nhôm cầu cách nhiệt Eurowindow', 
     'Hệ thống cửa nhôm tích hợp dải cầu cách nhiệt Polyamide Technoform Đức hạn chế truyền nhiệt tối đa.', 
     '2.0mm + Cầu cách nhiệt Polyamide', 'Phụ kiện cao cấp Roto / CMECH rãnh C', 'Gioăng trung tâm đa khoang EPDM cao cấp', 
     'Cách âm cách nhiệt hoàn hảo nhất, tiết kiệm 30% điện năng điều hòa nhiệt độ. Giá thành rất cao.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_Eurowindow.xlsx
# ─────────────────────────────────────────────
def create_eurowindow_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue Eurowindow"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHỰA uPVC & CỬA NHÔM CAO CẤP EUROWINDOW"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="0D47A1") # Xanh dương đậm Eurowindow
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="E3F2FD")
    for r_idx, row_vals in enumerate(EUROWINDOW_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 24, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_Eurowindow.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel Eurowindow saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel Eurowindow to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_Eurowindow.docx
# ─────────────────────────────────────────────
def create_eurowindow_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x0D,0x47,0xA1))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ THỐNG CỬA uPVC & CỬA NHÔM EUROWINDOW", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp hệ cửa nhựa lõi thép uPVC và nhôm định hình Eurowindow\nTích hợp phân tích ưu nhược điểm R&D và ma trận so sánh chéo phân khúc cao cấp", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU EUROWINDOW", bold=True, size=14, color=RGBColor(0x0D,0x47,0xA1), space_before=12)
    add_p(
        "Eurowindow là đơn vị đi tiên phong khai mở phân khúc cửa uPVC cao cấp tại Việt Nam từ năm 2002. "
        "Với nhà máy quy mô công nghiệp lớn hiện đại, hãng cung cấp trọn gói các giải pháp cửa nhựa lõi thép Koemmerling "
        "và cửa nhôm cao cấp Eurowindow (tự thiết kế định hình rãnh C chuẩn Châu Âu) đạt chứng nhận chất lượng quốc tế. "
        "Đặc trưng lớn nhất của Eurowindow là chất lượng đồng bộ cực cao từ profile nhựa/nhôm, phụ kiện Roto/G-U/Cmech "
        "cho đến gioăng đúc sẵn, được tin dùng trong hàng triệu biệt thự, lâu đài cao cấp và cao ốc văn phòng hạng sang trên toàn lãnh thổ Việt Nam."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D 5 HỆ CỬA EUROWINDOW", bold=True, size=14, color=RGBColor(0x0D,0x47,0xA1), space_before=12)
    
    # 1. EW uPVC Asia60
    add_p("1. EW uPVC Asia60 (Cửa nhựa uPVC Koemmerling lõi thép)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày profile: 2.2mm - 2.5mm vỏ uPVC + lõi thép gia cường 1.5mm chịu va đập.")
    add_bullet("Ưu điểm R&D: Khả năng chống truyền nhiệt và cách âm tuyệt đỉnh hàng đầu (nhờ cấu trúc rỗng nhựa uPVC triệt tiêu sóng âm), gioăng đúc TPE mềm êm kín khít khói bụi hoàn hảo.")
    add_bullet("Nhược điểm R&D: Bản khung nhựa to, thô dày, kém phóng khoáng về view nhìn so với hệ nhôm Slim hiện đại. Nhựa trắng có thể bị ố vàng nhẹ sau 20-30 năm mưa nắng.")
    
    # 2. EW EA55
    add_p("2. EW EA55 (Cửa nhôm mở quay EA55 cao cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 2.0mm tiêu chuẩn rãnh C.")
    add_bullet("Ưu điểm R&D: Thiết kế bo góc thẩm mỹ cực kỳ sang trọng quý phái, rãnh C Châu Âu lắp CMECH/Roto chính hãng, công nghệ sơn bột tĩnh điện độc quyền bảo hành 10-20 năm.")
    add_bullet("Nhược điểm R&D: Giá nhôm thanh và thành phẩm thi công trọn gói cao hơn hẳn nhôm Quảng Đông phổ thông.")

    # 3. EW EA70
    add_p("3. EW EA70 (Cửa nhôm lùa trượt EA70)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.6mm - 1.8mm.")
    add_bullet("Ưu điểm R&D: Hệ thống lùa trượt đóng mở êm nhẹ thích hợp lắp đặt làm cửa ra ban công chung cư và biệt thự nghỉ dưỡng giúp tối ưu không gian sinh hoạt.")
    add_bullet("Nhược điểm R&D: Độ kín khít chống bão kém hơn hệ quay EA55.")

    # 4. EW EA90
    add_p("4. EW EA90 (Vách mặt dựng Stick EA90)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.2mm - 3.0mm.")
    add_bullet("Ưu điểm R&D: Kết cấu xương chịu lực siêu dày chống gió bão giật ngoại thất tòa nhà cực tốt, sơn tĩnh điện PVDF bền màu muối mặn.")
    add_bullet("Nhược điểm R&D: Đòi hỏi thiết kế bản vẽ Shopdrawing chuẩn xác và thi công giàn giáo phức tạp ngoài công trình.")

    # 5. EW Thermal
    add_p("5. EW Thermal (Cửa nhôm cầu cách nhiệt Eurowindow)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm tích hợp dải cầu cách nhiệt Polyamide Technoform Đức.")
    add_bullet("Ưu điểm R&D: Cách âm và ngăn truyền nhiệt tuyệt hảo nhất, ngăn thất thoát hơi lạnh điều hòa giúp tiết kiệm 35% điện năng tiêu thụ.")
    add_bullet("Nhược điểm R&D: Giá thành lắp đặt thuộc phân khúc siêu đắt đỏ Luxury biệt thự nghỉ dưỡng.")

    # ── PHẦN SO SÁNH MA TRẬN ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN PHÂN KHÚC HỆ NHÔM EUROWINDOW EA55 VỚI ĐỐI THỦ CẠNH TRANH", bold=True, size=14, color=RGBColor(0x0D,0x47,0xA1), space_before=12)
    add_p(
        "Bảng ma trận so sánh chéo đặc tính kỹ thuật hệ nhôm mở quay Eurowindow EA55 "
        "với các dòng nhôm cùng phân khúc cao cấp trên thị trường xây dựng Việt Nam:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "Eurowindow EA55"
    hdr_c[2].text = "Civro AW55"
    hdr_c[3].text = "Maxpro R55"
    hdr_c[4].text = "Xingfa Quảng Đông 55"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ("Nguồn gốc xuất xứ", "Eurowindow Việt Nam", "Civro Đức / đùn ép TQ", "Nhật Bản / đùn ép VN", "Quảng Đông (Trung Quốc)"),
        ("Công nghệ sơn phủ", "Sơn tĩnh điện Eurowindow", "Sơn Matt cao cấp", "Mạ điện di Anodise ED", "Sơn tĩnh điện tiêu chuẩn"),
        ("Tiêu chuẩn rãnh phụ kiện", "Chuẩn rãnh C Châu Âu", "Chuẩn rãnh C Châu Âu", "Chuẩn rãnh C Châu Âu", "Rãnh 22mm truyền thống"),
        ("Thời gian bảo hành sơn", "Bảo hành 10-20 năm", "Bảo hành 20 năm", "Bảo hành 25 năm", "Bảo hành 5 năm"),
        ("Thương hiệu định vị", "Quốc gia (Việt Nam)", "Thượng lưu (Châu Âu)", "Cận cao cấp (Nhật Bản)", "Quốc dân thông dụng"),
        ("Phân khúc giá bán", "$$$$ (Cao cấp)", "$$$$$ (Siêu cao cấp)", "$$$ (Cận cao cấp)", "$$ (Tầm trung)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_Eurowindow.docx")
    try:
        doc.save(path)
        print(f"  >> Docx Eurowindow saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx Eurowindow to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone Eurowindow docx & xlsx in BASE directory...")
    print("=" * 60)
    create_eurowindow_xlsx()
    create_eurowindow_docx()
    print("=" * 60)
