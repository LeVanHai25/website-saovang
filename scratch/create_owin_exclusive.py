# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_owin_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM OWIN ĐỒNG BỘ CAO CẤP")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật chuyên đề cửa trượt quay không ray dưới và hệ cửa thủy lực phào nổi\nGiải pháp không gian thông minh độc đáo tiêu chuẩn CHLB Đức chất lượng vượt trội")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM OWIN", level=1)
    add_paragraph_with_spacing(doc, 
        "Owin là một trong những thương hiệu nhôm định hình cao cấp hàng đầu tại Việt Nam, sản xuất theo các "
        "tiêu chuẩn công nghệ cơ khí Đức. Owin nổi tiếng trên thị trường nhờ các giải pháp kiến trúc cửa 'độc đáo, "
        "chuyên biệt' và đi đầu xu hướng. Sản phẩm tiêu biểu nhất của Owin là hệ cửa mở trượt quay (Slide & Turn) "
        "không cần ray dưới sàn nhà, giúp tối ưu 100% diện tích không gian và loại bỏ hoàn toàn nguy cơ vấp ngã. "
        "Bên cạnh đó, Owin còn sở hữu các hệ nhôm thủy lực bản lớn tích hợp phào chỉ nổi nghệ thuật tân cổ điển "
        "vô cùng sang trọng. Hợp kim nhôm Owin mác 6063-T5 dày dặn kết hợp công nghệ xi mạ Anodized bọc ED bền bỉ, "
        "đảm bảo tính năng chống chịu thời tiết biển và an ninh tuyệt hảo cho các công trình cao cấp."
    )
    
    add_paragraph_with_spacing(doc, "Owin cung cấp các giải pháp hệ nhôm đặc chủng và thông dụng bao gồm:")
    
    # Table of 7 systems overview in Word
    table = doc.add_table(rows=8, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.6)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "Owin Trượt Quay", "Cửa trượt quay thông minh", "1.6mm - 3.5mm", "Cửa vừa lùa trượt vừa quay cánh xếp gọn góc 100%. Không ray dưới. Độc đáo nhất."),
        ("2", "Owin Thủy Lực", "Cửa thủy lực bản lớn phào nổi", "2.0mm", "Cửa đi bản lề sàn đại sảnh cánh chỉ nổi phào nghệ thuật tân cổ điển 120-180mm."),
        ("3", "Owin Hệ 55 Quay", "Cửa đi/sổ mở quay chuẩn Đức", "1.4mm - 2.0mm", "Cửa đi mở quay rãnh C Châu Âu chốt đa điểm kín khít (Mã C3332D, C3303D)."),
        ("4", "Owin Hệ 55 Lùa", "Cửa lùa trượt sổ/đi 55", "1.2mm - 1.4mm", "Cửa lùa trượt lồng cánh mỏng tiết kiệm diện tích (Mã C3208)."),
        ("5", "Owin Trượt Nâng", "Cửa trượt nâng Lift & Slide", "1.5mm - 2.0mm", "Cửa lùa trượt nâng hạ nén gioăng EPDM cách âm chống bão ban công (Mã C3329LS)."),
        ("6", "Owin Xếp Trượt", "Cửa đi xếp trượt gấp Owin", "1.8mm - 2.0mm", "Cửa xếp gấp trượt xếp gọn cánh chịu tải lớn cho biệt thự showroom."),
        ("7", "Owin Vách Ngăn", "Vách ngăn kính văn phòng", "1.4mm - 2.0mm", "Vách ngăn phòng thạch cao kính cường lực định hình khung nhôm Owin vững chãi.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM OWIN", level=1)
    
    owin_details = [
        ("1. Owin Cửa trượt quay thông minh (Slide & Turn) - Sản phẩm biểu tượng",
         "• Độ dày nhôm: 1.6mm - 2.0mm, các điểm ray treo và đố chịu lực đặc biệt dày từ 2.5mm - 3.5mm.\n"
         "• Cơ chế đóng mở: Kết hợp đồng thời cơ cấu trượt lùa và mở quay. Cửa mở 100% diện tích thông thủy dồn cánh về góc tường.\n"
         "• Đặc trưng: Thiết kế không có ray dưới sàn nhà, loại bỏ vấp ngã nguy hiểm, dễ vệ sinh, quét dọn nhà cửa và đẩy xe.\n"
         "• Phụ kiện: Hệ thống bánh xe treo, khóa tay nắm và chốt liên kết đặc chủng đồng bộ Owin chịu lực tải trọng nặng treo trên cực tốt.\n"
         "• Ứng dụng: Cửa chính mặt tiền nhà phố, cửa ngăn phòng khách - sảnh vườn."),
         
        ("2. Owin Cửa thủy lực cánh lớn đại sảnh phào nổi",
         "• Độ dày nhôm: 2.0mm bản cánh siêu rộng (120mm - 180mm x 60mm) chịu lực uốn nén cao.\n"
         "• Thiết kế phào nổi: Profile cánh tích hợp sẵn đường phào chỉ nổi tân cổ điển nghệ thuật, sơn Anodized màu Champagne/Bronze bóng đẹp sang trọng.\n"
         "• Cấu tạo: Bản lề sàn thủy lực âm sàn nặng chịu tải lớn tự động đóng mở giảm chấn êm ái.\n"
         "• Ứng dụng: Cửa đi chính đại sảnh biệt thự lâu đài, sảnh trung tâm showroom cao cấp."),
         
        ("3. Owin Hệ 55 Cửa mở quay rãnh C tiêu chuẩn Đức",
         "• Độ dày nhôm: 1.4mm cho cửa sổ, 2.0mm cho cửa đi (Thanh mã hiệu C3332D, C3303D).\n"
         "• Cấu tạo: Thiết kế rãnh C tiêu chuẩn Châu Âu giúp lắp khóa gạt đa điểm chìm tăng độ khít kín tối đa và chống cạy phá.\n"
         "• Ứng dụng: Cửa đi ban công 1 cánh, 2 cánh, cửa sổ mở quay lật thông gió."),
         
        ("4. Owin Hệ 55 Cửa trượt lùa tiết kiệm",
         "• Độ dày nhôm: 1.2mm - 1.4mm (Thanh C3208).\n"
         "• Công năng: Cửa sổ trượt lùa 2 cánh, cửa lùa ngăn phòng hẹp.\n"
         "• Ưu điểm: Mỏng nhẹ thanh thoát tối ưu diện tích kính lấy sáng."),
         
        ("5. Owin Cửa trượt nâng (Lift & Slide) cách âm tốt",
         "• Độ dày nhôm: 1.5mm - 2.0mm (Thanh C3329LS).\n"
         "• Cơ chế Lift & Slide: Khi đóng gạt hạ tay nắm nén ép gioăng cao su chặt sát ray phẳng mặt sàn cách âm cách nhiệt tốt chống gió bão biển.\n"
         "• Ứng dụng: Cửa trượt lùa ban công rộng penthouse biệt thự hướng biển."),
         
        ("6. Owin Cửa xếp trượt gấp nhiều cánh",
         "• Độ dày nhôm: 1.8mm - 2.0mm.\n"
         "• Công năng: Xếp lùa gấp dồn cánh mở thoáng lối ra hồ bơi bể cảnh biệt thự.\n"
         "• Ưu điểm: Bản cánh hộp to dập ép góc thủy lực ke nhôm dày cứng cáp trượt êm ái nhẹ nhàng."),
         
        ("7. Owin Vách Ngăn kính văn phòng",
         "• Độ dày nhôm: 1.4mm - 2.0mm.\n"
         "• Công năng: Khung nhôm bao giữ vách kính ngăn phòng họp, showroom hiện đại.\n"
         "• Ưu điểm: Kết cấu chắc chắn liên kết vít nở inox thẩm mỹ chống rung lắc giật ồn.")
    ]
    
    for title, desc in owin_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Materials and Gasket Specs ---
    add_heading_with_spacing(doc, "TIÊU CHUẨN KỸ THUẬT VÀ GIA CÔNG CỬA NHÔM OWIN", level=1)
    add_paragraph_with_spacing(doc, 
        "Thanh profile nhôm Owin được đùn ép từ phôi nhôm sạch mác hợp kim 6063-T5 đạt tiêu chuẩn chất lượng. "
        "Bề mặt thanh nhôm được xử lý hóa học và xi mạ Anodized bọc lớp bảo vệ Electrodeposition ED siêu cứng chống trầy xước "
        "và chống ăn mòn hóa học bởi sương muối axit tuyệt hảo.\n\n"
        "Yêu cầu gia công và lắp đặt:\n"
        "• Máy cắt nhôm góc 45 độ CNC chuẩn xác, ép góc thủy lực kết hợp keo góc PU chuyên dụng.\n"
        "• Gioăng cao su EPDM 2 thành phần chất lượng Đức bền bỉ cản gió cản nước mưa bão.\n"
        "• Vít bắt liên kết 100% Inox SUS304 chống gỉ sét ăn mòn điện hóa nhôm.\n"
        "• Keo bọt nở PU trám khe hở tường và keo silicon trung tính chống thấm ngoài trời."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ OWIN]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Owin.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Owin Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_owin_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Owin"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM ĐẶC CHỦNG CAO CẤP OWIN"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        'Phụ Kiện', 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Dòng Đặc Chủng (Signature)", "Owin Trượt Quay", "Cửa trượt quay thông minh không ray dưới", 
         "Cửa đi chính mở trượt kết hợp mở quay dồn cánh góc 100% mở rộng tối đa lối đi sảnh lớn.", 
         "1.6mm - 3.5mm", "Phụ kiện trượt quay Owin đồng bộ", "Gioăng EPDM kép chèn ép lực, silicon", "Hệ thống trượt quay không ray dưới sàn tránh vấp ngã cực kỳ thông minh."),
         
        ("2", "Dòng Thiết Kế (Luxury)", "Owin Thủy Lực", "Cửa thủy lực bản lớn phào nổi đại sảnh", 
         "Cửa đi sảnh lớn lắp bản lề sàn thủy lực cánh phào chỉ chỉ nổi tân cổ điển mạ vàng titan.", 
         "2.0mm", "Bản lề sàn Adler/Hafele đồng bộ, tay kéo", "Keo kết cấu chuyên dụng, gioăng kép", "Bản cánh rộng lớn 120-180mm x 60mm sang trọng lâu đài."),
         
        ("3", "Dòng Cao Cấp (Premium)", "Hệ 55 Mở quay rãnh C", "Cửa mở quay hệ rãnh C tiêu chuẩn Đức", 
         "Cửa đi & cửa sổ mở quay, mở hất lắp khóa gạt đa điểm chìm cản cạy phá chống nước tốt.", 
         "1.4mm - 2.0mm", "Phụ kiện rãnh C đồng bộ CMECH, Roto", "Gioăng EPDM 2 thành phần kín khít", "Sử dụng các thanh nhôm mã C3332D, C3303D dày dặn chắc chắn."),
         
        ("4", "Dòng Tiết Kiệm (Economy)", "Hệ 55 Trượt lùa", "Cửa trượt lùa 55 phổ thông", 
         "Cửa sổ và cửa đi trượt lùa mỏng nhẹ cho ô cửa ban công, phòng ngủ hẹp.", 
         "1.2mm - 1.4mm", "Bánh xe đơn & khóa sập Kinlong, Draho", "Gioăng lông, keo bọt nở PU", "Nhôm mã C3208 tối ưu hóa chi phí sản xuất nhanh."),
         
        ("5", "Dòng Cao Cấp (Premium)", "Owin Trượt Nâng", "Cửa trượt nâng Lift & Slide", 
         "Cửa đi trượt nâng hạ nén gioăng sát ray chống gió bão biển và cách âm tuyệt đối.", 
         "1.5mm - 2.0mm", "Phụ kiện nâng hạ Roto, CMECH, Owin", "Gioăng nén EPDM chèn phẳng sàn", "Nhôm mã C3329LS chuyên dụng trượt lướt êm nhẹ cánh kính lớn."),
         
        ("6", "Dòng Cao Cấp (Premium)", "Owin Xếp Trượt", "Cửa xếp gấp trượt xếp lùa cánh to", 
         "Cửa đi mở xếp trượt gấp nhiều cánh dồn góc mở rộng tối ưu không gian bể bơi biệt thự.", 
         "1.8mm - 2.0mm", "Phụ kiện trượt xếp Hopo, Roto chính hãng", "Gioăng EPDM đàn hồi co giãn cao", "Hệ ray treo trên chịu lực dập góc chắc chắn chống rung xệ."),
         
        ("7", "Dòng Dự Án (B2B)", "Owin Vách Ngăn", "Vách ngăn thạch cao kính Owin", 
         "Vách ngăn kính cường lực chia phòng họp, vách tĩnh lấy sáng showroom văn phòng.", 
         "1.4mm - 2.0mm", "Vít nở inox SUS304 liên kết", "Gioăng cao su đệm kính, keo silicon", "Cấu trúc vững chãi cách âm cách nhiệt tốt dễ thi công.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Đặc Chủng" in seg:
            seg_fill = fill_premium
        elif "Tiết Kiệm" in seg:
            seg_fill = fill_economy
        else:
            seg_fill = fill_standard
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Owin.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Owin Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_owin_docx()
    generate_owin_xlsx()
