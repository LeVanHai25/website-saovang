# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_soco_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM SOCO SYSTEM - ITALY LOOK")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp hệ thống cửa nhôm định hình rãnh C Châu Âu SOCO System\nỨng dụng công nghệ xử lý hóa hóa nhôm Anodized siêu cứng chống mài mòn muối biển vượt trội")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM SOCO SYSTEM (YONGXING)", level=1)
    add_paragraph_with_spacing(doc, 
        "SOCO System (Solution Construction) là dòng sản phẩm nhôm hệ cao cấp được thiết kế theo rãnh C Châu Âu và phong cách Italy (Ý) "
        "thanh lịch, lịch lãm, sản xuất đùn ép bởi tập đoàn nhôm Yongxing Việt Nam. Đặc trưng công nghệ nổi bật nhất của SOCO "
        "là ứng dụng công nghệ xử lý bề mặt nhôm mạ màu Anodized hiện đại: thanh nhôm được đưa qua các bể dung dịch axit hóa học "
        "kết hợp dòng điện phân tạo nên lớp oxit nhôm bọc ngoài cực kỳ dày cứng chống ăn mòn hóa cơ học. Nhôm Anode của SOCO đạt "
        "độ bền màu vượt trội trên 30 năm, chống trầy xước hiệu quả và hoàn toàn miễn dịch với tác hại của không khí muối biển mặn. "
        "SOCO sở hữu tông màu vàng Champagne lộng lẫy và bạc nhôm sang trọng, rãnh C Châu Âu tương thích phụ kiện cao cấp CMECH, Hopo."
    )
    
    add_paragraph_with_spacing(doc, "SOCO System cung cấp các phân hệ nhôm đùn ép đa dạng bao gồm:")
    
    # Table of 6 systems overview in Word
    table = doc.add_table(rows=7, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.6)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "SOCO 65 Quay", "Cửa mở quay phong cách Ý", "1.6mm - 2.0mm", "AW/AD 65 mở quay rãnh C Châu Âu cánh phẳng trơn không gân nổi sang trọng."),
        ("2", "SOCO 94 Lùa", "Cửa trượt lùa 2 ray", "1.4mm - 1.6mm", "Cửa đi/sổ trượt lùa ngang lồng cánh ray inox tròn đúc lướt êm ái."),
        ("3", "SOCO 120 Lùa", "Cửa trượt lùa cánh rộng", "2.0mm", "Cửa lùa trượt hoặc trượt nâng Lift & Slide bản rộng 120mm chịu tải gió penthouse."),
        ("4", "SOCO 180 Lùa", "Cửa trượt lùa 3 ray", "1.4mm - 2.0mm", "Cửa đi trượt lùa 3 ray mở rộng tối đa 2/3 diện tích lối ra sân vườn đại sảnh."),
        ("5", "SOCO X xếp trượt", "Cửa xếp gấp trượt xếp lùa", "1.8mm - 2.0mm", "Cửa xếp trượt dồn cánh ray treo trên chịu lực dập ke ép góc thủy lực cứng vững."),
        ("6", "SOCO Vách Kính", "Vách ngăn kính tĩnh SOCO", "1.6mm - 2.0mm", "Vách kính cố định liên kết gioăng EPDM kín khí cách âm cản bức xạ nóng tốt.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC PHÂN HỆ SOCO SYSTEM", level=1)
    
    soco_details = [
        ("1. SOCO 65 Cửa mở quay phong cách Italy look",
         "• Độ dày nhôm: 1.6mm cho cửa sổ hất, 2.0mm cho cửa đi mở quay chính.\n"
         "• Tên mã cụ thể: SOCO6501, SOCO6502, SOCO6504, SOCO6505, SOCO6506.\n"
         "• Điểm nhấn thiết kế: Khung bao bản dày 65mm. Rãnh C Châu Âu chèn khóa đa điểm chìm giúp tăng an ninh và cản cạy phá.\n"
         "• Thẩm mỹ: Thiết kế trơn phẳng phẳng lì không gân nổi mang đậm nét tối giản hiện đại của Ý."),
         
        ("2. SOCO 94 Cửa trượt lùa ngang phổ thông",
         "• Độ dày nhôm: 1.4mm - 1.6mm.\n"
         "• Tên mã cụ thể: SOCO9401, SOCO9402, SOCO9403, SOCO9405, SOCO9406, SOCO9407.\n"
         "• Cơ cấu trượt: Ray Inox đúc tròn chìm giúp các bánh xe kép chạy êm, giảm ma sát tối đa cản xệ cánh.\n"
         "• Ứng dụng: Cửa sổ lùa trượt, cửa phòng ngăn ban công vừa phải."),
         
        ("3. SOCO 120 Cửa đi trượt nâng Lift & Slide chịu bão gió",
         "• Độ dày nhôm: 2.0mm bản cánh to chịu mô-men xoắn lớn.\n"
         "• Tên mã cụ thể: SOCO12001, SOCO12002, SOCO12003, SOCO12004, SOCO12005.\n"
         "• Điểm nhấn: Thiết kế bản lớn 120mm thích hợp cho các biệt thự hướng biển hoặc ban công nhà cao tầng chịu áp lực bão giật cao."),
         
        ("4. SOCO 180 Cửa trượt lùa 3 ray mở rộng tối đa",
         "• Độ dày nhôm: 1.4mm - 2.0mm.\n"
         "• Tên mã cụ thể: SOCO18001, SOCO18002, SOCO18003, SOCO18004.\n"
         "• Ưu điểm vượt trội: Thiết kế bản to 180mm có 3 đường ray lướt cánh độc lập. Cho phép kéo dồn cánh mở toang tối đa 2/3 diện tích sảnh ra vào vườn ban công rộng lớn sảnh biệt thự, tạo tầm nhìn kính Panorama hoàn mỹ."),
         
        ("5. SOCO Hệ cửa đi xếp trượt xếp lùa cánh to",
         "• Độ dày nhôm: 1.8mm - 2.0mm.\n"
         "• Cơ chế xếp gấp: Trục bản lề chịu tải treo trên dập Ke góc ép thủy lực nén khít, dồn gấp xếp gọn 95% diện tích mặt tiền biệt thự showroom."),
         
        ("6. SOCO Vách ngăn kính cố định cách âm cách nhiệt",
         "• Độ dày nhôm: 1.6mm - 2.0mm.\n"
         "• Ưu điểm: Hệ gioăng cao su đàn hồi kép EPDM nhiều khoang kín chèn khít mép kính hộp Low-E chống nóng tối ưu.")
    ]
    
    for title, desc in soco_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Anodize technology and specs ---
    add_heading_with_spacing(doc, "CÔNG NGHỆ NHÔM ANODE SOCO VÀ BẢO HÀNH 30 NĂM", level=1)
    add_paragraph_with_spacing(doc, 
        "Nhôm SOCO đùn ép từ mác nhôm sạch đạt tiêu chuẩn độ cứng vật lý 6063-T6. "
        "Bề mặt thanh nhôm được xử lý hóa học trong bể dung dịch Acetone và Axit Sulfuric đậm đặc kết hợp dòng điện phân "
        "tạo nên lớp oxit nhôm bọc ngoài siêu cứng. Đây là công nghệ Anodized tiên tiến giúp nhôm chống chịu muối biển "
        "ăn mòn tốt nhất thế giới, bảo hành bề mặt bền màu lên đến 30 năm.\n\n"
        "Quy chuẩn liên kết phụ kiện gia công:\n"
        "• Máy cắt nhôm góc 45 độ CNC chuẩn xác, ke góc ép thủy lực keo PU nén khít.\n"
        "• Gioăng cao su đàn hồi kép EPDM cản gió nước giông tốt.\n"
        "• Phụ kiện rãnh C Châu Âu CMECH Mỹ, Hopo lắp chìm chốt âm đa điểm an toàn tuyệt đối chống cạy phá."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ SOCO]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Soco.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Soco Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_soco_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Soco"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury/design
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for standard
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM ANODIZED SOCO SYSTEM"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        'Phụ Kiện', 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Dòng Mở Quay (Italy look)", "SOCO 65 Quay", "Cửa mở quay rãnh C Châu Âu", 
         "Cửa đi & cửa sổ mở quay cánh phẳng trơn không gân nổi phong cách Ý tối giản sang trọng.", 
         "1.6mm - 2.0mm", "Phụ kiện rãnh C CMECH (Mỹ), Hopo, Sigico", "Gioăng EPDM kép nhiều khoang", "Bản khuôn bao dày 65mm vững vàng cách âm cực tốt."),
         
        ("2", "Dòng Trượt Lùa (Standard)", "SOCO 94 Lùa", "Cửa trượt lùa ngang 2 ray", 
         "Cửa đi/sổ trượt lùa lồng cánh ray inox tròn đúc lướt nhẹ êm hạn chế ma sát xệ cánh.", 
         "1.4mm - 1.6mm", "Bánh xe kép & chốt sập Hopo, Draho", "Gioăng lông cản bụi mịn PM2.5", "Giải pháp trượt lùa tiết kiệm diện tích tối đa nhà phố."),
         
        ("3", "Dòng Trượt Lùa (Premium)", "SOCO 120 Lùa", "Cửa trượt lùa cánh rộng chịu tải", 
         "Cửa lùa trượt hoặc trượt nâng Lift & Slide ban công chịu bão gió căn hộ penthouses.", 
         "2.0mm", "Phụ kiện nâng hạ Hopo, CMECH chính hãng", "Gioăng nén EPDM chèn phẳng sàn", "Bản khuôn bao rộng lớn 120mm chịu lực gió va đập cực khỏe."),
         
        ("4", "Dòng Đặc Biệt (Luxury)", "SOCO 180 Lùa", "Cửa trượt lùa 3 ray mở rộng 2/3", 
         "Cửa đi lùa trượt 3 ray 3 cánh, 6 cánh lướt êm mở toang lối đi chính ra vườn ban công rộng.", 
         "1.4mm - 2.0mm", "Bánh xe & khóa tay nắm chịu tải CMECH", "EPDM đệm cao su kép dẻo dai", "Bản lớn 180mm phong cách Ý, mở rộng tối đa tầm nhìn Panorama."),
         
        ("5", "Dòng Xếp Gấp (Premium)", "SOCO X xếp trượt", "Cửa xếp gấp lùa dồn cánh", 
         "Cửa đi xếp trượt lùa gấp gọn nhiều cánh dồn góc mở rộng 95% sảnh showroom biệt thự.", 
         "1.8mm - 2.0mm", "Phụ kiện trượt xếp Hopo, Bogo chịu tải treo", "Gioăng nén đàn hồi co giãn tốt", "Trục bản lề chịu tải ray treo trên dập ke góc thủy lực chắc chắn."),
         
        ("6", "Vách Kính (Standard)", "SOCO Vách Kính", "Vách ngăn kính tĩnh SOCO", 
         "Vách kính ngăn phòng sảnh showroom lấy sáng tự nhiên cách âm cản bức xạ tốt.", 
         "1.6mm - 2.0mm", "Vít nở inox SUS304 liên kết", "Gioăng cao su kẹp giữ, silicon Dow", "Màu vàng Champagne lộng lẫy bảo hành sơn Anode tới 30 năm.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Đặc Biệt" in seg or "Luxury" in seg:
            seg_fill = fill_premium
        elif "Trượt Lùa" in seg:
            seg_fill = fill_standard
        else:
            seg_fill = fill_economy
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Soco.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Soco Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_soco_xlsx()
    generate_soco_docx()
