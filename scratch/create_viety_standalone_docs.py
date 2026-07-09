# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng Việt Ý (Italumi):
1. Catalogue_HeCuaNhom_VietY.xlsx
2. ThuVien_HeCuaNhom_VietY.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

VIETY_SYSTEMS = [
    ('1', 'Tầm trung phổ thông', 'VI Italumi 55', 'Cửa đi & sổ mở quay Việt Ý 55', 
     'Cửa mở quay thế hệ mới bản 55mm vát mép nhẹ hiện đại, dùng cho cửa sổ ngoại thất và cửa thông phòng.', 
     '1.2mm - 1.4mm', 'Phụ kiện Kinlong / PMA / Huy Hoàng đồng bộ', 'Gioăng EPDM kép / Silicon Apollo A500', 
     'Màu sắc trang nhã, thiết kế vát cạnh cản nước dột tốt. Giá bình dân phù hợp nhà phố xây sẵn.'),
     
    ('2', 'Tầm trung phổ thông', 'VI Italumi 93', 'Cửa lùa trượt Việt Ý 93', 
     'Cửa đi và cửa sổ trượt lùa bản 93mm giúp tiết kiệm không gian diện tích đóng mở tối đa.', 
     '1.4mm - 1.6mm', 'Bánh xe đơn & khóa bán nguyệt đồng bộ', 'Gioăng chèn nỉ mịn chắn bụi nỉ / Apollo', 
     'Trượt êm dột nước ở mức khá, chi phí đầu tư nguyên liệu rất tiết kiệm.'),
     
    ('3', 'Tầm trung phổ thông', 'VI Italumi 450', 'Cửa mở quay Việt Ý 450', 
     'Cửa đi bản nhỏ thông dụng cho nhà WC, phòng kho hoặc cửa sổ lấy sáng phụ.', 
     '1.0mm - 1.2mm', 'Khóa tay gạt tròn / chốt phụ trong nước', 'Gioăng cao su chèn sập thường / Silicon', 
     'Chi phí siêu rẻ cho phân khúc nhà dân cỏ bình dân, dễ thi công dập ke góc nhảy ke ma thuật.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_VietY.xlsx
# ─────────────────────────────────────────────
def create_viety_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue Viet Y"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM ĐỊNH HÌNH VIỆT Ý (ITALUMI)"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="2E7D32") # Xanh lục sậm Italia
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="E8F5E9")
    for r_idx, row_vals in enumerate(VIETY_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_VietY.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel Việt Ý saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel Việt Ý to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_VietY.docx
# ─────────────────────────────────────────────
def create_viety_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x2E,0x7D,0x32))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM ĐỊNH HÌNH VIỆT Ý (ITALUMI)", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp hệ cửa nhôm định hình sản xuất tại nhà máy nhôm Việt Ý (Italumi)\nTích hợp phân tích ưu nhược điểm R&D và ma trận so sánh phân khúc nhà dân dụng phổ thông", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM VIỆT Ý (ITALUMI)", bold=True, size=14, color=RGBColor(0x2E,0x7D,0x32), space_before=12)
    add_p(
        "Nhôm Việt Ý (Italumi) được sản xuất tại nhà máy đùn ép nhôm quy mô lớn ở Việt Nam, đi đầu phân khúc "
        "nhôm hệ định hình phổ thông kinh tế. Với dây chuyền đùn ép công nghệ hiện đại chuyển giao, "
        "Việt Ý cung cấp các dòng sản phẩm nhôm hệ 55 vát mép nhẹ hiện đại, hệ 93 lùa trượt và hệ 450 kinh tế. "
        "Đặc trưng lớn nhất của nhôm Việt Ý là tối ưu hóa trọng lượng vật liệu để mang lại giá thành cửa thành phẩm "
        "cực kỳ rẻ, dễ thi công lắp ráp ke nhảy ke ma thuật nhanh chóng, phù hợp hoàn hảo với nhu cầu xây nhà phố thương mại "
        "xây bán hoặc các công trình ngân sách tiết kiệm."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D 3 HỆ CỬA NHÔM VIỆT Ý", bold=True, size=14, color=RGBColor(0x2E,0x7D,0x32), space_before=12)
    
    # 1. VI Italumi 55
    add_p("1. VI Italumi 55 (Cửa đi & sổ mở quay vát mép)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.4mm tiêu chuẩn hệ 55.")
    add_bullet("Ưu điểm R&D: Thiết kế mặt vát mép nhẹ giúp hạn chế đọng bụi dột nước tốt, thẩm mỹ thanh thoát giống Xingfa hiện đại nhưng giá mềm hơn hẳn, dễ mua ke nhảy gia công nhanh.")
    add_bullet("Nhược điểm R&D: Độ dày mỏng nên không khuyên dùng cho các công trình cao tầng chịu sức gió bão trực diện lớn.")
    
    # 2. VI Italumi 93
    add_p("2. VI Italumi 93 (Cửa lùa trượt kinh tế)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 1.6mm bản lùa 93mm.")
    add_bullet("Ưu điểm R&D: Trượt lướt êm ái, tiết kiệm diện tích tối đa khi mở cánh cho phòng ngủ hẹp hoặc lối đi hông căn nhà, chi phí phôi nhôm thấp.")
    add_bullet("Nhược điểm R&D: Ray lùa mỏng dễ bị cọ mòn xước xát nếu bị kẹt cát bẩn.")

    # 3. VI Italumi 450
    add_p("3. VI Italumi 450 (Cửa mở quay phụ)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.0mm - 1.2mm.")
    add_bullet("Ưu điểm R&D: Giá rẻ vô đối, thi công dập ke nhảy siêu nhanh không đòi hỏi máy móc ép góc đắt tiền, dùng tốt làm cửa vệ sinh, cửa kho hoặc ô thoáng.")
    add_bullet("Nhược điểm R&D: Độ cứng kết cấu yếu hơn hẳn hệ 55, dễ rung lắc cánh cửa nếu lắp ở cửa ngoại thất đón gió.")

    # ── PHẦN SO SÁNH MA TRẬN ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN PHÂN KHÚC HỆ NHÔM VIỆT Ý 55 VỚI ĐỐI THỦ CẠNH TRANH", bold=True, size=14, color=RGBColor(0x2E,0x7D,0x32), space_before=12)
    add_p(
        "Bảng ma trận so sánh chéo đặc tính kỹ thuật hệ nhôm Việt Ý Italumi 55 "
        "với các dòng nhôm cùng phân khúc phổ thông kinh tế trên thị trường:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "Việt Ý Italumi 55"
    hdr_c[2].text = "PMA 55 vát cạnh"
    hdr_c[3].text = "Việt Pháp 450"
    hdr_c[4].text = "Xingfa Việt Nam 55"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ("Nguồn gốc xuất xứ", "Nhà máy Việt Ý (VN)", "Việt Nam (PMA Vina)", "Liên doanh Việt - Pháp", "Nhà máy đùn ép VN"),
        ("Kiểu dáng thiết kế", "Mặt vát mép nhẹ hiện đại", "Mặt vát góc nẹp liền", "Sập rời bo góc tròn", "2 gân gia cường phẳng"),
        ("Độ dày cánh cửa đi", "1.2mm - 1.4mm", "1.2mm - 1.4mm", "1.2mm - 1.4mm", "1.2mm - 1.4mm"),
        ("Bản rộng khung bao", "Khung 55mm", "Khung 55mm", "Khung 450mm bản mỏng", "Khung 55mm"),
        ("Đặc tính nổi bật", "Giá cực tốt, phôi mềm", "Rất phổ biến miền Bắc", "Hệ kinh điển lâu năm", "Dễ mạo danh hàng nhập"),
        ("Phân khúc giá bán", "$ (Siêu tiết kiệm)", "$$ (Bình dân)", "$ (Siêu tiết kiệm)", "$$ (Bình dân)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_VietY.docx")
    try:
        doc.save(path)
        print(f"  >> Docx Việt Ý saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx Việt Ý to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone VietY docx & xlsx in BASE directory...")
    print("=" * 60)
    create_viety_xlsx()
    create_viety_docx()
    print("=" * 60)
