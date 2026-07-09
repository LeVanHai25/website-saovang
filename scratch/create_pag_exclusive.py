# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_pag_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM CẦU CÁCH NHIỆT PAG")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp hệ thống cửa nhôm cầu cách nhiệt và công nghệ hàn góc liền khối PAG\nGiải pháp kiến trúc cách âm cách nhiệt tiêu chuẩn Châu Âu tiết kiệm năng lượng điện tối đa")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM CẦU CÁCH NHIỆT PAG", level=1)
    add_paragraph_with_spacing(doc, 
        "PAG là một trong những thương hiệu cửa nhôm cầu cách nhiệt cao cấp hàng đầu được thiết kế theo tiêu chuẩn cơ khí "
        "và chất lượng Châu Âu. Điểm nhấn đột phá lớn nhất của nhôm PAG là kết cấu 3 lớp cách nhiệt vượt trội: gồm hai "
        "thanh profile nhôm hợp kim 6063-T6 cường độ cao liên kết với nhau thông qua dải cầu cách nhiệt Polyamide PA66 của Technoform. "
        "Ngoài ra, cửa nhôm PAG còn ứng dụng công nghệ hàn góc liền khối (Seamless welding) hiện đại, giúp loại bỏ hoàn toàn "
        "các khe hở ghép nối truyền thống, tăng tính thẩm mỹ nhẵn bóng và đảm bảo kín gió cách âm tới 40dB, kín nước mưa tuyệt hảo. "
        "Bề mặt thanh nhôm sơn phủ tĩnh điện của hãng AkzoNobel giúp chống chịu mặn tốt, cực kỳ thích hợp cho các công trình resort ven biển, "
        "biệt thự và nhà cao tầng hiện đại."
    )
    
    add_paragraph_with_spacing(doc, "Hệ thống sản phẩm cửa nhôm cầu cách nhiệt PAG tiêu biểu bao gồm:")
    
    # Table of 6 systems overview in Word
    table = doc.add_table(rows=7, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.6)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "PAG Hệ 60 Quay", "Cửa sổ mở quay cách nhiệt", "1.8mm", "Cửa sổ mở quay/hất dải cầu cách nhiệt PA66, hàn góc liền khối trơn phẳng thẩm mỹ cao."),
        ("2", "PAG Hệ 65 Quay", "Cửa đi mở quay cách nhiệt", "2.0mm", "Cửa đi mở quay rãnh C Châu Âu chốt đa điểm Hopo/Cmech chịu lực an toàn cao."),
        ("3", "PAG Hệ 83 Xếp", "Cửa đi xếp trượt bản 83", "2.0mm", "Cửa đi xếp gấp trượt nhiều cánh dồn góc chạy ray treo trên chịu lực cực êm nhẹ."),
        ("4", "PAG Hệ 115 Trượt", "Trượt nâng Lift & Slide 115", "2.0mm - 3.0mm", "Cửa đi lùa nâng ray inox chịu tải cánh lớn chống bão và cách âm ban công."),
        ("5", "PAG Hệ Lùa Sổ", "Cửa sổ trượt lùa cách nhiệt", "1.6mm - 1.8mm", "Cửa sổ trượt ngang phẳng rãnh chìm tiết kiệm diện tích cho nhà phố."),
        ("6", "PAG Vách Tĩnh", "Vách kính cách âm cách nhiệt", "1.8mm - 2.5mm", "Vách ngăn kính cố định lấy sáng sảnh văn phòng biệt thự sang trọng.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM PAG", level=1)
    
    pag_details = [
        ("1. PAG Hệ 60 Mở Quay - Cầu cách nhiệt & Hàn góc liền khối",
         "• Độ dày nhôm: 1.8mm.\n"
         "• Bản rộng khuôn bao: 60mm.\n"
         "• Công nghệ hàn góc liền khối: Sử dụng máy hàn nhiệt CNC liền mạch, mài phẳng nhẵn bóng bề mặt sơn tĩnh điện. Triệt tiêu hoàn toàn khe hở nứt góc ngấm nước mưa của các loại nhôm bắt vít ke thường.\n"
         "• Cách âm: Đạt tới 40dB nhờ hệ gioăng EPDM đa khoang kín khít.\n"
         "• Ứng dụng: Cửa sổ mở quay, cửa sổ mở hất ngoài trời."),
         
        ("2. PAG Hệ 65 Cửa đi mở quay - Phân khúc cao cấp",
         "• Độ dày nhôm: 2.0mm cứng vững bản lớn.\n"
         "• Bản khuôn bao: 65mm.\n"
         "• Phụ kiện: Tương thích hoàn hảo phụ kiện kim khí rãnh C Châu Âu của HOPO, CMECH, Sigico mang lại tuổi thọ mở trên 100.000 lần.\n"
         "• Ứng dụng: Cửa đi ban công chính, cửa ra vào sảnh biệt thự."),
         
        ("3. PAG Hệ 83 Cửa đi xếp trượt dồn góc (Folding Door)",
         "• Độ dày nhôm: 2.0mm.\n"
         "• Bản khuôn bao: 83mm vững chắc chịu lực treo cánh nặng.\n"
         "• Cơ chế trượt: Ray treo trên và bánh xe chịu lực dẫn hướng dưới chạy trơn tru, dồn gấp xếp gọn 95% lối ra sân vườn hồ bơi biệt thự."),
         
        ("4. PAG Hệ 115 Cửa đi trượt nâng Lift & Slide chống bão",
         "• Độ dày nhôm: 2.0mm - 3.0mm tùy thuộc điểm ray treo chịu tải trọng gió.\n"
         "• Bản khuôn bao: 115mm.\n"
         "• Đặc tính cơ học: Cơ chế Lift & Slide nâng hạ cánh kính trượt nhẹ như không. Khi hạ cánh hạ lực khóa ép gioăng nằm phẳng chèn sàn cách âm chống gió bão, cản ngấm nước mưa ngoài ban công căn hộ penthouse cao tầng.\n"
         "• Kính: Tương thích tốt kính hộp Double Low-E bơm khí trơ cách nhiệt tốt nhất."),
         
        ("5. PAG Cửa sổ trượt lùa cách nhiệt kinh tế",
         "• Độ dày nhôm: 1.6mm - 1.8mm.\n"
         "• Công năng: Cửa sổ lùa ngang 2 cánh, 4 cánh mỏng nhẹ phẳng thẩm mỹ.\n"
         "• Ưu điểm: Lướt nhẹ trên ray inox đúc tròn, tiết kiệm diện tích tối đa ô cửa ban công nhỏ."),
         
        ("6. PAG Vách kính cố định (Vách ngăn cách âm)",
         "• Độ dày nhôm: 1.8mm - 2.5mm.\n"
         "• Công năng: Vách ngăn văn phòng họp lớn, vách kính lấy sáng biệt thự.\n"
         "• Ưu điểm: Cầu cách nhiệt Polyamide giúp cản bức xạ nóng mùa hè từ bên ngoài giữ ổn định nhiệt độ phòng điều hòa.")
    ]
    
    for title, desc in pag_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- AkzoNobel Paint and PA66 specs ---
    add_heading_with_spacing(doc, "CÔNG NGHỆ SƠN AKZONOBEL VÀ TIÊU CHUẨN CẦU CÁCH NHIỆT PAG", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ thống thanh nhôm định hình PAG sử dụng hợp kim nhôm sạch mác 6063-T6 có độ bền và tính chịu tải tối ưu. "
        "Bề mặt thanh nhôm được sơn phủ bằng dòng sơn tĩnh điện cao cấp AkzoNobel (Hà Lan) kháng sương muối và axit "
        "tuyệt đối, bảo hành bề mặt bền màu lên đến 15 năm.\n\n"
        "Cấu trúc liên kết cầu cách nhiệt và hàn góc:\n"
        "• Dải cầu cách nhiệt Polyamide PA66 Technoform Đức giúp hệ thống giảm 30-40% sự thất thoát nhiệt năng máy lạnh.\n"
        "• Ke ép góc nhôm đúc dày và keo góc PU kết hợp công nghệ hàn góc nhiệt CNC tạo độ liền khối liền mạch.\n"
        "• Hệ gioăng cao su đệm kép EPDM đàn hồi cao kín nước tuyệt đối."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ PAG]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_PAG.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"PAG Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_pag_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue PAG"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM CẦU CÁCH NHIỆT HÀN GÓC PAG"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        'Phụ Kiện', 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Hàn Góc & Cách Nhiệt", "PAG Hệ 60 Quay", "Cửa sổ mở quay cách nhiệt hàn góc liền khối", 
         "Cửa sổ mở quay hoặc hất ngoài dải cầu cách nhiệt PA66 hàn nhiệt CNC nhẵn mịn phẳng gờ mép sơn.", 
         "1.8mm", "Phụ kiện Hopo đồng bộ rãnh C, Sigico", "Gioăng EPDM đa khoang cách âm tới 40dB", "Công nghệ hàn góc seamless cản nước hoàn hảo thẩm mỹ tuyệt đối."),
         
        ("2", "Hàn Góc & Cách Nhiệt", "PAG Hệ 65 Quay", "Cửa đi mở quay rãnh C cách nhiệt", 
         "Cửa đi chính mở quay ban công bản cánh rộng rãnh C Châu Âu chốt đa điểm Hopo an ninh.", 
         "2.0mm", "Tay nắm khóa gạt & bản lề Hopo, CMECH", "Gioăng EPDM kép dẻo dai, silicon Dow", "Dải cách nhiệt PA66 Technoform Đức cản nóng tiết kiệm điện năng."),
         
        ("3", "Dòng Thiết Kế (Premium)", "PAG Hệ 83 Xếp", "Cửa đi xếp trượt gấp bản 83mm", 
         "Cửa xếp gấp trượt dồn cánh mở rộng tối đa lối đi sảnh lớn bể bơi biệt thự.", 
         "2.0mm", "Trục bánh xe chịu treo trên Hopo, khóa đa điểm", "Gioăng cao su đệm đàn hồi cao", "Cơ cấu ray treo chịu lực treo trên cực êm chống rung lắc xệ."),
         
        ("4", "Dòng Siêu Chịu Lực (Luxury)", "PAG Hệ 115 Trượt", "Cửa trượt nâng Lift & Slide 115mm", 
         "Cửa đi lùa nâng hạ nén gioăng chịu bão cản áp lực gió cao tầng, ven biệt thự biển.", 
         "2.0mm - 3.0mm", "Phụ kiện nâng hạ Hopo, Sigico rãnh C", "Gioăng nén EPDM chèn phẳng sàn", "Tương thích tốt kính hộp Double Low-E dày bơm khí trơ."),
         
        ("5", "Cách Nhiệt (Standard)", "PAG Lùa Sổ", "Cửa sổ trượt lùa phẳng rãnh chìm", 
         "Cửa sổ lùa ngang 2 cánh, 4 cánh phẳng ray inox lướt nhẹ tiết kiệm diện tích.", 
         "1.6mm - 1.8mm", "Bánh xe đơn & khóa chốt sập Hopo, Draho", "Gioăng lông, keo bọt nở PU", "Thiết kế cách nhiệt cản bụi mịn PM2.5 vào trong nhà phố."),
         
        ("6", "Cầu Cách Nhiệt (Luxury)", "PAG Vách Tĩnh", "Vách kính cố định lấy sáng cách âm", 
         "Vách kính tĩnh ngăn phòng họp sảnh tòa nhà đố cầu cách nhiệt Technoform.", 
         "1.8mm - 2.5mm", "Ke liên kết chịu lực chuyên dụng, keo dán", "Gioăng cao su đa khoang, silicon Dow", "Giảm 30-40% sự thất thoát nhiệt lạnh máy điều hòa.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Hàn Góc" in seg:
            seg_fill = fill_premium
        elif "Siêu Chịu Lực" in seg or "Luxury" in seg:
            seg_fill = fill_premium
        else:
            seg_fill = fill_standard
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_PAG.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"PAG Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_pag_xlsx()
    generate_pag_docx()
