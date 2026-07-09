# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng PMI (Press Metal - Malaysia):
1. Catalogue_HeCuaNhom_PMI.xlsx
2. ThuVien_HeCuaNhom_PMI.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

PMI_SYSTEMS = [
    ('1', 'Tầm trung rãnh C', 'PMI PE45', 'Cửa đi & sổ mở quay PE45', 
     'Cửa đi, cửa sổ mở quay không cầu cách nhiệt rãnh C tiêu chuẩn Châu Âu gọn nhẹ, tối ưu thẩm mỹ.', 
     '1.4mm - 2.0mm', 'Phụ kiện rãnh C Châu Âu (Sigico, Hopo, CMECH)', 'Gioăng EPDM kép / Silicon Dow Corning', 
     'Thiết kế mỏng nhẹ, tối ưu chi phí nhưng vẫn đảm bảo chuẩn rãnh C Châu Âu cao cấp.'),
     
    ('2', 'Cầu cách nhiệt', 'PMI DPE56', 'Cửa sổ & đi cách nhiệt DPE56', 
     'Hệ thống cửa mở quay có dải cầu cách nhiệt Polyamide giúp cản nhiệt tối đa cho công trình hiện đại.', 
     '1.4mm - 2.0mm', 'Phụ kiện CMECH / Roto rãnh C cách nhiệt', 'Gioăng EPDM đa khoang trung tâm / Silicon Dow', 
     'Khả năng cách âm và cách nhiệt xuất sắc, giảm tiền điện điều hòa đến 30%.'),
     
    ('3', 'Tầm trung phổ thông', 'PMI P55D', 'Hệ Xingfa PMI P55D', 
     'Cửa đi, cửa sổ mở quay kiểu dáng Xingfa định hình bản 55mm của hãng Press Metal.', 
     '1.4mm - 2.0mm', 'Phụ kiện đồng bộ PMI / Kinlong rãnh 22', 'Gioăng EPDM đơn / Silicon Apollo A500', 
     'Thẩm mỹ giống Xingfa nhưng bề mặt sơn mượt mà đặc trưng của Press Metal Malaysia.'),
     
    ('4', 'Cao cấp rãnh C', 'PMI P55M', 'Cửa mở quay bản lớn P55M', 
     'Hệ thống cửa đi mở quay bản lớn dày dặn chuyên dùng cho cửa chính hoặc cửa thông phòng biệt thự.', 
     '2.0mm', 'Bản lề cối chịu lực Hopo / CMECH & khóa đa điểm', 'Gioăng EPDM đúc kép / Silicon Dow Corning', 
     'Cực kỳ cứng vững, chống rung chấn tốt. Cánh nặng hơn các hệ thường.'),
     
    ('5', 'Tầm trung rãnh C', 'PMI MP58A', 'Cửa đi sổ tiết kiệm năng lượng MP58A', 
     'Hệ nhôm tích hợp khoang rãnh C cải tiến giúp tăng tính thẩm mỹ và cản gió tốt.', 
     '1.4mm - 1.8mm', 'Phụ kiện Sigico / Hopo đồng bộ rãnh C', 'Gioăng cao su chèn nỉ / Silicon Sika', 
     'Đáp ứng tốt nhu cầu thiết kế hiện đại và tiết kiệm năng lượng phân khúc trung cấp.'),
     
    ('6', 'Cầu cách nhiệt', 'PMI MP75A', 'Cửa cầu cách nhiệt siêu nặng MP75A', 
     'Hệ cửa đi mở quay cầu cách nhiệt bản lớn cực dày, chống chịu khí hậu khắc nghiệt vùng biển.', 
     '2.0mm', 'Phụ kiện an ninh chống cạy RC2 Sobinco', 'Gioăng đúc xốp trung tâm kép / Dow', 
     'Chống bão gió giật và cản nóng hướng Tây hoàn hảo. Giá thành rất cao.'),
     
    ('7', 'Trượt lùa', 'PMI L100', 'Cửa đi trượt lùa bản 100mm', 
     'Cửa đi trượt lùa 2 ray hoặc 3 ray bản rộng 100mm cứng cáp, trượt êm.', 
     '1.4mm - 2.0mm', 'Bánh xe chịu lực & tay gạt Hopo', 'Gioăng EPDM chống nước chèn nỉ', 
     'Không gian đóng mở rộng thoáng, thích hợp cửa ra ban công căn hộ và biệt thự vườn.'),
     
    ('8', 'Trượt lùa', 'PMI LE70', 'Cửa lùa bản mỏng LE70', 
     'Cửa lùa trượt bản nhỏ gọn tinh tế, thẩm mỹ tối giản sang trọng.', 
     '1.4mm - 1.6mm', 'Bánh xe trượt nhẹ & chốt sập PMI', 'Gioăng EPDM kép / Apollo A500', 
     'Mảnh mai tinh tế, độ bền cao, chi phí bảo trì rất thấp.'),
     
    ('9', 'Lift & Slide', 'PMI L120C', 'Cửa đi trượt nâng L120C', 
     'Cửa trượt nâng hạ nén gioăng kín khít tuyệt đối bản rộng 120mm chịu tải cánh lớn.', 
     '2.0mm', 'Phụ kiện trượt nâng Sobinco (Bỉ) / CMECH tải trọng 300kg', 'Gioăng nén EPDM phẳng sàn', 
     'Vận hành trượt nâng êm ái, cách âm vượt trội cho penhouse hoặc resort ven biển.'),
     
    ('10', 'Lift & Slide', 'PMI DLM128', 'Cửa trượt nâng siêu rộng DLM128', 
     'Hệ trượt nâng bản lớn 128mm siêu dày dặn chuyên trị các ô cửa khẩu độ rộng lớn nhất biệt thự.', 
     '2.2mm', 'Phụ kiện trượt nâng siêu tải trọng CMECH/Hopo', 'Gioăng đúc nén đa khoang chịu nhiệt', 
     'Kết cấu chịu lực gió bão cao nhất, mở rộng tầm nhìn tối đa. Giá thành siêu đắt đỏ.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_PMI.xlsx
# ─────────────────────────────────────────────
def create_pmi_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue PMI"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM CAO CẤP PMI IMPORT MALAYSIA"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="C62828") # Màu đỏ sậm đặc trưng PMI
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="FFEBEE")
    for r_idx, row_vals in enumerate(PMI_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_PMI.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel PMI saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel PMI to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_PMI.docx
# ─────────────────────────────────────────────
def create_pmi_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xC6,0x28,0x28))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM CAO CẤP PMI (PRESS METAL MALAYSIA)", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp hệ cửa nhôm nhập khẩu từ Tập đoàn Press Metal Malaysia\nTích hợp phân tích ưu nhược điểm R&D và ma trận so sánh phân khúc thị trường", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM PMI (MALAYSIA)", bold=True, size=14, color=RGBColor(0xC6,0x28,0x28), space_before=12)
    add_p(
        "Nhôm cao cấp PMI (Press Metal Aluminum) là sản phẩm nhôm thanh định hình đùn ép nhập khẩu 100% từ Tập đoàn Press Metal Malaysia "
        "- tập đoàn sản xuất nhôm lớn nhất khu vực Đông Nam Á. "
        "Sản phẩm được chế tạo từ phôi nhôm sạch nhập khẩu nguyên khối chất lượng cực cao AL6063-T5, tạo nên thanh nhôm có bề mặt "
        "mịn phẳng đặc trưng và nước sơn tĩnh điện vô cùng bền bỉ đạt tiêu chuẩn kiểm định quốc tế AAMA, Qualicoat, GSB Châu Âu. "
        "Nhôm PMI rất nổi tiếng tại thị trường Việt Nam nhờ dải sản phẩm phong phú rãnh C Châu Âu cách âm cách nhiệt tốt "
        "như hệ PE45 không cầu, hệ DPE56 có cầu cách nhiệt, hệ mở quay P55D kiểu dáng Xingfa bản lớn cứng cáp "
        "cùng các dòng cửa trượt lùa và trượt nâng bản rộng (L100, L120C, DLM128) chịu bão gió ven biển đặc chủng."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D 10 HỆ NHÔM PMI", bold=True, size=14, color=RGBColor(0xC6,0x28,0x28), space_before=12)
    
    # 1. PE45
    add_p("1. PMI PE45 (Cửa mở quay rãnh C)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Thiết kế gọn nhẹ tinh tế, tích hợp sẵn rãnh C chuẩn Châu Âu lắp đặt phụ kiện CMECH/Hopo vô cùng dễ dàng, nước sơn tĩnh điện mịn màng bền bỉ.")
    add_bullet("Nhược điểm R&D: Không có cầu cách nhiệt nên hiệu năng cản nhiệt ở mức trung bình, cách âm khoảng 32-35dB.")
    
    # 2. DPE56
    add_p("2. PMI DPE56 (Cầu cách nhiệt)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Hệ cửa cầu cách nhiệt cao cấp dải Polyamide cách âm cách nhiệt hoàn hảo, thích hợp phòng ngủ biệt thự hướng Tây.")
    add_bullet("Nhược điểm R&D: Kỹ thuật sản xuất phức tạp, chi phí đầu tư phôi nhôm nhập khẩu rất cao.")

    # 3. P55D
    add_p("3. PMI P55D (Hệ Xingfa PMI)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Kiểu dáng giống Xingfa 55 thông dụng nhưng đùn ép bởi Press Metal Malaysia, nước sơn bóng mịn chất lượng cao, dễ gia công ép góc.")
    add_bullet("Nhược điểm R&D: Phân khúc giá cao hơn Xingfa Quảng Đông nhập khẩu một chút.")

    # 4. P55M
    add_p("4. PMI P55M (Cửa đi bản lớn rãnh C)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm bản dày.")
    add_bullet("Ưu điểm R&D: Cánh mở quay bản lớn dày dặn 2.0mm chịu lực gió tốt, thẩm mỹ vững chãi cho cửa đi mặt tiền.")
    add_bullet("Nhược điểm R&D: Trọng lượng cánh nặng đòi hỏi dùng bản lề chịu tải cao cấp chuyên dụng.")

    # 5. MP58A
    add_p("5. PMI MP58A (Tầm trung rãnh C tiết kiệm năng lượng)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 1.8mm.")
    add_bullet("Ưu điểm R&D: Khung cánh thiết kế phẳng tinh giản hiện đại, tối ưu dải kính lắp ghép kín khít cách âm tầm trung.")
    add_bullet("Nhược điểm R&D: Phụ kiện lắp đặt rãnh C đi kèm làm tăng giá thành tổng thể cửa.")

    # 6. MP75A
    add_p("6. PMI MP75A (Cầu cách nhiệt siêu nặng chịu bão)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm.")
    add_bullet("Ưu điểm R&D: Cách âm cách nhiệt siêu việt nhờ cầu Polyamide bản dày 75mm vững chãi, chống rung chấn bão giật ven biển rất tốt.")
    add_bullet("Nhược điểm R&D: Chỉ thích hợp phân khúc biệt thự Luxury cao cấp hoặc penhouse vì chi phí siêu đắt.")

    # 7. L100
    add_p("7. PMI L100 (Cửa trượt lùa bản 100mm)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Khung lùa bản 100mm cứng cáp trượt rất đầm chắc chắn, không lo rung giật khi đóng mở.")
    add_bullet("Nhược điểm R&D: Chiếm độ rộng ô tường lớn khi lắp đặt.")

    # 8. LE70
    add_p("8. PMI LE70 (Cửa lùa bản mỏng)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 1.6mm.")
    add_bullet("Ưu điểm R&D: Mảnh gọn tinh tế mở rộng tầm nhìn, trượt nhẹ nhàng, giá cả vô cùng cạnh tranh trong dòng trượt lùa nhập khẩu.")
    add_bullet("Nhược điểm R&D: Khả năng chắn nước ngập bão ngoài trời hạn chế.")

    # 9. L120C
    add_p("9. PMI L120C (Cửa trượt nâng bản lớn)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm.")
    add_bullet("Ưu điểm R&D: Cơ cấu trượt nâng hạ nén gioăng kín khít tuyệt đối cách âm tốt, bánh xe chịu lực cực tốt cánh lướt nhẹ nhàng.")
    add_bullet("Nhược điểm R&D: Chi phí phụ kiện nâng hạ nhập khẩu đắt đỏ.")

    # 10. DLM128
    add_p("10. PMI DLM128 (Cửa trượt nâng siêu tải trọng)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.2mm cực dày.")
    add_bullet("Ưu điểm R&D: Bản rộng 128mm chịu bão gió tối đa ngoại thất, lắp được cánh kích thước siêu rộng lấy trọn tầm nhìn panorama.")
    add_bullet("Nhược điểm R&D: Yêu cầu độ phẳng mặt nền sàn lắp dựng rất cao.")

    # ── PHẦN SO SÁNH MA TRẬN ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN PHÂN KHÚC HỆ PMI PE45 VỚI ĐỐI THỦ CẠNH TRANH", bold=True, size=14, color=RGBColor(0xC6,0x28,0x28), space_before=12)
    add_p(
        "Bảng ma trận so sánh chéo đặc tính kỹ thuật hệ nhôm PMI PE45 nhập khẩu Malaysia "
        "với các dòng nhôm rãnh C Châu Âu cao cấp cạnh tranh trực tiếp trên thị trường:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "PMI PE45"
    hdr_c[2].text = "Kogen Slim 50"
    hdr_c[3].text = "Soco 65"
    hdr_c[4].text = "Maxpro R55"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ("Nguồn gốc xuất xứ", "Malaysia", "Đức/Việt Nam", "Trung Quốc (Cao cấp)", "Nhật Bản (Cận cao cấp)"),
        ('Độ Dày (mm)', "1.4mm - 2.0mm", "1.6mm - 1.8mm", "1.6mm - 2.0mm", "1.4mm"),
        ("Hệ phụ kiện đi kèm", "Hopo / CMECH rãnh C", "Kogen / CMECH rãnh C", "Soco / CMECH rãnh C", "Maxpro / CMECH rãnh C"),
        ("Công nghệ hoàn thiện", "Sơn tĩnh điện AkzoNobel", "Sơn tĩnh điện hạt mịn", "Sơn tĩnh điện Fluorocarbon", "Mạ điện di Anodise ED"),
        ("Phân khúc giá bán", "$$$ (Cận cao cấp)", "$$$$ (Cao cấp)", "$$$ (Khá - Cao cấp)", "$$$ (Cận cao cấp)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_PMI.docx")
    try:
        doc.save(path)
        print(f"  >> Docx PMI saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx PMI to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone PMI docx & xlsx in BASE directory...")
    print("=" * 60)
    create_pmi_xlsx()
    create_pmi_docx()
    print("=" * 60)
