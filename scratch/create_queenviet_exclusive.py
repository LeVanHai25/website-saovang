# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_queenviet_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM EUROVN - QUEENVIET")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp hệ nhôm EuroVN vát cạnh, trượt quay, thủy lực và tủ nội thất Omega\nỨng dụng dây chuyền công nghệ sơn đứng tự động Wagner Thụy Sĩ bảo hành bề mặt 10 năm")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM EUROVN (QUEENVIET GROUP)", level=1)
    add_paragraph_with_spacing(doc, 
        "EuroVN là thương hiệu nhôm định hình cao cấp được sản xuất và phân phối bởi Tập đoàn QueenViet Group tại Việt Nam. "
        "Với triết lý nghiên cứu thực tế nhu cầu thị trường xây dựng nội địa, QueenViet đã đầu tư dây chuyền sản xuất đồng bộ quy mô lớn "
        "sử dụng phôi nhôm sạch mác 6063-T5 nguyên chất và hệ thống phun sơn tĩnh điện dạng đứng tự động của hãng Wagner (Thụy Sĩ). "
        "Nhờ đó, bề mặt sơn của EuroVN có độ mịn bóng đều tuyệt hảo và khả năng kháng tia UV, chống bong tróc bay màu tối đa (bảo hành tới 10 năm). "
        "EuroVN tiên phong kiến tạo các giải pháp độc đáo như cửa tích hợp chấn song nhôm đúc bảo vệ vững chãi đồng bộ, hệ trượt quay mở 100% "
        "không ray dưới sàn và nhôm nội thất Omega Deco giả gỗ tự nhiên giống thật tới 95% chống mối mọt ẩm mốc tuyệt đối."
    )
    
    add_paragraph_with_spacing(doc, "EuroVN (QueenViet) cung cấp 6 hệ nhôm và cấu kiện đặc sắc bao gồm:")
    
    # Table of 6 systems overview in Word
    table = doc.add_table(rows=7, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.6)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "EuroVN Trượt Quay", "Cửa trượt quay thông minh", "1.6mm - 2.8mm", "Cửa mở quay kết hợp mở trượt lùa dồn góc 100%. Không ray dưới sàn, đi lại dễ dàng."),
        ("2", "EuroVN Thủy Lực", "Cửa thủy lực bản lớn phào gỗ", "2.0mm", "Cửa đi bản lề sàn cánh to 120mm - 180mm soi hèm phào chỉ nghệ thuật tân cổ điển."),
        ("3", "Hệ Chấn Song", "Cửa tích hợp chấn song nhôm", "1.5mm - 2.0mm", "Khung cửa đi/sổ tích hợp sẵn các thanh chấn song nhôm bảo vệ đúc vững chắc chống trộm."),
        ("4", "EuroVN Xingfa", "Cửa mở quay hệ Xingfa", "1.4mm - 2.0mm", "Cửa đi/sổ mở quay có gân nổi gia cường chịu lực bề mặt kiểu Xingfa truyền thống."),
        ("5", "EuroVN VIP (Gold)", "Cửa nhôm cầu cách nhiệt", "1.6mm - 2.2mm", "Hệ nhôm cao cấp có dải cầu cách nhiệt Polyamide cách âm cản bức xạ nhiệt tối ưu."),
        ("6", "Nội Thất Omega", "Nhôm nội thất Omega Deco", "1.0mm - 1.2mm", "Nhôm hệ vân gỗ cao cấp giống gỗ thật 95% chuyên làm tủ bếp, tủ áo chống ẩm mọt.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM EUROVN", level=1)
    
    queenviet_details = [
        ("1. EuroVN Hệ trượt quay thông minh (Slide & Turn)",
         "• Độ dày nhôm: 1.6mm - 2.0mm, tại ray treo và các điểm lắp góc chịu lực dày tới 2.8mm.\n"
         "• Cơ chế: Kết hợp đồng thời cơ cấu trượt lùa và mở quay. Cho phép mở 100% diện tích lối đi.\n"
         "• Điểm nhấn: Thiết kế không ray dưới sàn tránh vấp ngã vướng chân, quét dọn dễ dàng. Bánh xe trượt treo dẫn hướng trên chịu lực đầm êm.\n"
         "• Ứng dụng: Cửa chính mặt tiền nhà phố, cửa sảnh ngăn phòng rộng lớn."),
         
        ("2. EuroVN Hệ cửa thủy lực cánh lớn giả gỗ sang trọng",
         "• Độ dày nhôm: 2.0mm cứng vững bản lớn cánh 120-180mm x 60mm.\n"
         "• Điểm nhấn: Cạnh nhôm soi hèm giật cấp, khuôn bao bản to ôm tường tích hợp phào chỉ nổi nghệ thuật tân cổ điển giả gỗ tự nhiên rất đẹp.\n"
         "• Ứng dụng: Cửa đi chính đại sảnh biệt thự lâu đài lắp bản lề sàn tự động đóng giảm chấn."),
         
        ("3. EuroVN Hệ cửa tích hợp chấn song nhôm đúc bảo vệ",
         "• Độ dày chấn song: 1.5mm - 2.0mm đúc rỗng tăng gờ cứng vững.\n"
         "• Ưu điểm vượt trội: Tích hợp trực tiếp thanh chấn song bảo vệ nhôm đúc đồng bộ chất liệu và màu sắc với khuôn bao cửa chính. Đảm bảo an toàn tuyệt đối cho trẻ nhỏ ở các tầng cao chung cư nhà phố, chống trộm đột nhập thay thế chấn song sắt thô cứng gỉ sét.\n"
         "• Ứng dụng: Cửa sổ mở quay, cửa đi ô thoáng phòng ngủ."),
         
        ("4. EuroVN Hệ Xingfa (XF-QueenViet) có gân nổi",
         "• Độ dày nhôm: 1.4mm cho cửa sổ, 2.0mm cho cửa đi mở quay ban công.\n"
         "• Công năng: Thiết kế thân nhôm có đường gân nổi hai bên gia cường chịu lực gió va đập giống hệ Xingfa Quảng Đông truyền thống."),
         
        ("5. EuroVN VIP (Gold) Cửa nhôm cầu cách nhiệt cao cấp",
         "• Độ dày nhôm: 1.6mm - 2.2mm.\n"
         "• Điểm nhấn: Có dải cầu cách nhiệt Polyamide Technoform đặt trong khoang hộp giúp cản bức xạ nhiệt làm mát phòng lạnh và cách âm cách nhiệt tốt.\n"
         "• Ứng dụng: Cửa phòng khách, cửa biệt thự ven biển resort hạng sang."),
         
        ("6. Nhôm nội thất cao cấp Omega Deco giả gỗ tự nhiên",
         "• Độ dày nhôm: 1.0mm - 1.2mm dẻo dai nhẹ dễ gia công.\n"
         "• Điểm nhấn: Bề mặt phủ film vân gỗ chân không công nghệ cao giống gỗ tự nhiên 95%. Không cong vênh, không mối mọt ẩm mốc, chịu nước tuyệt đối.\n"
         "• Ứng dụng: Gia công tủ bếp cao cấp, tủ quần áo phòng ngủ, tủ giày dép, tủ trang trí nội thất.")
    ]
    
    for title, desc in queenviet_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Materials and Gasket Specs ---
    add_heading_with_spacing(doc, "TIÊU CHUẨN KỸ THUẬT VÀ GIA CÔNG CỦA QUEENVIET", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ thống nhôm thanh định hình EuroVN đùn ép từ mác nhôm sạch 6063-T5 đảm bảo giới hạn bền kéo lý tưởng. "
        "Bề mặt thanh nhôm được xử lý hóa học phốt-phát bám sơn và phủ sơn tĩnh điện trên dây chuyền đứng Wagner (Thụy Sĩ). "
        "QueenViet áp dụng chế độ bảo hành bề mặt sơn phủ lên tới 10 năm chống bong tróc, bay màu mốc rỉ.\n\n"
        "Quy chuẩn liên kết phụ kiện gia công:\n"
        "• Ép góc thủy lực máy CNC góc 45 độ, sử dụng ke góc dày kết hợp keo PU góc chuyên dụng khít kín.\n"
        "• Hệ gioăng cao su đàn hồi kép EPDM tăng độ khít chống ngấm nước mưa giông.\n"
        "• Vít liên kết 100% Inox SUS304 chống rỉ sét ăn mòn điện hóa tiếp xúc nhôm.\n"
        "• Keo silicon trung tính chống nước ngoài trời Dow Corning."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ EUROVN]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_EuroVN.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"EuroVN Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_queenviet_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue EuroVN"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM VÀ NỘI THẤT EUROVN QUEENVIET"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        'Phụ Kiện', 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Dòng Thông Minh", "EuroVN Trượt Quay", "Cửa trượt quay mở rộng 100%", 
         "Cửa đi chính mở trượt kết hợp mở quay dồn cánh dẹp góc không có ray dưới tránh vấp ngã chân.", 
         "1.6mm - 2.8mm", "Phụ kiện trượt quay EuroVN chính hãng, Janus", "Gioăng EPDM kép dẻo dai, silicon", "Bản nhôm dày dặn ray treo trên dẫn hướng chịu lực cực êm."),
         
        ("2", "Dòng Thiết Kế", "EuroVN Thủy Lực", "Cửa thủy lực bản lớn phào gỗ nghệ thuật", 
         "Cửa đi sảnh lớn bản lề sàn mở 2 chiều bản cánh phào nổi soi hèm nghệ thuật tân cổ điển.", 
         "2.0mm", "Bản lề sàn Adler/Hafele đồng bộ chính hãng", "Keo kết cấu chuyên dụng, gioăng kép", "Bản cánh rộng lớn 120-180mm giả gỗ tự nhiên rất lộng lẫy."),
         
        ("3", "Dòng An Ninh", "Hệ Chấn Song", "Cửa tích hợp chấn song bảo vệ nhôm", 
         "Cửa sổ và đi tích hợp sẵn các thanh chấn song nhôm đúc bảo vệ an toàn chống trộm đột nhập.", 
         "1.5mm - 2.0mm", "Tay nắm khóa gạt đa điểm, phụ kiện chấn song", "Gioăng cao su chèn, silicon", "Đồng bộ màu sắc khung bao, thay thế chấn song sắt thô gỉ."),
         
        ("4", "Dòng Phổ Thông", "EuroVN Xingfa", "Cửa mở quay Xingfa EuroVN", 
         "Cửa đi & cửa sổ mở quay, hất có gân tăng cứng kiểu Xingfa Quảng Đông phổ biến dễ gia công.", 
         "1.4mm - 2.0mm", "Phụ kiện rãnh 22 Kinlong, Draho chính hãng", "Gioăng EPDM kép, silicon Dow", "Sử dụng phôi nhôm sạch 6063-T5 đùn ép tại nhà máy Việt Nam."),
         
        ("5", "Cầu Cách Nhiệt", "EuroVN VIP (Gold)", "Cửa nhôm cầu cách nhiệt cao cấp", 
         "Cửa đi & sổ có dải cầu cách nhiệt Polyamide Technoform giúp cách âm cản nhiệt chống nóng tốt.", 
         "1.6mm - 2.2mm", "Phụ kiện rãnh C Châu Âu CMECH, Sobinco", "Gioăng EPDM đa khoang chất lượng cao", "Chuyên biệt lắp biệt thự, penthouses, resort hướng biển mặn."),
         
        ("6", "Nhôm Nội Thất", "Nội Thất Omega", "Tủ nội thất nhôm Omega Deco", 
         "Gia công tủ bếp, tủ quần áo, tủ giày dép giả gỗ tự nhiên giống thật tới 95%.", 
         "1.0mm - 1.2mm", "Bản lề hơi hơi giảm chấn, tay nắm Anode", "Gioăng kẹp kính, keo liên kết", "Không mối mọt, không cong vênh nứt nẻ, chịu nước tuyệt đối.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Thông Minh" in seg or "Thiết Kế" in seg:
            seg_fill = fill_premium
        elif "Nội Thất" in seg:
            seg_fill = fill_standard
        else:
            seg_fill = fill_economy
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_EuroVN.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"EuroVN Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_queenviet_xlsx()
    generate_queenviet_docx()
