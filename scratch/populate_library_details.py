# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

base_dir = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"
master_dir = os.path.join(base_dir, "THƯ VIỆN HỆ NHÔM SAO VÀNG")

# Helper to generate basic DXF file containing a drawing
def generate_dxf(file_path, width=100.0, height=50.0, layer_name="Profile_Nhom"):
    dxf_content = f"""0
SECTION
2
HEADER
9
$ACADVER
1
AC1015
0
ENDSEC
0
SECTION
2
TABLES
0
TABLE
2
LTYPE
70
1
0
LTYPE
2
CONTINUOUS
70
0
3
Solid line
72
65
73
0
40
0.0
0
ENDTAB
0
TABLE
2
LAYER
70
2
0
LAYER
2
0
70
0
62
7
6
CONTINUOUS
0
LAYER
2
{layer_name}
70
0
62
1
6
CONTINUOUS
0
ENDTAB
0
ENDSEC
0
SECTION
2
BLOCKS
0
ENDSEC
0
SECTION
2
ENTITIES
0
LINE
8
{layer_name}
10
0.0
20
0.0
30
0.0
11
{width}
21
0.0
31
0.0
0
LINE
8
{layer_name}
10
{width}
20
0.0
30
0.0
11
{width}
21
{height}
31
0.0
0
LINE
8
{layer_name}
10
{width}
20
{height}
30
0.0
11
0.0
21
{height}
31
0.0
0
LINE
8
{layer_name}
10
0.0
20
{height}
30
0.0
11
0.0
21
0.0
31
0.0
0
ENDSEC
0
EOF"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(dxf_content)

# Populate details in folders
def populate_all():
    print("Populating CAD library...")
    cad_dirs = {
        "Cửa đi": ("Ban_ve_mau_Cua_di.dxf", 120.0, 220.0, "CAD_Cua_di"),
        "Cửa sổ": ("Ban_ve_mau_Cua_so.dxf", 100.0, 120.0, "CAD_Cua_so"),
        "Cửa lùa": ("Ban_ve_mau_Cua_lua.dxf", 180.0, 200.0, "CAD_Cua_lua"),
        "Cửa Slim": ("Ban_ve_mau_Cua_Slim.dxf", 150.0, 220.0, "CAD_Cua_Slim"),
        "Vách mặt dựng": ("Ban_ve_mau_Vach_mat_dung.dxf", 250.0, 300.0, "CAD_Vach_Mat_Dung")
    }
    
    for sub, (dxf_name, w, h, layer) in cad_dirs.items():
        sub_path = os.path.join(master_dir, "03_Thư viện CAD", sub)
        # Generate DXF
        generate_dxf(os.path.join(sub_path, dxf_name), w, h, layer)
        # Write technical instruction guide
        with open(os.path.join(sub_path, "huong_dan_cad.txt"), "w", encoding="utf-8") as f:
            f.write(f"QUY CHUẨN VẼ CAD HỆ CỬA: {sub.upper()}\n"
                    f"================================================\n"
                    f"1. Tỷ lệ bản vẽ (Scale): Vẽ tỷ lệ 1:1 trong Model Space, Block thể hiện 1:10 hoặc 1:5.\n"
                    f"2. Quản lý Layer:\n"
                    f"   - Layer '{layer}': Dùng nét mảnh để thể hiện khung cánh.\n"
                    f"   - Layer 'NET_THAY': Thể hiện các đường biên dạng của nhôm.\n"
                    f"   - Layer 'NET_KHUAT': Thể hiện đường bao khuất chèn bê tông hoặc gioăng cao su.\n"
                    f"   - Layer 'DIM_TEXT': Cho kích thước và ghi chú chữ.\n"
                    f"3. Quy chuẩn Hatch nét kính:\n"
                    f"   - Kính cường lực đơn: Hatch nét xiên 45 độ liền mảnh.\n"
                    f"   - Kính hộp chân không: Hatch 2 nét kính cách nhau 12mm chèn dải keo silicon đen.\n")

    # 04. Thư viện Profile Nhôm
    print("Populating Profile library...")
    profile_path = os.path.join(master_dir, "04_Thư viện Profile Nhôm")
    generate_dxf(os.path.join(profile_path, "Profile_Khung_Bao_He_65.dxf"), 65.0, 55.0, "Profile_Khung_Bao")
    generate_dxf(os.path.join(profile_path, "Profile_Canh_He_65.dxf"), 65.0, 75.0, "Profile_Canh")
    with open(os.path.join(profile_path, "huong_dan_profile.txt"), "w", encoding="utf-8") as f:
        f.write("QUY CHUẨN TRA CỨU HỆ PROFILE NHÔM ĐỊNH HÌNH\n"
                "================================================\n"
                "1. Tất cả profile đùn ép phải thể hiện chuẩn độ dày thực tế (1.4mm - 2.0mm).\n"
                "2. Rãnh chèn gioăng EPDM phải đúng kích thước hình học 5.5mm - 6.2mm.\n"
                "3. Rãnh rơ-le phụ kiện rãnh C phải tuân thủ chuẩn Euro-Groove 15/20mm.\n"
                "4. Xem mặt cắt mẫu đính kèm: Profile_Khung_Bao_He_65.dxf và Profile_Canh_He_65.dxf.")

    # 05. Thư viện Phụ kiện
    print("Populating Hardware library...")
    hardware_path = os.path.join(master_dir, "05_Thư viện Phụ kiện")
    generate_dxf(os.path.join(hardware_path, "Ban_le_3D_Cmech.dxf"), 40.0, 90.0, "Phu_kien_Ban_le")
    generate_dxf(os.path.join(hardware_path, "Khoa_Da_Diem_Hopo.dxf"), 20.0, 200.0, "Phu_kien_Khoa")
    with open(os.path.join(hardware_path, "huong_dan_phu_kien.txt"), "w", encoding="utf-8") as f:
        f.write("QUY CHUẨN PHỤ KIỆN KIM KHÍ ĐỒNG BỘ\n"
                "================================================\n"
                "1. Bản lề 3D / Bản lề ma sát: Phải kiểm tra góc mở tối đa (90 - 180 độ) để chống va vào tường.\n"
                "2. Thân khóa đa điểm: Rãnh truyền động phải khớp với thanh truyền lực chuyển động của tay nắm.\n"
                "3. Bánh xe chịu lực: Bánh xe inox đúc đơn/kép chịu tải cánh sổ lùa từ 80kg, cánh đi lùa từ 150kg - 400kg.\n"
                "4. Xem bản vẽ mẫu: Ban_le_3D_Cmech.dxf và Khoa_Da_Diem_Hopo.dxf.")

    # 06. Hướng dẫn Lắp đặt (Word DOCX)
    print("Populating Installation guide...")
    install_path = os.path.join(master_dir, "06_Hướng dẫn Lắp đặt")
    doc_inst = Document()
    doc_inst.add_heading("HƯỚNG DẪN QUY TRÌNH LẮP ĐẶT CỬA NHÔM KÍNH DỰ ÁN SAO VÀNG", level=1)
    
    install_steps = [
        "Bước 1: Khảo sát đo đạc ô chờ trước khi lắp đặt (kiểm tra cốt 0.0, độ dốc cốt nền).",
        "Bước 2: Định vị khung bao vào tường bằng vít nở inox SUS304 chuyên dụng (khoảng cách vít < 600mm).",
        "Bước 3: Lắp đặt kính hộp / kính cường lực vào cánh cửa bằng nẹp sập kính và nêm nhựa đàn hồi chống sệ.",
        "Bước 4: Bơm keo bọt nở PU trương nở xung quanh khe hở giữa khung nhôm và tường (đợi khô cắt phẳng).",
        "Bước 5: Bơm keo liên kết chống thấm nước mưa ngoài trời (silicon trung tính Dow Corning / Apollo A500/A600).",
        "Bước 6: Căn chỉnh bản lề, phụ kiện tay nắm khóa cửa đạt độ trơn tru và đóng kín khít.",
        "Bước 7: Vệ sinh bóc băng keo bảo vệ và nghiệm thu bàn giao chủ đầu tư."
    ]
    for step in install_steps:
        doc_inst.add_paragraph(step)
        
    doc_inst.save(os.path.join(install_path, "Huong_dan_lap_dat_chi_tiet.docx"))

    # 07. Hướng dẫn Gia công (Word DOCX)
    print("Populating Fabrication manual...")
    fab_path = os.path.join(master_dir, "07_Hướng dẫn Gia công")
    doc_fab = Document()
    doc_fab.add_heading("HƯỚNG DẪN QUY TRÌNH SẢN XUẤT GIA CÔNG CỬA NHÔM KÍNH TẠI XƯỞNG", level=1)
    
    fab_steps = [
        "Quy trình cắt nhôm: Sử dụng máy cắt 2 đầu CNC đạt độ chính xác góc cắt 45 độ và 90 độ tuyệt đối.",
        "Quy trình ép góc nhôm: Sử dụng ke góc nhôm đúc dày kết hợp máy ép góc thủy lực lực ép 15-20 Bar, bơm keo góc PU chuyên dụng.",
        "Quy trình khoét lỗ phụ kiện: Sử dụng máy chép hình CNC khoét lỗ khóa, tay nắm và lỗ thoát nước mưa đúng bản vẽ kỹ thuật.",
        "Quy trình chèn gioăng cao su: Đi gioăng EPDM đa khoang liên tục xung quanh rãnh nhôm, nối góc bằng keo lưu hóa.",
        "Quy trình kiểm tra QA/QC tại xưởng trước khi đóng gói xuất xưởng."
    ]
    for step in fab_steps:
        doc_fab.add_paragraph(step)
        
    doc_fab.save(os.path.join(fab_path, "Huong_dan_gia_cong_san_xuat.docx"))

    # 09. Hồ sơ Dự án Tham khảo
    print("Populating Project references...")
    project_path = os.path.join(master_dir, "09_Hồ sơ Dự án Tham khảo")
    with open(os.path.join(project_path, "Ho_so_du_an_biet_thu_mau.txt"), "w", encoding="utf-8") as f:
        f.write("HỒ SƠ DỰ ÁN BIỆT THỰ MẪU SAO VÀNG VILLA\n"
                "================================================\n"
                "1. Địa điểm xây dựng: Khu đô thị Vinhomes Riverside, Hà Nội.\n"
                "2. Quy mô dự án: Biệt thự đơn lập 3 tầng nổi, 1 tầng hầm.\n"
                "3. Giải pháp cửa nhôm kính áp dụng:\n"
                "   - Tầng 1 (Đại sảnh): Hệ cửa thủy lực Owin bản lớn 180mm mạ nhôm Anode màu Champagne phối kính hộp Low-E.\n"
                "   - Tầng 2 & 3 (Phòng ngủ): Hệ cửa sổ mở quay cầu cách nhiệt Civro AW65 rãnh C Châu Âu cách âm chống ồn.\n"
                "   - Lối ra bể bơi sân vườn: Cửa đi trượt nâng Lift & Slide Soco 120 Anodized trượt siêu rộng.\n"
                "   - Phòng thay đồ & bếp nội thất: Cửa nhôm Slim lùa treo không ray dưới sàn giảm chấn tự hãm OPK.\n"
                "4. Quy chuẩn kính: Kính hộp Hải Long 6Temper + 12Argon + 6Low-E bơm khí trơ cản bức xạ UV.")

    # 10. Hình ảnh & Video
    print("Populating Images and Videos...")
    media_path = os.path.join(master_dir, "10_Hình ảnh & Video")
    with open(os.path.join(media_path, "Danh_sach_video_huong_dan.txt"), "w", encoding="utf-8") as f:
        f.write("DANH SÁCH VIDEO HƯỚNG DẪN KỸ THUẬT VÀ LẮP ĐẶT\n"
                "================================================\n"
                "1. Video [01_Huong_dan_lap_giam_chan_Slim.mp4]: Mô phỏng hành trình piston thủy lực tự hãm cánh cửa lùa Slim.\n"
                "2. Video [02_Huong_dan_han_goc_seamless_PAG.mp4]: Hướng dẫn vận hành máy hàn nhiệt CNC liền khối không vết ghép.\n"
                "3. Video [03_Huong_dan_lap_cua_truot_quay_Owin.mp4]: Hướng dẫn căn chỉnh bản lề xoay định vị ray treo trên cửa trượt quay.\n"
                "4. Video [04_Huong_dan_bom_keo_ silicon_chong_tham.mp4]: Quy chuẩn đi đường keo silicon ngoài trời bóng mịn không răng cưa.\n"
                "5. Tất cả video được lưu hành nội bộ trên ổ đĩa mạng công ty.")

    print("All directories populated successfully!")

if __name__ == "__main__":
    populate_all()
