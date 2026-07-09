# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_viralwindow_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM VIRALWINDOW CHUYÊN SÂU")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp 15 hệ cửa nhôm VRA, VRE, VR Chuyên dụng và VRX Cách nhiệt\nPhiên bản dành cho Kiến trúc sư, Kỹ sư mặt dựng & Ban Giám đốc")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Table of Contents Placeholder & Brand Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU VIRALWINDOW", level=1)
    add_paragraph_with_spacing(doc, 
        "Viralwindow là một trong những thương hiệu cửa nhôm cao cấp hàng đầu tại thị trường Việt Nam, "
        "được nghiên cứu và phát triển theo các tiêu chuẩn kỹ thuật khắt khe của Châu Âu. Hệ thống sản phẩm "
        "của hãng sử dụng phôi nhôm sạch nhập khẩu mác 6063-T5/T6 có cấu trúc khoang trống dày dặn cùng "
        "các gờ gia cường chịu lực lớn. Điểm vượt trội của Viralwindow là sử dụng công nghệ sơn phủ KCC Hàn Quốc "
        "bảo hành lên tới 20 năm chống ăn mòn muối biển và hệ gioăng EPDM 2-3 lớp kép đồng bộ cùng hệ thống phụ kiện "
        "chính hãng tiêu chuẩn rãnh C Châu Âu (Cmech, Hopo, Roto, Viral đồng bộ) tạo độ kín khít và cách âm hoàn hảo."
    )
    
    add_paragraph_with_spacing(doc, "Tài liệu này tổng hợp chi tiết toàn bộ 15 hệ nhôm của Viralwindow được chia thành 4 nhóm phân khúc sản phẩm cốt lõi:")
    
    # Table of 15 systems overview in Word
    table = doc.add_table(rows=16, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", 'Nhóm Hệ', "Mã Hệ", 'Tên Hệ', 'Loại Sản Phẩm']
    col_widths = [Inches(0.5), Inches(1.5), Inches(1.0), Inches(2.0), Inches(2.5)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "Hệ VRA (Châu Á)", "VRA55", "Cửa đi & Cửa sổ mở quay VRA55", "Cửa đi mở quay, cửa sổ mở quay, sổ mở hất, vách kính cố định (Cánh phẳng khung, kính lõm)"),
        ("2", "Hệ VRA (Châu Á)", "VRA63", "Cửa xếp trượt VRA63", "Cửa đi mở xếp trượt góc hoặc xếp gấp phẳng"),
        ("3", "Hệ VRA (Châu Á)", "VRA64", "Cửa sổ trượt VRA64", "Cửa sổ mở trượt lùa 2 cánh, 4 cánh"),
        ("4", "Hệ VRA (Châu Á)", "VRA94 / VRA80", "Cửa trượt VRA94", "Cửa đi mở trượt lùa 2 ray"),
        ("5", "Hệ VRA (Châu Á)", "VRA120", "Cửa trượt VRA120", "Cửa đi mở trượt lùa 3 ray"),
        ("6", "Hệ VRE (Châu Âu)", "VRE55", "Cửa mở quay tiêu chuẩn VRE55", "Cửa sổ mở quay, mở hất hoặc cửa đi bản nhỏ"),
        ("7", "Hệ VRE (Châu Âu)", "VRE65", "Cửa mở quay Premium VRE65", "Cửa đi mở quay, cửa sổ mở quay lật (Hệ rãnh C tiêu chuẩn Châu Âu, cánh ôm khung)"),
        ("8", "Hệ VRE (Châu Âu)", "VRE75", "Cửa xếp gấp VRE75", "Cửa đi mở xếp trượt cao cấp"),
        ("9", "Hệ VRE (Châu Âu)", "VRE77", "Cửa xếp gấp ẩn bản lề VRE77", "Cửa đi xếp trượt Luxury (Bản lề giấu kín trong khe nhôm)"),
        ("10", "Hệ VRE (Châu Âu)", "VRE94", "Cửa sổ trượt VRE94", "Cửa sổ mở trượt lùa bản dày tiêu chuẩn Châu Âu"),
        ("11", "Hệ VRE (Châu Âu)", "VRE120", "Cửa trượt nâng Lift & Slide VRE120", "Cửa đi mở trượt lùa 2 ray, có phụ kiện nâng hạ cánh"),
        ("12", "Hệ VRE (Châu Âu)", "VRE180", "Cửa lùa siêu khổ VRE180", "Cửa đi mở trượt lùa 3 ray tích hợp lưới chống côn trùng"),
        ("13", "Hệ VR Chuyên Dụng", "VR100", "Cửa mở trượt quay VR100", "Cửa đi mở trượt kết hợp mở quay (Slide & Turn), không ray dưới"),
        ("14", "Hệ VR Chuyên Dụng", "VR150", "Cửa thủy lực đại sảnh VR150", "Cửa đi bản lề sàn (thủy lực), bản cánh lớn phào nổi"),
        ("15", "Hệ Cầu Cách Nhiệt", "VRX75", "Cửa nhôm cầu cách nhiệt VRX75", "Cửa đi và cửa sổ mở quay đa khoang có cầu Polyamide")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- PHẦN I: HỆ VRA (CHÂU Á) ---
    add_heading_with_spacing(doc, "PHẦN I: HỆ VRA (PHÂN KHÚC CHÂU Á - HIỆN ĐẠI & TỐI GIẢN)", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ VRA của Viralwindow được thiết kế theo xu hướng thiết kế Á Đông hiện đại. "
        "Đặc trưng lớn nhất là thiết kế phẳng mặt ngoài (khung và cánh bằng phẳng không có đường gân nổi) "
        "kết hợp các gờ kính lõm vào trong tạo độ phẳng mướt, thanh thoát cho ngôi nhà. Phân khúc này tối ưu hóa "
        "trọng lượng nhôm nhưng vẫn bảo đảm cấu trúc cơ học dập góc thủy lực và chịu tải gió rất tốt."
    )
    
    vra_systems = [
        ("VRA55 - Cửa đi & Cửa sổ mở quay", 
         "• Độ dày nhôm: 1.4mm - 2.0mm.\n"
         "• Bản khuôn bao: 55mm.\n"
         "• Công năng: Phù hợp làm cửa đi 1 cánh, 2 cánh mở quay phòng ngủ, cửa toilet, cửa ra ban công; cửa sổ mở quay hoặc hất ngoài trời.\n"
         "• Phụ kiện: Đồng bộ Viral, Draho, Hopo chính hãng.\n"
         "• Hệ gioăng: EPDM 2 lớp độ bền cao, đàn hồi tốt chống ồn."),
         
        ("VRA63 - Cửa xếp trượt gấp phổ thông", 
         "• Độ dày nhôm: 1.6mm - 1.8mm.\n"
         "• Cơ chế hoạt động: Xếp gấp cánh gọn gàng sang 1 hoặc 2 bên, mở rộng không gian 90%.\n"
         "• Thiết kế đặc biệt: Các ray trượt và bánh xe chịu tải cao, trượt êm phẳng mặt sàn.\n"
         "• Ứng dụng: Cửa ra ban công, cửa garage ô tô, lối ra hồ bơi nhà dân dụng tầm trung."),
         
        ("VRA64 - Cửa sổ mở trượt lùa", 
         "• Độ dày nhôm: 1.2mm - 1.4mm giúp tối ưu hóa chi phí cho dự án.\n"
         "• Công năng: Cửa sổ lùa 2 cánh, 4 cánh trượt ngang, chống đập gió mưa tuyệt hảo.\n"
         "• Phụ kiện: Khóa sập tự động, bánh xe đơn chịu mòn chạy trên ray inox tròn."),
         
        ("VRA94 / VRA80 - Cửa đi trượt lùa 2 ray", 
         "• Độ dày nhôm: 1.4mm - 1.8mm.\n"
         "• Bản khuôn bao: 94mm.\n"
         "• Công năng: Làm cửa đi lùa phòng khách ra ban công, cửa lùa sảnh vườn biệt thự.\n"
         "• Điểm nhấn: Vận hành rất nhẹ nhàng nhờ hệ bánh xe kép đồng bộ chạy trên ray inox chống xệ cánh."),
         
        ("VRA120 - Cửa đi trượt lùa 3 ray", 
         "• Độ dày nhôm: 1.8mm - 2.0mm.\n"
         "• Bản khuôn bao: 120mm.\n"
         "• Công năng: Thiết kế 3 ray 3 cánh lùa song song, cho phép mở rộng tới 2/3 diện tích ô cửa.\n"
         "• Điểm nhấn: Phù hợp cho công trình hướng view rộng lớn cần không gian thoáng đãng tối đa.")
    ]
    
    for title, desc in vra_systems:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- PHẦN II: HỆ VRE (CHÂU ÂU) ---
    add_heading_with_spacing(doc, "PHẦN II: HỆ VRE (PHÂN KHÚC CHÂU ÂU - CÁNH ÔM KHUNG SANG TRỌNG)", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ VRE đại diện cho triết lý thiết kế cửa nhôm cao cấp Châu Âu. Cấu tạo cánh cửa thiết kế ôm khít "
        "lên mặt ngoài khung bao, tạo gờ chỉ nổi sang trọng giúp cửa có chiều sâu, tăng cường tối đa khả năng cách âm "
        "và ngăn nước mưa tuyệt đối. Tích hợp rãnh C tiêu chuẩn Châu Âu lắp đặt phụ kiện cao cấp (Cmech, Roto, Hopo) "
        "chống xệ cánh và cạy phá an toàn bậc nhất."
    )
    
    vre_systems = [
        ("VRE55 - Cửa mở quay tiêu chuẩn Châu Âu", 
         "• Độ dày nhôm: 1.4mm - 1.6mm.\n"
         "• Công năng: Phù hợp cho cửa sổ mở quay, cửa sổ mở hất hoặc cửa đi thông phòng bản nhỏ mỏng nhẹ tinh tế.\n"
         "• Phụ kiện: Rãnh C tiêu chuẩn lắp khóa đa điểm Cmech, Sigico."),
         
        ("VRE65 - Cửa mở quay Premium Châu Âu", 
         "• Độ dày nhôm: 1.6mm - 2.0mm cứng cáp.\n"
         "• Công năng: Chuyên dùng cho cửa đi mở quay chính, cửa sổ quay lật (Tilt & Turn) góc hất 15 độ thông gió.\n"
         "• Đặc trưng: Thiết kế cánh ôm khung tăng độ dày đệm khí, kết hợp hệ gioăng EPDM 3 tầng kín khít cách âm tuyệt đối."),
         
        ("VRE75 - Cửa xếp gấp Premium", 
         "• Độ dày nhôm: 1.8mm - 2.2mm bản cánh lớn.\n"
         "• Công năng: Hệ thống cửa xếp trượt cao cấp dành cho biệt thự, sảnh nghỉ khách sạn nghỉ dưỡng.\n"
         "• Phụ kiện: Đồng bộ bánh xe chịu lực treo trên và dẫn hướng dưới của hãng CMECH/Hopo vận hành siêu êm."),
         
        ("VRE77 - Cửa xếp gấp ẩn bản lề Luxury", 
         "• Thiết kế độc quyền: Bản lề xếp cánh được giấu kín hoàn toàn bên trong khe rãnh nhôm khi đóng cửa.\n"
         "• Ưu điểm thẩm mỹ: Mặt cửa phẳng tuyệt đối không lộ con lề thô cồng kềnh, tránh bụi bẩn bám vào bản lề.\n"
         "• Độ an toàn: Chống cạy phá góc lề từ bên ngoài tuyệt hảo. Phân khúc Luxury siêu cao cấp biệt thự nghỉ dưỡng."),
         
        ("VRE94 - Cửa sổ mở trượt lùa Châu Âu", 
         "• Độ dày nhôm: 1.6mm - 1.8mm.\n"
         "• Công năng: Cửa sổ trượt lùa bản rộng dày dặn, chịu lực rung bão gió giật vùng ven biển rất tốt.\n"
         "• Thiết kế: Cạnh góc vuông sắc nét phong cách tối giản Châu Âu."),
         
        ("VRE120 - Cửa trượt nâng Lift & Slide", 
         "• Cơ chế Lift & Slide: Khi gạt tay nắm khóa, toàn bộ cánh cửa sẽ được nhấc lên khỏi ray để trượt nhẹ nhàng; "
         "khi khóa cửa, cánh cửa hạ xuống ép chặt gioăng xuống sàn cách âm, cách nhiệt hoàn toàn.\n"
         "• Trọng lượng cánh: Chịu tải kính hộp nặng lên đến 350kg.\n"
         "• Ứng dụng: Cửa lùa chính biệt thự hướng vườn hoặc ban công penthouse."),
         
        ("VRE180 - Cửa lùa siêu khổ Jumbo tích hợp lưới", 
         "• Cấu tạo 3 ray tích hợp: 2 ray trượt cánh kính lớn và 1 ray tích hợp sẵn lưới chống côn trùng inox 304 siêu bền.\n"
         "• Tiện ích: Đón gió tự nhiên mát mẻ mà không lo muỗi, ruồi; bảo vệ an toàn trẻ nhỏ.\n"
         "• Bản nhôm: Bản nhôm dày 2.0mm - 2.5mm cho khẩu độ kính cực đại.")
    ]
    
    for title, desc in vre_systems:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- PHẦN III: HỆ VR CHUYÊN DỤNG ---
    add_heading_with_spacing(doc, "PHẦN III: HỆ VR CHUYÊN DỤNG (GIẢI PHÁP ĐỘC ĐÁO & ẤN TƯỢNG)", level=1)
    
    special_systems = [
        ("VR100 - Cửa mở trượt quay thông minh (Slide & Turn)", 
         "• Cơ chế vận hành sáng tạo: Kết hợp độc đáo giữa cửa lùa trượt và cửa mở quay. Cánh cửa được trượt lùa dồn "
         "về một bên sau đó quay 90 độ mở ra hoàn toàn lối đi.\n"
         "• Đặc trưng: Thiết kế không có ray dưới sàn nhà, loại bỏ hoàn toàn nguy cơ vấp ngã, dễ dàng quét dọn nhà cửa "
         "và đẩy xe nôi, xe lăn.\n"
         "• Phụ kiện: Đồng bộ hệ bánh xe treo dẫn hướng trên đặc chủng bền bỉ của Viral.\n"
         "• Ứng dụng: Phù hợp làm cửa thông phòng khách - bếp, cửa đi mặt tiền nhà ống cần mở tối đa diện tích."),
         
        ("VR150 - Cửa thủy lực đại sảnh Luxury", 
         "• Cấu tạo: Cánh cửa nhôm bản siêu lớn 136mm x 65mm dày 2.0mm chịu lực cực tốt, lắp đặt bản lề sàn thủy lực nặng.\n"
         "• Thiết kế phào nổi: Profile nhôm tích hợp phào chỉ nổi nghệ thuật tân cổ điển kết hợp mạ màu Gold/Champagne sang trọng.\n"
         "• Tương thích: Lắp đặt kính cường lực dày 10mm - 12mm hoặc kính hộp nan đồng trang trí nghệ thuật.\n"
         "• Ứng dụng: Cửa chính đại sảnh biệt thự, lâu đài hoặc trung tâm hội nghị tiệc cưới lớn.")
    ]
    
    for title, desc in special_systems:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- PHẦN IV: HỆ CẦU CÁCH NHIỆT ---
    add_heading_with_spacing(doc, "PHẦN IV: HỆ CẦU CÁCH NHIỆT CAO CẤP", level=1)
    
    thermal_system = [
        ("VRX75 - Cửa nhôm cầu cách nhiệt đa khoang", 
         "• Cấu trúc Profile nhôm: Cấu tạo 3 khoang độc lập với dải cầu cách nhiệt Polyamide PA66 GF25 chèn xốp "
         "nằm ở trung tâm ngăn cản tuyệt đối dòng truyền nhiệt cơ học từ ngoài vào trong.\n"
         "• Hiệu quả năng lượng: Giảm truyền nhiệt tới 40% - 50%, giúp tiết kiệm điện năng điều hòa lên đến 30% hàng năm.\n"
         "• Tương thích kính: Thiết kế tối ưu cho kính hộp Low-E cản nhiệt dày 24mm - 39mm bơm khí trơ Argon.\n"
         "• Vị trí khuyên dùng: Hướng Tây chịu nắng trực tiếp, biệt thự vùng núi Sapa, Đà Lạt hoặc phòng karaoke cách âm.")
    ]
    
    for title, desc in thermal_system:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    # Document Footer Note
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ VIRALWINDOW]", before=200, italic=True)
    
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Viralwindow.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Viralwindow Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_viralwindow_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Viralwindow"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_vra = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # VRA Accent (Gold/Beige)
    fill_vre = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # VRE Accent (Light Green)
    fill_vr = PatternFill(start_color="EAF2F8", end_color="EAF2F8", fill_type="solid") # VR Accent (Light Blue)
    fill_vrx = PatternFill(start_color="F5EEF8", end_color="F5EEF8", fill_type="solid") # VRX Accent (Light Purple)
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM ĐỒNG BỘ CAO CẤP VIRALWINDOW"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        "Phụ Kiện Đồng Bộ", 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Hệ VRA (Châu Á)", "VRA55", "Cửa đi & Cửa sổ mở quay VRA55", 
         "Cửa đi mở quay, cửa sổ mở quay, sổ mở hất, vách kính cố định (Cánh phẳng khung, kính lõm).", 
         "1.4mm - 2.0mm", "Viral đồng bộ, Draho, Hopo", "Gioăng EPDM, Keo Dow Corning", "Thiết kế hiện đại phẳng mặt ngoài nhôm sạch."),
         
        ("2", "Hệ VRA (Châu Á)", "VRA63", "Cửa xếp trượt VRA63", 
         "Cửa đi mở xếp trượt góc hoặc xếp gấp phẳng, tối ưu góc mở.", 
         "1.6mm - 1.8mm", "Phụ kiện trượt xếp Viral đồng bộ", "Gioăng EPDM chống nước, ke ép góc", "Bánh xe chịu lực chịu tải tốt chạy êm."),
         
        ("3", "Hệ VRA (Châu Á)", "VRA64", "Cửa sổ trượt VRA64", 
         "Cửa sổ mở trượt lùa 2 cánh, 4 cánh trượt ngang nhẹ nhàng.", 
         "1.2mm - 1.4mm", "Bánh xe đơn & khóa sập Viral", "Gioăng lông, keo bọt nở PU", "Giải pháp cửa sổ trượt kinh tế tối ưu diện tích."),
         
        ("4", "Hệ VRA (Châu Á)", "VRA94 / VRA80", "Cửa trượt VRA94", 
         "Cửa đi mở trượt lùa 2 ray có ray inox tròn dẫn hướng.", 
         "1.4mm - 1.8mm", "Bánh xe kép & Khóa đa điểm Viral", "EPDM bọc chèn mép kính", "Nhôm độ dày vừa, trượt êm không xệ."),
         
        ("5", "Hệ VRA (Châu Á)", "VRA120", "Cửa trượt VRA120", 
         "Cửa đi mở trượt lùa 3 ray mở rộng tới 2/3 không gian lối đi.", 
         "1.8mm - 2.0mm", "Phụ kiện khóa lùa đồng bộ Viral", "Gioăng EPDM đệm khít kín nước", "Phù hợp cửa đi lớn view ban công rộng."),
         
        ("6", "Hệ VRE (Châu Âu)", "VRE55", "Cửa mở quay tiêu chuẩn VRE55", 
         "Cửa sổ mở quay, mở hất hoặc cửa đi bản nhỏ (Cánh ôm khung gờ chỉ nổi).", 
         "1.4mm - 1.6mm", "Viral đồng bộ rãnh C, Cmech", "Gioăng EPDM 2 lớp kép dẻo dai", "Thiết kế gờ chỉ ôm khít chống thấm gió."),
         
        ("7", "Hệ VRE (Châu Âu)", "VRE65", "Cửa mở quay Premium VRE65", 
         "Cửa đi mở quay, cửa sổ mở quay lật (Hệ rãnh C tiêu chuẩn Châu Âu, cánh ôm khung).", 
         "1.6mm - 2.0mm", "Phụ kiện đa điểm Cmech, Roto", "EPDM đệm chèn 3 tầng kín khít", "Hệ cửa cách âm cao cấp nhất phân hệ VRE."),
         
        ("8", "Hệ VRE (Châu Âu)", "VRE75", "Cửa xếp gấp VRE75", 
         "Cửa đi mở xếp trượt xếp gấp bản rộng phân khúc cao cấp.", 
         "1.8mm - 2.2mm", "Phụ kiện trượt xếp Cmech, Roto", "Gioăng EPDM đặc biệt co giãn cực tốt", "Mở rộng tối ưu, lắp cho resort biệt thự biển."),
         
        ("9", "Hệ VRE (Châu Âu)", "VRE77", "Cửa xếp gấp ẩn bản lề VRE77", 
         "Cửa đi xếp trượt Luxury, bản lề được giấu kín hoàn toàn trong khe nhôm.", 
         "1.8mm - 2.2mm", "Bản lề ẩn & Khóa Cmech Luxury", "Gioăng EPDM kép chèn ép lực kín", "Không lộ bản lề tăng thẩm mỹ tối đa, chống trộm."),
         
        ("10", "Hệ VRE (Châu Âu)", "VRE94", "Cửa sổ trượt VRE94", 
         "Cửa sổ mở trượt lùa bản dày chịu lực gió bão tiêu chuẩn Châu Âu.", 
         "1.6mm - 1.8mm", "Khóa sập rãnh C, bánh xe chịu lực Cmech", "Gioăng nén EPDM chống mưa tạt bão", "Kết cấu góc vuông chịu lực gió bão cao."),
         
        ("11", "Hệ VRE (Châu Âu)", "VRE120", "Cửa trượt nâng Lift & Slide VRE120", 
         "Cửa đi mở trượt lùa 2 ray, có phụ kiện cơ cấu nâng hạ cánh kính hộp lớn.", 
         "2.0mm", "Phụ kiện nâng hạ CMECH, Roto chính hãng", "EPDM gioăng chèn nén chặt khi đóng", "Trượt siêu nhẹ cho cánh lớn nặng 350kg."),
         
        ("12", "Hệ VRE (Châu Âu)", "VRE180", "Cửa lùa siêu khổ VRE180", 
         "Cửa đi mở trượt lùa 3 ray tích hợp sẵn lưới chống muỗi/côn trùng inox 304.", 
         "2.0mm - 2.5mm", "Khóa đa điểm & lưới inox đồng bộ Viral", "Gioăng kép EPDM ôm kính", "Đón gió mát tự nhiên, ngăn muỗi côn trùng hiệu quả."),
         
        ("13", "Hệ VR Chuyên Dụng", "VR100", "Cửa mở trượt quay VR100", 
         "Cửa đi mở trượt kết hợp mở quay (Slide & Turn), thiết kế không ray dưới sàn.", 
         "1.6mm - 2.0mm", "Phụ kiện trượt quay đồng bộ Viral", "Gioăng EPDM ôm khít cánh giảm chấn", "Không ray dưới tránh vấp ngã, mở rộng 100% diện tích."),
         
        ("14", "Hệ VR Chuyên Dụng", "VR150", "Cửa thủy lực đại sảnh VR150", 
         "Cửa đi bản lề sàn (thủy lực), cánh lớn phào nổi nghệ thuật sang trọng.", 
         "2.0mm", "Bản lề sàn Adler/Hafele đồng bộ", "Keo kết cấu Dow Corning chịu uốn", "Bản nhôm siêu rộng 136mm kết hợp phào nổi nghệ thuật."),
         
        ("15", "Hệ Cầu Cách Nhiệt", "VRX75", "Cửa nhôm cầu cách nhiệt VRX75", 
         "Cửa đi và cửa sổ mở quay đa khoang có cầu polyamide PA66 cản nhiệt.", 
         "2.0mm", "Phụ kiện đa điểm Cmech rãnh C, Roto", "Gioăng EPDM đa khoang đệm chèn ép nước", "Dải cách nhiệt chèn xốp, lắp với kính hộp Low-E.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "VRA" in seg:
            seg_fill = fill_vra
        elif "VRE" in seg:
            seg_fill = fill_vre
        elif "Chuyên Dụng" in seg:
            seg_fill = fill_vr
        else:
            seg_fill = fill_vrx
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Viralwindow.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Viralwindow Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_viralwindow_docx()
    generate_viralwindow_xlsx()