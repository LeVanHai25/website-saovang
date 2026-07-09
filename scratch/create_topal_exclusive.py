# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_topal_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM TOPAL - AUSTDOOR")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp hệ thống cửa nhôm định hình đồng bộ Topal Prima, Slima, XFAD\nĐạt các tiêu chuẩn Châu Âu khắt khe về độ cách âm EN140 và độ kín nước tuyệt hảo EN1027")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM TOPAL (AUSTDOOR GROUP)", level=1)
    add_paragraph_with_spacing(doc, 
        "Topal là thương hiệu nhôm định hình cao cấp thuộc Tập đoàn Austdoor - tập đoàn hàng đầu Việt Nam về giải pháp cửa sổ và cửa đi. "
        "Với hệ thống 6 nhà máy quy mô lớn trải dài từ Bắc vào Nam (như nhà máy Mỹ Hào 100.000m², Nhơn Trạch 50.000m²) đạt công nghệ hiện đại, "
        "nhôm Topal đáp ứng năng lực cung ứng dồi dào trên 100.000 tấn/năm. Điểm nổi bật nhất của dòng sản phẩm Topal là thiết kế đồng bộ rãnh C Châu Âu "
        "cho phép lắp ráp phụ kiện không cần bắt vít cồng kềnh, tối ưu thẩm mỹ trơn phẳng hiện đại và tăng khả năng cách âm cách nhiệt vượt trội. "
        "Hệ thống sản phẩm của Topal đã vượt qua các bài kiểm chứng khắt khe và đạt tiêu chuẩn Châu Âu EN1027 (Kín nước tối đa) và EN140 (Cách âm tối đa)."
    )
    
    add_paragraph_with_spacing(doc, "Topal cung cấp các giải pháp nhôm định hình từ phân khúc phổ thông đến cao cấp:")
    
    # Table of 8 systems overview in Word
    table = doc.add_table(rows=9, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.6)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "Topal Prima Mở Quay", "Cửa mở quay Prima cao cấp", "1.3mm - 2.0mm", "Cửa đi & cửa sổ mở quay cánh trơn phẳng 63mm rãnh C Châu Âu. Cao cấp nhất."),
        ("2", "Topal Prima Trượt Lùa", "Cửa lùa trượt nâng Prima", "1.9mm - 2.0mm", "Cửa đi trượt nâng Lift & Slide chịu bão, cách âm cách nhiệt tối ưu."),
        ("3", "Topal Slima Mở Quay", "Cửa mở quay Slima trung cấp", "1.2mm - 1.5mm", "Cửa đi & sổ mở quay mỏng nhẹ cánh phẳng không gân tối giản thanh lịch."),
        ("4", "Topal Slima Trượt Lùa", "Cửa sổ trượt lùa Slima", "1.2mm - 1.5mm", "Cửa sổ trượt lùa mỏng gọn gàng tiết kiệm diện tích cho nhà phố."),
        ("5", "Topal XFAD Mở Quay", "Cửa mở quay Xingfa Topal", "1.4mm - 2.0mm", "Cửa đi/sổ có gân tăng cứng bề mặt theo biên dạng Xingfa Quảng Đông."),
        ("6", "Topal XFAD Trượt Lùa", "Cửa trượt lùa Xingfa Topal", "1.4mm - 2.0mm", "Cửa lùa trượt lồng cánh hệ Xingfa truyền thống của Topal."),
        ("7", "Topal XF Xếp Trượt", "Cửa đi xếp trượt gấp XF", "1.8mm - 2.0mm", "Cửa xếp gấp trượt xếp lùa hệ Xingfa của Topal cho lối mở rộng."),
        ("8", "Mặt Dựng Prima", "Vách mặt dựng đứng Prima", "2.0mm - 3.0mm", "Mặt dựng kính liên kết hệ mặt tiền các showroom tòa nhà văn phòng cao cấp.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM TOPAL", level=1)
    
    topal_details = [
        ("1. Topal Prima Cửa mở quay - Phân khúc cao cấp",
         "• Độ dày nhôm: 1.3mm - 1.4mm cho cửa sổ, 1.9mm - 2.0mm cho cửa đi.\n"
         "• Bản rộng khuôn bao: 63mm. Thiết kế cánh phẳng trơn không gân nổi hiện đại Châu Âu.\n"
         "• Điểm nhấn: Nẹp kính bo tròn mềm mại tạo tính mỹ thuật cao. Rãnh C tiêu chuẩn Châu Âu đồng bộ phụ kiện CMECH, Roto không cần khoan bắt vít thanh nhôm.\n"
         "• Khả năng cách âm: Đạt tiêu chuẩn EN140 cách âm tuyệt đối khỏi tiếng ồn ngoài phố.\n"
         "• Ứng dụng: Cửa ra vào ban công biệt thự, resort cao cấp."),
         
        ("2. Topal Prima Cửa lùa trượt nâng (Lift & Slide)",
         "• Độ dày nhôm: 1.9mm - 2.0mm bản cánh lớn cực kỳ chắc chắn.\n"
         "• Đặc tính: Cơ chế trượt nâng Lift & Slide, bánh xe chịu lực đẩy cánh kính nặng lướt êm ái trên ray inox tròn đúc. Khi hạ khóa cửa nén ép gioăng sát mặt sàn cản gió và chống nước tối đa.\n"
         "• Khả năng cản nước: Đạt tiêu chuẩn EN1027 chống rò rỉ nước mưa ở áp lực bão gió lớn."),
         
        ("3. Topal Slima Cửa mở quay - Phân khúc trung cấp tối giản",
         "• Độ dày nhôm: 1.2mm - 1.5mm.\n"
         "• Thiết kế tối giản: Bản cánh và khung bao mỏng gọn gàng, tăng diện tích lấy sáng cho ô kính.\n"
         "• Công năng: Phù hợp làm cửa sổ mở quay, cửa sổ mở hất và cửa đi 1 cánh phòng ngủ căn hộ chung cư hiện đại."),
         
        ("4. Topal Slima Cửa trượt lùa gọn gàng",
         "• Độ dày nhôm: 1.2mm - 1.5mm.\n"
         "• Công năng: Cửa sổ trượt lùa 2 cánh, 4 cánh mỏng nhẹ tinh tế.\n"
         "• Điểm nhấn: Tối ưu chi phí sản xuất và thi công lắp ráp nhanh cho nhà phố."),
         
        ("5. Topal XFAD Cửa mở quay - Biên dạng Xingfa chính hãng",
         "• Độ dày nhôm: 1.4mm cho cửa sổ, 2.0mm cho cửa đi.\n"
         "• Thiết kế Xingfa: Thân nhôm có đường gân nổi hai bên gia cường tăng độ cứng chống uốn xoắn.\n"
         "• Công năng: Cửa đi chính mở quay ra ngoài ban công, cửa sổ hất giống hệ Xingfa Quảng Đông.\n"
         "• Phụ kiện: Tương thích hệ phụ kiện rãnh 22 phổ thông dễ tìm trên thị trường."),
         
        ("6. Topal XFAD Cửa trượt lùa Xingfa",
         "• Độ dày nhôm: 1.4mm - 2.0mm.\n"
         "• Công năng: Cửa đi trượt lùa lồng cánh 2 ray, 3 ray dồn góc.\n"
         "• Ứng dụng: Cửa lùa ban công lớn nhà ống hoặc chung cư phân khúc phổ thông."),
         
        ("7. Topal XF Cửa xếp trượt gấp nhiều cánh",
         "• Độ dày nhôm: 1.8mm - 2.0mm.\n"
         "• Công năng: Cửa đi xếp lùa gấp xếp gọn nhiều cánh dồn góc mở rộng tối đa lối ra sân vườn đại sảnh.\n"
         "• Ưu điểm: Phụ kiện trục bản lề chịu tải treo trên chắc chắn vận hành ổn định."),
         
        ("8. Mặt Dựng Topal Prima cách âm cách nhiệt tốt",
         "• Độ dày nhôm: 2.0mm - 3.0mm tùy uốn mô-men xoắn lớn.\n"
         "• Công năng: Vách dựng đứng kính cường lực khổ lớn mặt tiền cho showroom văn phòng cao cấp.\n"
         "• Ưu điểm: Đạt độ kín khít chống thấm và cách nhiệt rất tốt nhờ dải gioăng đệm EPDM đa khoang bọc mép cấu trúc.")
    ]
    
    for title, desc in topal_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Materials and Gasket Specs ---
    add_heading_with_spacing(doc, "TIÊU CHUẨN KỸ THUẬT VÀ CHÍNH SÁCH BẢO HÀNH AUSTDOOR", level=1)
    add_paragraph_with_spacing(doc, 
        "Nhôm thanh định hình Topal được sản xuất từ hợp kim nhôm chất lượng cao đạt mác 6063-T5. "
        "Toàn bộ quy trình sơn tĩnh điện được kiểm soát tự động hóa đáp ứng tiêu chuẩn bền màu của Hiệp hội Sơn kiến trúc Hoa Kỳ AAMA. "
        "Topal cam kết chính sách bảo hành bề mặt sơn phủ chống bay màu và rỉ sét từ 5 năm cho dòng phổ thông "
        "đến 10 năm cho dòng cao cấp Prima.\n\n"
        "Quy chuẩn liên kết cơ khí và gioăng chèn:\n"
        "• Hệ gioăng cao su EPDM 2 thành phần chất lượng Châu Âu đàn hồi cao giúp kín gió kín nước tuyệt đối.\n"
        "• Ke ép góc nhôm đúc dày kết hợp bơm keo góc PU chuyên dụng ép góc khít 90 độ.\n"
        "• Phụ kiện đồng bộ Topal hoặc phụ kiện châu Âu cao cấp rãnh C mang lại tuổi thọ cơ học đóng mở bền bỉ vượt thời gian."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ TOPAL]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Topal.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Topal Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_topal_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Topal"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM ĐỒNG BỘ CAO CẤP TOPAL AUSTDOOR"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        'Phụ Kiện', 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Dòng Cao Cấp (Luxury)", "Topal Prima mở quay", "Cửa mở quay Prima cánh trơn phẳng 63mm", 
         "Cửa đi mở quay, cửa sổ mở quay hoặc hất rãnh C tiêu chuẩn Châu Âu cánh trơn không gân nẹp tròn.", 
         "1.3mm - 2.0mm", "Phụ kiện Prima đồng bộ rãnh C, Cmech", "Gioăng EPDM cao cấp, keo Dow Corning", "Đạt chuẩn cách âm EN140, kín nước EN1027 Châu Âu."),
         
        ("2", "Dòng Cao Cấp (Luxury)", "Topal Prima trượt lùa", "Cửa lùa trượt nâng Prima", 
         "Cửa đi trượt nâng Lift & Slide ray inox chịu lực va đập gió bão lớn ngoài ban công.", 
         "1.9mm - 2.0mm", "Phụ kiện nâng hạ Roto, CMECH, Prima", "EPDM đệm chèn ép lực kín khí", "Khả năng cản nước gió bão bão giật ven biển rất tốt."),
         
        ("3", "Dòng Trung Cấp (Slima)", "Topal Slima mở quay", "Cửa mở quay Slima mỏng nhẹ", 
         "Cửa sổ mở quay, cửa sổ hất và cửa đi 1 cánh phòng ngủ căn hộ hiện đại tối giản.", 
         "1.2mm - 1.5mm", "Phụ kiện Slima đồng bộ rãnh C, Sigico", "Gioăng EPDM kép dẻo dai", "Khung cánh phẳng trơn không gân thanh thoát tăng diện tích kính."),
         
        ("4", "Dòng Trung Cấp (Slima)", "Topal Slima trượt lùa", "Cửa lùa trượt sổ/đi Slima", 
         "Cửa sổ và đi trượt lùa mỏng nhẹ tiết kiệm diện tích cho nhà phố.", 
         "1.2mm - 1.5mm", "Bánh xe & khóa sập Slima đồng bộ", "Gioăng lông cản bụi chống ồn", "Tối ưu hóa chi phí sản xuất và lắp đặt."),
         
        ("5", "Dòng Phổ Thông (XFAD)", "Topal XFAD mở quay", "Cửa mở quay Xingfa Topal", 
         "Cửa đi/sổ mở quay bản gân tăng cứng bề mặt giống Xingfa Quảng Đông.", 
         "1.4mm - 2.0mm", "Phụ kiện rãnh 22 phổ thông Kinlong, Draho", "Gioăng cao su EPDM kép", "Thiết kế gá lắp gân chịu lực bão gió tốt phổ biến."),
         
        ("6", "Dòng Phổ Thông (XFAD)", "Topal XFAD trượt lùa", "Cửa trượt lùa Xingfa Topal", 
         "Cửa đi lùa trượt lồng cánh hệ Xingfa truyền thống của Topal.", 
         "1.4mm - 2.0mm", "Bánh xe kép & chốt âm Kinlong, Draho", "Gioăng lông, keo bọt nở PU", "Giải pháp cửa trượt lùa phổ thông kinh tế."),
         
        ("7", "Dòng Phổ Thông (XF)", "Topal XF xếp trượt", "Cửa xếp gấp hệ Xingfa Topal", 
         "Cửa xếp trượt lùa gấp gọn nhiều cánh dồn góc mở tối đa lối đi sảnh biệt thự.", 
         "1.8mm - 2.0mm", "Phụ kiện trượt xếp Roto, Hopo chính hãng", "Gioăng EPDM đệm đàn hồi cao", "Ray treo trên dẫn hướng dưới chắc chắn chịu lực."),
         
        ("8", "Dòng Dự Án (B2B)", "Mặt Dựng Prima", "Vách mặt dựng kính Prima", 
         "Vách dựng đứng kính cường lực khổ lớn cho showroom văn phòng cao cấp.", 
         "2.0mm - 3.0mm", "Ke liên kết mặt tiền chịu tải lực chuyên dụng", "Gioăng EPDM nhiều khoang, keo dán Dow", "Cách âm cách nhiệt và cản gió bão lớn tuyệt đối.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Phổ Thông" in seg:
            seg_fill = fill_economy
        elif "Cao Cấp" in seg or "Luxury" in seg:
            seg_fill = fill_premium
        else:
            seg_fill = fill_standard
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Topal.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Topal Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_topal_docx()
    generate_topal_xlsx()
