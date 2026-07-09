# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng Kogen:
1. Catalogue_HeCuaNhom_Kogen.xlsx
2. ThuVien_HeCuaNhom_Kogen.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_Kogen.xlsx
# ─────────────────────────────────────────────
def create_kogen_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue Kogen"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM CAO CẤP PHONG CÁCH ĐỨC KOGEN"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="0D2240")
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    # Kogen Data
    kogen_rows = [
        ('1', 'Xếp Trượt (Luxury)', 'Kogen Slim 50/68', 'Cửa xếp lùa ẩn bản lề Slim', 
         'Cửa đi xếp trượt Slim ẩn bản lề độc đáo (Hidden Hinge), khi đóng không nhìn thấy bản lề, giải phóng tối đa tầm nhìn.', 
         '1.6mm', 'Phụ kiện rãnh C đồng bộ Kogen ẩn bản lề', 'Gioăng EPDM đa khoang + Silicon Dow', 
         'Bề mặt sơn phủ hạt mịn bóng mờ phong cách Bauhaus tối giản sang trọng.'),
         
        ('2', 'Slim Lùa (Premium)', 'Kogen Slim Lùa', 'Cửa trượt lùa Slim Panorama', 
         'Cửa đi trượt lùa siêu mỏng cánh trơn phẳng, ray inox SUS316L trượt nhẹ êm ái, kính cường lực khổ lớn lấy trọn view.', 
         '2.0mm', 'Ray inox chịu tải & Bánh xe chuyên dụng Kogen', 'Gioăng EPDM kép chèn ép lực chìm', 
         'Thiết kế cánh siêu mảnh tăng diện tích kính tối đa hướng ngoại.'),
         
        ('3', 'Casement (Premium)', 'Kogen 60 Đa khoang', 'Cửa đi + cửa sổ mở quay hệ 60', 
         'Cửa sổ/đi mở quay vào hoặc ra ngoài, kết cấu đa khoang tăng cứng cách âm cách nhiệt tốt, chống rung chấn.', 
         '1.6mm', 'Bản lề rãnh C CMECH & khóa đa điểm Hopo', 'Gioăng EPDM hàng không 3 lớp', 
         'Giải pháp hệ rãnh C Châu Âu đồng bộ phù hợp nhà phố phong cách Đức hiện đại.')
    ]
    
    fill_row = PatternFill(fill_type="solid", fgColor="EDF7ED")
    for r_idx, row_vals in enumerate(kogen_rows, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_Kogen.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel Kogen saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel Kogen to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_Kogen.docx
# ─────────────────────────────────────────────
def create_kogen_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x0D,0x22,0x40))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM CAO CẤP KOGEN", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu nghiên cứu R&D, thuyết minh giải pháp kỹ thuật hệ cửa nhôm cao cấp phong cách Bauhaus (Đức)\nỨng dụng kết cấu đa khoang rãnh C Châu Âu và công nghệ bản lề ẩn độc đáo", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM CAO CẤP KOGEN (PHONG CÁCH ĐỨC)", bold=True, size=14, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    add_p(
        "Nhôm Kogen là dòng cửa nhôm hệ cao cấp được phát triển bởi Kenwin Group tại Việt Nam, lấy cảm hứng từ triết lý kiến trúc Bauhaus của Đức: "
        "tối giản, tinh tế, loại bỏ những chi tiết rườm rà để tập trung hoàn toàn vào công năng và độ bền bỉ cơ học. "
        "Điểm đặc trưng làm nên vị thế của Kogen chính là hệ thống xếp gấp rãnh C tích hợp dải ẩn bản lề độc đáo (Hidden Hinge), "
        "khi đóng cửa lại toàn bộ bản lề sẽ nằm gọn trong rãnh nhôm tạo nên một bề mặt nhôm phẳng khít sang trọng. "
        "Với kết cấu nhôm đa khoang chịu lực vượt trội, kết hợp lớp sơn tĩnh điện bề mặt mịn hạt mịn cao cấp bảo vệ thanh nhôm trước muối biển, "
        "Kogen đang trở thành lựa chọn hàng đầu cho các công trình thiết kế hiện đại như biệt thự, penthouses, resort cao cấp tại Việt Nam."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM KOGEN", bold=True, size=14, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    
    add_p("1. Kogen Hệ cửa xếp lùa ẩn bản lề Slim 50 / 68", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.6mm cứng vững tối ưu tải trọng.")
    add_bullet("Đặc tính bản lề ẩn (Hidden Hinge): Khi đóng cửa bản lề ẩn hoàn toàn trong rãnh cánh giúp tăng tính thẩm mỹ, chống bụi bẩn và ngăn nước mưa mài mòn bản lề.")
    add_bullet("Thiết kế cánh Slim mỏng: Bản rộng cánh siêu nhỏ giúp lấy trọn view bên ngoài, mở rộng không gian tối đa.")
    add_bullet("Ứng dụng: Lối ra hồ bơi biệt thự, ngăn sảnh phòng khách vườn, lối ra ban công penthouses.")
    
    add_p("2. Kogen Hệ cửa trượt lùa Slim Panorama", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm cứng vững.")
    add_bullet("Thiết kế siêu mỏng: Khung và cánh mảnh trơn nhẵn mang phong cách Minimalist Châu Âu.")
    add_bullet("Ray lùa đặc chủng: Sử dụng ray bằng inox SUS316L kết hợp bánh xe chịu tải trọng lớn lướt nhẹ cực êm.")
    add_bullet("Ứng dụng: Cửa kính ban công lớn biệt thự, penthouse đón ánh sáng và không gian panorama.")
    
    add_p("3. Kogen 60 Hệ cửa mở quay đa khoang rãnh C", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.6mm vách nhôm đa khoang.")
    add_bullet("Cấu trúc đa khoang: Tăng khả năng chịu tải trọng gió rung, cách âm vượt trội lên đến 35-38dB.")
    add_bullet("Rãnh C chuẩn Châu Âu: Tương thích hoàn hảo với phụ kiện cao cấp CMECH (Mỹ), Hopo, Sigico.")
    add_bullet("Ứng dụng: Cửa sổ mở quay, cửa đi thông phòng nhà phố hiện đại, biệt thự cao cấp.")

    # Bảng tóm tắt
    add_p("BẢNG TỔNG HỢP THÔNG SỐ KỸ THUẬT HỆ NHÔM KOGEN", bold=True, size=12, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    
    hdr = tbl.rows[0].cells
    hdr[0].text = "STT"
    hdr[1].text = "Mã Hệ"
    hdr[2].text = 'Tên Hệ'
    hdr[3].text = 'Độ Dày (mm)'
    hdr[4].text = "Loại Sản Phẩm & Công Năng Tiêu Biểu"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    kogen_tbl_data = [
        ("1", "Kogen Slim 50/68", "Cửa xếp lùa ẩn bản lề Slim", "1.6mm", "Cửa đi xếp lùa ẩn bản lề, cánh cực mỏng phong cách hiện đại."),
        ("2", "Kogen Slim Lùa", "Cửa trượt lùa Slim Panorama", "2.0mm", "Cửa lùa phẳng panorama cánh mảnh ray inox trượt êm."),
        ("3", "Kogen 60", "Cửa mở quay đa khoang hệ 60", "1.6mm", "Cửa sổ/đi mở quay kết cấu đa khoang cách âm rãnh C Châu Âu.")
    ]
    
    for row in kogen_tbl_data:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph()
    add_p("Tài liệu được bảo mật bởi Phòng R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_Kogen.docx")
    try:
        doc.save(path)
        print(f"  >> Docx Kogen saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx Kogen to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone KOGEN docx & xlsx in BASE directory...")
    print("=" * 60)
    create_kogen_xlsx()
    create_kogen_docx()
    print("=" * 60)
