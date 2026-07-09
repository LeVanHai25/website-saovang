# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_civro_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM CẦU CÁCH NHIỆT CIVRO")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp hệ thống cửa nhôm cầu cách nhiệt Technoform BAUTEC siêu cao cấp\nỨng dụng công nghệ xử lý bề mặt Matte Electrodeposition (MED) độc quyền từ CHLB Đức")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU THƯƠNG HIỆU NHÔM SIÊU CAO CẤP CIVRO (ĐỨC)", level=1)
    add_paragraph_with_spacing(doc, 
        "Civro GmbH là nhà phát triển hệ thống cửa nhôm, vách mặt dựng siêu cao cấp có trung tâm R&D đặt tại Würzburg, "
        "Đức và hoạt động toàn cầu từ năm 1993. Triết lý thiết kế của CIVRO dựa trên ba nền tảng cốt lõi: an ninh bảo mật, "
        "bảo vệ môi trường và phát triển bền vững. CIVRO tiên phong đập tan giới hạn cách nhiệt truyền thống bằng việc ứng dụng "
        "dải cầu cách nhiệt Polyamide Technoform BAUTEC (Đức) giữa hai khoang nhôm để ngăn chặn hoàn toàn truyền nhiệt. "
        "Ngoài ra, thanh profile CIVRO được hoàn thiện bề mặt bằng công nghệ Matte Electrodeposition (MED) độc quyền cho nước sơn "
        "mịn mờ chống xước, kháng muối mặn muối biển cấp độ cao nhất. CIVRO chính là giải pháp kiến trúc thượng lưu không thể "
        "thiếu trong các căn biệt thự triệu đô, penthouses đẳng cấp và resort 6 sao tại Việt Nam."
    )
    
    add_paragraph_with_spacing(doc, "CIVRO cung cấp hệ thống sản phẩm cách âm cách nhiệt tích hợp cao bao gồm:")
    
    # Table of 6 systems overview in Word
    table = doc.add_table(rows=7, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.6)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "AW55 / AW65 / AW75", "Cửa sổ cầu cách nhiệt mở quay", "1.4mm - 1.8mm", "AW55IN/OU, AW65IN/OU, AW75IN mở vào/ra cách âm 45dB nhờ 3 lớp gioăng hàng không."),
        ("2", "AD55 / AD65 / AD75", "Cửa đi mở quay cách nhiệt", "2.0mm", "AD55IN/OU, AD65IN/OU, AD75IN/OU đi chính rãnh C khóa đa điểm an ninh RC2 chống trộm."),
        ("3", "CSD80 trượt nâng", "Trượt nâng bản mỏng 80mm", "1.8mm - 2.2mm", "Cửa trượt nâng Lift & Slide tối giản khung bao tăng ô kính panorama."),
        ("4", "CSD130 trượt nâng", "Trượt nâng siêu chịu lực 130", "2.0mm - 3.0mm", "Cửa đi lùa nâng Lift & Slide cánh lớn chịu tải 400kg, tích hợp kính hộp cách âm."),
        ("5", "CSD xếp trượt", "Cửa xếp gấp trượt cánh lớn", "2.0mm", "Cửa xếp gấp nhiều cánh dồn góc chạy ray treo trên chịu lực cực êm nhẹ."),
        ("6", "Vách mặt dựng", "Vách dựng kính cách nhiệt", "2.5mm - 3.0mm", "Vách dựng kính khổ lớn liên kết đố cầu cách nhiệt Technoform cản bức xạ mặt trời.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM CIVRO", level=1)
    
    civro_details = [
        ("1. Civro Hệ cửa sổ mở quay cách nhiệt AW55 / AW65 / AW75",
         "• Độ dày nhôm: 1.4mm cho cửa hất ngoài, 1.8mm cho cửa đi kèm.\n"
         "• Tên mã cụ thể: AW55IN (mở vào), AW55OU (mở ra ngoài), AW65IN/OU, AW75IN.\n"
         "• Dải cách nhiệt: Technoform BAUTEC đặt trong lòng khoang nhôm cản khí nóng xâm nhập.\n"
         "• Hệ gioăng: 3 lớp gioăng cao su EPDM tiêu chuẩn ngành hàng không cách âm đạt tới 45dB, cản bụi mịn PM2.5.\n"
         "• Ứng dụng: Cửa sổ phòng ngủ, phòng làm việc biệt thự cần yên tĩnh tuyệt đối."),
         
        ("2. Civro Hệ cửa đi mở quay an ninh AD55 / AD65 / AD75",
         "• Độ dày nhôm: 2.0mm cứng vững.\n"
         "• Phụ kiện: Tương thích phụ kiện cao cấp Sobinco (Bỉ) hoặc CMECH (Mỹ) chốt đa điểm chống cạy phá đạt chuẩn EN RC2.\n"
         "• Ứng dụng: Cửa đi chính phòng khách biệt thự, cửa lối phụ ban công."),
         
        ("3. Civro Hệ trượt nâng Lift & Slide CSD80 tối giản",
         "• Độ dày nhôm: 1.8mm - 2.2mm.\n"
         "• Bản khuôn bao: 80mm gọn gàng.\n"
         "• Công năng: Cửa lùa trượt nâng hạ ray Inox tròn êm nhẹ phẳng sàn.\n"
         "• Thẩm mỹ: Triết lý Minimalist thu hẹp tối đa tiết diện khung nhôm bao để tăng tầm nhìn panorama bể bơi, sân vườn."),
         
        ("4. Civro Hệ trượt nâng siêu chịu lực CSD130 cánh lớn",
         "• Độ dày nhôm: 2.0mm - 3.0mm chịu uốn va đập cực lớn.\n"
         "• Bản khuôn bao: 130mm vững chãi chịu mô-men xoắn gió lớn bão giật ven biển.\n"
         "• Sức tải bánh xe: Đạt tới 400kg/cánh kính, cho phép sản xuất những cánh cửa trượt khổ khổng lồ cao đến 3.5m rộng 2.5m.\n"
         "• Kính áp dụng: Phù hợp lắp kính hộp Double Glazing Low-E bơm khí trơ cách nhiệt tốt nhất."),
         
        ("5. Civro Hệ cửa xếp gấp trượt nhiều cánh",
         "• Độ dày nhôm: 2.0mm bản cánh to dập thủy lực chắc chắn.\n"
         "• Cơ chế trượt: Ray treo chịu lực trên và dẫn hướng ray dưới, đóng mở xếp gọn dồn góc 95% mở toang sảnh biệt thự rộng mở ra sân vườn."),
         
        ("6. Civro Vách mặt dựng cầu cách nhiệt ngoài trời",
         "• Độ dày nhôm: 2.5mm - 3.0mm tùy tải trọng gió đứng.\n"
         "• Điểm nhấn: Đố đứng tích hợp dải Technoform cản nhiệt cản ánh nắng mặt trời bức xạ nhiệt truyền vào làm mát máy lạnh bên trong tòa nhà kính penthouse.")
    ]
    
    for title, desc in civro_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Materials and Gasket Specs ---
    add_heading_with_spacing(doc, "CÔNG NGHỆ BỀ MẶT MED & CHẤT LƯỢNG SIÊU CAO CẤP", level=1)
    add_paragraph_with_spacing(doc, 
        "Điểm độc đáo tạo nên danh tiếng CIVRO là công nghệ phủ bề mặt Matte Electrodeposition (MED) độc quyền. "
        " MED tạo ra một màng sơn điện di mờ cực mịn bảo vệ chống muối biển ăn mòn trầy xước, không bám bẩn vân tay "
        "và ngăn ngừa khúc xạ ánh sáng chói mắt bảo vệ thị lực. Toàn bộ thanh nhôm CIVRO có tuổi thọ vận hành trên 30 năm "
        "đảm bảo không lão hóa bay màu sơn dưới bức xạ mặt trời.\n\n"
        "Quy chuẩn vật tư phụ chèn ép góc:\n"
        "• Ke góc nhôm đúc nguyên chất liên kết vít Inox kết hợp bơm keo góc PU của Đức chống nước tuyệt hảo.\n"
        "• Hệ gioăng cao su EPDM 3 lớp tiêu chuẩn hàng không độ đàn hồi nén dẻo dai dài lâu.\n"
        "• Phụ kiện rãnh C đồng bộ cao cấp lắp đặt âm chìm tăng độ an ninh bảo mật chống cạy cửa."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ CIVRO]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Civro.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Civro Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_civro_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Civro"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM CẦU CÁCH NHIỆT SIÊU CAO CẤP CIVRO (ĐỨC)"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        'Phụ Kiện', 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Cầu Cách Nhiệt (Luxury)", "AW55IN / AW55OU", "Cửa sổ mở quay cách nhiệt hệ 55", 
         "Cửa sổ mở quay vào trong (IN) hoặc ra ngoài (OU) cầu cách nhiệt dải Polyamide Technoform cách âm tới 45dB.", 
         "1.4mm - 1.8mm", "Phụ kiện Sobinco (Bỉ) rãnh C đồng bộ", "Gioăng EPDM hàng không 3 lớp", "Sơn mờ phủ MED độc quyền Đức chống vân tay trầy xước."),
         
        ("2", "Cầu Cách Nhiệt (Luxury)", "AW65IN / AW65OU", "Cửa sổ mở quay cách nhiệt hệ 65", 
         "Cửa sổ mở quay hất rãnh C Châu Âu bản dày 65mm lắp đặt kính hộp dày cách nhiệt tốt.", 
         "1.4mm - 1.8mm", "Phụ kiện CMECH (Mỹ) chốt đa điểm âm", "Gioăng cao su EPDM 3 lớp", "Chịu áp lực bão gió cấp độ cao cực kỳ kín khít."),
         
        ("3", "Cầu Cách Nhiệt (Luxury)", "AW75IN", "Cửa sổ mở quay cánh lớn hệ 75", 
         "Cửa sổ mở quay vào trong có lưới chống côn trùng tích hợp, cách nhiệt cách âm tối đa.", 
         "1.8mm", "Phụ kiện Sobinco, tay nắm đặc chủng", "Gioăng cao su đúc đa khoang", "Đạt chuẩn an ninh RC2 chống cạy khóa đột nhập."),
         
        ("4", "Cầu Cách Nhiệt (Luxury)", "AD55 / AD65 / AD75", "Cửa đi mở quay cách nhiệt đại sảnh", 
         "Cửa đi mở quay 1 cánh, 2 cánh, 4 cánh rãnh C Châu Âu bản cánh to chắc chắn vững chãi.", 
         "2.0mm", "Bản lề chịu tải CMECH & khóa đa điểm", "Gioăng EPDM kép chèn ép lực chìm", "Lối đi chính sang trọng lâu đài biệt thự sân vườn lớn."),
         
        ("5", "Trượt Nâng (Minimalist)", "CSD80", "Cửa trượt lùa nâng bản mỏng 80mm", 
         "Cửa đi trượt nâng Lift & Slide bản mỏng thu gọn khuôn bao tăng diện tích ô kính panorama.", 
         "1.8mm - 2.2mm", "Bánh xe nâng & tay khóa CMECH", "Gioăng nén EPDM phẳng mặt sàn", "Cơ cấu trượt nâng gạt hạ nén gioăng chống gió bão biển."),
         
        ("6", "Trượt Nâng (Luxury)", "CSD130", "Cửa trượt nâng siêu cánh rộng 130mm", 
         "Cửa đi trượt lùa nâng Lift & Slide chịu tải trọng cánh 400kg siêu khổ, cao đến 3.5m rộng 2.5m.", 
         "2.0mm - 3.0mm", "Phụ kiện nâng hạ Sobinco, khóa RC2", "Gioăng EPDM nén chặt cách âm tốt", "Chuyên biệt lắp kính hộp Double Low-E dày penthouses."),
         
        ("7", "Cầu Cách Nhiệt (Luxury)", "CSD xếp trượt", "Cửa xếp gấp lùa cách nhiệt dồn góc", 
         "Cửa đi xếp trượt nhiều cánh trượt xếp gọn gàng mở toang 95% diện tích không gian.", 
         "2.0mm", "Phụ kiện xếp trượt chịu tải CMECH bệ treo", "Gioăng nén đàn hồi co giãn nhiệt", "Cơ cấu trượt treo cực khỏe êm nhẹ không rung lắc."),
         
        ("8", "Mặt Dựng Cầu Cách Nhiệt", "Mặt dựng đố đứng đứng cách nhiệt", "Vách mặt dựng đứng kính Panorama", 
         "Vách kính ngăn ngoài trời cho biệt thự kính Panorama đố đứng dải Polyamide cản nhiệt.", 
         "2.5mm - 3.0mm", "Ke liên kết chịu lực chuyên dụng Đức", "Silicon kết cấu Dow, gioăng đa khoang", "Giải pháp kính chống nóng bức xạ mặt trời tối ưu điện năng.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Minimalist" in seg:
            seg_fill = fill_standard
        else:
            seg_fill = fill_premium
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Civro.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Civro Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_civro_docx()
    generate_civro_xlsx()
