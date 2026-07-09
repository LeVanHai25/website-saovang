# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# Auto-install openpyxl and python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. GENERATE EXCLUSIVE WORD DOCUMENT
def generate_slim_docx():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM KÍNH SLIM TỐI GIẢN")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tài liệu kỹ thuật tổng hợp hệ cửa lùa trượt treo giảm chấn thủy lực và cửa mở quay Slim\nGiải pháp kiến trúc nội ngoại thất tối giản (Minimalism) tăng tối đa tầm nhìn Panorama")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Introduction ---
    add_heading_with_spacing(doc, "GIỚI THIỆU HỆ CỬA NHÔM VIỀN SIÊU MỎNG SLIM", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ cửa nhôm Slim là giải pháp cửa kính khung nhôm tiên phong đại diện cho phong cách thiết kế tối giản hiện đại (Minimalism). "
        "Điểm đặc thù nhận diện lớn nhất của hệ Slim là thiết kế viền thanh cánh nhôm siêu mỏng (viền chỉ rộng 16mm - 25mm, "
        "giúp thu nhỏ tiết diện khung tới 60% so với hệ nhôm Xingfa thông thường) nhằm tối đa hóa diện tích lấy sáng và mở rộng tầm nhìn "
        "kính Panorama không giới hạn. Khung nhôm tuy siêu mảnh nhưng được chế tạo từ hợp kim nhôm 6063-T5 dày dặn (1.2mm - 2.0mm) "
        "đảm bảo độ cứng cơ học cao. Đi kèm cửa lùa Slim là bộ piston giảm chấn thủy lực đóng mở tự động hãm tốc độ đóng chậm lại nhẹ nhàng, "
        "triệt tiêu hoàn toàn rung chấn va đập nứt vỡ kính và chống kẹt tay tuyệt hảo cho trẻ em."
    )
    
    add_paragraph_with_spacing(doc, "Hệ cửa nhôm Slim được phân loại theo công năng hoạt động chi tiết:")
    
    # Table of 6 systems overview in Word
    table = doc.add_table(rows=7, cols=5)
    table.alignment = 1 # Center
    
    headers = ["STT", "Mã Phân Hệ", 'Tên Hệ', 'Độ Dày (mm)', "Loại Sản Phẩm & Công Năng Tiêu Biểu"]
    col_widths = [Inches(0.5), Inches(1.3), Inches(2.0), Inches(1.2), Inches(2.5)]
    
    hdr_row = table.rows[0]
    for idx, name in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    systems_data = [
        ("1", "Slim Lùa Treo", "Cửa lùa treo giảm chấn", "1.6mm - 2.0mm", "Cửa đi lùa trượt treo ray trên không có ray dưới sàn, tránh vấp ngã chân, giảm chấn đóng mở 2 chiều."),
        ("2", "Slim Ray Dẹt", "Cửa lùa ray âm/ray dẹt", "1.6mm", "Cửa đi lùa nhiều cánh liên kết lướt êm ái trên ray inox siêu phẳng mỏng dẹt nằm sát mặt sàn gỗ/đá."),
        ("3", "Slim Ngoại Thất", "Cửa lùa Slim ngoài trời", "2.0mm", "Hệ cánh slim gia cường chịu bão, dập gioăng EPDM nhiều lớp chống nước ngấm và dùng kính hộp cách nhiệt."),
        ("4", "Slim Mở Quay", "Cửa mở quay Slim tối giản", "1.2mm - 1.4mm", "Cửa sổ hất viền mỏng hoặc cửa đi 1 cánh phòng ngủ thông phòng thanh lịch."),
        ("5", "Slim Pocket", "Cửa lùa âm tường Pocket", "1.6mm", "Cửa lùa trượt giấu cánh chìm hoàn toàn vào khoang hộp vách tường thạch cao tối ưu diện tích hẹp."),
        ("6", "Slim Vách Tĩnh", "Vách ngăn thạch cao kính Slim", "1.2mm", "Vách tĩnh phân vùng phòng ngủ - làm việc thẩm mỹ tối giản sang trọng.")
    ]
    
    for row_idx, s_data in enumerate(systems_data, start=1):
        row = table.rows[row_idx]
        bg_color = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(s_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if col_idx in [0, 1, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    doc.add_page_break()
    
    # --- Chi tiết các hệ nhôm ---
    add_heading_with_spacing(doc, "CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC PHÂN HỆ SLIM", level=1)
    
    slim_details = [
        ("1. Cửa lùa treo Slim giảm chấn thủy lực không ray dưới (Nội thất)",
         "• Độ dày nhôm: 1.6mm - 2.0mm, đố cánh đứng siêu mảnh rộng 16mm.\n"
         "• Cơ chế trượt: Ray treo trên bắt chắc chắn vào trần thạch cao/bê tông gia cố thép hộp. Hoàn toàn không có ray dưới sàn nhà giúp việc lau chùi và quét dọn thuận tiện, không cản trở xe lăn/xe đẩy.\n"
         "• Giảm chấn 2 chiều: Tích hợp bộ Piston hãm thủy lực âm chìm trong thanh ray, khi cửa chạy gần mép sẽ tự động hãm phanh chậm lại kéo nhẹ khít kín khung.\n"
         "• Ứng dụng: Cửa phòng thay đồ, cửa ngăn phòng bếp tránh mùi bay ra phòng khách."),
         
        ("2. Cửa lùa Slim liên kết ray dẹt (Trượt nhẹ cánh rộng)",
         "• Độ dày nhôm: 1.6mm.\n"
         "• Bản ray dưới: Ray Inox tròn dẹt nổi lên chỉ 3mm hoặc âm chìm phẳng bằng mặt sàn. Cửa lùa 3 cánh, 4 cánh liên kết kéo cánh đầu tiên thì các cánh sau tự chạy theo nhẹ nhàng.\n"
         "• Ứng dụng: Cửa ngăn phân chia phòng ngủ lớn hoặc sảnh rộng ban công căn hộ."),
         
        ("3. Cửa lùa Slim ngoại thất cản bão chịu nước (Ngoại thất)",
         "• Độ dày nhôm: 2.0mm cứng vững bản lớn hơn (viền nhôm rộng 25mm - 35mm).\n"
         "• Đặc thù: Cửa đi lùa ngoài trời cần chống thấm nước mưa và gió bão. Profile thiết kế các rãnh chèn dải gioăng cao su EPDM 3 tầng kín khít.\n"
         "• Kính: Lắp được kính hộp Low-E cách âm cách nhiệt tốt dày tới 20mm - 24mm."),
         
        ("4. Cửa mở quay Slim tối giản viền mỏng",
         "• Độ dày nhôm: 1.2mm - 1.4mm.\n"
         "• Bản cánh: Bản cánh mỏng bo viền kính. Sử dụng phụ kiện bản lề ẩn giấu chìm chốt đa điểm thẩm mỹ cao.\n"
         "• Ứng dụng: Cửa thông phòng ngủ 1 cánh, cửa sổ mở quay hất chung cư tối giản."),
         
        ("5. Cửa lùa âm tường Slim (Pocket Door)",
         "• Độ dày nhôm: 1.6mm.\n"
         "• Cơ cấu Pocket: Cửa trượt lùa đẩy ngang giấu ẩn chìm 100% cánh vào trong hộp vách tường thạch cao hoặc vách gỗ nội thất, giải phóng hoàn toàn lối đi rộng mở.\n"
         "• Ứng dụng: Phòng tắm hẹp, phòng giặt quần áo căn hộ studio tối giản."),
         
        ("6. Vách ngăn kính khung Slim tĩnh",
         "• Độ dày nhôm: 1.2mm mỏng nhẹ.\n"
         "• Công năng: Vách ngăn kính tĩnh cố định bọc viền nhôm đen mờ hoặc vàng titan sang trọng lấy ánh sáng xuyên thấu toàn bộ căn nhà.")
    ]
    
    for title, desc in slim_details:
        add_heading_with_spacing(doc, title, level=2)
        add_paragraph_with_spacing(doc, desc)
        
    doc.add_page_break()

    # --- Hardware and Installation Specs ---
    add_heading_with_spacing(doc, "YÊU CẦU KỸ THUẬT LẮP ĐẶT VÀ PHỤ KIỆN HỆ SLIM", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ cửa Slim là dòng cửa kỹ thuật có viền nhôm rất mảnh nên độ an toàn phụ thuộc rất lớn vào kết cấu phụ kiện "
        "và tay nghề gia công lắp dựng:\n\n"
        "1. Gia cố trần treo: Đối với hệ lùa treo không ray dưới, bắt buộc đơn vị thi công phải gia cố dầm sắt hộp "
        "hoặc bê tông trên trần thạch cao chịu tải trọng treo nặng của toàn bộ cánh cửa kính cường lực (từ 80kg - 150kg).\n"
        "2. Bộ giảm chấn thủy lực (Soft closing damper): Phải lắp đặt bộ hãm thủy lực đồng bộ cao cấp (như OPK, Papo, Heckler) "
        "đảm bảo hành trình giảm tốc 2 đầu mở và đóng êm ái.\n"
        "3. Kính: Bắt buộc dùng kính cường lực Temper an toàn dày từ 8mm - 10mm mài cạnh bóng nhẵn.\n"
        "4. Bề mặt nhôm: Khuyên dùng xi mạ Anodized xước kim loại màu vàng titan, đen mờ bền bỉ không bám bẩn vân tay."
    )
    
    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU KHẢO SÁT CHUYÊN ĐỀ SLIM]", before=200, italic=True)
    
    # Save Word document
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\ThuVien_HeCuaNhom_Slim.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Slim Exclusive DOCX successfully saved to {output_path}")

# 2. GENERATE EXCLUSIVE EXCEL CATALOGUE
def generate_slim_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Catalogue Slim"
    
    # Show gridlines
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid") # Deep Navy
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Light Green for standard
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Light Gold for luxury/design
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Light Pink for standard
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_title = Alignment(horizontal="center", vertical="center")
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)

    # Title block
    ws.merge_cells("A1:I2")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC PHÂN HỆ CỬA NHÔM KÍNH HỆ SLIM TỐI GIẢN"
    title_cell.font = font_title
    title_cell.alignment = align_title
    
    # Headers
    headers = [
        "STT", 
        'Nhóm Hệ', 
        'Mã Hệ', 
        'Tên Hệ', 
        'Loại Sản Phẩm', 
        'Độ Dày (mm)',
        "Phụ Kiện Chuyên Dụng", 
        'Gioăng & Keo', 
        'Đặc Điểm Kỹ Thuật'
    ]
    
    for col_idx, h_text in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h_text
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border

    # Data Rows
    data = [
        ("1", "Trượt Treo (Nội thất)", "Slim Lùa Treo", "Cửa lùa trượt treo không ray dưới sàn", 
         "Cửa đi lùa trượt treo trên, đóng mở giảm chấn thủy lực 2 đầu chống va đập, đi lại êm ái.", 
         "1.6mm - 2.0mm", "Phụ kiện bánh xe treo chịu tải OPK, Papo", "Gioăng cao su đệm kính chèn mép", "Trần treo thạch cao bắt buộc phải gia cố sắt hộp chịu lực nặng."),
         
        ("2", "Ray Dẹt (Nội thất)", "Slim Ray Dẹt", "Cửa lùa trượt ray dẹt/âm sàn", 
         "Cửa lùa nhiều cánh liên kết chạy ray dẹt nổi 3mm hoặc ray âm phẳng nền, giảm chấn đóng mở.", 
         "1.6mm", "Bánh xe dưới chịu lực & bộ dẫn hướng OPK", "Gioăng cao su, keo silicon chèn", "Kéo 1 cánh thì các cánh sau chạy theo liên động nhẹ nhàng."),
         
        ("3", "Ngoại Thất (Premium)", "Slim Ngoại Thất", "Cửa lùa Slim cách âm ngoài trời", 
         "Cửa đi trượt lùa ngoài trời viền cánh rộng hơn dập gioăng EPDM chống bão chịu nước ngấm.", 
         "2.0mm", "Phụ kiện khóa tay nắm dẹt đa điểm, bánh xe", "Gioăng EPDM kép dẻo dai chống mặn", "Phù hợp lắp kính hộp Double Low-E 20mm-24mm cách nhiệt tốt."),
         
        ("4", "Mở Quay (Tối giản)", "Slim Mở Quay", "Cửa mở quay sổ/đi Slim", 
         "Cửa đi thông phòng ngủ 1 cánh, cửa sổ mở quay hất viền mỏng bo ô kính Panorama.", 
         "1.2mm - 1.4mm", "Bản lề ẩn giấu chìm & khóa chốt đa điểm dẹt", "Gioăng EPDM kép khít kín tốt", "Tăng diện tích truyền sáng tối đa, thẩm mỹ tinh tế."),
         
        ("5", "Trượt Âm (Nội thất)", "Slim Pocket Door", "Cửa lùa âm tường Pocket", 
         "Cửa trượt đẩy ngang giấu ẩn chìm 100% cánh vào khoang tường thạch cao tối ưu diện tích hẹp.", 
         "1.6mm", "Bánh xe treo tải trọng nặng & ray hộp âm", "Gioăng cao su đệm cánh", "Giải pháp thông minh cho phòng tắm hẹp, căn hộ studio."),
         
        ("6", "Vách Kính (Minimalist)", "Slim Vách Tĩnh", "Vách kính ngăn phòng khung Slim", 
         "Vách kính ngăn phòng họp, phòng thay đồ, vách lấy sáng trang trí bọc viền nhôm siêu mỏng.", 
         "1.2mm", "Liên kết ke góc nhôm & vít inox chìm", "Gioăng cao su kẹp giữ kính", "Sơn Anode xước màu vàng titan, đen mịn không bám vân tay.")
    ]
    
    # Populate rows
    for row_idx, row_data in enumerate(data, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = row_data[1]
        
        # Color segment cell differently based on system group
        if "Ngoại Thất" in seg:
            seg_fill = fill_premium
        elif "Trượt Treo" in seg or "Trượt Âm" in seg:
            seg_fill = fill_premium
        else:
            seg_fill = fill_standard
            
        for col_idx, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            # Apply color fills
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            # Alignment rules
            if col_idx in [1, 2, 3, 6]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

    # Add a heavy bottom border to the last data row
    last_row_idx = len(data) + 4
    for col_idx in range(1, 10):
        ws.cell(row=last_row_idx, column=col_idx).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom\Catalogue_HeCuaNhom_Slim.xlsx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Slim Exclusive XLSX successfully saved to {output_path}")

if __name__ == "__main__":
    generate_slim_xlsx()
    generate_slim_docx()
