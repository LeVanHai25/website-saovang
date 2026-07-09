# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng Xingfa (Xingfa Quảng Đông Nhập khẩu & Xingfa Việt Nam):
1. Catalogue_HeCuaNhom_Xingfa.xlsx
2. ThuVien_HeCuaNhom_Xingfa.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

XINGFA_SYSTEMS = [
    ('1', 'Nhập khẩu tem đỏ', 'Xingfa Guangdong 55', 'Cửa đi & sổ mở quay XF55', 
     'Cửa sổ và cửa đi mở quay tem đỏ Quảng Đông chính hãng nhập khẩu rãnh 22 thông dụng nhất.', 
     '1.4mm - 2.0mm', 'Phụ kiện Kinlong / Draho / 3H đồng bộ', 'Gioăng EPDM kép / Silicon Apollo A500', 
     'Hệ nhôm quốc dân, độ cứng cao nhờ thiết kế 2 đường gân gia cường đặc trưng.'),
     
    ('2', 'Nhập khẩu tem đỏ', 'Xingfa Guangdong 93', 'Cửa đi lùa trượt XF93', 
     'Cửa đi trượt lùa bản dày 2.0mm hệ 93 2 ray hoặc 3 ray vô cùng vững chắc.', 
     '2.0mm', 'Bánh xe chịu lực kép & khóa chữ D Kinlong', 'Gioăng EPDM kép chèn nỉ chống nước', 
     'Chịu gió bão cực tốt, vận hành đầm chắc chắn. Nhưng chiếm nhiều diện tích ray dưới.'),
     
    ('3', 'Nhập khẩu tem đỏ', 'Xingfa Guangdong 55 Slide', 'Cửa sổ lùa trượt XF55', 
     'Cửa sổ trượt lùa bản mỏng hệ 55 giúp tối ưu chi phí nguyên vật liệu.', 
     '1.2mm - 1.4mm', 'Bánh xe đơn & khóa sập Kinlong', 'Gioăng chèn nỉ / Silicon Apollo', 
     'Phù hợp cửa sổ các chung cư, nhà phố diện tích hẹp không thể mở quay.'),
     
    ('4', 'Nhập khẩu tem đỏ', 'Xingfa Guangdong 63', 'Cửa xếp trượt gấp XF63', 
     'Cửa xếp trượt gấp bifold từ 3 cánh đến hơn 10 cánh giúp mở rộng tối đa không gian.', 
     '1.5mm - 3.0mm', 'Bánh xe treo chịu lực & bản lề liên kết Kinlong', 'Gioăng chèn nỉ đa lớp / Silicon Dow', 
     'Lấy thoáng 99% diện tích ô cửa biệt thự vườn, đòi hỏi kỹ thuật lắp đặt căn chỉnh bản lề cao.'),
     
    ('5', 'Nhập khẩu tem đỏ', 'Xingfa Guangdong 65', 'Vách mặt dựng Stick XF65', 
     'Vách kính mặt dựng Stick bao che tòa nhà khung xương nổi hoặc xương chìm.', 
     '2.5mm - 3.0mm', 'Ke liên kết & Bulong kết cấu inox SUS304', 'Silicon kết cấu chống thấm Dow Corning 791/895', 
     'Chịu lực gió bão cao tầng tốt, lấy sáng tự nhiên tối đa.'),
     
    ('6', 'Cao cấp Class A', 'Xingfa Class A', 'Hệ nhôm rãnh C Class A', 
     'Dòng nhôm Xingfa cao cấp không gân, thiết kế phẳng rãnh C Châu Âu mạ điện di Anodized ED.', 
     '2.0mm - 2.2mm', 'Phụ kiện CMECH / Roto rãnh C Châu Âu', 'Gioăng EPDM đa khoang trung tâm / Silicon Dow', 
     'Xi mạ ED bền màu sơn trên 25 năm chống muối mặn ven biển. Giá thành rất cao.'),
     
    ('7', 'Việt Nam sản xuất', 'Xingfa Việt Nam 55', 'Cửa mở quay XF55 Việt Nam', 
     'Cửa đi & cửa sổ mở quay sản xuất trong nước bởi các nhà máy đùn nhôm nội địa, tem xanh/tem xám.', 
     '1.2mm - 1.4mm', 'Phụ kiện PMA / Kinlong / Huy Hoàng rãnh 22', 'Gioăng EPDM đơn / Silicon Apollo', 
     'Chi phí rẻ, nguồn hàng nhanh gọn phù hợp các dự án nhà phố bình dân.'),
     
    ('8', 'Việt Nam sản xuất', 'Xingfa Việt Nam 93', 'Cửa lùa trượt XF93 Việt Nam', 
     'Cửa đi lùa trượt 2 ray/3 ray đùn ép trong nước giúp tối ưu giá thành.', 
     '1.4mm - 1.6mm', 'Bánh xe đơn & khóa sập trong nước', 'Gioăng chèn nỉ mỏng / Apollo', 
     'Giá thành rẻ hơn hàng nhập khẩu tem đỏ 30%, kết cấu mỏng hơn.'),
     
    ('9', 'Việt Nam sản xuất', 'Xingfa Việt Nam 63', 'Cửa xếp trượt gấp XF63 Việt Nam', 
     'Hệ cửa xếp trượt gấp nội địa phục vụ công trình dân dụng giá rẻ.', 
     '1.4mm - 1.6mm', 'Phụ kiện xếp trượt Kinlong / Draho', 'Gioăng EPDM kép / Silicon Apollo', 
     'Kinh tế, dễ mua. Chịu tải bản bánh xe treo ở mức vừa phải.'),
     
    ('10', 'Việt Nam sản xuất', 'Xingfa Việt Nam 65', 'Vách mặt dựng Stick XF65 Việt Nam', 
     'Vách kính mặt dựng Stick trong nước chuyên dùng cho showroom, tòa nhà thấp và trung tầng.', 
     '2.0mm', 'Ke kết cấu chịu lực trong nước', 'Silicon kết cấu chống thấm Apollo', 
     'Chi phí rẻ, lắp dựng nhanh chóng, không tốn thời gian chờ hàng nhập khẩu.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_Xingfa.xlsx
# ─────────────────────────────────────────────
def create_xingfa_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue Xingfa"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM XINGFA NHẬP KHẨU & VIỆT NAM"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="D32F2F") # Màu đỏ đặc trưng Xingfa
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="FFEBEE")
    for r_idx, row_vals in enumerate(XINGFA_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_Xingfa.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel Xingfa saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel Xingfa to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_Xingfa.docx
# ─────────────────────────────────────────────
def create_xingfa_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xD3,0x2F,0x2F))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM XINGFA NHẬP KHẨU & VIỆT NAM", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp hệ cửa nhôm định hình Xingfa Quảng Đông và Xingfa Việt Nam\nTích hợp phân tích ưu nhược điểm R&D và ma trận so sánh chéo phân khúc", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM XINGFA", bold=True, size=14, color=RGBColor(0xD3,0x2F,0x2F), space_before=12)
    add_p(
        "Nhôm định hình hệ Xingfa là thương hiệu cửa nhôm 'quốc dân' vô cùng nổi tiếng và chiếm thị phần lớn nhất tại Việt Nam. "
        "Dòng nhôm Xingfa Quảng Đông (tem đỏ) đùn ép bởi Công ty TNHH Nhôm Xingfa Quảng Đông nhập khẩu chính hãng "
        "đã tạo nên vị thế vững chắc nhờ hợp kim nhôm chất lượng cao AL6063-T5, kết cấu 2 gân gia cường cứng vững xuất sắc "
        "và công nghệ sơn phủ tĩnh điện ngoài trời vượt trội đạt chuẩn quốc tế Qualicoat, GSB Châu Âu. "
        "Tại Việt Nam, nhằm tối ưu chi phí cho các công trình trung cấp và phổ thông, các nhà máy đùn nhôm nội địa "
        "cũng đã đùn ép thành công hệ nhôm Xingfa Việt Nam với các độ dày linh hoạt từ 1.2mm - 1.6mm "
        "tạo nên sự phong phú tuyệt vời cho phân khúc thị trường cửa nhôm kính."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D 10 HỆ NHÔM XINGFA", bold=True, size=14, color=RGBColor(0xD3,0x2F,0x2F), space_before=12)
    
    # 1. Xingfa Guangdong 55
    add_p("1. Xingfa Guangdong Hệ 55 (Cửa đi & sổ mở quay nhập khẩu)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm (sổ) - 2.0mm (đi) rãnh 22mm.")
    add_bullet("Ưu điểm R&D: Cực kỳ cứng chắc nhờ kết cấu gân gia cường, màu sơn đẹp bền màu đạt chuẩn Qualicoat, phụ kiện Kinlong/Draho rất đồng bộ và phổ biến nhất thị trường.")
    add_bullet("Nhược điểm R&D: Nhiều hàng giả hàng nhái trên thị trường, rãnh 22mm truyền thống không lắp được phụ kiện rãnh C cao cấp.")
    
    # 2. Xingfa Guangdong 93
    add_p("2. Xingfa Guangdong Hệ 93 (Cửa lùa trượt nhập khẩu)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm.")
    add_bullet("Ưu điểm R&D: Bản nhôm dày 2.0mm cực kỳ vững chãi, chống rung rung chấn gió bão rất tốt, thiết kế trượt nhẹ đầm chắc.")
    add_bullet("Nhược điểm R&D: Thanh ray dưới nhôm đúc nổi có thể gây vấp chân khi đi lại nếu không thiết kế âm sàn.")

    # 3. Xingfa Guangdong 55 Slide
    add_p("3. Xingfa Guangdong Hệ 55 Slide (Cửa sổ lùa nhập khẩu)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.4mm.")
    add_bullet("Ưu điểm R&D: Gọn nhẹ tiết kiệm không gian, chi phí lắp đặt rẻ hơn hệ 93, thích hợp làm cửa sổ lùa.")
    add_bullet("Nhược điểm R&D: Kết cấu nhẹ nên không lắp cánh đi quá rộng cao.")

    # 4. Xingfa Guangdong 63
    add_p("4. Xingfa Guangdong Hệ 63 (Cửa xếp trượt gấp nhập khẩu)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.5mm - 3.0mm.")
    add_bullet("Ưu điểm R&D: Mở thông suốt 99% khẩu độ ô cửa, bánh xe treo chịu tải lớn giúp gấp xếp cánh êm ái thích hợp biệt thự sân vườn lớn.")
    add_bullet("Nhược điểm R&D: Kỹ thuật căn chỉnh khe hở bản lề phức tạp, phụ kiện xếp treo khá đắt đỏ.")

    # 5. Xingfa Guangdong 65
    add_p("5. Xingfa Guangdong Hệ 65 (Vách mặt dựng Stick nhập khẩu)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.5mm - 3.0mm.")
    add_bullet("Ưu điểm R&D: Xương chịu lực dày cản gió bão tuyệt đối, bề mặt sơn tĩnh điện chống tia UV tốt thích hợp bao ngoài các tòa nhà cao tầng showroom.")
    add_bullet("Nhược điểm R&D: Thi công giàn giáo đu dây bên ngoài phức tạp.")

    # 6. Xingfa Class A
    add_p("6. Xingfa Class A (Hệ rãnh C Châu Âu xi mạ ED cao cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm - 2.2mm.")
    add_bullet("Ưu điểm R&D: Bản phẳng không gân cực kỳ sang trọng hiện đại, chuẩn rãnh C Châu Âu cách âm cản nhiệt tốt lắp CMECH/Roto đồng bộ, xi mạ ED bảo hành 25 năm muối mặn.")
    add_bullet("Nhược điểm R&D: Chi phí nguyên liệu và phụ kiện đi kèm phân khúc Luxury siêu đắt.")

    # 7. Xingfa Việt Nam 55
    add_p("7. Xingfa Việt Nam Hệ 55 (Cửa mở quay trong nước)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.4mm.")
    add_bullet("Ưu điểm R&D: Giá rẻ hơn hàng nhập Quảng Đông 30-40%, sẵn nguồn hàng đùn ép trong nước, dễ nhảy ke ép góc gia công nhanh.")
    add_bullet("Nhược điểm R&D: Độ dày mỏng nên chống rung gió bão kém hơn, nước sơn tĩnh điện bảo hành ngắn hơn hàng nhập khẩu.")

    # 8. Xingfa Việt Nam 93
    add_p("8. Xingfa Việt Nam Hệ 93 (Cửa lùa trượt nội địa)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 1.6mm.")
    add_bullet("Ưu điểm R&D: Giải pháp cửa đi trượt lùa vô cùng kinh tế cho nhà phố, nhà xưởng văn phòng vừa và nhỏ.")
    add_bullet("Nhược điểm R&D: Ray trượt mỏng hơn dễ mài mòn, độ kín khít chắn dột mưa gió trung bình.")

    # 9. Xingfa Việt Nam 63
    add_p("9. Xingfa Việt Nam Hệ 63 (Cửa xếp gấp nội địa)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm - 1.6mm.")
    add_bullet("Ưu điểm R&D: Lắp dựng nhanh chóng tiết kiệm đầu tư, phụ kiện bình dân.")
    add_bullet("Nhược điểm R&D: Càng về lâu dài dễ sệ cánh xếp gấp nếu tần suất mở ra vào lớn.")

    # 10. Xingfa Việt Nam 65
    add_p("10. Xingfa Việt Nam Hệ 65 (Vách kính Stick nội địa)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm.")
    add_bullet("Ưu điểm R&D: Chi phí rẻ, sẵn hàng cắt gia công nhanh cho các vách ngăn kính văn phòng showroom trong nhà hoặc thấp tầng.")
    add_bullet("Nhược điểm R&D: Độ chịu lực mạ chống rỉ kém hơn hàng nhập khẩu đắt tiền.")

    # ── PHẦN SO SÁNH MA TRẬN ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN PHÂN KHÚC HỆ XINGFA 55 VỚI ĐỐI THỦ CẠNH TRANH", bold=True, size=14, color=RGBColor(0xD3,0x2F,0x2F), space_before=12)
    add_p(
        "Bảng ma trận so sánh chéo đặc tính kỹ thuật hệ nhôm Xingfa Quảng Đông 55 nhập khẩu tem đỏ "
        "với các dòng nhôm hệ 55 thông dụng trên thị trường Việt Nam giúp hỗ trợ tư vấn bán hàng:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "Xingfa Quảng Đông 55"
    hdr_c[2].text = "PMA 55 vát cạnh"
    hdr_c[3].text = "Việt Pháp 450"
    hdr_c[4].text = "Xingfa Việt Nam 55"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ("Nguồn gốc xuất xứ", "Quảng Đông (Trung Quốc)", "Việt Nam (PMA Vina)", "Liên doanh Việt - Pháp", "Đùn ép trong nước (VN)"),
        ("Độ dày cánh cửa đi", "2.0mm cực chắc khỏe", "1.2mm - 1.4mm", "1.2mm - 1.4mm", "1.2mm - 1.4mm"),
        ("Kiểu dáng thiết kế", "2 gân gia cường nổi", "Mặt vát góc nẹp liền", "Sập rời bo góc", "2 gân gia cường nổi"),
        ("Hệ phụ kiện phù hợp", "Kinlong / Draho rãnh 22", "PMA / Kinlong rãnh 22", "Chốt gạt Việt Pháp", "Kinlong / PMA rãnh 22"),
        ("Mức độ tin dùng", "Rất cao (Quốc dân)", "Cao (Kinh tế dân dụng)", "Trung bình (Phổ thông)", "Khá cao (Nhà phố rẻ)"),
        ("Phân khúc giá bán", "$$$ (Tầm trung)", "$$ (Bình dân)", "$ (Siêu tiết kiệm)", "$$ (Bình dân)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_Xingfa.docx")
    try:
        doc.save(path)
        print(f"  >> Docx Xingfa saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx Xingfa to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone Xingfa docx & xlsx in BASE directory...")
    print("=" * 60)
    create_xingfa_xlsx()
    create_xingfa_docx()
    print("=" * 60)
