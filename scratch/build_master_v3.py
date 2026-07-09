# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 4 tài liệu Master hoàn chỉnh
Phiên bản 3.0 - Dữ liệu thực tế từ catalogue các hãng
"""
import os, sys, subprocess

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

for pkg in ["python-docx", "openpyxl"]:
    try:
        __import__(pkg.replace("-",""))
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from openpyxl import Workbook
from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                              GradientFill)
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.filters import AutoFilter

# ─────────────────────────────────────────────
# HÀM LƯU TỆP AN TOÀN CHỐNG LOCK FILE
# ─────────────────────────────────────────────
def safe_save_excel(wb, path):
    try:
        wb.save(path)
        print(f"  >> Saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_DANHGIA_MOI.xlsx")
        wb.save(alt_path)
        print(f"  [!] CANH BAO: Tep '{os.path.basename(path)}' dang duoc mo trong Excel.")
        print(f"  >> Da luu tam sang: {alt_path}")

def safe_save_docx(doc, path):
    try:
        doc.save(path)
        print(f"  >> Saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_DANHGIA_MOI.docx")
        doc.save(alt_path)
        print(f"  [!] CANH BAO: Tep '{os.path.basename(path)}' dang duoc mo trong Word.")
        print(f"  >> Da luu tam sang: {alt_path}")

# ─────────────────────────────────────────────
# THƯ MỤC GỐC
# ─────────────────────────────────────────────
BASE  = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"
ROOT  = os.path.join(BASE, "THƯ VIỆN HỆ NHÔM SAO VÀNG")

SUBDIRS = [
    "05_Thư viện CAD/Cửa đi",
    "05_Thư viện CAD/Cửa sổ",
    "05_Thư viện CAD/Cửa lùa",
    "05_Thư viện CAD/Cửa Slim",
    "05_Thư viện CAD/Vách mặt dựng",
    "05_Thư viện CAD/Lan can & Mái kính",
    "06_Thư viện Profile Nhôm",
    "07_Thư viện Phụ kiện",
    "08_Hướng dẫn Lắp đặt",
    "09_Hướng dẫn Gia công",
    "10_Catalogue theo Hãng",
    "11_Hồ sơ Dự án Tham khảo",
    "12_Hình ảnh & Video",
]

for d in SUBDIRS:
    os.makedirs(os.path.join(ROOT, d), exist_ok=True)

# ─────────────────────────────────────────────
# CƠ SỞ DỮ LIỆU THỰC TẾ  (nghiên cứu từ catalogue các hãng)
# ─────────────────────────────────────────────
# Cột: ID, Nhóm, Mã Hệ, Hãng, QG, TC, Loại Cửa, Kiểu Mở,
#       Loại Profile, Công Năng, Ứng Dụng,
#       Dày Nhôm(mm), Kính Thích(mm), Kính Tối Đa(mm),
#       Kích Thước Tối Đa (W×H mm), Tải Trọng(kg),
#       Chuẩn PK, Gioăng/Keo, Bề Mặt,
#       Ưu Điểm, Nhược Điểm,
#       CAD, SD, Cat, HDLD, Tình Trạng, %, Web, Link Cat
ROWS = [
    # ══════ VIỆT PHÁP SHAL (LIÊN DOANH / VIỆT NAM) ══════
    ("VP-001","Phổ thông","Việt Pháp 450","Việt Pháp","Việt Nam","TCVN",
     "Cửa đi","Mở quay","Khung bản sập rời 450mm + Cánh","Nội/Ngoại thất","Nhà cấp 4, nhà phố giá rẻ, nhà xưởng",
     1.2,"Kính đơn 5-10mm","10mm","1100×2200",80,"Chốt gạt Việt Pháp / Kinlong rãnh 22","Gioăng EPDM đơn / Silicon Apollo","Sơn tĩnh điện bóng/mờ",
     "Giá thành rẻ, phổ biến rộng rãi, cực kỳ dễ gia công bằng máy dập thường","Bản cánh mỏng chống rung gió yếu, cách âm trung bình",
     "✓","✗","✓","✓","Hoàn thành",85,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-viet-phap-shal-18.html"),

    ("VP-002","Phổ thông","Việt Pháp 4400","Việt Pháp","Việt Nam","TCVN",
     "Cửa sổ","Mở quay + Mở hất","Khung bản nhỏ 4400mm + Cánh","Nội/Ngoại thất","Cửa sổ phụ nhà dân, vách ngăn phòng",
     1.1,"Kính đơn 5-10mm","10mm","1000×1400",60,"Khớp chốt Việt Pháp / Kinlong","Gioăng cao su chèn sập / Apollo","Sơn tĩnh điện bóng/mờ",
     "Bản nhôm mỏng nhẹ, chi phí siêu rẻ, dễ gia công nhanh chóng","Độ kín nước và chắn rung gió bão kém",
     "✓","✗","✓","✓","Hoàn thành",85,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-viet-phap-shal-18.html"),

    ("VP-003","Phổ thông","Việt Pháp 2600","Việt Pháp","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Trượt lùa (Sliding)","Khung lùa 2600mm 2 ray","Ngoại thất","Căn hộ chung cư, nhà phố bình dân",
     1.1,"Kính đơn 5-10mm","10mm","2400×2000",70,"Bánh xe đơn trượt & khóa sập Việt Pháp","Gioăng chèn nỉ mỏng","Sơn tĩnh điện bóng/mờ",
     "Giá thành rẻ, tiết kiệm 100% diện tích không gian đóng mở","Kết cấu mỏng nhẹ, trượt không đầm tay, ray nhôm dễ mòn",
     "✓","✗","✓","✓","Hoàn thành",85,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-viet-phap-shal-18.html"),

    ("VP-004","Tầm trung","Shal 55","Việt Pháp","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay","Khung phẳng rãnh C 55mm","Nội/Ngoại thất","Nhà phố tầm trung, văn phòng công ty",
     1.2,"Kính đơn 6-12mm / Kính hộp","20mm","1400×2400",100,"Phụ kiện PMA/Draho rãnh C Châu Âu","Gioăng EPDM kép / Silicon Apollo","Sơn tĩnh điện AkzoNobel",
     "Thiết kế phẳng phiu giống Xingfa hiện đại, khắc phục thẩm mỹ hệ 450","Phân khúc cạnh tranh khốc liệt với nhôm cỏ trong nước",
     "✓","✓","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-viet-phap-shal-18.html"),

    ("VP-005","Mặt dựng","Việt Pháp 1100","Việt Pháp","Việt Nam","TCVN",
     "Vách mặt dựng","Cố định (Fixed Curtain Wall)","Xương mặt dựng Stick 1100","Ngoại thất","Showroom, tòa nhà văn phòng trung tầng",
     1.5,"Kính cường lực 10-12mm / Kính hộp","24mm","Không giới hạn",400,"Ke chịu lực & bulong inox SUS304","Silicon kết cấu chịu thời tiết","Sơn tĩnh điện bảo hành 10 năm",
     "Thi công lắp dựng ngoài công trường đơn giản, chi phí hợp lý","Độ kín nước phụ thuộc lớn vào chất lượng đi keo tại công trình",
     "✓","✓","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-viet-phap-shal-18.html"),

    ("VP-006","Cao cấp Premium","Shal Premium","Việt Pháp","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay","Khung rãnh C chuẩn Châu Âu","Ngoại thất","Biệt thự, căn hộ cao cấp ven biển",
     1.6,"Kính đơn/hộp 8-32mm","32mm","1600×2800",150,"Phụ kiện CMECH / Roto rãnh C Châu Âu","Gioăng EPDM đa khoang / Silicon Dow","Sơn tĩnh điện bảo hành 30 năm",
     "Chất lượng sơn cao cấp bảo hành 30 năm chống muối biển, liên doanh Pháp","Giá thành cao phân khúc biệt thự Luxury",
     "✓","✓","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-viet-phap-shal-18.html"),

    # ══════ PMA (VIỆT NAM) ══════
    ("PM-001","Phổ thông Vát cạnh","PMA 55","PMA","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất","Khung vát cạnh 55mm + Cánh","Nội/Ngoại thất","Nhà phố phổ thông, chung cư, trường học",
     1.2,"5-12mm","12mm","1200×2200",80,"Rãnh 22mm đồng bộ","Gioăng EPDM đơn / Silicon Apollo","Sơn tĩnh điện bóng/mờ",
     "Giá thành hợp lý, dễ gia công, linh kiện phổ thông, phù hợp nhà phố","Không phải rãnh C, cách âm thấp hơn hệ cao cấp, không phù hợp cánh quá lớn",
     "✓","✗","✓","✓","Hoàn thành",85,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-002","Cao cấp Platinum","PMA Platinum","PMA","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay","Phào kép nghệ thuật ôm khít khuôn cửa","Nội/Ngoại thất","Biệt thự vườn, biệt thự phố cao cấp",
     1.4,"Kính đơn 8-12mm, kính hộp","24mm","1600×2800",150,"Rãnh C Châu Âu đồng bộ","Gioăng EPDM đa khoang / Silicon Dow","Sơn tĩnh điện cao cấp AkzoNobel",
     "Cánh phẳng, thiết kế hiện đại, phụ kiện đồng bộ, độ kín nước cao","Chi phí cao, yêu cầu gia công chính xác",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-003","Rãnh C Classic","PMA Classic 58","PMA","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất","Khung rãnh C 58mm + Cánh","Nội/Ngoại thất","Căn hộ chung cư cao cấp, văn phòng",
     1.3,"Kính đơn/hộp 5-28mm","28mm","1400×2600",100,"Rãnh C Châu Âu","Gioăng EPDM đa khoang / Silicon Sika","Sơn tĩnh điện bảo hành 10-15 năm",
     "Tương thích phụ kiện rãnh C Châu Âu, cách âm tốt","Giá trung cấp",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-004","Rãnh C Classic","PMA 65","PMA","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay","Khung 65mm + Cánh bản lớn","Ngoại thất","Ứng dụng công nghiệp nặng, kết cấu chịu tải lớn",
     2.0,"Kính đơn/hộp 6-30mm","30mm","1500×2600",120,"Rãnh C Châu Âu","Gioăng EPDM đa khoang / Sika","Sơn tĩnh điện bảo hành 10-15 năm",
     "Bản dày chắc khỏe hơn hệ 58, thích ứng tốt cho ứng dụng chịu lực vặn lớn","Trọng lượng lớn hơn, giá nhôm thanh cao hơn",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-005","Rãnh C Classic","PMA 75","PMA","Việt Nam","TCVN",
     "Cửa đi","Mở quay","Khung 75mm dày dặn","Ngoại thất","Hệ thống đóng tự động, khu vận chuyển công nghiệp nặng, biệt thự mặt phố chịu gió lớn",
     2.5,"Kính hộp đến 32mm","32mm","1800×2800",150,"Rãnh C Châu Âu","Gioăng EPDM kép chèn / Silicon Dow","Sơn tĩnh điện bảo hành 15 năm",
     "Khả năng chịu áp lực tải trọng siêu nặng, chịu bão cấp cao ngoại thất","Chi phí đầu tư cao, nặng",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-006","Trượt lùa Trung cấp","PMA 93","PMA","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa (Sliding)","Khung lùa 93mm 2 ray/3 ray","Ngoại thất","Căn hộ, văn phòng, nhà phố",
     1.4,"Kính đơn/hộp 6-20mm","20mm","3000×2400",100,"Bánh xe trượt PMA đồng bộ","Gioăng EPDM kép / Apollo A500","Sơn tĩnh điện bóng/mờ",
     "Giá bình dân, tối ưu diện tích không gian đóng mở","Độ kín khít cách âm ở mức trung bình",
     "✓","✗","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-007","Trượt lùa Trung cấp","PMA 95","PMA","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa (Sliding)","Khung lùa 95mm dày dặn","Ngoại thất","Nhà phố, căn hộ cao cấp",
     1.6,"Kính đơn/hộp 6-24mm","24mm","4000×2400",120,"Bánh xe Hopo / PMA chịu tải","Gioăng EPDM kép / Silicon Dow","Sơn tĩnh điện AkzoNobel",
     "Cứng vững chống gió giật tốt hơn lùa 93, trượt đầm tay","Chi phí phụ kiện cao hơn",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-008","Trượt lùa Cao cấp","PMA 115","PMA","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa 3 ray có lưới chống muỗi","Khung 3 ray 115mm + ray lưới","Ngoại thất","Biệt thự vườn, biệt thự ven sông",
     1.8,"Kính đơn/hộp 6-24mm","24mm","5000×2600",150,"Ray inox chặn nỉ / bánh xe chịu lực","Gioăng EPDM chèn lưới","Sơn tĩnh điện AkzoNobel",
     "Tích hợp sẵn ray lưới chống côn trùng đồng bộ, thoáng mát","Chiếm nhiều diện tích độ dày tường khi lắp",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-009","Lift & Slide Cao cấp","PMA Lift & Slide","PMA","Việt Nam","TCVN",
     "Cửa đi","Trượt nâng (Lift & Slide)","Khung trượt nâng bản dày","Ngoại thất","Biệt thự sát biển, penthouses, resort",
     2.0,"Kính hộp cách âm 24-32mm","32mm","6000×2800",300,"Bánh xe nâng hạ & tay gạt Sobinco/Hopo","Gioăng nén phẳng mặt sàn","Sơn tĩnh điện bảo hành 15 năm",
     "Nén gioăng kín khít tuyệt đối cách âm tốt, trượt nâng chịu tải lớn 300kg","Chi phí phụ kiện nâng hạ nhập khẩu rất đắt",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-010","Slim Ngoại thất Cao cấp","PMA Slim","PMA","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa Slim Panorama (Sliding)","Khung cánh siêu mỏng ray inox","Nội/Ngoại thất","Vách ngăn phòng khách, cửa ban công biệt thự",
     1.6,"Kính cường lực 8-12mm / Kính hộp","24mm","4000×2600",100,"Ray inox SUS316L + bánh xe giảm chấn","Gioăng chèn nỉ mỏng","Sơn tĩnh điện mờ",
     "Thiết kế tối giản cánh cực mảnh, tăng 95% diện tích kính lấy view","Chống bão gió ngoài trời hạn chế, phù hợp trong nhà",
     "✓","✗","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-011","Cầu cách nhiệt Cao cấp","PMA Thermal 65","PMA","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay (Casement)","Khung cách nhiệt 65mm có cầu PA66","Ngoại thất","Biệt thự, phòng ngủ hướng Tây nắng nóng",
     1.4,"Kính hộp Low-E cách nhiệt 24-32mm","32mm","1400×2400",120,"Phụ kiện CMECH / Roto rãnh C Châu Âu","Gioăng EPDM đa khoang trung tâm / Silicon Dow","Sơn tĩnh điện bảo hành 15 năm",
     "Cầu cách nhiệt dải Polyamide cản nóng 30%, cách âm 38dB vượt trội","Gia công góc cắt cần ke keo cấu trúc phức tạp",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-012","Cầu cách nhiệt Siêu cao cấp","PMA Thermal 75","PMA","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay (Casement)","Khung cách nhiệt bản lớn 75mm","Ngoại thất","Biệt thự Luxury, resort cao cấp",
     2.0,"Kính hộp cách âm hộp 28-36mm","36mm","1600×2800",150,"Phụ kiện Sobinco (Bỉ) chống cạy RC2","Gioăng đúc xốp trung tâm","Sơn tĩnh điện bảo hành 15 năm AkzoNobel",
     "Khả năng cách âm tối đa đạt trên 40dB, cản nhiệt cao nhất thị trường","Giá thành đắt đỏ phân khúc Luxury",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-013","Curtain Wall Stick","PMA Curtain Wall","PMA","Việt Nam","TCVN",
     "Vách mặt dựng","Cố định (Fixed Curtain Wall)","Profile xương mặt dựng Stick","Ngoại thất","Tòa nhà văn phòng, showroom mặt tiền",
     2.0,"Kính hộp phản quang / Low-E 24-32mm","32mm","Không giới hạn",500,"Kẹp kính kết cấu chìm/nổi","Silicon kết cấu Dow Corning 895","Anodized, PVDF",
     "Chịu áp lực gió giật bão tốt, thẩm mỹ hiện đại lấy sáng tối đa","Thi công giàn giáo cao tầng bên ngoài phức tạp",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-014","Phổ thông","PMA Railing","PMA","Việt Nam","TCVN",
     "Lan can","Lắp ghép cố định (Railing)","Profile lan can hợp kim nhôm","Ngoại thất","Lan can ban công biệt thự, chung cư",
     1.5,"Không dùng kính / Kính cường lực 10-12mm","12mm","Không giới hạn",100,"Ốc vít inox SUS304 lắp ghép modul","Gioăng đệm nhựa chặn chân","Sơn tĩnh điện AkzoNobel",
     "Không rỉ sét muối mặn han rỉ, lắp đặt nhanh gọn, đồng bộ biệt thự","Kiểu dáng định hình sẵn ít tùy biến hoa văn cổ điển",
     "✓","✗","✓","✓","Hoàn thành",75,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-015","Phổ thông","PMA Partition","PMA","Việt Nam","TCVN",
     "Vách ngăn","Cố định + Cửa lùa văn phòng","Profile vách ngăn mỏng","Nội thất","Văn phòng làm việc, vách ngăn nhà xưởng",
     1.0,"Kính đơn 5-10mm / Tấm alu nhôm","10mm","Không giới hạn",50,"Ke nhảy / Ke ma thuật liên kết góc","Gioăng chèn nỉ / Silicon Apollo","Sơn tĩnh điện bóng/mờ",
     "Lắp đặt tháo dỡ cực nhanh, chi phí siêu rẻ phân chia phòng làm việc","Khả năng cách âm thấp, kết cấu chịu tải kém",
     "✓","✗","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-016","Phổ thông","PMA 50","PMA","Việt Nam","TCVN",
     "Cửa sổ + Cửa đi nhẹ + Vách ngăn","Mở quay + Cố định","Khung nhôm 50mm nhỏ gọn","Nội/Ngoại thất","Cửa phòng tắm, vách toilet, cửa sổ phụ nhà dân",
     1.2,"Kính đơn 5-10mm / Kính chèn nỉ","10mm","1000×2000",60,"Phụ kiện PMA đồng bộ nhẹ","Gioăng EPDM đơn / Silicon Apollo","Sơn tĩnh điện bóng/mờ",
     "Nhỏ gọn linh hoạt, chi phí cực kỳ tối ưu, dễ lắp ráp","Chịu gió ngoài trời yếu, không chịu tải kính hộp lớn",
     "✓","✗","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-017","Phổ thông","PMA 63","PMA","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ lớn","Mở quay","Khung bao tăng cứng","Ngoại thất","Khuôn viên cửa sổ lớn chịu gió, cửa cổng sân vườn biệt thự",
     1.5,"Kính đơn/hộp 6-24mm","24mm","1400×2400",100,"Phụ kiện PMA đồng bộ chịu tải","Gioăng EPDM kép chống dột","Sơn tĩnh điện AkzoNobel",
     "Khung cánh lớn chắc khỏe chống rung chống võng, thích ứng gió giật","Cần kỹ thuật lắp ráp chống sệ cánh cửa lớn",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),

    ("PM-018","Phổ thông","PMA 72","PMA","Việt Nam","TCVN",
     "Cửa đi kích thước lớn","Mở quay","Khung bao dày bản lớn","Ngoại thất","Cửa chính mở rộng, cửa ban công khẩu độ lớn biệt thự biển",
     2.0,"Kính cường lực 10-15mm / Kính hộp","28mm","1800×2800",150,"Bản lề cối chịu lực & khóa đa điểm","Gioăng đúc xốp kép / Sika","Sơn tĩnh điện AkzoNobel bảo hành 15 năm",
     "Thanh profile cực dày chống rung chấn, vững chãi đóng mở khẩu độ rộng","Cánh nặng nề, chi phí phụ kiện kim khí lớn",
     "✓","✓","✓","✓","Hoàn thành",80,"https://pma.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html"),


    # ══════ KOGEN (ĐỨC/VIỆT NAM) ══════
    ("KG-001","Xếp trượt Cao cấp","Kogen Slim 50/68","Kogen","Đức/Việt Nam","EN Châu Âu",
     "Cửa đi","Xếp trượt ẩn bản lề (Hidden Hinge)","Khung mỏng rãnh C ẩn bản lề","Ngoại thất","Biệt thự, penthouse, resort",
     1.6,"Kính đơn 8-12mm, kính hộp","28mm","5000×2800",80,"Rãnh C Châu Âu đồng bộ Kogen","Gioăng EPDM đa khoang / Silicon Dow","Sơn tĩnh điện hạt mịn chống UV",
     "Bản lề ẩn hoàn toàn khi đóng tạo bề mặt phẳng tinh tế Bauhaus","Gia công cần dưỡng chuyên dụng khoét bản lề ẩn",
     "✓","✓","✓","✓","Hoàn thành",85,"https://kogen.vn",""),

    ("KG-002","Slim Ngoại thất Cao cấp","Kogen Slim Lùa","Kogen","Đức/Việt Nam","EN Châu Âu",
     "Cửa đi","Trượt lùa (Sliding Panorama)","Khung cánh siêu mỏng ray inox","Ngoại thất","Biệt thự hướng biển, ban công panorama",
     2.0,"Kính cường lực 10-15mm, kính hộp","32mm","6000×3000",150,"Ray inox SUS316L + Bánh xe chịu tải","Gioăng EPDM kép / Dow Corning","Sơn tĩnh điện hạt mịn màu xám ghi",
     "Cánh siêu mảnh tối đa hóa tầm nhìn panorama, ray inox trượt rất nhẹ","Chống gió bão cấp trung bình",
     "✓","✓","✓","✓","Hoàn thành",85,"https://kogen.vn",""),

    ("KG-003","Anodized Cao cấp","Kogen 60 Đa khoang","Kogen","Đức/Việt Nam","EN",
     "Cửa đi + Cửa sổ","Mở quay (Casement)","Khung 60mm đa khoang rãnh C","Ngoại thất","Nhà phố hiện đại, biệt thự",
     1.6,"Kính đơn/hộp 5-28mm","28mm","1400×2600",100,"Rãnh C Châu Âu","Gioăng EPDM đa khoang / Silicon Sika","Sơn tĩnh điện hạt mịn màu ghi",
     "Cấu trúc đa khoang chịu lực cực tốt, cách âm 35dB, rãnh C Châu Âu","Không có cầu cách nhiệt trong hệ này",
     "✓","✓","✓","✓","Hoàn thành",80,"https://kogen.vn",""),

    # ══════ VIRALWINDOW (VR) ══════
    ("VR-001","Anodized Phổ thông","VRA55","Viralwindow","Việt Nam","TCVN 5826",
     "Cửa đi + Cửa sổ","Mở quay (Casement)","Khung 55mm + Cánh","Nội/Ngoại thất","Nhà phố, chung cư phổ thông",
     1.4,"5-12mm","12mm","1200×2400",80,"Rãnh 22mm","Gioăng EPDM đơn / Keo PU + Silicon","Sơn KCC 20 năm – đa màu",
     "Giá cạnh tranh, đa dạng màu sắc, phổ biến miền Nam","Profile mỏng, không phù hợp cửa rộng lớn",
     "✓","✓","✓","✓","Hoàn thành",90,"https://viralwindow.com",""),

    ("VR-002","Anodized Trung cấp","VRA64","Viralwindow","Việt Nam","TCVN 5826",
     "Cửa sổ","Trượt lùa (Sliding)","Khung 66×42mm + Cánh 70×28mm","Ngoại thất","Nhà phố, chung cư",
     1.3,"5-10mm","10mm","2400×1800",80,"Rãnh 22mm","Gioăng EPDM / Silicon","Sơn tĩnh điện đa màu",
     "Trượt nhẹ, giá tốt","Kín khít trung bình","✓","✗","✓","✓","Hoàn thành",80,"https://viralwindow.com",""),

    ("VR-003","Anodized Trung cấp","VRA94","Viralwindow","Việt Nam","TCVN 5826",
     "Cửa đi","Trượt lùa (Sliding)","Khung 94×44mm + Cánh 80×34mm","Ngoại thất","Nhà phố, biệt thự",
     1.4,"6.38-30mm","30mm","4000×2600",150,"Rãnh C (Euro Groove)","Gioăng EPDM kép / Silicon Dow Corning","Sơn KCC, anodized",
     "Phù hợp cửa lùa rộng, cánh chứa kính hộp","Cần bảo trì bánh xe","✓","✓","✓","✓","Hoàn thành",85,"https://viralwindow.com",""),

    ("VR-004","Anodized Cao cấp","VRE65 (Cửa sổ)","Viralwindow","Việt Nam","TCVN",
     "Cửa sổ","Mở quay + Mở hất (Casement/Tilt-Turn)","Khung 65×62mm + Cánh 90×75mm","Ngoại thất","Biệt thự, nhà phố cao cấp",
     1.8,"6.38-37mm","37mm","1600×2800",150,"Rãnh C Châu Âu","Gioăng EPDM đa khoang / Dow Corning","Anodized ED, sơn AkzoNobel",
     "Hệ Châu Âu chất lượng cao, phụ kiện CMECH nguyên bộ","Giá cao hơn VRA","✓","✓","✓","✓","Hoàn thành",90,"https://viralwindow.com",""),

    ("VR-005","Anodized Cao cấp","VRE65 (Cửa đi)","Viralwindow","Việt Nam","TCVN",
     "Cửa đi","Mở quay (Casement)","Khung 65×70mm + Cánh 97×75mm","Ngoại thất","Biệt thự, nhà phố cao cấp",
     1.8,"6.38-37mm","37mm","1400×2800",180,"Rãnh C Châu Âu","Gioăng EPDM đa khoang / Dow Corning","Anodized ED, PVDF",
     "Cứng vững, chứa kính hộp lớn, profile thanh lịch","Giá trung-cao","✓","✓","✓","✓","Hoàn thành",90,"https://viralwindow.com",""),

    ("VR-006","Anodized Cao cấp","VRE120","Viralwindow","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa 2 ray (Sliding)","Khung 120×38mm + Cánh 47×95mm","Ngoại thất","Biệt thự, ban công rộng",
     2.0,"6.38-34mm","34mm","5000×2600",200,"Rãnh C Châu Âu","Gioăng EPDM kép / Silicon Dow Corning","Anodized ED chống muối biển",
     "Trượt lùa 2 ray, cánh rộng, kính hộp lớn, bền bỉ biển","Nặng, cần ray thẳng tuyệt đối","✓","✓","✓","✓","Hoàn thành",85,"https://viralwindow.com",""),

    ("VR-007","Anodized Cao cấp","VRE180","Viralwindow","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa 3 ray (Sliding)","Khung 180mm + Cánh 3 ray","Ngoại thất","Biệt thự, resort cao cấp",
     2.0,"Kính hộp đến 36mm","36mm","7000×2800",250,"Rãnh C Châu Âu","Gioăng EPDM 3 lớp / Silicon","Anodized ED",
     "Mở 3 cánh song song, không gian rộng tối đa","Chi phí lắp đặt cao","✓","✓","✓","✓","Hoàn thành",80,"https://viralwindow.com",""),

    ("VR-008","Thủy lực Siêu cao cấp","VRE Thủy lực","Viralwindow","Việt Nam","TCVN",
     "Cửa đi","Thủy lực (Hydraulic / Zero Gravity)","Khung siêu nặng + Piston thủy lực","Ngoại thất","Biệt thự siêu sang, villa resort",
     2.0,"Kính hộp 36-50mm / Kính cường lực 15-19mm","19mm","3500×3500",800,"Thủy lực Biloba/DICTATOR","Keo SG + EPDM giả thép","Anodized ED vàng 24K, đồng đỏ",
     "Cửa siêu rộng, đóng mở bằng một ngón tay nhờ piston thủy lực","Giá 300tr-1 tỷ/bộ, thi công phức tạp","✓","✓","✓","✓","Hoàn thành",85,"https://viralwindow.com",""),

    ("VR-009","Lift & Slide Cao cấp","VRE Lift Slide","Viralwindow","Việt Nam","TCVN",
     "Cửa đi","Trượt nâng (Lift & Slide)","Khung + Cánh + Ray đặc biệt","Ngoại thất","Biệt thự, lối ra sân vườn",
     2.0,"Kính hộp 24-48mm","19mm","6000×2800",400,"Roto NT / Siegenia LSS","Gioăng EPDM đa khoang / Dow Corning","Anodized, sơn tĩnh điện",
     "Kín hoàn toàn khi đóng, mở nhẹ dù cánh cực nặng","Giá cao, ray thẳng tuyệt đối","✓","✓","✓","✓","Hoàn thành",80,"https://viralwindow.com",""),

    ("VR-010","Xếp trượt Cao cấp","VRE Bi-fold","Viralwindow","Việt Nam","TCVN",
     "Cửa đi","Xếp trượt (Bi-fold Folding)","Cánh xếp gấp + Bản lề xếp + Ray trên","Ngoại thất","Biệt thự, nhà hàng, lối ra sân",
     1.8,"Kính cường lực 8-12mm","12mm","5000×2600",40,"Roto Fold / Hopo Bi-fold","Gioăng EPDM / Silicon","Sơn tĩnh điện, anodized",
     "Mở toàn bộ không gian, xếp gọn về một phía","Kín khít kém nhất, cần bảo trì","✓","✗","✓","✓","Đang nghiên cứu",60,"https://viralwindow.com",""),

    # ══════ CIVRO ══════
    ("CV-001","Anodized Trung cấp","AW55 IN/OU","Civro","Việt Nam","EN 12207/12210",
     "Cửa sổ","Mở quay vào/ra (Casement IN/OUT)","Khung 55mm cầu PA66","Ngoại thất","Nhà phố, chung cư trung cấp",
     1.5,"5-10mm","10mm","1200×2200",80,"Rãnh C Euro Groove","Gioăng EPDM 3 lớp / Silicon Dow","MED xi mạ độc quyền + Sơn tĩnh điện",
     "Hệ châu Âu giá Việt Nam, xi mạ MED không phai màu","Ít tùy chọn màu sắc hơn hệ 65","✓","✓","✓","✓","Hoàn thành",85,"https://civrowindow.com",""),

    ("CV-002","Anodized Cao cấp","AW65 IN/OU","Civro","Việt Nam","EN 12207/12210",
     "Cửa sổ","Mở quay vào/ra (Casement IN/OUT)","Khung 65mm cầu PA66","Ngoại thất","Biệt thự, văn phòng cao cấp",
     1.8,"8-14mm / Kính hộp 20-32mm","14mm","1600×2800",150,"Rãnh C Châu Âu – tương thích Roto/Siegenia/Maco","Gioăng EPDM đa khoang 3 lớp / Dow Corning 795","MED xi mạ 2 màu trong/ngoài, PVDF",
     "Cứng vững, cách âm 38dB, bảo mật RC2, phụ kiện Châu Âu đầy đủ","Giá trung-cao, cần thợ chuyên nghiệp","✓","✓","✓","✓","Hoàn thành",90,"https://civrowindow.com",""),

    ("CV-003","Anodized Siêu cao cấp","AW75 IN/OU","Civro","Việt Nam","EN 12207/12210",
     "Cửa sổ","Mở quay + Mở hất (Casement + Awning)","Khung 75mm cầu PA66 dày","Ngoại thất","Biệt thự cao cấp, khách sạn resort 5 sao",
     2.0,"10-14mm / Kính hộp 24-36mm","16mm","2000×3000",200,"Rãnh C Châu Âu cao cấp","Gioăng EPDM 4 lớp / Dow Corning 791","MED xi mạ 2 màu, PVDF 2 lớp, wood finish",
     "Profile dày nhất dòng Civro, cách âm vượt trội, Passive House chuẩn Châu Âu","Giá rất cao, thi công phức tạp","✓","✓","✓","✓","Hoàn thành",85,"https://civrowindow.com",""),

    ("CV-004","Anodized Cao cấp","AD55 IN/OU","Civro","Việt Nam","EN 14351-1",
     "Cửa đi","Mở quay vào/ra (Casement IN/OUT)","Khung cửa đi 55mm cầu PA66","Ngoại thất","Nhà phố, văn phòng",
     1.5,"5-10mm","10mm","1200×2400",180,"Rãnh C Euro Groove","Gioăng EPDM 3 lớp / Silicon Dow","MED xi mạ độc quyền",
     "Cửa đi hệ Châu Âu giá tốt, cách âm tốt","Ít hệ chuyên dụng hơn AD65","✓","✓","✓","✓","Hoàn thành",80,"https://civrowindow.com",""),

    ("CV-005","Anodized Cao cấp","AD65 IN/OU","Civro","Việt Nam","EN 14351-1",
     "Cửa đi","Mở quay vào/ra (Casement IN/OUT)","Khung cửa đi 65mm cầu PA66","Ngoại thất","Biệt thự, văn phòng cao cấp",
     1.8,"6-14mm / Kính hộp 20-32mm","14mm","1400×2800",180,"Rãnh C Châu Âu","Gioăng EPDM đa khoang / Dow Corning","MED xi mạ độc quyền, PVDF",
     "Cứng vững, cách âm vượt trội, phù hợp cửa đi đơn 2 cánh lớn","Giá cao","✓","✓","✓","✓","Hoàn thành",90,"https://civrowindow.com",""),

    ("CV-006","Anodized Siêu cao cấp","AD65 IN_HB/OU_HB","Civro","Việt Nam","EN 14351-1",
     "Cửa đi","Mở quay – Đố ngang lớn (Panorama Door)","Khung 65mm + Đố ngang bản rộng","Ngoại thất","Biệt thự siêu cao cấp, sảnh khách sạn",
     2.0,"6-14mm / Kính hộp lớn","14mm","2000×3200",250,"Rãnh C đặc biệt panorama","Gioăng EPDM 4 lớp / Dow Corning 795","MED xi mạ độc quyền, PVDF 2 lớp",
     "Đố ngang siêu rộng tạo cửa panorama, thẩm mỹ đỉnh cao","Thiết kế kết cấu phức tạp","✓","✓","✓","✓","Đang nghiên cứu",70,"https://civrowindow.com",""),

    ("CV-007","Anodized Siêu cao cấp","AD75 IN/OU","Civro","Việt Nam","EN 14351-1",
     "Cửa đi","Mở quay vào/ra (Casement IN/OUT)","Khung cửa đi 75mm siêu dày cầu PA66","Ngoại thất","Biệt thự siêu sang, resort 5 sao",
     2.0,"10-16mm / Kính hộp lớn","16mm","2000×3200",250,"Rãnh C Châu Âu siêu cao cấp","Gioăng EPDM 5 lớp / Dow Corning","MED xi mạ, PVDF 2 lớp, wood-look",
     "Profile 75mm dày nhất, cách âm cao nhất dòng Civro","Giá rất cao","✓","✓","✓","✓","Hoàn thành",85,"https://civrowindow.com",""),

    ("CV-008","Trượt lùa Cao cấp","CSD110","Civro","Việt Nam","EN 13830",
     "Cửa đi","Trượt lùa (Sliding)","Khung CSD 110mm + Cánh + Ray inox","Ngoại thất","Biệt thự, lối ra ban công sân hồ bơi",
     2.0,"10-15mm / Kính hộp 24-36mm","15mm","4000×2600",200,"Ray inox SUS316L","Gioăng EPDM kép / Silicon cấu trúc","MED xi mạ, PVDF",
     "Phù hợp biệt thự biển, chịu muối, trượt êm ray inox","Nặng, cần tường rộng để cánh trượt vào","✓","✓","✓","✓","Hoàn thành",80,"https://civrowindow.com",""),

    ("CV-009","Trượt lùa Cao cấp","CSD120","Civro","Việt Nam","EN 13830",
     "Cửa đi","Trượt lùa (Sliding)","Khung CSD 120mm + Cánh + Ray inox","Ngoại thất","Biệt thự cao cấp, resort",
     2.0,"10-15mm / Kính hộp 24-36mm","15mm","5000×2800",250,"Ray inox SUS316L","Gioăng EPDM kép / Silicon cấu trúc","MED xi mạ, PVDF",
     "Cánh lùa rộng hơn CSD110, phù hợp không gian mở rộng","Chi phí cao","✓","✓","✓","✓","Hoàn thành",80,"https://civrowindow.com",""),

    ("CV-010","Lift & Slide Siêu cao cấp","Civro Lift Slide","Civro","Việt Nam","EN 13830/12208",
     "Cửa đi","Trượt nâng (Lift & Slide)","Khung CSD cách nhiệt + Phụ kiện Lift Slide","Ngoại thất","Biệt thự siêu sang, biệt thự biển, penthouse",
     2.0,"Kính hộp 28-52mm / Kính 3 lớp","19mm","7000×3000",500,"Siegenia Patio Life / Roto Patio Fold","Gioăng EPDM cầu cách nhiệt / Dow Corning","MED xi mạ, PVDF 2 lớp",
     "Kín hoàn toàn, cách âm tốt nhất dòng lùa, mở nhẹ dù 500kg","Giá cao nhất dòng lùa","✓","✓","✓","✓","Hoàn thành",85,"https://civrowindow.com",""),

    ("CV-011","Pivot Siêu cao cấp","Civro Pivot","Civro","Việt Nam","EN 14351-1",
     "Cửa đi","Xoay trục (Pivot)","Cánh xoay trục + Pivot floor spring","Ngoại thất","Biệt thự siêu sang, villa resort",
     2.0,"Kính cường lực 15-22mm","22mm","2000×4000",500,"FritsJurgens / Dorma Pivot","Keo silicon kết cấu / Gioăng cạnh kín","MED xi mạ cao cấp, PVDF",
     "Thẩm mỹ cực cao, xoay 180°, cánh siêu rộng một tấm kính","Giá cao, thi công phức tạp","✓","✗","✓","✗","Đang nghiên cứu",55,"https://civrowindow.com",""),

    ("CV-012","Xếp trượt Cao cấp","Civro Folding","Civro","Việt Nam","EN 14351-1",
     "Cửa đi","Xếp trượt (Bi-fold / Folding)","Cánh xếp gấp + Ray trên đặc biệt","Ngoại thất","Biệt thự, nhà hàng, quán cà phê, resort",
     1.8,"Kính cường lực 8-12mm","12mm","5000×2800",40,"Roto Fold / OPK Folding","Gioăng EPDM / Silicon","MED xi mạ, sơn tĩnh điện",
     "Mở toàn bộ không gian, xếp gọn hoàn toàn","Kín khít kém nhất","✓","✗","✓","✓","Đang nghiên cứu",60,"https://civrowindow.com",""),

    ("CV-013","Curtain Wall Stick","CCW50 Stick","Civro","Việt Nam","EN 13830",
     "Vách mặt dựng","Cố định (Fixed – Stick System)","Đố đứng + Đố ngang + Nẹp kính lắp tại chỗ","Ngoại thất – Mặt dựng","Văn phòng, TTTM, khách sạn",
     2.0,"Kính hộp 20-36mm / Kính 3 lớp 44mm","36mm","Không giới hạn",500,"Hệ Stick Châu Âu","Silicon kết cấu SG2000 / Sika","Anodized, PVDF",
     "Linh hoạt thiết kế, thi công từng thanh tại chỗ, dễ sửa chữa","Thi công chậm hơn Unitized","✓","✓","✓","✓","Đang nghiên cứu",65,"https://civrowindow.com",""),

    # ══════ MAXPRO JP ══════
    ("MP-001","Anodized Phổ thông","R55","Maxpro JP","Việt Nam","TCVN / JIS A 4702",
     "Cửa đi + Cửa sổ","Mở quay (Casement)","Khung 55mm + Cánh","Nội/Ngoại thất","Nhà phố, chung cư, nhà xưởng",
     1.4,"5-10mm","10mm","1200×2200",80,"Rãnh C / Rãnh 22","Gioăng EPDM 2 khoang / Silicon trung tính","Anodized ED chuẩn Nhật JIS H8602 – bền 40 năm",
     "Anodize ED bền màu 40 năm, phụ kiện CMECH nguyên bộ, tỷ lệ bản rộng Nhật Bản tinh tế","Phân phối ít ở miền Bắc","✓","✓","✓","✓","Hoàn thành",85,"https://maxprojp.com.vn",""),

    ("MP-002","Anodized Trung cao cấp","R65","Maxpro JP","Việt Nam","TCVN / JIS A 4702",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất (Casement + Awning)","Khung 65mm + Cánh","Ngoại thất","Biệt thự, nhà phố cao cấp, văn phòng hạng B",
     1.6,"5-12mm / Kính hộp 20-28mm","12mm","1600×2800",150,"Rãnh C Euro Groove","Gioăng EPDM đa khoang / Silicon Sika 11FC","Anodized ED Nhật bền màu, sơn PVDF",
     "Hệ phụ kiện Nhật tích hợp sẵn, profile thanh lịch hơn hệ Trung Quốc","Giá cao hơn hàng TQ 20-30%","✓","✓","✓","✓","Hoàn thành",85,"https://maxprojp.com.vn",""),

    ("MP-003","Anodized Siêu cao cấp","R65 Plus","Maxpro JP","Việt Nam","TCVN / JIS",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất (Casement + Awning)","Khung 65mm Plus + thêm gioăng giữa","Ngoại thất","Biệt thự cao cấp",
     1.6,"5-20mm / Kính hộp","20mm","1600×2800",150,"Rãnh C Châu Âu","Gioăng EPDM 3 khoang (cải tiến) / Silicon","Anodized ED Nhật",
     "Cải tiến thêm gioăng giữa, kín khít hơn R65 thường","Giá cao hơn R65","✓","✓","✓","✓","Hoàn thành",82,"https://maxprojp.com.vn",""),

    ("MP-004","Tân cổ điển Cao cấp","R70","Maxpro JP","Việt Nam","TCVN / JIS",
     "Cửa đi","Mở quay – Vát cạnh tân cổ điển","Khung 70mm vát cạnh mạnh","Ngoại thất","Biệt thự tân cổ điển",
     2.0,"5-15mm","15mm","1400×2800",180,"Rãnh C Châu Âu","Gioăng EPDM / Silicon","Anodized ED, sơn gold champagne, vân gỗ",
     "Profile vát cạnh đặc trưng tân cổ điển, không trùng hệ khác","Không phù hợp phong cách hiện đại","✓","✓","✓","✓","Hoàn thành",80,"https://maxprojp.com.vn",""),

    ("MP-005","Tân cổ điển Siêu cao cấp","R83","Maxpro JP","Việt Nam","TCVN / JIS",
     "Cửa đi + Cửa vòm + Cửa sổ","Mở quay + Mở hất","Khung 83mm vát cạnh + Cánh vòm cung","Ngoại thất","Biệt thự cổ điển Pháp, biệt thự vòm",
     1.8,"5-25mm","25mm","2000×3500",200,"Rãnh C Châu Âu","Gioăng EPDM đa khoang / Dow Corning","Anodized, sơn vàng gold, vân gỗ",
     "Hệ tân cổ điển đỉnh cao, làm được cửa vòm cung phức tạp","Chỉ phù hợp phong cách cổ điển","✓","✓","✓","✓","Hoàn thành",80,"https://maxprojp.com.vn",""),

    ("MP-006","Thủy lực Siêu cao cấp","R200 Thủy lực","Maxpro JP","Việt Nam","TCVN",
     "Cửa đi","Thủy lực (Hydraulic Door)","Khung 200mm siêu nặng + Hệ thủy lực","Ngoại thất","Biệt thự siêu cao cấp, cửa đại sảnh",
     2.0,"Kính cường lực 15-22mm","22mm","3000×3500",800,"Thủy lực đặc chủng Dorma/Biloba","Keo silicon kết cấu / EPDM giả thép","Anodized ED, sơn RAL đặc chủng",
     "Bản cánh 200mm siêu rộng, đóng mở êm tuyệt đối bằng hệ thủy lực cân bằng","Giá rất cao, gia cố nền sàn đặc biệt","✓","✓","✓","✓","Hoàn thành",75,"https://maxprojp.com.vn",""),

    ("MP-007","Trượt lùa Trung cấp","SW55","Maxpro JP","Việt Nam","TCVN",
     "Cửa sổ","Trượt lùa (Sliding Window)","Khung SW 55mm + Cánh trượt","Ngoại thất","Nhà phố phổ thông",
     1.4,"5-10mm","10mm","2400×1800",80,"Rãnh 22 / Rãnh C","Gioăng EPDM / Silicon","Anodized ED Nhật, sơn tĩnh điện",
     "Trượt êm, giá cạnh tranh, Anodized ED bền 40 năm","Kín khít trung bình","✓","✗","✓","✓","Hoàn thành",75,"https://maxprojp.com.vn",""),

    ("MP-008","Trượt lùa Cao cấp","SD83","Maxpro JP","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa phẳng (Sliding Door)","Khung SD 83mm + Cánh","Ngoại thất","Nhà phố, biệt thự",
     1.6,"8-12mm / Kính hộp 20-28mm","12mm","4000×2400",150,"Ray inox SUS304","Gioăng EPDM / Dow Corning","Anodized ED Nhật",
     "Trượt êm hệ Nhật, Anodized ED bền 40 năm","Kín khít kém hơn Lift Slide","✓","✓","✓","✓","Hoàn thành",80,"https://maxprojp.com.vn",""),

    ("MP-009","Trượt lùa Siêu cao cấp","SD115 Chống bão","Maxpro JP","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa chịu bão (Storm Resistant Sliding)","Khung SD 115mm + Cánh siêu dày","Ngoại thất – Biển","Biệt thự biển, resort duyên hải chịu bão cấp 14",
     2.5,"Kính hộp chống bão 24-36mm","15mm","5000×2600",300,"Ray inox SUS316L chống muối biển","Gioăng EPDM cầu cách nhiệt / Silicon Dow Corning","Anodized ED chống muối biển AS 1734-5",
     "Chịu gió bão cấp 14, kín nước EN 12208 Cls 9A, chịu muối biển lâu dài","Giá rất cao, nặng","✓","✓","✓","✓","Hoàn thành",80,"https://maxprojp.com.vn",""),

    ("MP-010","Xếp trượt Cao cấp","SFD80","Maxpro JP","Việt Nam","TCVN",
     "Cửa đi","Xếp trượt (Folding Sliding Door)","Cánh xếp gấp 80mm + Bản lề xếp","Ngoại thất","Biệt thự, nhà hàng, không gian mở",
     1.8,"Kính cường lực 8-12mm","12mm","5000×2800",40,"OPK Bi-fold / Roto Fold","Gioăng EPDM / Silicon","Anodized ED Nhật, sơn tĩnh điện",
     "Mở toàn bộ không gian Nhật, bền bỉ phụ kiện Nhật","Kín khít thấp nhất trong hệ lùa","✓","✓","✓","✓","Đang nghiên cứu",65,"https://maxprojp.com.vn",""),

    ("MP-011","Lam chắn nắng","GW65","Maxpro JP","Việt Nam","TCVN",
     "Lam chắn nắng","Cố định / Xoay (Fixed / Motorized Louvres)","Profile lam 65mm xoay","Ngoại thất","Mặt ngoài công trình, ban công, sân thượng",
     1.6,"Không dùng kính","N/A","Không giới hạn",50,"Hệ xoay cơ / Mô tơ tự động","Không cần keo","Anodized, sơn tĩnh điện chịu UV",
     "Chắn nắng điều tiết ánh sáng, thẩm mỹ hiện đại mặt công trình","Cần bảo trì mô tơ định kỳ","✓","✗","✓","✓","Đang nghiên cứu",55,"https://maxprojp.com.vn",""),

    # ══════ SOCO ══════
    ("SC-001","Anodized Cao cấp","Soco 65","Soco","Việt Nam","EN Châu Âu",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất (Casement + Awning)","Khung 65mm rãnh C","Ngoại thất","Biệt thự, nhà phố cao cấp",
     1.6,"5-12mm / Kính hộp 20-28mm","12mm","1600×2800",150,"Rãnh C Châu Âu – Cmech/Sigico/Bogo","Gioăng EPDM đa khoang / Silicon","Anodized vàng ánh kim Champagne đặc trưng Soco",
     "Màu Champagne đặc trưng sang trọng, rãnh C Châu Âu đầy đủ phụ kiện","Ít phân phối ngoài miền Nam","✓","✓","✓","✓","Hoàn thành",80,"https://nhomsocowindow.com",""),

    ("SC-002","Trượt lùa Trung cấp","Soco 94","Soco","Việt Nam","EN Châu Âu",
     "Cửa đi","Trượt lùa 2 ray (Sliding)","Khung 94mm + Ray inox","Ngoại thất","Nhà phố, biệt thự",
     1.6,"6-12mm","12mm","4000×2600",150,"Ray inox SUS304","Gioăng EPDM / Silicon","Anodized Champagne, sơn tĩnh điện",
     "Trượt êm ray inox, màu Champagne sang trọng","Kín khít trung bình","✓","✓","✓","✓","Hoàn thành",78,"https://nhomsocowindow.com",""),

    ("SC-003","Trượt lùa Cao cấp","Soco 120","Soco","Việt Nam","EN Châu Âu",
     "Cửa đi","Trượt lùa 2 ray (Sliding)","Khung 120mm + Ray inox SUS316","Ngoại thất","Biệt thự, ban công rộng",
     1.8,"6-15mm / Kính hộp 24-36mm","15mm","5000×2800",200,"Ray inox SUS316L","Gioăng EPDM kép / Dow Corning","Anodized Champagne ED",
     "Phù hợp biệt thự biển, ray inox 316 chống muối","Giá cao","✓","✓","✓","✓","Hoàn thành",80,"https://nhomsocowindow.com",""),

    ("SC-004","Trượt lùa Siêu cao cấp","Soco 180","Soco","Việt Nam","EN Châu Âu",
     "Cửa đi","Trượt lùa 3 ray (Sliding)","Khung 180mm + Ray inox SUS316","Ngoại thất","Biệt thự, resort cao cấp",
     2.0,"Kính hộp đến 36mm","15mm","7000×2800",250,"Ray inox SUS316L","Gioăng EPDM 3 lớp / Dow Corning","Anodized Champagne ED chống muối biển",
     "Cửa lùa 3 ray, mở rộng tối đa, chống muối biển","Chi phí cao","✓","✓","✓","✓","Hoàn thành",75,"https://nhomsocowindow.com",""),

    ("SC-005","Lift & Slide Siêu cao cấp","Soco LS100","Soco","Việt Nam","EN 13830/12208",
     "Cửa đi","Trượt nâng (Lift & Slide)","Khung nặng đặc biệt + Ray Lift Slide","Ngoại thất","Biệt thự siêu sang, penthouse, resort",
     2.5,"Kính hộp 32-60mm / Kính 3 lớp","22mm","8000×3200",700,"Roto Patio Life / Siegenia Titan LSS","Gioăng EPDM cầu cách nhiệt 4 lớp","Anodized ED biển chống ăn mòn, PVDF",
     "Lift Slide rộng nhất Việt Nam (8m), nâng êm 700kg, kín khít EN 12208 Cls 9A","Giá 500tr – 2 tỷ/bộ","✓","✓","✓","✓","Hoàn thành",90,"https://nhomsocowindow.com",""),

    ("SC-006","Xếp trượt","Soco Folding","Soco","Việt Nam","EN",
     "Cửa đi","Xếp trượt (Folding/Bi-fold)","Cánh xếp gấp + Ray trên","Ngoại thất","Biệt thự, nhà hàng",
     1.6,"Kính cường lực 8-12mm","12mm","5000×2800",40,"OPK Folding / Roto Fold","Gioăng EPDM / Silicon","Anodized Champagne",
     "Mở toàn bộ, màu Champagne đặc trưng","Kín khít kém nhất","✓","✗","✓","✓","Đang nghiên cứu",60,"https://nhomsocowindow.com",""),

    ("SC-007","Vách kính Cố định","Soco Fix Vách","Soco","Việt Nam","EN 13830",
     "Vách kính","Cố định (Fixed Glass Partition/Facade)","Profile vách 65mm cố định","Ngoại thất + Nội thất","Văn phòng, TTTM, công trình",
     1.6,"Kính hộp 20-32mm","32mm","Không giới hạn",500,"Hệ Stick Châu Âu","Silicon kết cấu Dow Corning","Anodized Champagne, sơn PVDF",
     "Thiết kế linh hoạt, cùng hệ màu Champagne đặc trưng Soco","Thi công phụ thuộc kiểm soát tại công trường","✓","✓","✓","✓","Đang nghiên cứu",60,"https://nhomsocowindow.com",""),

    ("SC-008","Curtain Wall Unitized","Soco UCW Unitized","Soco","Việt Nam","EN 13830/AAMA 501",
     "Vách mặt dựng","Panel đúc sẵn (Unitized Panel)","Panel hoàn chỉnh đúc sẵn tại xưởng Soco","Ngoại thất – Tòa nhà cao tầng","Tòa nhà văn phòng hạng A, khách sạn 5 sao 30+ tầng",
     2.5,"Kính hộp Low-E 28-48mm / Kính 3 lớp 60mm","48mm","Không giới hạn (module)",1000,"Hệ Unitized panel + crane","Silicon kết cấu Dow Corning AAMA chuẩn","Anodized, PVDF 70%",
     "Thi công nhanh 3-4 ngày/tầng, chất lượng đồng đều xưởng, chuẩn AAMA tòa nhà cao tầng","Chi phí mold cao, không linh hoạt sửa đổi tại công trường","✓","✓","✓","✓","Đang nghiên cứu",65,"https://nhomsocowindow.com",""),

    # ══════ PAG ══════
    ("PG-001","Cầu cách nhiệt Trung cấp","PAG 60","PAG","Việt Nam","EN 14351-1",
     "Cửa sổ + Cửa đi","Mở quay (Casement)","Khung 60mm cầu PA66 GF25","Ngoại thất – Công trình xanh","Nhà phố cao cấp, chung cư xanh",
     1.6,"Kính hộp Low-E 20-28mm","14mm","1200×2400",100,"Rãnh C cách nhiệt","Gioăng EPDM cách nhiệt 2 lớp / Dow Corning 795","AkzoNobel chống muối biển – bảo hành 25 năm",
     "Uf < 2.5 W/(m²K), giải pháp cách nhiệt giá tốt nhất Việt Nam","Profile hẹp hơn hệ 65","✓","✓","✓","✓","Hoàn thành",85,"https://nhompag.com.vn",""),

    ("PG-002","Cầu cách nhiệt Cao cấp","PAG 80 Lùa","PAG","Việt Nam","EN 14351-1",
     "Cửa sổ","Trượt lùa (Sliding) + Chống côn trùng","Khung cầu PA66 80mm + Lưới chắn côn trùng","Ngoại thất","Nhà phố, biệt thự ở vùng nhiều muỗi, côn trùng",
     1.6,"Kính hộp Low-E 20-28mm","14mm","2400×1800",100,"Rãnh C cách nhiệt","Gioăng EPDM cách nhiệt / Dow Corning","AkzoNobel bền 25 năm",
     "Kết hợp cách nhiệt + lưới chắn côn trùng EPDM, tiện lợi khí hậu nhiệt đới","Kín khít cửa lùa trung bình","✓","✗","✓","✓","Hoàn thành",75,"https://nhompag.com.vn",""),

    ("PG-003","Xếp trượt Cao cấp","PAG 83 Folding","PAG","Việt Nam","EN 14351-1",
     "Cửa đi","Xếp trượt cầu cách nhiệt (Thermal Bi-fold)","Cánh xếp gấp 83mm cầu cách nhiệt","Ngoại thất","Biệt thự cao cấp",
     1.8,"Kính cường lực 8-12mm","12mm","5000×2800",40,"Rãnh C cách nhiệt + Roto Fold","Gioăng EPDM cách nhiệt / Dow Corning","AkzoNobel chống muối biển",
     "Xếp trượt CÓ cầu cách nhiệt, hiếm có trên thị trường Việt Nam","Kín khít không hoàn toàn","✓","✗","✓","✓","Đang nghiên cứu",65,"https://nhompag.com.vn",""),

    ("PG-004","Trượt nâng Cao cấp","PAG 120","PAG","Việt Nam","EN 14351-1",
     "Cửa đi","Trượt nâng + Lùa vuông góc (Lift Slide + Inline Slide)","Khung cầu PA66 + Phụ kiện Lift đặc biệt","Ngoại thất","Biệt thự cao cấp, resort",
     1.8,"Kính hộp Low-E 24-36mm","16mm","6000×2800",300,"Siegenia Patio / Roto Patio Life","Gioăng EPDM cầu cách nhiệt / Dow Corning","AkzoNobel chống muối biển",
     "Kết hợp Lift Slide VÀ Pocket Slide lùa vuông góc – không giới hạn số cánh","Thiết kế kỹ lưỡng, giá cao","✓","✓","✓","✓","Đang nghiên cứu",65,"https://nhompag.com.vn",""),

    ("PG-005","Lift & Slide Siêu cao cấp","PAG 125","PAG","Việt Nam","EN 14351-1",
     "Cửa đi","Trượt nâng (Lift & Slide)","Khung cầu PA66 + Phụ kiện LSS đặc biệt","Ngoại thất","Biệt thự siêu sang, resort",
     2.0,"Kính hộp Low-E 24-44mm","18mm","7000×3000",400,"Siegenia Titan LSS / Roto Patio Life","Gioăng EPDM cầu cách nhiệt 4 lớp / Dow Corning","AkzoNobel bền 25 năm, anodized 2 màu",
     "Lift Slide CÓ cầu cách nhiệt PAG đầu tiên Việt Nam, cách âm 42dB, Uf < 1.8 W/(m²K)","Giá cao, sản xuất theo đơn","✓","✓","✓","✓","Đang nghiên cứu",65,"https://nhompag.com.vn",""),

    ("PG-006","Lift & Slide Siêu cao cấp","PAG 200","PAG","Việt Nam","EN 14351-1",
     "Cửa đi","Trượt nâng bản lớn (Heavy Lift & Slide)","Khung cầu PA66 200mm siêu dày","Ngoại thất","Biệt thự siêu sang, penthouse",
     2.0,"Kính hộp Low-E 28-52mm","20mm","8000×3200",400,"Roto Patio Life đặc biệt","Gioăng EPDM cầu cách nhiệt 4 lớp","AkzoNobel bền 25 năm",
     "Lift Slide bản rộng nhất PAG, tải trọng 400kg/cánh, cách nhiệt kép","Giá dự án cao","✓","✓","✓","✓","Đang nghiên cứu",60,"https://nhompag.com.vn",""),

    ("PG-007","Hàn góc liền khối","PAG Seamless","PAG","Việt Nam","EN 14351-1",
     "Cửa sổ + Cửa đi","Mở quay + Hàn góc CNC (Seamless Welded)","Khung hàn nhiệt liền khối CNC – không vết ghép","Ngoại thất","Biệt thự siêu cao cấp, villa resort",
     2.0,"Kính hộp Low-E 24-44mm / Kính 3 lớp","16mm","1400×2800",150,"Rãnh C cách nhiệt đặc biệt","Gioăng EPDM 4 lớp / Dow Corning 791","Anodized tách màu siêu mịn – không vết ghép góc",
     "Góc hàn nhiệt liền khối hoàn toàn không vết ghép – thẩm mỹ cao nhất, kín tuyệt đối","Cần máy hàn CNC đặc chủng, +30% so với ép góc","✓","✓","✓","✓","Hoàn thành",85,"https://nhompag.com.vn",""),

    ("PG-008","Chống côn trùng","PAG 140","PAG","Việt Nam","EN",
     "Cửa đi","Lùa chống côn trùng (Insect Screen Door)","Khung đặc biệt + Lưới chắn côn trùng EPDM","Ngoại thất","Nhà phố, biệt thự ở vùng nhiều muỗi",
     1.4,"Không dùng kính (lưới inox 304)","N/A","2400×2400",60,"Rãnh C / Lưới inox SUS304","Không cần keo","Anodized, sơn tĩnh điện",
     "Chắn côn trùng, thông gió tối đa, phù hợp khí hậu nhiệt đới","Không cách âm, không cách nước","✓","✗","✓","✓","Hoàn thành",65,"https://nhompag.com.vn",""),

    # ══════ OWIN ══════
    ("OW-001","Anodized Phổ thông","Owin Hệ 55","Owin","Việt Nam","EN Tiêu chuẩn Đức",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất","Khung 55mm Anodized-ED Luxanode","Ngoại thất","Nhà phố, văn phòng",
     1.4,"5-12mm","12mm","1200×2200",80,"Rãnh C / Kinlong / Cmech","Gioăng EPDM / Silicon trung tính","Anodized-ED Luxanode chuẩn Đức – chống phai màu",
     "Anodized-ED Luxanode chuẩn Đức bền bỉ, hệ phụ kiện Kinlong/Cmech nguyên bộ","Ít đa dạng hệ chuyên dụng","✓","✓","✓","✓","Hoàn thành",75,"",""),

    ("OW-002","Anodized Cao cấp","Owin Hệ 65","Owin","Việt Nam","EN Tiêu chuẩn Đức",
     "Cửa đi + Cửa sổ","Mở quay + Trượt lùa","Khung 65mm Anodized-ED","Ngoại thất","Biệt thự, nhà phố cao cấp",
     2.0,"5-15mm / Kính hộp 20-28mm","15mm","1600×2800",150,"Rãnh C / Cmech / Roto","Gioăng EPDM đa khoang / Silicon","Anodized-ED Luxanode",
     "Profile cứng, Anodized-ED Luxanode bền bỉ, đa dạng ứng dụng","Ít hệ chuyên dụng hơn Civro","✓","✓","✓","✓","Hoàn thành",75,"",""),

    ("OW-003","Trượt lùa Cao cấp","Owin Hệ 100","Owin","Việt Nam","EN Tiêu chuẩn Đức",
     "Cửa đi","Trượt lùa + Xếp gấp (Sliding + Folding)","Khung 100mm + Cánh đặc biệt","Ngoại thất","Biệt thự, ban công rộng",
     2.0,"Kính hộp 20-32mm","15mm","5000×2800",200,"Rãnh C / Cmech","Gioăng EPDM kép / Silicon","Anodized-ED Luxanode",
     "Hệ lùa + xếp gấp 100mm linh hoạt","Chi phí cao","✓","✓","✓","✓","Đang nghiên cứu",65,"",""),

    ("OW-004","Thủy lực Siêu cao cấp","Owin HL180","Owin","Việt Nam","TCVN",
     "Cửa đi","Thủy lực 2 chiều (Hydraulic Pivot)","Khung siêu nặng 180×60mm cánh + Piston thủy lực","Ngoại thất","Biệt thự siêu sang, sảnh khách sạn 5 sao",
     2.5,"Kính cường lực/laminated 15-22mm","22mm","5000×4000",1500,"Thủy lực Biloba / Floor Spring Dorma TS98","Keo silicon kết cấu / Gioăng tổng hợp","Anodized-ED Luxanode, sơn đặc chủng vàng/đồng/titan",
     "Cửa siêu rộng một cánh (đến 5m), đóng mở bằng một ngón tay, tự đóng piston thủy lực","Giá 200tr-vài tỷ, gia cố nền đặc biệt","✓","✓","✓","✓","Hoàn thành",85,"",""),

    ("OW-005","Trượt quay Cao cấp","Owin SR Trượt quay","Owin","Việt Nam","TCVN",
     "Cửa đi","Trượt quay (Sliding Rotating / Sliding Pivot)","Khung + Cánh + Ray treo trên đặc biệt","Ngoại thất","Biệt thự, nhà phố sang trọng, phòng trưng bày",
     2.0,"Kính cường lực 10-15mm","15mm","2000×3000",150,"Slido Dorma / Häfele trượt quay","Gioăng EPDM / Silicon","Anodized-ED Luxanode, sơn mờ",
     "Không ray dưới sàn, đóng mở 180° trơn tru, thẩm mỹ cao","Tải trọng cánh hạn chế, lắp đặt chính xác cao","✓","✓","✓","✓","Hoàn thành",80,"",""),

    ("OW-006","Vách kính","Owin Vách kính","Owin","Việt Nam","EN",
     "Vách kính","Cố định (Fixed Glass)","Profile vách ngăn kính","Nội/Ngoại thất","Văn phòng, công trình",
     1.6,"Kính đơn/hộp 8-24mm","24mm","Không giới hạn",400,"Hệ Stick Owin","Silicon kết cấu","Anodized-ED Luxanode",
     "Anodized-ED bền bỉ, phù hợp vách kính thương mại","Ít linh hoạt hơn hệ Civro","✓","✓","✓","✓","Đang nghiên cứu",60,"",""),

    ("OW-007","Lan can & Chấn song","Owin Chấn song","Owin","Việt Nam","TCVN 9392",
     "Lan can + Chấn song","Cố định (Fixed Railing)","Profile chấn song nhôm + Thanh đứng","Ngoại thất + Nội thất","Ban công, cầu thang, hành lang",
     1.2,"Không dùng kính","N/A","Không giới hạn",80,"Phụ kiện chấn song đặc chủng","Không cần keo","Anodized-ED, sơn tĩnh điện",
     "Bền bỉ, đa dạng kiểu dáng, Anodized-ED chống ăn mòn","Không cách âm/nhiệt","✓","✗","✓","✓","Đang nghiên cứu",55,"",""),

    # ══════ TOPAL (Austdoor) ══════
    ("TP-001","Tân cổ điển Cao cấp","Topal Prima 55+","Topal","Việt Nam","EN Châu Âu / Rãnh C",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất + Trượt lùa","Cánh trơn không gân, bản khung 63mm","Ngoại thất","Biệt thự, căn hộ cao cấp",
     1.6,"Kính đơn 6.38mm đến kính hộp đầy đủ","12mm","1600×2800",120,"Rãnh C Châu Âu đồng bộ Topal/Austdoor","Gioăng EPDM đa khoang / Silicon","Sơn tĩnh điện, ép vân gỗ, anodized – bảo hành 5 năm",
     "Cánh trơn không gân – profile thanh lịch nhất dòng Topal, rãnh C Châu Âu đồng bộ, cách âm ~28dB","Chỉ rõ ràng phong cách hiện đại, ít mẫu cổ điển","✓","✓","✓","✓","Hoàn thành",80,"https://topal.vn",""),

    ("TP-002","Anodized Mỏng hiện đại","Topal Slima 48","Topal","Việt Nam","EN Châu Âu / Rãnh C",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất","Cánh mảnh 48mm, rãnh C, tối giản hiện đại","Ngoại thất","Nhà phố, chung cư cao cấp phong cách hiện đại",
     1.2,"Kính đơn/hộp","10mm","1400×2600",100,"Rãnh C Châu Âu Topal/Austdoor","Gioăng EPDM / Silicon","Sơn tĩnh điện bột, anodized bạc",
     "Profile siêu mỏng 48mm – nhìn nhiều kính nhất trong dòng Topal, tối giản phong cách Châu Âu, cách âm ~28dB","Tải trọng giới hạn do profile mỏng","✓","✓","✓","✓","Hoàn thành",78,"https://topal.vn",""),

    ("TP-003","Phổ thông Rãnh 22","Topal XFAD 55","Topal","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay + Trượt lùa","Profile Xingfa 55mm rãnh 22","Ngoại thất phổ thông","Nhà phố, tối ưu chi phí",
     1.4,"Kính đơn/hộp 5-12mm","12mm","1200×2400",80,"Rãnh 22mm phổ thông","Gioăng EPDM / Silicon","Sơn tĩnh điện, đa màu",
     "Giá cạnh tranh nhất dòng Topal, phổ biến dự án ngân sách trung bình","Profile rãnh 22 kém phụ kiện Châu Âu","✓","✗","✓","✓","Hoàn thành",70,"https://topal.vn",""),

    ("TP-004","Phổ thông tiết kiệm","Topal XFEC 55","Topal","Việt Nam","TCVN",
     "Cửa sổ + Cửa đi","Mở quay","Profile phổ thông dày 1.0-1.4mm","Nội/Ngoại thất phổ thông","Công trình ngân sách thấp, nhà ở phổ thông",
     1.1,"Kính đơn 5-10mm","10mm","1000×2000",60,"Rãnh 22mm","Gioăng EPDM / Silicon","Sơn tĩnh điện",
     "Giá rẻ nhất dòng Topal, phù hợp nhà ở bình dân","Profile mỏng nhất, không phù hợp yêu cầu cao","✓","✗","✓","✓","Hoàn thành",60,"https://topal.vn",""),

    # ══════ EUROVN ══════
    ("EV-001","Anodized Trung cao cấp","EuroVN Gold","EuroVN","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay + Mở hất","Khung vát cạnh + Sơn Wagner Thụy Sỹ","Ngoại thất","Biệt thự, nhà phố",
     1.2,"5-12mm","12mm","1400×2600",100,"Rãnh C","Gioăng EPDM / Silicon","Sơn Wagner Thụy Sỹ + ép vân gỗ hàn túi – bảo hành 10 năm",
     "Sơn Wagner Thụy Sỹ bền màu 10 năm, vát cạnh tạo hiệu ứng 3D, nhôm 6063T5 >98% nguyên chất","Profile mỏng 1.2mm","✓","✓","✓","✓","Hoàn thành",72,"",""),

    ("EV-002","Anodized Cao cấp","EuroVN VIP","EuroVN","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ + Trượt","Đa kiểu mở","Khung vát cạnh + dập gân chìm, nhiều hệ","Ngoại thất","Biệt thự, công trình cao cấp",
     1.4,"5-15mm / Kính hộp 20-28mm","15mm","1600×2800",120,"Rãnh C","Gioăng EPDM đa khoang / Dow Corning","Sơn Wagner Thụy Sỹ, ép vân gỗ hàn túi tự động",
     "Hệ VIP toàn diện nhất dòng EuroVN, sơn Wagner bền, vát cạnh & gân chìm tăng cứng","Thị phần nhỏ hơn Civro","✓","✓","✓","✓","Hoàn thành",72,"",""),

    ("EV-003","Trượt quay Cao cấp","EuroVN Trượt quay","EuroVN","Việt Nam","TCVN",
     "Cửa đi","Trượt quay (Sliding Rotating)","Khung + Cánh + Ray đặc biệt 2 trong 1","Ngoại thất","Biệt thự, tối ưu không gian",
     1.6,"Kính cường lực 10-15mm","15mm","2000×2800",150,"Slido / Häfele trượt quay","Gioăng EPDM / Silicon","Sơn Wagner Thụy Sỹ",
     "Trượt quay 180° tối ưu không gian hành lang, thẩm mỹ cao","Giá cao hơn EuroVN thường","✓","✓","✓","✓","Hoàn thành",68,"",""),

    ("EV-004","Thủy lực Siêu cao cấp","EuroVN Thủy lực","EuroVN","Việt Nam","TCVN",
     "Cửa đi","Thủy lực 2 chiều (Hydraulic Door)","Khung 140-180mm cánh + Hệ thủy lực QueenViet","Ngoại thất","Biệt thự siêu cao cấp, showroom",
     1.8,"Kính cường lực 15-19mm","19mm","3000×3500",800,"Thủy lực đặc chủng QueenViet Group","Keo silicon kết cấu / EPDM","Sơn Wagner Thụy Sỹ đặc chủng",
     "Cửa thủy lực QueenViet Group nổi tiếng, tự đóng êm, bản cánh 140-180mm siêu rộng","Giá cao, ít nhà phân phối","✓","✓","✓","✓","Hoàn thành",70,"",""),

    ("EV-005","Cầu cách nhiệt Cao cấp","EuroVN Thermal","EuroVN","Việt Nam","EN 14351-1",
     "Cửa đi + Cửa sổ","Mở quay + Trượt","Khung + dải Polyamide cách nhiệt","Ngoại thất – Công trình xanh","Biệt thự, tiết kiệm năng lượng",
     1.6,"Kính hộp Low-E 24-36mm","16mm","1400×2800",120,"Rãnh C cách nhiệt","Gioăng EPDM cách nhiệt / Dow Corning","Sơn Wagner Thụy Sỹ 2 màu trong ngoài",
     "Cách nhiệt + sơn Wagner Thụy Sỹ bền 10 năm, tiết kiệm điện điều hòa","Thị phần nhỏ ở mảng cách nhiệt","✓","✓","✓","✓","Đang nghiên cứu",60,"",""),

    # ══════ YANGLI ══════
    ("YL-001","Phổ thông","Yangli 55 Rãnh C","Yangli","Trung Quốc/Việt Nam","QCVN 16:2019 / ISO 9001",
     "Cửa đi + Cửa sổ","Mở quay (Casement)","Khung 55mm vát cạnh 3D + Gân tăng cứng","Ngoại thất","Nhà phố, chung cư",
     1.3,"5-10mm","10mm","1200×2400",80,"Rãnh C","Gioăng EPDM / Silicon trung tính","Anodized bạc, sơn tĩnh điện bột – vát cạnh 3D",
     "Thiết kế vát cạnh tạo hiệu ứng 3D đặc trưng Yangli, giá cạnh tranh, QCVN","Chất lượng không đồng đều tùy lô hàng","✓","✗","✓","✓","Hoàn thành",70,"https://yangli.vn",""),

    ("YL-002","Phổ thông","Yangli 60/65 Rãnh C","Yangli","Trung Quốc/Việt Nam","QCVN 16:2019",
     "Cửa đi + Cửa sổ","Mở quay + Trượt lùa","Khung 60-65mm rãnh C vát cạnh","Ngoại thất","Nhà phố, biệt thự trung cấp",
     1.3,"5-12mm / Kính hộp 20-28mm","12mm","1600×2800",100,"Rãnh C","Gioăng EPDM / Silicon","Anodized bạc, sơn bột",
     "Rãnh C tiêu chuẩn, giá trung bình, đa dạng mẫu mã","Ít hệ chuyên dụng","✓","✗","✓","✓","Hoàn thành",68,"https://yangli.vn",""),

    ("YL-003","Tân cổ điển","Yangli-XF Tân cổ điển","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Cửa đi + Cửa sổ","Mở quay + Trượt lùa","Profile Xingfa mô phỏng, đa hệ","Ngoại thất","Nhà phố phong cách",
     1.3,"5-12mm","12mm","1400×2600",100,"Rãnh C / Rãnh 22","Gioăng EPDM / Silicon","Sơn bột vân gỗ, champagne vàng",
     "Đa dạng phong cách, giá tốt","Chất lượng kém hơn VRE/Civro","✓","✗","✓","✓","Hoàn thành",65,"https://yangli.vn",""),

    ("YL-004","Thủy lực","Yangli Thủy lực","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Cửa đi","Thủy lực 2 chiều","Bản cánh lớn + Hệ thủy lực","Ngoại thất","Biệt thự, showroom",
     1.6,"Kính cường lực 15-19mm","19mm","3000×3500",800,"Thủy lực đặc chủng","Silicon kết cấu","Sơn tĩnh điện đặc chủng",
     "Cửa thủy lực giá tốt hơn thương hiệu Việt Nam","Chất lượng thủy lực kém hơn Owin/Maxpro","✓","✗","✓","✓","Đang nghiên cứu",60,"https://yangli.vn",""),

    ("YL-005","Trượt lùa","Yangli Lùa 95/97","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Cửa đi + Cửa sổ","Trượt lùa (Sliding)","Khung 95-97mm + Ray inox","Ngoại thất","Nhà phố, biệt thự trung cấp",
     1.4,"5-12mm","12mm","4000×2400",150,"Ray inox SUS304","Gioăng EPDM / Silicon","Anodized, sơn tĩnh điện",
     "Ray inox bền, giá cạnh tranh","Kín khít thấp hơn Lift Slide","✓","✗","✓","✓","Hoàn thành",68,"https://yangli.vn",""),

    ("YL-006","Trượt lùa Rộng","Yangli Lùa 120","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Cửa đi","Trượt lùa 2 ray (Sliding)","Khung 120mm + Ray inox","Ngoại thất","Biệt thự",
     1.5,"Kính hộp 20-28mm","15mm","5000×2600",200,"Ray inox SUS304","Gioăng EPDM kép / Silicon","Anodized, sơn bột",
     "Cánh lùa rộng 2 ray, giá tốt","Kín khít kém hơn Lift Slide","✓","✗","✓","✓","Hoàn thành",65,"https://yangli.vn",""),

    ("YL-007","Trượt lùa 3 Ray","Yangli Lùa 180","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Cửa đi","Trượt lùa 3 ray (Sliding)","Khung 180mm + Ray inox","Ngoại thất","Biệt thự, resort",
     1.6,"Kính hộp 20-28mm","15mm","7000×2800",250,"Ray inox SUS304","Gioăng EPDM 3 lớp / Silicon","Anodized, sơn bột",
     "Mở 3 cánh, giá tốt nhất cho hệ 3 ray","Kín khít kém hơn VRE180","✓","✗","✓","✓","Đang nghiên cứu",62,"https://yangli.vn",""),

    ("YL-008","Slim hiện đại","Yangli Slim 40/45","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Cửa đi + Cửa sổ","Mở quay (Minimal Slim)","Profile siêu mỏng 40-45mm","Ngoại thất hiện đại","Nhà phố hiện đại, căn hộ tối giản",
     1.1,"Kính đơn 5-10mm","10mm","1200×2400",60,"Rãnh C mỏng","Gioăng EPDM mỏng / Silicon","Anodized bạc siêu mỏng, đen mờ",
     "Profile siêu mỏng 40-45mm – tỷ lệ kính/nhôm tối ưu nhất","Tải trọng rất giới hạn","✓","✗","✓","✓","Đang nghiên cứu",60,"https://yangli.vn",""),

    ("YL-009","Xếp trượt","Yangli Folding 80","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Cửa đi","Xếp trượt (Folding 80)","Cánh xếp gấp 80mm","Ngoại thất","Biệt thự, nhà hàng",
     1.5,"Kính cường lực 8-12mm","12mm","5000×2800",40,"Hopo Bi-fold / OPK","Gioăng EPDM / Silicon","Anodized, sơn bột",
     "Xếp trượt giá tốt","Kín khít kém nhất","✓","✗","✓","✓","Đang nghiên cứu",58,"https://yangli.vn",""),

    ("YL-010","Mặt dựng","Yangli Mặt dựng 52/65","Yangli","Trung Quốc/Việt Nam","QCVN",
     "Vách mặt dựng","Cố định (Stick System)","Module mặt dựng 52-65mm lắp tại chỗ","Ngoại thất – Công trình","Tòa nhà thương mại, công trình công nghiệp",
     1.8,"Kính hộp 20-32mm","32mm","Không giới hạn",600,"Hệ Stick Yangli","Silicon kết cấu Sika","Anodized, sơn bột",
     "Chi phí thấp nhất cho mặt dựng, phù hợp dự án ngân sách trung bình","Kém tiêu chuẩn Châu Âu","✓","✗","✓","✓","Đang nghiên cứu",55,"https://yangli.vn",""),

    # ══════ SCHÜCO – GERMANY ══════
    ("SH-001","Casement Đức Cao cấp","AWS 65","Schüco","Đức","EN 14351-1 / EN 12207 Cls4 / CE",
     "Cửa sổ + Cửa đi","Mở quay + Tilt-Turn (Casement + Tilt-Turn)","Khung 65mm cầu cách nhiệt PA66 GF25","Ngoại thất – Công trình xanh","Biệt thự, khách sạn 5 sao, văn phòng hạng A",
     2.0,"Kính hộp Low-E 24-36mm / Kính 3 lớp 44mm","18mm","1800×3000",200,"Rãnh C Euro Groove – tương thích Roto/Siegenia/Maco","Gioăng EPDM 4 lớp / Dow Corning 795","Anodized tách màu, PVDF 70%, wood-look, RAL đặc chủng",
     "Uf=1.3 W/(m²K), cách âm Rw 45dB, bền 50 năm, phụ kiện AvanTec nguyên bộ Châu Âu","Giá cao hơn VN 3-5 lần, nhập khẩu dài ngày","✓","✓","✓","✓","Hoàn thành",90,"https://www.schueco.com",""),

    ("SH-002","Casement Đức Passive House","AWS 75.SI","Schüco","Đức","EN 14351-1 / PHI Passive",
     "Cửa sổ + Cửa đi","Mở quay + Tilt-Turn","Khung 75mm cầu cách nhiệt PA66 GF25 siêu dày","Ngoại thất – Passive House","Nhà thụ động (Passive House), tòa nhà Zero Energy",
     2.0,"Kính 3 lớp 44-52mm / Kính chân không triple","20mm","2000×3500",250,"Rãnh C Châu Âu đặc biệt","Gioăng EPDM 5 lớp cách nhiệt siêu cao cấp","Anodized 2 màu, PVDF 2 lớp",
     "Uf=0.97 W/(m²K) – đạt Passive House PHI, cách âm 52dB, bền 60 năm","Giá cực cao, thị trường VN rất hạn chế","✓","✓","✓","✓","Đang nghiên cứu",60,"https://www.schueco.com",""),

    ("SH-003","Sliding Đức Cao cấp","ADS 70.HI","Schüco","Đức","EN 12207/12208",
     "Cửa đi","Trượt lùa cách nhiệt (Thermal Sliding)","Khung cánh cách nhiệt ADS 70","Ngoại thất","Biệt thự, văn phòng cao cấp",
     2.0,"Kính hộp 24-40mm","18mm","6000×2800",400,"Schüco AvanTec Slide","Gioăng EPDM 4 lớp cách nhiệt / Dow Corning","Anodized, PVDF",
     "Uf=1.3 W/(m²K), cách âm Rw 44dB, bền vững cao","Giá nhập khẩu cao","✓","✓","✓","✓","Đang nghiên cứu",55,"https://www.schueco.com",""),

    ("SH-004","Lift Slide Đức Siêu cao cấp","ASS 77 PD.HI","Schüco","Đức","EN 13830/12208",
     "Cửa đi","Trượt nâng cách nhiệt (Thermal Lift & Slide)","Khung cầu cách nhiệt + Cánh nặng Lift Slide","Ngoại thất – Siêu biệt thự","Penthouse, biệt thự siêu sang, resort 6 sao",
     2.5,"Kính 3 lớp 48-60mm","22mm","10000×3500",800,"Schüco AvanTec đặc chủng","Gioăng EPDM cầu cách nhiệt 5 lớp","Anodized, PVDF 70%, gỗ ốp mặt trong",
     "Lift Slide rộng nhất thế giới (10m), cách nhiệt kép Uf<1.5 W/(m²K)","Giá 2-10 tỷ VND/bộ","✓","✓","✓","✓","Đang nghiên cứu",55,"https://www.schueco.com",""),

    ("SH-005","Curtain Wall Đức","FW 50+","Schüco","Đức","EN 13830 / CWCT",
     "Vách mặt dựng","Cố định (Stick/Unitized)","Module đố đứng đố ngang cấu trúc Schüco","Ngoại thất – Mặt dựng","Tòa nhà văn phòng hạng A, khách sạn quốc tế",
     2.5,"Kính hộp 24-50mm / Kính 3 lớp 60mm","50mm","Không giới hạn",1000,"Hệ Stick + Unitized AvanTec CW","Silicon kết cấu Dow 983 / Sika SG-500","Anodized, PVDF 70% AkzoNobel, mill finish",
     "Tiêu chuẩn tòa nhà thế giới, bền 60 năm, tối ưu tải gió động đất","Giá nhập khẩu cực cao","✓","✓","✓","✓","Đang nghiên cứu",60,"https://www.schueco.com",""),

    # ══════ REYNAERS – BELGIUM ══════
    ("RY-001","Casement Bỉ Cao cấp","CS 60","Reynaers","Bỉ","EN 14351-1/12207 CE",
     "Cửa sổ + Cửa đi","Mở quay + Tilt-Turn","Khung 60mm cầu cách nhiệt PA66","Ngoại thất","Biệt thự, khách sạn 5 sao, văn phòng cao cấp",
     1.8,"Kính hộp Low-E 24-36mm / Kính 3 lớp 44mm","16mm","1800×3000",180,"Rãnh C Euro Groove","Gioăng EPDM đa khoang 3 lớp / Dow Corning","Anodized tách màu, PVDF 70%",
     "Uf=1.5 W/(m²K), thiết kế Bỉ sang trọng, bền 50 năm","Giá nhập khẩu cao, giao hàng lâu","✓","✓","✓","✓","Đang nghiên cứu",65,"https://www.reynaers.com",""),

    ("RY-002","Casement Bỉ Passive","CS 77","Reynaers","Bỉ","EN 14351-1 / PHI Passive",
     "Cửa sổ + Cửa đi","Mở quay + Tilt-Turn","Khung 77mm cầu cách nhiệt siêu dày","Ngoại thất – Passive House","Biệt thự siêu sang, Zero Energy Building",
     2.0,"Kính 3 lớp 44-56mm","20mm","2000×3500",250,"Rãnh C Châu Âu siêu cao cấp","Gioăng EPDM 5 lớp","Anodized 2 màu, PVDF 2 lớp",
     "Uf=1.0 W/(m²K), cách âm 48dB, Passive House chuẩn Bỉ","Giá cực cao, thị trường VN rất hạn chế","✓","✓","✓","✓","Đang nghiên cứu",50,"https://www.reynaers.com",""),

    ("RY-003","Lift Slide Bỉ","CP 155-LS","Reynaers","Bỉ","EN 13830",
     "Cửa đi","Trượt nâng (Lift & Slide)","Khung cách nhiệt + Cánh nặng Lift Slide Reynaers","Ngoại thất – Siêu biệt thự","Biệt thự cao cấp Châu Âu, penthouse",
     2.5,"Kính 3 lớp 44-56mm","20mm","8000×3200",600,"Phụ kiện Reynaers LSS","Gioăng EPDM cầu cách nhiệt 4 lớp","Anodized tách màu, PVDF",
     "Uf<1.5 W/(m²K), Lift Slide cách nhiệt chuẩn Bỉ hàng đầu","Nhập khẩu lâu, giá cực cao","✓","✓","✓","✓","Đang nghiên cứu",50,"https://www.reynaers.com",""),

    ("RY-004","Curtain Wall Bỉ","CW 50","Reynaers","Bỉ","EN 13830",
     "Vách mặt dựng","Cố định (Stick/Unitized)","Module CW Reynaers Châu Âu","Ngoại thất – Tòa nhà cao tầng","Tòa nhà cao tầng Châu Âu",
     2.0,"Kính hộp 24-48mm","48mm","Không giới hạn",1000,"Hệ Stick/Unitized Reynaers","Silicon kết cấu Dow/Sika","Anodized, PVDF",
     "Chuẩn Châu Âu, bền 50 năm","Nhập khẩu, giá cao","✓","✓","✓","✓","Đang nghiên cứu",50,"https://www.reynaers.com",""),

    # ══════ XINGFA – CHINA ══════
    ("XF-001","Nhập khẩu tem đỏ","Xingfa Guangdong 55","Xingfa","Trung Quốc","GB/T 8478",
     "Cửa sổ + Cửa đi","Mở quay (Casement)","Khung 55mm tem đỏ Quảng Đông","Ngoại thất","Nhà phố, biệt thự phổ thông, chung cư",
     1.4,"5-10mm / Kính hộp đến 24mm","24mm","1200×2200",80,"Phụ kiện Kinlong / Draho rãnh 22","Gioăng EPDM kép / Silicon Apollo","Sơn tĩnh điện chất lượng cao",
     "Cực kỳ cứng chắc nhờ gân gia cường, màu sơn đẹp bền màu đạt chuẩn Qualicoat","Nhiều hàng giả hàng nhái, rãnh 22 không lắp được phụ kiện rãnh C",
     "✓","✗","✓","✓","Hoàn thành",90,"https://www.xingfa-alu.com","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-nhap-khau-quang-dong-1.html"),

    ("XF-002","Nhập khẩu tem đỏ","Xingfa Guangdong 93","Xingfa","Trung Quốc","GB/T 8478",
     "Cửa đi","Trượt lùa (Sliding)","Khung trượt 93mm 2 ray/3 ray","Ngoại thất","Căn hộ, văn phòng, biệt thự",
     2.0,"Kính đơn/hộp 6-24mm","24mm","4000×2400",150,"Bánh xe chịu lực kép & khóa chữ D Kinlong","Gioăng EPDM kép chèn nỉ","Sơn tĩnh điện bóng/mờ",
     "Thanh nhôm dày 2.0mm cực kỳ vững chãi, chịu bão gió giật ven biển rất tốt","Thanh ray dưới đúc nổi gây vấp chân khi đi lại nếu không âm sàn",
     "✓","✗","✓","✓","Hoàn thành",85,"https://www.xingfa-alu.com","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-nhap-khau-quang-dong-1.html"),

    ("XF-003","Nhập khẩu tem đỏ","Xingfa Guangdong 55 Slide","Xingfa","Trung Quốc","GB/T 8478",
     "Cửa sổ","Trượt lùa (Sliding)","Khung lùa bản mỏng 55mm","Ngoại thất","Chung cư, nhà phố diện tích hẹp",
     1.2,"Kính đơn 5-10mm","10mm","1800×1400",70,"Bánh xe đơn & khóa sập Kinlong","Gioăng chèn nỉ / Silicon Apollo","Sơn tĩnh điện mờ",
     "Giải pháp cửa sổ lùa cực kỳ gọn nhẹ tiết kiệm diện tích mở cánh, giá rẻ","Kết cấu mỏng nhẹ không lắp được cánh cửa đi lớn",
     "✓","✗","✓","✓","Hoàn thành",80,"https://www.xingfa-alu.com","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-nhap-khau-quang-dong-1.html"),

    ("XF-004","Nhập khẩu tem đỏ","Xingfa Guangdong 63","Xingfa","Trung Quốc","GB/T 8478",
     "Cửa đi","Xếp trượt gấp (Bifold)","Khung lùa xếp bản dày rãnh C","Ngoại thất","Biệt thự sân vườn, lối ra ban công penhouse lớn",
     1.5,"Kính đơn/hộp 6-28mm","28mm","6000×2800",200,"Bánh xe treo & bản lề liên kết Kinlong","Gioăng chèn nỉ đa lớp / Silicon Dow","Sơn tĩnh điện bóng/mờ",
     "Mở rộng tối đa 99% khẩu độ ô cửa lấy thoáng tối đa, sang trọng biệt thự","Kỹ thuật căn chỉnh bản lề bánh xe phức tạp, phụ kiện đắt",
     "✓","✗","✓","✓","Hoàn thành",85,"https://www.xingfa-alu.com","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-nhap-khau-quang-dong-1.html"),

    ("XF-005","Nhập khẩu tem đỏ","Xingfa Guangdong 65","Xingfa","Trung Quốc","GB/T 8478",
     "Vách mặt dựng","Cố định (Fixed Curtain Wall)","Xương mặt dựng Stick bản 65mm","Ngoại thất","Tòa nhà văn phòng cao tầng, showroom ô tô mặt tiền lớn",
     2.5,"Kính cường lực / Kính hộp phản quang","32mm","Không giới hạn",500,"Ke chịu lực kết cấu & bulong inox","Silicon kết cấu Dow Corning 791/895","Anodized, PVDF 3 lớp",
     "Xương chịu lực cực dày cản gió bão lớn cao tầng, mạ bề mặt bền bỉ","Thi công lắp dựng leo giàn giáo bên ngoài nguy hiểm phức tạp",
     "✓","✓","✓","✓","Hoàn thành",85,"https://www.xingfa-alu.com","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-nhap-khau-quang-dong-1.html"),

    ("XF-006","Cao cấp Class A","Xingfa Class A","Xingfa","Trung Quốc/VN","GB/T & EN",
     "Cửa đi + Cửa sổ","Mở quay","Khung phẳng rãnh C Châu Âu mạ ED","Ngoại thất","Biệt thự Luxury, resort ven biển cao cấp",
     2.0,"Kính hộp Low-E cách âm cách nhiệt","32mm","1600×2800",150,"Phụ kiện CMECH / Roto rãnh C Châu Âu","Gioăng EPDM đa khoang trung tâm","Mạ điện di Anodized ED bảo hành 25 năm",
     "Bản nhôm phẳng không gân hiện đại, xi mạ ED chống muối mặn ăn mòn muối biển","Giá thành nhôm thanh và phụ kiện đi kèm thuộc hàng đắt nhất",
     "✓","✓","✓","✓","Hoàn thành",80,"https://www.xingfa-alu.com","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-nhap-khau-quang-dong-1.html"),

    ("XF-007","Việt Nam sản xuất","Xingfa Việt Nam 55","Nhiều nhà máy","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Mở quay","Khung 55mm đùn ép trong nước","Nội/Ngoại thất","Nhà phố dân dụng phổ thông, văn phòng giá rẻ",
     1.2,"Kính đơn 5-10mm / Kính dán","10mm","1200×2200",80,"Phụ kiện PMA / Kinlong rãnh 22","Gioăng EPDM đơn / Silicon Apollo","Sơn tĩnh điện nội địa",
     "Chi phí cực kỳ kinh tế rẻ hơn hàng nhập 35%, sẵn hàng nhanh chóng","Kết cấu mỏng, chống gió bão giật yếu hơn hàng nhập khẩu Quảng Đông",
     "✓","✗","✓","✓","Hoàn thành",85,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-viet-nam-10.html"),

    ("XF-008","Việt Nam sản xuất","Xingfa Việt Nam 93","Nhiều nhà máy","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa (Sliding)","Khung trượt 93mm trong nước","Ngoại thất","Cửa đi thông phòng văn phòng, nhà xưởng",
     1.4,"Kính đơn 5-12mm","12mm","3000×2400",100,"Bánh xe đơn & khóa sập trong nước","Gioăng chèn nỉ mỏng / Apollo","Sơn tĩnh điện mờ",
     "Giải pháp cửa đi trượt lùa vô cùng kinh tế giúp tiết kiệm không gian","Ray trượt dễ bị mòn xước, trượt không đầm chắc tay",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-viet-nam-10.html"),

    ("XF-009","Việt Nam sản xuất","Xingfa Việt Nam 63","Nhiều nhà máy","Việt Nam","TCVN",
     "Cửa đi","Xếp trượt gấp (Bifold)","Khung xếp gấp đùn ép trong nước","Ngoại thất","Nhà dân trung cấp, cửa ra ban công rẻ",
     1.4,"Kính đơn/dán 6-12mm","12mm","4500×2400",120,"Phụ kiện xếp gấp Kinlong / Draho nội địa","Gioăng EPDM kép / Silicon Apollo","Sơn tĩnh điện",
     "Chi phí đầu tư rẻ, dễ mua sẵn hàng thi công nhanh chóng","Khung mỏng dễ sệ cánh nếu số lượng cánh trên 5 cánh",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-viet-nam-10.html"),

    ("XF-010","Việt Nam sản xuất","Xingfa Việt Nam 65","Nhiều nhà máy","Việt Nam","TCVN",
     "Vách mặt dựng","Cố định (Stick)","Profile xương mặt dựng Stick nội địa","Ngoại thất","Showroom, vách mặt tiền nhà phố trung tầng",
     2.0,"Bulong & ke chịu lực nội địa","24mm","Không giới hạn",350,"Chốt kết cấu thép mạ kẽm","Silicon Apollo kết cấu chống nước","Sơn tĩnh điện thường",
     "Rút ngắn thời gian đặt hàng nhập khẩu, giá rẻ, dễ lắp ráp Stick","Chỉ phù hợp tòa nhà thấp tầng, chống rung chấn trung bình",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-viet-nam-10.html"),

    # ══════ PMI (MALAYSIA) ══════
    ("PI-001","Tầm trung rãnh C","PMI PE45","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi + Cửa sổ","Mở quay","Khung rãnh C 45mm + Cánh","Ngoại thất","Căn hộ chung cư, nhà phố trung cấp",
     1.4,"Kính đơn/hộp 6-24mm","24mm","1400×2400",100,"Phụ kiện rãnh C Châu Âu (Sigico, Hopo)","Gioăng EPDM kép / Silicon Dow","Sơn tĩnh điện AkzoNobel",
     "Khung rãnh C chuẩn Châu Âu, dễ gia công lắp phụ kiện cao cấp, nước sơn mịn","Không có cầu cách nhiệt nên cản nhiệt trung bình",
     "✓","✓","✓","✓","Hoàn thành",85,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-002","Cầu cách nhiệt","PMI DPE56","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi + Cửa sổ","Mở quay","Khung cầu Polyamide 56mm","Ngoại thất","Biệt thự, resort, phòng ngủ hướng Tây nắng nóng",
     1.4,"Kính hộp Low-E cách nhiệt 24-32mm","32mm","1400×2600",120,"Phụ kiện CMECH / Roto rãnh C cách nhiệt","Gioăng EPDM đa khoang trung tâm","Sơn tĩnh điện cao cấp",
     "Khả năng cách âm cách nhiệt tuyệt vời, giảm truyền nhiệt phòng điều hòa 30%","Chi phí phôi nhôm cầu cách nhiệt nhập khẩu đắt đỏ",
     "✓","✓","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-003","Tầm trung phổ thông","PMI P55D","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi + Cửa sổ","Mở quay","Khung bao rãnh 22 kiểu Xingfa","Ngoại thất","Nhà phố dân dụng, biệt thự vườn phổ thông",
     1.4,"Kính đơn/hộp 6-20mm","20mm","1200×2200",85,"Phụ kiện PMA / Kinlong rãnh 22","Gioăng EPDM đơn / Apollo","Sơn tĩnh điện mờ",
     "Thẩm mỹ vát góc kiểu Xingfa dễ bán, bề mặt nước sơn mịn đẹp đặc trưng PMI","Phân khúc giá nhôm thanh hơi cao hơn Xingfa TQ",
     "✓","✗","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-004","Cao cấp rãnh C","PMI P55M","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi","Mở quay","Khung bao bản cánh 55mm dày dặn","Nội/Ngoại thất","Cửa chính biệt thự, cửa đi thông phòng",
     2.0,"Kính cường lực 10-15mm / Kính hộp","24mm","1600×2800",150,"Bản lề cối chịu lực Hopo / CMECH","Gioăng EPDM kép / Silicon Dow","Sơn tĩnh điện bảo hành 15 năm",
     "Cực kỳ cứng vững, chống rung chấn, an toàn chống đột nhập tốt","Trọng lượng cánh nặng, đòi hỏi bản lề chịu tải tốt",
     "✓","✓","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-005","Tầm trung rãnh C","PMI MP58A","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi + Cửa sổ","Mở quay","Khung phẳng rãnh C cải tiến","Nội/Ngoại thất","Chung cư cao cấp, văn phòng công ty",
     1.4,"Kính đơn/hộp 6-24mm","24mm","1400×2600",100,"Phụ kiện Sigico / Hopo rãnh C","Gioăng chèn nỉ / Silicon Sika","Sơn tĩnh điện mờ",
     "Đáp ứng tốt tiêu chuẩn thiết kế hiện đại phẳng phiu, tiết kiệm năng lượng","Giá phụ kiện rãnh C cao hơn rãnh thường",
     "✓","✓","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-006","Cầu cách nhiệt","PMI MP75A","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi","Mở quay","Khung cầu cách nhiệt bản 75mm dày","Ngoại thất","Biệt thự cao cấp ven biển, penhouse chịu bão",
     2.0,"Kính hộp Low-E cách âm cách nhiệt","32mm","1800×2800",150,"Phụ kiện an ninh chống cạy RC2 Sobinco","Gioăng đúc xốp kép trung tâm","Sơn tĩnh điện bảo hành 15 năm",
     "Cách âm và cách nhiệt vượt trội, chống bão gió giật ven biển rất tốt","Giá thành cực cao ở phân khúc siêu cao cấp",
     "✓","✓","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-007","Trượt lùa","PMI L100","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi","Trượt lùa (Sliding)","Khung trượt lùa 100mm 2 ray/3 ray","Ngoại thất","Nhà phố, căn hộ có ban công rộng",
     1.4,"Kính đơn/hộp 6-24mm","24mm","4000×2400",120,"Bánh xe chịu lực & tay gạt Hopo","Gioăng EPDM kép chèn nỉ","Sơn tĩnh điện mờ",
     "Khung bản 100mm cứng cáp trượt rất đầm chắc chắn, tối ưu không gian","Chiếm diện tích độ rộng tường lắp khuôn lớn",
     "✓","✗","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-008","Trượt lùa","PMI LE70","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa lùa","Trượt lùa bản mỏng (Sliding)","Khung lùa bản nhỏ 70mm gọn nhẹ","Nội/Ngoại thất","Vách ngăn phòng, cửa lùa phụ biệt thự",
     1.4,"Kính cường lực 8-12mm","12mm","3000×2200",80,"Bánh xe trượt nhẹ & chốt sập PMI","Gioăng chèn nỉ mỏng / Apollo","Sơn tĩnh điện mờ",
     "Thanh nhôm mảnh mai tinh tế, đóng mở nhẹ nhàng, chi phí bảo trì rất thấp","Không chịu tải bão lớn ngoài trời, cản nước mưa kém hơn bản 100",
     "✓","✗","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-009","Lift & Slide","PMI L120C","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi","Trượt nâng (Lift & Slide)","Khung trượt nâng bản 120mm","Ngoại thất","Resort cao cấp, biệt thự sân vườn lớn",
     2.0,"Kính hộp cách âm cách nhiệt 24-32mm","32mm","6000×2800",300,"Bánh xe nâng hạ & tay gạt Sobinco/Hopo","Gioăng nén EPDM phẳng sàn","Sơn tĩnh điện bảo hành 15 năm",
     "Trượt nâng hạ nén gioăng kín khít tuyệt đối cách âm tốt, lướt êm ái","Giá bộ phụ kiện nâng hạ nhập khẩu đắt đỏ",
     "✓","✓","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    ("PI-010","Lift & Slide","PMI DLM128","PMI","Malaysia","AAMA, Qualicoat",
     "Cửa đi","Trượt nâng siêu tải (Lift & Slide)","Khung trượt nâng 128mm cực dày","Ngoại thất","Biệt thự Luxury diện tích mặt lớn, penthouses",
     2.2,"Kính hộp Low-E dải rộng","36mm","7000×3000",350,"Phụ kiện trượt nâng siêu tải Hopo/CMECH","Gioăng đúc đa khoang chịu lực","Sơn tĩnh điện bảo hành 15 năm",
     "Kết cấu vững chãi chịu gió bão ven biển tốt nhất, mở rộng khẩu độ cánh lớn","Chi phí cực kỳ đắt đỏ, đòi hỏi nền sàn phẳng tuyệt đối khi lắp",
     "✓","✓","✓","✓","Hoàn thành",80,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html"),

    # ══════ YKK AP – JAPAN ══════
    ("YK-001","Casement Nhật","YKK WS55","YKK AP","Nhật Bản","JIS A 4702 / JIS A 4706",
     "Cửa sổ + Cửa đi","Mở quay + Mở hất (Casement + Awning)","Khung 55mm chuẩn Nhật + Cánh đặc chủng Nhật","Ngoại thất","Biệt thự, chung cư cao cấp Nhật và Việt Nam",
     1.6,"5-12mm / Kính hộp 20-28mm","12mm","1600×2800",150,"Phụ kiện YKK AP nguyên bộ Nhật","Gioăng EPDM Nhật tiêu chuẩn JIS","Anodized Nhật bạc tự nhiên – bền 40 năm chuẩn Nhật",
     "Kiểm soát chất lượng chuẩn Nhật JIS, độ bền vượt trội, thẩm mỹ tinh tế","Giá cao hơn hàng VN 2-3 lần","✓","✓","✓","✓","Đang nghiên cứu",65,"https://www.ykkap.co.jp",""),

    ("YK-002","Sliding Nhật","YKK SL90","YKK AP","Nhật Bản","JIS A 4706",
     "Cửa đi","Trượt lùa (Sliding)","Khung + Ray inox Nhật tiêu chuẩn JIS","Ngoại thất","Biệt thự, chung cư cao cấp",
     1.8,"Kính cường lực 10-15mm / Kính hộp 20-28mm","15mm","4000×2400",200,"Ray inox SUS304 Nhật","Gioăng EPDM Nhật","Anodized Nhật, sơn bột Nhật",
     "Trượt cực êm chuẩn Nhật JIS, bánh xe chịu tải cao","Giá cao, phân phối hạn chế VN","✓","✓","✓","✓","Đang nghiên cứu",60,"https://www.ykkap.co.jp",""),

    # ══════ HONDALEX (JAPAN / LONG VAN GROUP) ══════
    ("HD-001","Anodized Cao cấp","Hondalex LV 60","Hondalex","Nhật Bản/VN","JIS H4100",
     "Cửa đi","Mở quay","Khung bản dày 60mm + Cánh","Ngoại thất","Biệt thự cao cấp, khách sạn mặt tiền biển",
     1.8,"Kính cường lực 10-15mm / Kính hộp","24mm","1600×2800",150,"Phụ kiện Hopo / CMECH / Long Vân rãnh C","Gioăng EPDM Nhật Bản / Silicon Dow","Mạ điện di Anodize ED bền trên 40 năm",
     "Lớp bề mặt Anodize ED bóng mịn chống ăn mòn cực cao, bền trên 40 năm chuẩn JIS Nhật Bản","Giá thành nhôm thanh xi mạ ED và phụ kiện đi kèm khá cao, kén thợ",
     "✓","✓","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-hondalex-longvangroup-50.html"),

    ("HD-002","Anodized Trung cấp","Hondalex LV 34","Hondalex","Nhật Bản/VN","JIS H4100",
     "Cửa đi","Mở quay","Khung bản nhỏ gọn + Cánh","Nội thất","Cửa thông phòng biệt thự, cửa đi căn hộ chung cư",
     1.2,"Kính đơn 5-10mm","10mm","1200×2400",90,"Phụ kiện đồng bộ Long Vân rãnh 22","Gioăng cao su chèn sập / Silicon Apollo","Mạ điện di Anodize ED",
     "Thẩm mỹ gọn nhẹ tinh giản kiểu Nhật Bản, giá thành hợp lý, dễ sản xuất","Bản nhôm mỏng không lắp được kính hộp quá dày cách âm",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-hondalex-longvangroup-50.html"),

    ("HD-003","Anodized Trung cấp","Hondalex LV 70","Hondalex","Nhật Bản/VN","JIS H4100",
     "Cửa đi","Mở quay","Khung bản 70mm tối ưu nhẹ","Ngoại thất","Căn hộ, văn phòng làm việc, nhà phố",
     1.2,"Kính đơn 5-12mm / Kính hộp","18mm","1400×2400",100,"Phụ kiện chốt gạt Long Vân / Kinlong","Gioăng EPDM / Silicon Apollo","Mạ điện di Anodize ED",
     "Trọng lượng tối ưu giúp vận hành đóng mở nhẹ tay êm ái, xi mạ chống xước rất tốt","Độ chắn gió giật ngoại thất kém hơn bản LV 60 dày",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-hondalex-longvangroup-50.html"),

    ("HD-004","Anodized Cao cấp","Hondalex LV 60 Window","Hondalex","Nhật Bản/VN","JIS H4100",
     "Cửa sổ","Mở hất + Mở quay","Khung sổ 60mm chống dột","Ngoại thất","Cửa sổ nhà cao tầng, biệt thự ven biển",
     1.5,"Kính đơn/hộp 6-24mm","24mm","1400×1600",120,"Bản lề chữ A & tay gạt đa điểm rãnh C","Gioăng EPDM đa lớp / Silicon Dow","Mạ điện di Anodize ED",
     "Cửa sổ chống dột nước mưa cực tốt nhờ hệ gioăng EPDM đúc Nhật, chịu bão tốt","Trọng lượng nặng hơn các dòng cửa sổ nhôm thường",
     "✓","✓","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-hondalex-longvangroup-50.html"),

    ("HD-005","Anodized Phổ thông","Hondalex LV 38 Window","Hondalex","Nhật Bản/VN","JIS H4100",
     "Cửa sổ","Mở hất gọn nhẹ","Khung sổ 38mm bản mỏng","Nội/Ngoại thất","Ô thoáng, cửa sổ phụ nhà dân bình dân",
     0.9,"Kính đơn 5-10mm","10mm","800×1200",50,"Chốt bật & phụ kiện nhẹ Long Vân","Gioăng cao su chèn sập / Apollo","Mạ điện di Anodize ED",
     "Chi phí siêu tiết kiệm cho công trình dân dụng, gia công dập ke góc cực nhanh","Độ bền chịu tải yếu, không lắp được ở các vị trí cao chịu bão lớn",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-hondalex-longvangroup-50.html"),

    ("HD-006","Anodized Trung cấp","Hondalex LV 76","Hondalex","Nhật Bản/VN","JIS H4100",
     "Cửa đi + Cửa sổ","Trượt lùa (Sliding)","Khung lùa 76mm 2 ray","Ngoại thất","Căn hộ chung cư, cửa ra ban công biệt thự",
     1.2,"Kính đơn/hộp 6-20mm","20mm","3000×2400",100,"Bánh xe trượt Long Vân & khóa bán nguyệt","Gioăng chèn nỉ mịn chống rung nước","Mạ điện di Anodize ED",
     "Trượt lướt êm nhẹ phong cách Nhật Bản, tiết kiệm diện tích tối đa, chắn nước tốt","Cần vệ sinh ray trượt định kỳ để tránh cát kẹt",
     "✓","✗","✓","✓","Hoành thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-hondalex-longvangroup-50.html"),

    # ══════ EUROWINDOW (uPVC & ALUMIUM) ══════
    ("EW-001","Nhựa lõi thép uPVC","EW uPVC Asia60","Eurowindow","Đức/Việt Nam","TCVN & EN",
     "Cửa sổ + Cửa đi","Mở quay + Mở hất","Profile nhựa uPVC Koemmerling + Lõi thép","Ngoại thất","Biệt thự, căn hộ chung cư cao cấp, phòng thu âm",
     2.2,"Kính đơn/hộp 6-28mm","28mm","1400×2400",100,"Phụ kiện đồng bộ Roto / G-U / Eurowindow","Gioăng TPE đa lớp đúc sẵn","Bề mặt nhựa trắng bóng / xám xi măng",
     "Khả năng cách âm đứng đầu các loại hệ cửa, không bị mài mòn rỉ sét, chịu gió tốt","Khung nhựa to dày thô hơn nhôm, nhựa trắng có thể ố vàng nhẹ sau 20 năm",
     "✓","✓","✓","✓","Hoàn thành",90,"https://eurowindow.biz","https://nhomkinhdaiphuc.com/catalogue/catalogue-cua-nhua-loi-thep-eurowindow-36.html"),

    ("EW-002","Nhôm Cao cấp","EW EA55","Eurowindow","Việt Nam","TCVN & EN",
     "Cửa sổ + Cửa đi","Mở quay","Khung nhôm EA55 rãnh C Châu Âu","Ngoại thất","Biệt thự hiện đại, khu nghỉ dưỡng, căn hộ sang trọng",
     1.4,"Kính đơn/hộp 6-32mm","32mm","1600×2800",130,"Phụ kiện Roto / CMECH / Eurowindow rãnh C","Gioăng EPDM đúc 3 lớp / Silicon Dow","Sơn tĩnh điện bảo hành lên tới 10-20 năm",
     "Thiết kế bo tròn thẩm mỹ sang trọng, rãnh C chuẩn Châu Âu lắp phụ kiện cao cấp","Giá thành thi công lắp đặt trọn gói cao hơn Xingfa nhập khẩu",
     "✓","✓","✓","✓","Hoàn thành",85,"https://eurowindow.biz","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-eurowindow-28.html"),

    ("EW-003","Nhôm Cao cấp","EW EA70","Eurowindow","Việt Nam","TCVN",
     "Cửa đi + Cửa sổ","Trượt lùa (Sliding)","Khung lùa bản mỏng EA70","Ngoại thất","Căn hộ chung cư, cửa ra ban công biệt thự",
     1.6,"Kính đơn/hộp 6-24mm","24mm","3000×2400",100,"Bánh xe chịu lực & khóa đa điểm Eurowindow","Gioăng chèn nỉ mịn kép chống rít nước","Sơn tĩnh điện chống UV",
     "Hoạt động lùa trượt vô cùng êm nhẹ, tiết kiệm không gian diện tích tối đa","Độ kín khít chắn bão giật ngoại thất kém hơn hệ quay EA55",
     "✓","✓","✓","✓","Hoàn thành",80,"https://eurowindow.biz","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-eurowindow-28.html"),

    ("EW-004","Nhôm Cao cấp","EW EA90","Eurowindow","Việt Nam","TCVN & ASTM",
     "Vách mặt dựng","Cố định (Stick Curtain Wall)","Profile vách dựng Stick bản EA90","Ngoại thất","Tòa nhà văn phòng, showroom thương mại mặt tiền rộng",
     2.2,"Bulong kết cấu nở & phụ kiện chịu lực Eurowindow","32mm","Không giới hạn",400,"Ke chịu lực và bulong kết cấu chuyên dụng","Silicon kết cấu chuyên dụng Dow Corning","Sơn tĩnh điện ngoài trời cao cấp",
     "Xương chịu tải lực gió bão siêu việt, thẩm mỹ kính liền dải sang trọng","Chi phí thiết kế, gia công lắp dựng lớn đòi hỏi kỹ thuật cao",
     "✓","✓","✓","✓","Hoàn thành",85,"https://eurowindow.biz","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-eurowindow-28.html"),

    ("EW-005","Cầu cách nhiệt Cao cấp","EW Thermal","Eurowindow","Đức/Việt Nam","EN & TCVN",
     "Cửa đi + Cửa sổ","Mở quay","Khung có dải cầu cách nhiệt Polyamide","Ngoại thất","Biệt thự nghỉ dưỡng cao cấp, penthouse, vùng lạnh/nóng ẩm",
     2.0,"Kính hộp cách âm cách nhiệt Low-E","36mm","1600×2800",150,"Phụ kiện cao cấp Roto / CMECH rãnh C","Gioăng trung tâm đa khoang EPDM","Sơn tĩnh điện / Anodized",
     "Khả năng cách nhiệt và cách âm siêu việt nhất, giảm 35% điện năng điều hòa","Giá thành thuộc hàng đắt đỏ nhất trong các dòng sản phẩm",
     "✓","✓","✓","✓","Hoàn thành",80,"https://eurowindow.biz","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-eurowindow-28.html"),

    # ══════ VIỆT Ý (ITALUMI) ══════
    ("VI-001","Tầm trung phổ thông","VI Italumi 55","Việt Ý","Việt Nam","TCVN",
     "Cửa sổ + Cửa đi","Mở quay","Khung 55mm vát mép nhẹ","Ngoại thất","Căn hộ, văn phòng, nhà phố dân dụng",
     1.2,"Kính đơn 5-10mm / Kính dán","10mm","1200×2200",80,"Phụ kiện Kinlong / PMA / Huy Hoàng rãnh 22","Gioăng EPDM kép / Silicon Apollo","Sơn tĩnh điện bóng mịn ngoài trời",
     "Màu sắc trang nhã, thiết kế vát cạnh cản nước dột tốt, phôi nhôm dẻo dai dễ gia công","Kết cấu nhôm khá mỏng, chống bão giật ngoại thất ở mức trung bình",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html"),

    ("VI-002","Tầm trung phổ thông","VI Italumi 93","Việt Ý","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa (Sliding)","Khung trượt 93mm 2 ray","Ngoại thất","Căn hộ chung cư, cửa ra ban công hông nhà",
     1.4,"Kính đơn 5-12mm","12mm","3000×2400",100,"Bánh xe đơn & khóa bán nguyệt đồng bộ","Gioăng chèn nỉ mịn kép chắn bụi","Sơn tĩnh điện mờ",
     "Trượt ngang tiết kiệm diện tích đóng mở cánh tối đa, giá thành phôi nhôm rẻ","Ray lùa mỏng dễ bị cọ xước nếu bị kẹt cát bẩn",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html"),

    ("VI-003","Tầm trung phổ thông","VI Italumi 450","Việt Ý","Việt Nam","TCVN",
     "Cửa đi thông phòng","Mở quay","Khung 450mm bản mỏng kinh tế","Nội thất","Toilet nhà dân bình dân, ô thoáng, cửa phòng kho",
     1.0,"Kính đơn 5-8mm / Panel đặc","8mm","1000×2000",60,"Khóa tay gạt tròn & bản lề trong nước","Gioăng cao su chèn sập thường","Sơn tĩnh điện thường",
     "Chi phí siêu tiết kiệm, thi công gia công nhanh chóng không cần máy móc phức tạp","Độ cứng kết cấu yếu, kỵ lắp ngoài trời nắng mưa trực diện",
     "✓","✗","✓","✓","Hoàn thành",80,"https://vietphapshal.com.vn","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html"),

    # ══════ HỆ SLIM ══════
    ("SL-001","Slim Nội thất Phổ thông","Slim 45 Nội thất","Nhiều hãng VN","Việt Nam","TCVN",
     "Cửa đi + Thông phòng","Trượt lùa giảm chấn (Soft Close)","Khung bản mỏng 45mm + Ray ẩn trên trần","Nội thất","Phòng ngủ, WC, phòng thay đồ, tủ quần áo",
     1.2,"Kính cường lực 8-12mm","12mm","2000×2800",80,"Ray treo Hettich/Häfele/Slido – 80kg","Gioăng mềm đệm nhẹ / không cần keo","Anodized bạc, vàng champagne, đen xước, trắng mờ",
     "Không ray dưới sàn, giảm chấn tự hãm, lắp đặt nhanh, thẩm mỹ cực cao","Chỉ nội thất, không chống thấm nước mưa","✓","✗","✓","✓","Hoàn thành",85,"",""),

    ("SL-002","Slim Nội thất Cao cấp","Slim 70 Liên động","Nhiều hãng VN","Việt Nam","TCVN",
     "Cửa đi + Vách ngăn phòng","Trượt liên động đa cánh (Multi-panel Sliding)","Khung mỏng 70mm + Ray đôi + Cánh kính lớn","Nội thất","Phòng khách lớn, văn phòng open space, phòng hội thảo",
     1.4,"Kính cường lực 10-15mm","15mm","6000×3000",120,"Ray đôi Slido Dorma / Häfele 50kg","Gioăng mềm kín cạnh / không keo","Anodized, đen mờ, champagne vàng, đồng antique",
     "Mở toàn bộ vách thông suốt, liên động kéo nhẹ 3-4 cánh cùng lúc","Chỉ nội thất, giá cao","✓","✓","✓","✓","Đang nghiên cứu",70,"",""),

    ("SL-003","Slim Ngoại thất Cao cấp","Slim 60 Ngoại thất","Nhiều hãng VN","Việt Nam","TCVN",
     "Cửa đi","Trượt lùa ngoại thất (Outdoor Slim Sliding)","Khung mỏng kép 60mm + Ray kép chịu lực","Ngoại thất","Biệt thự, ban công, lối ra sân vườn nhỏ",
     1.6,"Kính cường lực 10-15mm / Kính hộp 20-28mm","15mm","3000×2800",150,"Ray thép cường lực + Bánh xe inox","Gioăng EPDM kín nước / Dow Corning","Anodized ED, sơn chống UV",
     "Profile mỏng – nhìn nhiều kính hơn nhôm","Kín nước kém hơn cửa đi thông thường","✓","✗","✓","✓","Đang nghiên cứu",65,"",""),

    ("SL-004","Slim Steel Look","Steel Look Slim","Vitrocsa/Local","Thụy Sỹ/Việt Nam","EN 12207",
     "Cửa đi + Cửa sổ + Vách kính","Mở quay + Cố định","Khung profile giả thép nghệ thuật (Steel Look)","Ngoại thất cao cấp","Biệt thự Loft/Industrial, showroom, nhà hàng",
     3.0,"Kính cường lực 10-15mm / Kính hộp 20-28mm","15mm","3000×4000",200,"Phụ kiện ẩn đặc chủng","Gioăng kín EPDM / Silicon đen","Sơn đen nhám, đen bóng, sơn giả thép cát mờ",
     "Thẩm mỹ cực cao phong cách công nghiệp, profile đậm nét như khung thép","Giá cao, nặng, ít nhà cung cấp VN","✓","✗","✓","✗","Đang nghiên cứu",50,"",""),

    ("SL-005","Slim Pocket Âm tường","Slim Pocket Sliding","Các hãng","Châu Âu/VN","EN",
     "Cửa đi + Thông phòng","Trượt âm tường (Pocket Sliding)","Khung nhỏ + Cánh lùa vào trong tường","Nội thất","WC, tủ quần áo âm, phòng ngủ, bếp",
     1.2,"Kính cường lực 8-12mm / Panel đặc","12mm","1200×2800",60,"Ray treo Barn Door / Slido âm tường Häfele","Gioăng mềm cạnh / không keo","Anodized bạc, trắng mờ, đen mờ",
     "Tối ưu không gian, không cần diện tích xoay cánh","Cần tường đủ dày, cách âm kém","✗","✗","✓","✗","Đang nghiên cứu",45,"",""),

    # ══════ HỆ CHUYÊN DỤNG ══════
    ("SP-001","Cửa Pivot Xoay trục","Pivot Door","Civro/Owin/Local","Việt Nam","TCVN",
     "Cửa đi","Xoay trục (Pivot)","Cánh xoay trục giữa + Pivot floor/ceiling spring","Ngoại thất","Biệt thự siêu sang, sảnh resort, showroom",
     2.0,"Kính cường lực 15-22mm","22mm","2000×4000",500,"Dorma / FritsJurgens / Biloba Pivot","Gioăng cạnh kín / Silicon kết cấu","Anodized ED, sơn RAL đặc chủng",
     "Thẩm mỹ đỉnh cao, mở 180° tinh tế, cánh siêu rộng một tấm kính nguyên","Giá cao, thi công phức tạp","✓","✗","✓","✗","Đang nghiên cứu",55,"",""),

    ("SP-002","Cửa Bi-fold Xếp trượt","Bi-fold/Folding","Civro/Maxpro/Soco","Việt Nam","EN",
     "Cửa đi","Xếp trượt (Bi-fold / Folding)","Cánh xếp gấp + Bản lề xếp + Ray trên","Ngoại thất","Nhà hàng, quán cà phê, biệt thự mở",
     1.8,"Kính cường lực 8-12mm","12mm","5000×2800",40,"Roto Fold / OPK Bi-fold / Hopo","Gioăng EPDM đường viền / Silicon","Sơn tĩnh điện, anodized, champagne",
     "Mở toàn bộ không gian, xếp gọn về một phía, phù hợp nhà hàng mở","Kín khít kém nhất, mỗi cánh tải trọng giới hạn","✓","✗","✓","✓","Đang nghiên cứu",60,"",""),

    ("SP-003","Cửa tự động Trượt","Auto Sliding Door","Dorma/Besam/Local","Đức/Việt Nam","EN 16005",
     "Cửa đi","Tự động trượt (Auto Sliding)","Khung + Cánh kính + Cảm biến mắt thần","Thương mại","Siêu thị, TTTM, bệnh viện, khách sạn sảnh lớn",
     1.6,"Kính an toàn laminated 8-12mm","12mm","3000×2800",150,"Hệ truyền động AC/DC + Ray dẫn hướng","Gioăng mềm / Màn quét không chạm","Anodized bạc, đen, sơn tĩnh điện",
     "Hoàn toàn tự động, tiện lợi người tay bận hoặc xe lăn, chuẩn EN 16005","Cần bảo trì cảm biến và động cơ định kỳ","✓","✗","✓","✓","Đang nghiên cứu",55,"",""),

    ("SP-004","Cửa chống cháy","Fire Door EI60/EI90","Hệ PCCC","VN/Nhật","TCVN 9374/EN 1634",
     "Cửa đi","Mở quay (Fire Rated Casement)","Khung thép + Cánh nhôm/thép chống cháy","An toàn PCCC","Hành lang thoát hiểm, buồng thang bộ, phòng kỹ thuật",
     2.0,"Kính chống lửa 30-90 phút (Fire Glass)","25mm","1200×2400",100,"Bản lề chống cháy + Door Closer chống cháy","Gioăng bảo ôn tự trương nở khi cháy","Sơn Intumescent chịu nhiệt",
     "Ngăn lan cháy EI60/EI90, bắt buộc theo QCVN 06, bền 30 năm","Chỉ lắp nơi thoát hiểm, không trang trí","✓","✓","✓","✓","Đang nghiên cứu",55,"",""),

    ("SP-005","Mái kính Skylight","Skylight Giếng trời","Soco/Local","Việt Nam","TCVN",
     "Mái kính","Cố định / Mở nghiêng (Vented Skylight)","Khung nhôm chống thấm + Kính nghiêng","Mái công trình","Giếng trời biệt thự, mái hành lang, sảnh TTTM",
     2.0,"Kính cường lực laminated 10-15mm / Kính hộp Low-E","15mm","Không giới hạn",200,"Hệ thanh đỡ mái nhôm đặc chủng","Silicon kết cấu + Chống thấm dốc mái","Anodized, sơn bột chịu UV",
     "Đưa ánh sáng tự nhiên vào sâu trong nhà, tiết kiệm điện chiếu sáng","Cần thiết kế thoát nước kỹ, vệ sinh định kỳ","✓","✗","✓","✓","Đang nghiên cứu",55,"",""),

    ("SP-006","Lan can kính","Glass Railing Frameless","Maxpro/Owin/Local","Việt Nam","TCVN 9392",
     "Lan can","Cố định (Frameless Glass Railing)","Đế chân âm sàn inox + Kính cường lực đứng","Nội/Ngoại thất","Ban công biệt thự, cầu thang nội thất, hành lang cao tầng",
     3.0,"Kính cường lực đứng 12-19mm / Kính laminated 10+10mm","19mm","Không giới hạn",80,"Đế âm sàn inox SUS304/316 + Thanh nhôm đỡ trên","Keo silicon kết cấu đế kính","Anodized bạc, inox bóng, đen mờ",
     "Không gian mở, ánh sáng xuyên suốt, thẩm mỹ hiện đại","Cần bảo trì mối kết cấu kính, dễ bám bẩn","✓","✗","✓","✓","Đang nghiên cứu",55,"",""),

    ("SP-007","Pergola Nhôm Cơ điện","Pergola Motorized Louvres","Local/Ý","Việt Nam/Ý","TCVN",
     "Mái che + Pergola","Lá chớp xoay tự động (Motorized Aluminum Louvers)","Khung nhôm + Lá chớp xoay góc","Ngoại thất","Biệt thự sân vườn, nhà hàng ngoài trời, resort",
     2.5,"Không dùng kính (lá chớp nhôm đặc)","N/A","Không giới hạn",50,"Mô tơ AC bánh răng + Remote / Cảm biến mưa","Gioăng EPDM kín nước","Anodized ED, PVDF chịu UV, màu RAL",
     "Điều tiết ánh sáng và mưa tự động bằng remote hoặc cảm biến mưa","Giá cao, cần bảo trì mô tơ định kỳ","✓","✗","✓","✓","Đang nghiên cứu",50,"",""),

    ("SP-008","Vách ngăn văn phòng","Office Partition Glass","Nhiều hãng","Việt Nam","TCVN",
     "Vách ngăn","Cố định + Cửa trượt (Fixed + Sliding)","Profile vách ngăn kính mỏng","Nội thất","Văn phòng, phòng họp, coworking space",
     1.2,"Kính cường lực 10-12mm / Kính đục mờ","12mm","Không giới hạn",60,"Häfele / Hettich trượt văn phòng","Gioăng cạnh mềm / Silicon nội thất","Anodized bạc, đen mờ, champagne, trắng mờ",
     "Phân chia không gian linh hoạt, ánh sáng xuyên qua, thẩm mỹ hiện đại","Cách âm kém hơn tường xây","✓","✗","✓","✓","Đang nghiên cứu",50,"",""),

    ("SP-009","Cửa Tilt-Turn Mở lật","Tilt-Turn Window","Civro/PAG/Schüco","Châu Âu/VN","EN 14351-1",
     "Cửa sổ + Cửa đi","Mở quay + Mở lật (Casement + Tilt-Turn)","Khung + Cánh mở lật đa vị trí","Ngoại thất","Căn hộ cao tầng, văn phòng an toàn",
     1.8,"Kính hộp Low-E 24-36mm","16mm","1400×2600",120,"Rãnh C Euro Groove – Roto/Winkhaus Tilt-Turn","Gioăng EPDM đa khoang / Dow Corning","Anodized, PVDF",
     "Tilt (hé nghiêng thông gió an toàn) hoặc Turn (quay vào trong lau kính dễ dàng)","Phụ kiện đắt hơn Casement thông thường","✓","✓","✓","✓","Đang nghiên cứu",65,"",""),

    ("SP-010","Curtain Wall Unitized","UCW Unitized Panel","Soco/Maxpro/Civro","Việt Nam","EN 13830/AAMA 501",
     "Vách mặt dựng","Panel đúc sẵn (Unitized Panel)","Panel hoàn chỉnh lắp bằng cần trục","Ngoại thất – Tòa nhà cao tầng","Văn phòng hạng A, khách sạn 5 sao, cao tầng 30+",
     2.5,"Kính hộp Low-E 28-48mm / Kính 3 lớp 60mm","48mm","Không giới hạn",1000,"Hệ Unitized panel + crane","Silicon kết cấu AAMA","Anodized, PVDF 70%",
     "Thi công 3-4 ngày/tầng, chất lượng đồng đều chuẩn xưởng","Chi phí mold cao, không linh hoạt sửa đổi","✓","✓","✓","✓","Đang nghiên cứu",65,"",""),

    ("SP-011","Curtain Wall Spider","Spider Glass Facade","Hệ chuyên dụng","Đức/Việt Nam","EN 13830",
     "Vách mặt dựng đặc biệt","Cố định mặt kính điểm (Spider Connector)","Khung kết cấu thép ẩn + Cần nhện inox SUS316","Ngoại thất – Công trình biểu tượng","Sảnh khách sạn, hội nghị, biểu tượng",
     0.0,"Kính cường lực đặc biệt 15-22mm có lỗ cán khoan","22mm","Không giới hạn",2000,"Cần nhện inox SUS316 Spider 4 cánh","Silicon kết cấu Dow 895","Kính Low-E / Kính in màu kỹ thuật số",
     "Tối ưu thẩm mỹ mặt kính thuần túy, trong suốt hoàn toàn, không thấy frame nhôm","Kết cấu thép phụ phức tạp, giá cực cao","✓","✗","✓","✗","Đang nghiên cứu",40,"",""),

    ("SP-012","Canopy Nhôm","Aluminum Canopy","Local","Việt Nam","TCVN",
     "Mái che","Cố định (Fixed Canopy)","Khung nhôm profile + PC trong/Kính nghiêng","Ngoại thất","Sảnh đón, lối vào biệt thự, nhà hàng, cửa hàng",
     2.0,"PC trong suốt 10mm / Kính cường lực 8-10mm","10mm","Không giới hạn",100,"Hệ Stick nhôm đặc chủng","Dow Corning chống thấm + Vít inox SUS304","Anodized, sơn tĩnh điện, champagne",
     "Che nắng mưa sảnh đón, thẩm mỹ ngoại thất đẹp, lắp đặt nhanh","Cần thiết kế thoát nước mái","✓","✗","✓","✓","Đang nghiên cứu",55,"",""),

    ("SP-013","Smart Door Cửa thông minh","Smart Fingerprint Door","Hệ Smart","Việt Nam","TCVN",
     "Cửa đi","Mở quay + Khóa thông minh (Smart Casement)","Khung + Cánh + Module điện tử Smart Lock","Ngoại thất","Nhà thông minh Smart Home, biệt thự, căn hộ",
     1.8,"Kính cường lực 10-12mm","12mm","1200×2800",120,"Khóa vân tay + Face ID + Động cơ DC remote","Gioăng EPDM + module chống nước IP54","Anodized bạc, champagne, đen mờ",
     "Mở khóa bằng vân tay/khuôn mặt/điện thoại, tích hợp Smart Home","Cần nguồn điện dự phòng, chi phí lắp đặt cao","✗","✗","✓","✗","Đang nghiên cứu",40,"",""),

    ("SP-014","Cửa chống đạn","Bullet Resistant BR3/BR6","Hệ chuyên dụng","Đức/Mỹ","EN 1063/UL 752",
     "Cửa đi","Mở quay (Bullet Resistant)","Khung nhôm đặc biệt + Tấm thép gia cố","An ninh","Ngân hàng, đại sứ quán, phòng VIP, phòng máy chủ",
     3.0,"Kính chống đạn Lexan/Pilkington Protect 25-50mm","50mm","1200×2400",300,"Khóa an ninh đặc chủng chống phá","Gioăng cao su chịu lực","Sơn thép + anodized nhôm",
     "Chịu đạn BR3/BR6, hội tụ an ninh và thẩm mỹ","Giá rất cao, ít công ty thi công VN","✗","✗","✓","✗","Đang nghiên cứu",35,"",""),

    ("SP-015","Sunroom Phòng kính","Aluminum Sunroom","Local/Châu Âu","Châu Âu/Việt Nam","EN",
     "Vách + Mái","Vách cố định + Mái nghiêng (Sunroom)","Khung nhôm + Vách kính + Mái kính","Ngoại thất","Biệt thự sân vườn, resort, nhà hàng garden",
     2.0,"Kính hộp Low-E 24-36mm / Kính chống nóng mái","15mm","Không giới hạn",200,"Hệ Stick nhôm đặc chủng Sunroom","Silicon kết cấu + Keo chống thấm mái","Anodized, PVDF, trắng/champagne/đen RAL",
     "Không gian trong/ngoài kết nối quanh năm, lọc UV, chắn mưa","Nhiệt độ cao nếu không thiết kế thông gió tốt","✓","✗","✓","✓","Đang nghiên cứu",45,"",""),

    # ══════ KOSO – ABS DOORS (K KUMOVINA HÀN QUỐC) ══════
    ("KO-001","Nhựa ABS Hàn Quốc","KOS Flat","KOS","Hàn Quốc/VN","TCVN/Korean Standard",
     "Cửa đi thông phòng","Cửa phẳng (Flat Interior Door)","Tấm nhựa ABS 1.4mm ép chân không + Lõi Honeycomb","Nội thất","Căn hộ chung cư, phòng ngủ biệt thự, nhà phố",
     0.0,"Kính đơn hoặc Panel đặc (không dùng kính)","N/A","1000×2200",50,"Phụ kiện khóa tay gạt tròn / tay gạt KOS & Bản lề inox","Gioăng cao su chèn giảm chấn / Silicon Apollo","Phủ Deco Sheet vân gỗ bóng mịn, chống trầy xước",
     "Kháng nước 100%, chống ẩm mốc mối mọt cong vênh tuyệt đối, cánh nhẹ tránh xệ cánh bản lề","Chỉ lắp nội thất trong nhà, không chịu được nắng mưa ngoài trời hay lực chống trộm mạnh",
     "✗","✗","✓","✓","Hoàn thành",85,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html"),

    ("KO-002","Nhựa ABS Hàn Quốc","KOS Panel","KOS","Hàn Quốc/VN","TCVN/Korean Standard",
     "Cửa đi thông phòng","Cửa dập pano chìm nổi","Tấm nhựa ABS dập pano nổi + Lõi Honeycomb","Nội thất","Phòng ngủ phong cách tân cổ điển, truyền thống",
     0.0,"Panel đặc nhựa ABS dập pano","N/A","1000×2200",50,"Khóa tay gạt dài KOS & bản lề lá inox SUS304","Gioăng giảm chấn chèn sẵn trên khung","Deco Sheet vân gỗ dập pano nghệ thuật",
     "Kiểu dáng pano cổ điển ấm cúng y hệt gỗ tự nhiên, không sợ mối mọt ẩm ngập nước","Không uốn được vòm hay các ô kính dạng tròn phức tạp",
     "✗","✗","✓","✓","Hoàn thành",85,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html"),

    ("KO-003","Nhựa ABS Hàn Quốc","KOS Glass","KOS","Hàn Quốc/VN","TCVN/Korean Standard",
     "Cửa đi toilet/phòng tắm","Cửa phẳng kết hợp ô kính mờ","Tấm nhựa ABS kết hợp ô nẹp kính lấy sáng","Nội thất ẩm ướt","Toilet chung cư, phòng tắm biệt thự, phòng ngủ có ô quan sát",
     0.0,"Kính cường lực mờ 5-8mm","8mm","1000×2200",50,"Kính mờ và khóa tròn toilet KOS","Gioăng nẹp kính mút mềm / Silicon nội thất","Deco Sheet kết hợp ô kính lấy sáng",
     "Lấy sáng tự nhiên rất tốt, chống dột tràn nước toilet hoàn hảo, chống mốc mọt","Cách âm kém hơn do ghép ô kính",
     "✗","✗","✓","✓","Hoàn thành",85,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html"),

    ("KO-004","Nhựa ABS Hàn Quốc","KOS Line","KOS","Hàn Quốc/VN","TCVN/Korean Standard",
     "Cửa đi thông phòng","Cửa phẳng chạy chỉ trang trí nhôm","Tấm nhựa ABS phẳng chạy phay rãnh nẹp chỉ nhôm","Nội thất hiện đại","Thông phòng chung cư, phòng làm việc hiện đại",
     0.0,"Panel đặc chạy chỉ trang trí","N/A","1000×2200",50,"Khóa tay gạt dài / Khóa điện tử vân tay + Chỉ nhôm","Gioăng giảm chấn / Silicon nội thất","Deco Sheet vân gỗ chạy chỉ nhôm bạc/vàng nổi bật",
     "Điểm nhấn chỉ nhôm bạc hoặc vàng vô cùng sang trọng cá tính hiện đại, tương thích tốt khóa điện tử","Giá thành thanh nẹp gia công cao hơn cửa phẳng thường",
     "✗","✗","✓","✓","Hoàn thành",85,"https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html"),
]

# ─────────────────────────────────────────────
# TIÊU ĐỀ CỘT (30 CỘT)
# ─────────────────────────────────────────────
HEADERS = [
    "ID", "Nhóm phân khúc", "Mã Hệ (System Code)", "Hãng Sản Xuất",
    "Quốc Gia", "Tiêu Chuẩn", "Loại Cửa", "Kiểu Mở",
    "Loại Profile", "Công Năng", "Ứng Dụng Công Trình",
    "Độ Dày Nhôm (mm)", "Kính Tương Thích", "Kính Tối Đa",
    "Kích Thước Tối Đa (W×H)", "Tải Trọng Cánh (kg)",
    "Chuẩn Phụ Kiện", "Gioăng & Keo Lắp Đặt", "Bề Mặt & Đặc Tính",
    "Ưu Điểm", "Nhược Điểm",
    "Có CAD", "Có Shopdrawing", "Có Catalogue", "Có HD Lắp Đặt",
    "Tình Trạng", "Hoàn Thiện (%)", "Link Website", "Link Catalogue PDF", "Link Bản Vẽ CAD",
]

# Nhóm → màu nền
GROUP_COLORS = {
    "Nhập khẩu tem đỏ":       "FFCDD2",
    "Việt Nam sản xuất":      "F5F5F5",
    "Nhựa ABS Hàn Quốc":      "FFF9C4",
    "Nhựa lõi thép uPVC":     "E8F5E9",
    "Nhôm Cao cấp":           "E0F7FA",
    "Tầm trung rãnh C":       "FFE0B2",
    "Cầu cách nhiệt":         "81C784",
    "Tầm trung phổ thông":    "FFF3E0",
    "Cao cấp rãnh C":         "E0F7FA",
    "Trượt lùa":              "BBDEFB",
    "Lift & Slide":           "E1BEE7",
    "Tầm trung":              "FFF3E0",
    "Mặt dựng":               "E0F2F1",
    "Cao cấp Premium":        "E8EAF6",
    "Phổ thông Vát cạnh":     "ECEFF1",
    "Cao cấp Platinum":       "FFECB3",
    "Rãnh C Classic":         "E0F7FA",
    "Anodized Phổ thông":     "FFF3CD",
    "Anodized Trung cấp":     "FFE082",
    "Anodized Cao cấp":       "D4AF37",
    "Anodized Siêu cao cấp":  "C8A000",
    "Tân cổ điển Cao cấp":    "E8D5B7",
    "Tân cổ điển Siêu cao cấp":"D4B896",
    "Cầu cách nhiệt Trung cấp":"C8E6C9",
    "Cầu cách nhiệt Cao cấp":  "81C784",
    "Cầu cách nhiệt Siêu cao cấp":"4CAF50",
    "Trượt lùa Trung cấp":     "BBDEFB",
    "Trượt lùa Cao cấp":       "64B5F6",
    "Trượt lùa Siêu cao cấp":  "1565C0",
    "Lift & Slide Cao cấp":    "E1BEE7",
    "Lift & Slide Siêu cao cấp":"9C27B0",
    "Thủy lực Siêu cao cấp":   "FF8A65",
    "Xếp trượt Cao cấp":       "A5D6A7",
    "Slim Nội thất Phổ thông":  "F8BBD9",
    "Slim Nội thất Cao cấp":    "F48FB1",
    "Slim Ngoại thất Cao cấp":  "EC407A",
    "Slim Steel Look":          "424242",
    "Slim Pocket Âm tường":     "757575",
    "Curtain Wall Stick":       "B2EBF2",
    "Curtain Wall Unitized":    "00ACC1",
    "Phổ thông":                "F5F5F5",
    "Phổ thông TQ":             "EEEEEE",
    "Lam chắn nắng":            "FFF9C4",
    "Casement Đức Cao cấp":     "D7CCC8",
    "Casement Đức Passive House":"A1887F",
    "Sliding Đức Cao cấp":      "BCAAA4",
    "Lift Slide Đức Siêu cao cấp":"795548",
    "Curtain Wall Đức":         "4E342E",
    "Casement Bỉ Cao cấp":      "C5CAE9",
    "Casement Bỉ Passive":      "7986CB",
    "Lift Slide Bỉ":            "3949AB",
    "Curtain Wall Bỉ":          "1A237E",
    "Casement Nhật":            "FCE4EC",
    "Sliding Nhật":             "F48FB1",
    "Phổ thông TQ":             "FAFAFA",
}

def get_fill(group):
    hex_c = GROUP_COLORS.get(group, "FFFFFF")
    return PatternFill(fill_type="solid", fgColor=hex_c)

thin_b = Border(
    left=Side(style="thin",color="DDDDDD"),
    right=Side(style="thin",color="DDDDDD"),
    top=Side(style="thin",color="DDDDDD"),
    bottom=Side(style="thin",color="DDDDDD"),
)

# ─────────────────────────────────────────────
# TẠO FILE 1: 01_Sổ tay Kỹ thuật Tổng hợp.docx
# ─────────────────────────────────────────────
def create_handbook():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    def h1(text):
        p = doc.add_heading(text, level=1)
        p.runs[0].font.color.rgb = RGBColor(0x0D,0x22,0x40)
        p.runs[0].font.size = Pt(15)
        p.runs[0].font.name = "Arial"

    def h2(text):
        p = doc.add_heading(text, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x70,0x5D,0x30)
        p.runs[0].font.size = Pt(13)
        p.runs[0].font.name = "Arial"

    def h3(text):
        p = doc.add_heading(text, level=3)
        p.runs[0].font.color.rgb = RGBColor(0x4B,0x55,0x63)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.name = "Arial"

    def para(text, bold=False, italic=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name  = "Arial"
        run.font.size  = Pt(size)
        run.bold       = bold
        run.italic     = italic
        return p

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # ── BÌA ──────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("THƯ VIỆN HỆ NHÔM SAO VÀNG")
    run.font.name = "Arial Black"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0D,0x22,0x40)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("SỔ TAY KỸ THUẬT TỔNG HỢP")
    run2.font.name = "Arial"
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x70,0x5D,0x30)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Phiên bản: 3.0  |  Năm: 2026  |  Phòng R&D – Sao Vàng Group").font.size = Pt(10)
    doc.add_page_break()

    # ── TẬP 01: GIỚI THIỆU ──────────────────
    h1("TẬP 01 – GIỚI THIỆU VÀ TỔNG QUAN")
    h2("1.1 Mục đích Thư viện")
    para("Bộ tài liệu này được Phòng R&D Sao Vàng xây dựng nhằm hệ thống hóa toàn bộ kiến thức kỹ thuật về các hệ cửa nhôm kính từ hơn 15 thương hiệu lớn trong nước và quốc tế, phục vụ công tác tư vấn, thiết kế, và thi công.")

    h2("1.2 Cấu trúc Bộ Tài Liệu (4 Master Documents)")
    for item in [
        "01_Sổ tay Kỹ thuật Tổng hợp.docx – Tổng quan, phân loại hệ nhôm, sơ đồ hệ thống, so sánh hãng",
        "02_Cơ sở Dữ liệu Gốc.xlsx – Database phẳng tra cứu đầy đủ 100+ hệ nhôm, 30 cột thông số kỹ thuật",
        "03_Danh mục Catalogue theo Hãng.xlsx – Danh mục catalogue, CAD, BIM, hướng dẫn lắp đặt từng hãng",
        "04_Bảng theo dõi Tiến độ.xlsx – Trạng thái thu thập dữ liệu từng hệ, từng hãng",
    ]:
        bullet(item)

    h2("1.3 Phân loại Phân khúc Nhôm")
    segments = {
        "Nhôm Phổ thông (Economy)":   "Yangli, Xingfa, Topal XFEC – Nhôm 6063T5 dày 1.0-1.4mm, sơn tĩnh điện, phụ kiện rãnh 22",
        "Nhôm Trung cấp (Mid-range)": "VRA55, Maxpro R55, Topal XFAD, Soco 65 – Nhôm 1.4-1.6mm, rãnh C, phụ kiện Châu Âu",
        "Nhôm Cao cấp (Premium)":     "VRE65, Civro AW65/AD65, Maxpro R65+, PAG 60/65 – Nhôm 1.6-2.0mm, Anodized ED, cầu cách nhiệt",
        "Nhôm Siêu cao cấp (Luxury)": "Civro AW75/Lift Slide, Owin HL180, Soco LS100, Schüco, Reynaers – Nhôm 2.0-3.0mm",
        "Nhôm Nhập khẩu Châu Âu":     "Schüco (Đức), Reynaers (Bỉ), Technal (Pháp), Aluprof (Ba Lan), WICONA (Đức)",
        "Nhôm Nhật Bản":              "YKK AP – Tiêu chuẩn JIS, Anodize ED bền 40 năm",
    }
    for k, v in segments.items():
        bullet(f"{k}: {v}")

    doc.add_page_break()

    # ── TẬP 02: CỬA ĐI ──────────────────────
    h1("TẬP 02 – HỆ CỬA ĐI (DOORS)")
    systems_door = [
        ("Hệ Mở quay – Casement 55-65-75",
         ["Công năng: Cửa đi đơn (1 cánh) hoặc kép (2 cánh) mở vào hoặc mở ra",
          "Hãng có hệ: Viralwindow (VRA55/65, VRE65), Civro (AD55/65/75), Maxpro JP (R55/65/70/83), Soco 65, Topal Prima, EuroVN, Yangli, Xingfa",
          "Ứng dụng: Cửa đi chính biệt thự, nhà phố, văn phòng, chung cư",
          "Kích thước tối đa điển hình: 1600×2800mm mỗi cánh",
          "Kính: Đơn 5-14mm, Kính hộp Low-E 20-36mm"]),
        ("Hệ Mở quay Tân cổ điển – R70, R83",
         ["Công năng: Cửa đi tân cổ điển – profile vát cạnh, bản khung rộng (70-83mm)",
          "Hãng: Maxpro JP (R70, R83), EuroVN VIP, Topal Prima",
          "Ứng dụng: Biệt thự cổ điển Pháp, penthouse tân cổ điển, cửa cung vòm",
          "Đặc điểm: Làm được cánh vòm cung (arch) và cánh kích thước siêu lớn"]),
        ("Hệ Thủy lực – Hydraulic / Zero Gravity",
         ["Công năng: Cửa đi siêu nặng tự đóng êm nhờ piston thủy lực cân bằng",
          "Hãng: Owin HL180, Maxpro R200, Viralwindow Thủy lực, EuroVN Thủy lực, Yangli Thủy lực",
          "Ứng dụng: Biệt thự siêu sang, sảnh khách sạn 5 sao, showroom xe sang",
          "Kích thước: đến 5000×4000mm / Tải trọng: đến 1500kg/cánh",
          "Phụ kiện thủy lực: Biloba, DICTATOR, Dorma TS98 Floor Spring"]),
        ("Hệ Trượt quay – Sliding Rotating",
         ["Công năng: Cửa lùa kết hợp xoay 180° – không ray dưới sàn",
          "Hãng: Owin SR, EuroVN Trượt quay, Viralwindow VRSD",
          "Ứng dụng: Biệt thự, nhà phố sang trọng, phòng trưng bày",
          "Phụ kiện: Dorma Slido, Häfele Sliding-Pivot"]),
        ("Hệ Pivot – Xoay trục",
         ["Công năng: Cửa xoay quanh trục trung tâm hoặc lệch tâm đứng",
          "Hãng: Civro Pivot, và custom các hãng",
          "Ứng dụng: Biệt thự siêu sang, sảnh resort, lối vào đặc biệt",
          "Phụ kiện pivot: FritsJurgens (Hà Lan), Dorma, Biloba",
          "Kích thước: đến 2000×4000mm, tải trọng đến 500kg"]),
        ("Hệ Bi-fold / Folding – Xếp trượt",
         ["Công năng: Cửa xếp gấp trượt sang bên, mở toàn bộ không gian",
          "Hãng: Civro Folding, Maxpro SFD80, PAG 83, Soco Folding, Viralwindow Bi-fold",
          "Ứng dụng: Nhà hàng, quán cà phê, lối ra sân biệt thự",
          "Phụ kiện: Roto Fold, OPK Bi-fold, Hopo"]),
        ("Hệ Tự động – Auto Sliding Door",
         ["Công năng: Cửa tự động trượt khi phát hiện người qua cảm biến hồng ngoại hoặc radar",
          "Hãng: Dorma Automatics, Besam (Assa Abloy), các nhà lắp đặt nội địa",
          "Ứng dụng: Siêu thị, TTTM, bệnh viện, sảnh khách sạn, nhà ga",
          "Chuẩn: EN 16005 – An toàn cửa tự động"]),
    ]
    for name, pts in systems_door:
        h2(name)
        for pt in pts:
            bullet(pt)

    doc.add_page_break()

    # ── TẬP 03: CỬA SỔ ──────────────────────
    h1("TẬP 03 – HỆ CỬA SỔ (WINDOWS)")
    systems_window = [
        ("Hệ Mở quay – Casement Window",
         ["Hãng: Viralwindow VRE65, Civro AW55/65/75, Maxpro R65, PAG 60, Soco 65, Schüco AWS65/75",
          "Ứng dụng: Cửa sổ mở vào hoặc mở ra, phổ biến nhất cho nhà phố và biệt thự",
          "Kích thước: Tối đa 1800×3000mm/cánh"]),
        ("Hệ Mở hất – Awning / Top Hung",
         ["Công năng: Cánh mở hất lên trên – che mưa khi mở, thông gió tốt",
          "Hãng: Viralwindow VRAW, Civro AW65 OUT, Maxpro R65",
          "Ứng dụng: Phòng vệ sinh, bếp, nhà kho, nhà xưởng"]),
        ("Hệ Mở lật – Tilt-Turn",
         ["Công năng: 2 vị trí – Tilt (hé nghiêng thông gió an toàn) + Turn (mở quay vào trong)",
          "Hãng: Civro AW65 Tilt-Turn, PAG, Schüco AWS, Reynaers CS",
          "Ứng dụng: Chung cư cao tầng, văn phòng an toàn, khách sạn",
          "Ưu điểm: Lau kính từ bên trong, thông gió an toàn không rơi cánh"]),
        ("Hệ Trượt lùa – Sliding Window",
         ["Hãng: Viralwindow VRA64, VRA94, Maxpro SW55/SW65/SW83, Soco 94/120, Yangli Lùa 95/97",
          "Ứng dụng: Cửa sổ phổ thông nhà phố, chung cư, ban công nhỏ",
          "Ưu điểm: Giá tốt, lắp đặt nhanh, phù hợp mở rộng không gian nhỏ"]),
        ("Hệ Cố định – Fixed Light",
         ["Công năng: Không mở, chỉ lấy ánh sáng",
          "Hãng: Tất cả các hãng đều có panel cố định đi kèm",
          "Ứng dụng: Kết hợp với cửa mở, vách kính, mái giếng trời"]),
    ]
    for name, pts in systems_window:
        h2(name)
        for pt in pts:
            bullet(pt)

    doc.add_page_break()

    # ── TẬP 04: SLIM ────────────────────────
    h1("TẬP 04 – HỆ CỬA SLIM")
    h2("4.1 Slim Nội thất – Interior Slim")
    for pt in [
        "Profile siêu mỏng 40-70mm, trọng tải cánh 60-120kg, ray giảm chấn Hettich/Häfele/Slido Dorma",
        "Không ray dưới sàn – thẩm mỹ sàn hoàn toàn liền mạch",
        "Gioăng cạnh mềm đơn giản, không chống thấm nước mưa ngoài trời",
        "Ứng dụng: Phòng ngủ, WC, phòng thay đồ, tủ quần áo, văn phòng open space",
        "Trọng tải: 60-120kg/cánh theo ray được chọn (Slido 50, Slido 80kg, Slido 120kg...)",
    ]:
        bullet(pt)

    h2("4.2 Slim Ngoại thất – Outdoor Slim")
    for pt in [
        "Profile 60mm kép chịu lực, gioăng EPDM chống thấm nước mưa",
        "Ray thép cường lực + Bánh xe inox SUS304",
        "Ứng dụng: Lối ra ban công, sân vườn nhỏ biệt thự",
        "Tải trọng: 120-150kg/cánh, kính hộp được 20-28mm",
    ]:
        bullet(pt)

    h2("4.3 Slim Steel Look – Giả thép nghệ thuật")
    for pt in [
        "Profile nhôm sơn đặc chủng giả thép (powder coat black, grey matte, dark bronze)",
        "Hãng: Vitrocsa (Thụy Sỹ) – thương hiệu steel look cao cấp nhất thế giới; các hãng local Việt Nam đang phát triển",
        "Ứng dụng: Biệt thự phong cách Loft/Industrial, showroom, nhà hàng nghệ thuật",
        "Kích thước: đến 3000×4000mm, kính cường lực 10-15mm",
    ]:
        bullet(pt)

    h2("4.4 Slim Pocket – Lùa âm tường")
    for pt in [
        "Cánh trượt hoàn toàn vào bên trong tường khi mở",
        "Ứng dụng: WC, phòng ngủ, bếp – tiết kiệm diện tích tối đa",
        "Yêu cầu: Tường đủ dày để chứa chiều rộng cánh cửa",
        "Hệ ray: Häfele Pocket Door, Barn Door âm tường Slido",
    ]:
        bullet(pt)

    doc.add_page_break()

    # ── TẬP 05: CẦU CÁCH NHIỆT ──────────────
    h1("TẬP 05 – HỆ CẦU CÁCH NHIỆT (THERMAL BREAK)")
    h2("5.1 Nguyên lý Cầu cách nhiệt")
    para("Profile nhôm cầu cách nhiệt gồm 2 lớp nhôm ngăn cách bởi thanh Polyamide PA66 GF25 (Technoform Đức hoặc tương đương). Thanh PA66 có hệ số dẫn nhiệt λ ≈ 0.25 W/(m·K) so với nhôm λ = 200 W/(m·K) – giảm dẫn nhiệt đến 800 lần, loại bỏ hiện tượng đọng sương trong nhà.")
    h2("5.2 Các hãng nhôm cầu cách nhiệt tại Việt Nam")
    thermal_brands = {
        "PAG":    "PAG 60, PAG 80, PAG 83, PAG 120, PAG 125, PAG 200 – Cầu cách nhiệt PA66 + AkzoNobel chống muối biển 25 năm",
        "Civro":  "Civro AW/AD 55/65/75 IN/OUT – Cầu cách nhiệt Technoform Đức, MED xi mạ, đạt Passive House",
        "Maxpro": "Maxpro TB65 – Cầu cách nhiệt cải tiến Nhật cho khí hậu nhiệt đới",
        "Topal":  "Topal Prima – Profile cầu cách nhiệt nằm trong hệ rãnh C",
        "EuroVN": "EuroVN Thermal – Dải Polyamide + Sơn Wagner Thụy Sỹ",
    }
    for k, v in thermal_brands.items():
        bullet(f"{k}: {v}")

    h2("5.3 Hệ số nhiệt Uf so sánh các hãng")
    uf_data = [
        ("Schüco AWS 75.SI (Đức)", "0.97 W/(m²K)", "Passive House PHI"),
        ("Reynaers CS 77 (Bỉ)", "1.0 W/(m²K)", "Passive House PHI"),
        ("Schüco AWS 65 (Đức)", "1.3 W/(m²K)", "BREAM / LEED Platinum"),
        ("Reynaers CS 60 (Bỉ)", "1.5 W/(m²K)", "LEED Gold"),
        ("Civro AW75 IN (VN)", "1.6 W/(m²K)", "LEED Silver/Gold"),
        ("PAG 65TB (VN)", "1.8 W/(m²K)", "Công trình xanh tiêu chuẩn"),
        ("PAG 60 (VN)", "2.5 W/(m²K)", "Công trình tiết kiệm năng lượng cơ bản"),
        ("Nhôm thường không cầu cách nhiệt", ">5.0 W/(m²K)", "Không đạt tiêu chuẩn xanh"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Hãng / Hệ", "Uf W/(m²K)", "Chứng chỉ xanh"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row in uf_data:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()
    doc.add_page_break()

    # ── TẬP 06: CURTAIN WALL ────────────────
    h1("TẬP 06 – HỆ VÁch MẶT DỰNG (CURTAIN WALL)")
    for sub in [
        ("Hệ Stick – Lắp ghép tại công trường",
         ["Lắp từng thanh đố đứng (mullion) và đố ngang (transom) tại chỗ",
          "Hãng: Civro CCW50, Maxpro MPCW50, Soco Vách Fix, Yangli Mặt dựng 52/65",
          "Ưu điểm: Chi phí thấp, linh hoạt sửa đổi kiến trúc",
          "Nhược điểm: Thi công chậm, chất lượng phụ thuộc thực địa"]),
        ("Hệ Unitized – Panel đúc sẵn tại xưởng",
         ["Panel hoàn chỉnh (kính + nhôm) đúc tại xưởng, lắp bằng cần trục",
          "Hãng: Soco UCW Unitized, Maxpro Unitized, Schüco FW 50+ Unitized",
          "Ưu điểm: Thi công cực nhanh (3-4 ngày/tầng), chất lượng đồng đều tuyệt đối",
          "Nhược điểm: Chi phí mold cao, không linh hoạt sửa đổi tại chỗ"]),
        ("Hệ Spider – Mặt kính điểm Cần nhện",
         ["Kính đặc biệt có lỗ cán khoan + Cần nhện (Spider Fitting) inox SUS316",
          "Khung kết cấu thép ẩn phía sau – nhìn từ ngoài chỉ thấy kính thuần túy",
          "Ứng dụng: Sảnh khách sạn biểu tượng, trung tâm hội nghị",
          "Keo Silicon kết cấu: Dow Corning 895, Sika SikaSil SG-500"]),
        ("Hệ Cáp căng – Cable Tensile Facade",
         ["Hệ dây cáp inox SUS316 chịu lực căng thẳng đứng",
          "Kính điểm hoặc kính gương gắn lên cáp qua cần nhện",
          "Ứng dụng: Công trình nghệ thuật biểu tượng, sân bay, trung tâm văn hóa"]),
    ]:
        h2(sub[0])
        for pt in sub[1]:
            bullet(pt)

    doc.add_page_break()

    # ── TẬP 07: LAN CAN ─────────────────────
    h1("TẬP 07 – HỆ LAN CAN (RAILINGS)")
    for pt in [
        "Lan can nhôm đúc biệt thự: Profile nhôm đúc khuôn trang trí hoa văn phong cách cổ điển, sơn tĩnh điện vàng/đồng",
        "Lan can kính Frameless: Đế âm sàn inox SUS316 chôn bê tông + Kính cường lực đứng 12-19mm, không trụ giữa",
        "Lan can nhôm Slim: Profile nhôm mỏng 20×20-40×40mm làm trụ và thanh ngang, kết hợp kính hoặc cáp inox",
        "Lan can âm sàn – U-channel: Rãnh nhôm chữ U chôn âm sàn đỡ kính cường lực đứng, không thấy đế bắt vít",
        "Chuẩn TCVN 9392 – Lan can cầu thang: Chiều cao tối thiểu 1100mm cho công trình cao tầng",
    ]:
        bullet(pt)

    doc.add_page_break()

    # ── TẬP 08: MÁI KÍNH ────────────────────
    h1("TẬP 08 – HỆ MÁI KÍNH (SKYLIGHTS & ROOFING)")
    for pt in [
        "Skylight – Mái giếng trời: Khung nhôm nghiêng chứa kính hộp Low-E 20-28mm, góc nghiêng 15-45° để tự thoát nước mưa",
        "Pergola nhôm lá chớp xoay: Lá chớp nhôm đặc dày 80-120mm, mô tơ xoay 0-135° theo remote hoặc cảm biến mưa tự động, phù hợp nhà hàng sân vườn resort",
        "Canopy sảnh đón: Khung nhôm + tấm PC trong suốt hoặc kính cường lực 8-10mm nghiêng 3-5°, che mưa không cản ánh sáng",
        "Sunroom – Phòng kính ngoài trời: Hệ kết hợp vách kính (Casement + Cố định) và mái kính nghiêng tạo phòng kính đa năng quanh năm",
        "Mái kính phẳng: Hệ kính hộp nằm ngang chịu lực (Structural Glazing Roof) – cần kính laminated đặc chủng an toàn chống vỡ rơi xuống",
    ]:
        bullet(pt)

    doc.add_page_break()

    # ── TẬP 09: PHỤ KIỆN ────────────────────
    h1("TẬP 09 – HỆ PHỤ KIỆN (HARDWARE)")
    hardware = {
        "Bản lề (Hinges)": [
            "Bản lề cửa sổ Casement: CMECH 3D 7mm – Điều chỉnh 3 chiều X/Y/Z, khả năng mang tải 80-120kg/cánh",
            "Bản lề Pivot Floor Spring: Dorma BTS80, Geze, Biloba – Dùng cho cửa đi thủy lực siêu nặng",
            "Bản lề xếp Bi-fold: Roto Fold, OPK – Cho phép cánh gấp lại 180°",
        ],
        "Khóa (Locks)": [
            "Khóa đa điểm (Multi-point Lock): Roto NT, Siegenia, Winkhaus – 3-5 điểm khóa phân tán tải cửa",
            "Tay nắm (Handle): Hoppe (Đức), Roto (Đức), Giesse (Ý) – Tích hợp cơ khóa đa điểm",
            "Khóa thông minh: Hafele, Yale Smart Lock, VinSmart – Vân tay, thẻ từ, Bluetooth, App",
        ],
        "Ray và Bánh xe (Tracks & Rollers)": [
            "Ray inox SUS304: Cho cửa lùa nội địa, trọng tải 80-200kg",
            "Ray inox SUS316L: Chống muối biển, dùng công trình ven biển",
            "Phụ kiện Lift Slide: Roto Patio Life, Siegenia Titan LSS – Nâng cánh 1-2mm khi mở, hạ kín khi đóng",
            "Bánh xe giảm chấn: OPK, Kinlong – Tự hãm cuối hành trình, chống đập mạnh",
        ],
        "Gioăng cao su (Seals)": [
            "EPDM đơn khoang: Cho cửa phổ thông, chống bụi và mưa nhẹ",
            "EPDM đa khoang (2-4 khoang): Cho hệ cao cấp – chống ồn và kín nước tối ưu",
            "EPDM cầu cách nhiệt: Chứa túi khí hoặc foam cách nhiệt, phối hợp cầu PA66",
            "Gioăng trương nở (Intumescent): Tự phồng to khi nhiệt độ > 120°C – dùng cửa chống cháy",
        ],
    }
    for cat, items in hardware.items():
        h2(cat)
        for item in items:
            bullet(item)

    doc.add_page_break()

    # ── TẬP 10: PROFILE ─────────────────────
    h1("TẬP 10 – GIẢI PHẪU PROFILE NHÔM")
    h2("10.1 Thành phần mặt cắt profile cửa mở quay Casement")
    for pt in [
        "Thanh Khung bao (Frame): Gắn cố định vào tường – chịu lực toàn bộ cửa",
        "Thanh Cánh (Sash): Khung bao quanh kính – di chuyển khi mở/đóng",
        "Thanh Đố đứng (Mullion): Phân chia các ô kính trong cùng một cửa",
        "Thanh Đố ngang (Transom): Ngăn chia kính theo chiều ngang",
        "Nẹp kính (Glazing Bead): Giữ kính vào cánh bằng keo Silicone hoặc nẹp ép",
        "Rãnh chèn gioăng (Seal groove): Khe 5.5-6.2mm cho gioăng EPDM",
        "Rãnh C (C-slot): Rãnh 15/20mm chuẩn Euro Groove cho phụ kiện Châu Âu",
        "Cầu cách nhiệt PA66 (Thermal Break): Thanh Polyamide ngăn giữa 2 lớp nhôm",
    ]:
        bullet(pt)

    h2("10.2 Tiêu chuẩn nhôm nguyên chất đầu vào")
    for pt in [
        "Hợp kim 6063-T5: Al + Mg + Si, kéo dài tốt, bề mặt đẹp – dùng phổ thông",
        "Hợp kim 6063-T6: Nhiệt xử lý T6 tăng cứng, chịu lực tốt hơn T5 – dùng cao cấp",
        "Hợp kim 6060-T6: Civro dùng – hàm lượng Mg cao hơn, bề mặt anodized tốt hơn 6063",
        "Độ dày thực tế: Profile phổ thông 1.0-1.2mm, trung cấp 1.4-1.6mm, cao cấp 1.8-2.5mm",
    ]:
        bullet(pt)

    doc.add_page_break()

    # ── TẬP 11: SHOPDRAWING ─────────────────
    h1("TẬP 11 – QUY CHUẨN SHOPDRAWING")
    for pt in [
        "Tỷ lệ bản vẽ: 1:1 trong Model Space AutoCAD, xuất PDF 1:10 hoặc 1:5",
        "Layer quản lý: NET_THAY (đường biên dạng nhôm) – NET_KHUAT (khuất bê tông/gioăng) – DIM_TEXT (kích thước chữ) – KÍNH (hatch nét kính)",
        "Hatch nét kính đơn: ANSI31 45° nét mảnh – Hatch kính hộp: 2 lớp hatch song song cách nhau 12-14mm",
        "Ký hiệu cắt mặt cửa: Ký hiệu theo tiêu chuẩn ISO 7519 và TCVN 4608",
        "Mô tả profile kèm: Mã hãng, độ dày nhôm, màu sắc anodize/sơn, loại gioăng",
        "Bố cục trang in: Khung tên dự án, số tờ, ngày xuất, người vẽ, người duyệt",
    ]:
        bullet(pt)

    doc.add_page_break()

    # ── TẬP 12: LẮP ĐẶT ────────────────────
    h1("TẬP 12 – QUY TRÌNH LẮP ĐẶT NGOÀI CÔNG TRƯỜNG")
    steps = [
        "B1 – Khảo sát ô chờ: Kiểm tra độ vuông, cốt 0.0, độ dốc cốt nền, độ phẳng tường, sai số ≤ 3mm",
        "B2 – Định vị khung bao: Vít nở inox SUS304 M8×80mm, khoảng cách ≤ 600mm, chêm nhựa chống rung giữa nhôm và bê tông",
        "B3 – Lắp kính: Nêm nhựa EPDM chống sệ cánh, keo silicon kết cấu Dow Corning 795 hoặc Sika, không bơm keo silicone trực tiếp lên nhôm thô",
        "B4 – Bơm keo PU: Keo bọt nở polyurethane lấp đầy khe hở nhôm–tường, đợi khô 60 phút, cắt phẳng",
        "B5 – Bơm keo silicon ngoại thất: Silicon trung tính (neutral cure) Apollo A500/A600 hoặc Dow Corning 791, tạo đường keo đều 8-12mm, góc 45°, dùng dao cạo tạo phẳng",
        "B6 – Căn chỉnh phụ kiện: Bản lề 3D căn X/Y/Z, khóa đa điểm vận hành êm, bánh xe cân tải không nghiêng cánh",
        "B7 – Nghiệm thu: Bóc băng keo bảo vệ profile, kiểm tra độ kín (test nước phun), vận hành thử cửa đóng mở 10 lần liên tục, bàn giao biên bản",
    ]
    for s in steps:
        bullet(s)

    doc.add_page_break()

    # ── TẬP 13: SO SÁNH HÃNG ────────────────
    h1("TẬP 13 – BẢNG SO SÁNH MA TRẬN CÁC HÃNG")
    h2("13.1 Ma trận hệ nhôm theo hãng")
    brands = ["VR","Civro","Maxpro","Soco","PAG","Owin","Topal","EuroVN","Schüco","Reynaers"]
    sys_list = [
        ("Hệ 55 Mở quay",   ["✓","✓","✓","✓","✗","✓","✓","✓","✗","✗"]),
        ("Hệ 65 Mở quay",   ["✓","✓","✓","✓","✓","✓","✓","✓","✓","✓"]),
        ("Hệ 75 Mở quay",   ["✗","✓","✗","✗","✗","✗","✗","✗","✓","✓"]),
        ("Cầu cách nhiệt",  ["✗","✓","✓","✗","✓","✗","✓","✓","✓","✓"]),
        ("Lift & Slide",    ["✓","✓","✗","✓","✓","✗","✗","✗","✓","✓"]),
        ("Thủy lực",        ["✓","✗","✓","✗","✗","✓","✗","✓","✗","✗"]),
        ("Bi-fold/Folding", ["✓","✓","✓","✓","✓","✓","✗","✗","✓","✓"]),
        ("Pivot",           ["✗","✓","✗","✗","✗","✗","✗","✗","✓","✓"]),
        ("Tân cổ điển R83", ["✗","✗","✓","✗","✗","✗","✗","✗","✗","✗"]),
        ("Curtain Wall",    ["✗","✓","✓","✓","✗","✗","✗","✗","✓","✓"]),
        ("Slim Nội thất",   ["✗","✗","✗","✗","✗","✗","✓","✗","✗","✗"]),
    ]
    tbl2 = doc.add_table(rows=1, cols=len(brands)+1)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    hdr2[0].text = "Hệ nhôm"
    hdr2[0].paragraphs[0].runs[0].bold = True
    for i, b in enumerate(brands):
        hdr2[i+1].text = b
        hdr2[i+1].paragraphs[0].runs[0].bold = True
    for sys_name, checks in sys_list:
        row2 = tbl2.add_row().cells
        row2[0].text = sys_name
        for i, c in enumerate(checks):
            row2[i+1].text = c
    doc.add_paragraph()
    h2("13.2 Khuyến nghị chọn hệ theo loại công trình")
    recs = {
        "Nhà phố ngân sách 1-2 tỷ": "Topal XFAD 55 hoặc Yangli 55 – Giá cạnh tranh, đủ chức năng",
        "Nhà phố cao cấp 5-15 tỷ": "VRE65 hoặc Maxpro R65 hoặc Soco 65 – Cân bằng chất lượng và giá",
        "Biệt thự 20-50 tỷ": "Civro AW65/AD65 hoặc PAG 65TB – Cầu cách nhiệt, phụ kiện Châu Âu nguyên bộ",
        "Biệt thự siêu sang 50-200 tỷ+": "Civro Lift Slide + PAG 200 hoặc Soco LS100 + Owin HL180 thủy lực",
        "Tòa nhà văn phòng hạng A": "Schüco FW 50+ CW (nhập khẩu) hoặc Soco UCW Unitized + Civro CCW50 Stick",
        "Resort/Khách sạn 5 sao biển": "Maxpro SD115 chống bão hoặc Soco SD90 SUS316L + Owin HL180 thủy lực sảnh",
        "Biệt thự thông minh": "VRE65 hoặc Civro AW65 + Smart Lock Hafele/Yale tích hợp hệ thống KNX/Zigbee",
        "Công trình xanh LEED": "Civro AW65 IN cầu cách nhiệt + PAG 65TB + Kính hộp Low-E 32mm",
    }
    for k, v in recs.items():
        bullet(f"{k}: {v}")

    path = os.path.join(ROOT, "01_Sổ tay Kỹ thuật Tổng hợp.docx")
    safe_save_docx(doc, path)

# ─────────────────────────────────────────────
# TẠO FILE 2: 02_Cơ sở Dữ liệu Gốc.xlsx
# ─────────────────────────────────────────────
def create_database():
    wb = Workbook()
    # Xóa sheet mặc định ban đầu
    default_sheet = wb.active
    wb.remove(default_sheet)
    
    # ─────────────────────────────────────────────
    # SHEET 1: 01_MASTER_DATABASE
    # ─────────────────────────────────────────────
    ws_master = wb.create_sheet("01_MASTER_DATABASE")
    ws_master.views.sheetView[0].showGridLines = True
    
    hdr_fill = PatternFill(fill_type="solid", fgColor="0D2240")
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    all_headers = HEADERS + ["Link Bản Vẽ CAD"]
    ws_master.row_dimensions[1].height = 45
    for col_idx, h in enumerate(all_headers, 1):
        cell = ws_master.cell(row=1, column=col_idx, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    for row_idx, row in enumerate(ROWS, 2):
        row_data = list(row) + [""]
        # Tự động điền link catalogue Maxpro JP nếu trống
        if row_data[3] == "Maxpro JP" and not row_data[28]:
            row_data[28] = "https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-maxpro-7.html"
            
        group = row[1]
        fill = get_fill(group)
        ws_master.row_dimensions[row_idx].height = 35
        
        for col_idx, val in enumerate(row_data, 1):
            cell = ws_master.cell(row=row_idx, column=col_idx, value=val)
            cell.fill = fill
            cell.font = Font(name="Arial", size=9)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if col_idx == 27: # Hoàn thiện %
                cell.alignment = Alignment(horizontal="center", vertical="center")
                if isinstance(val, int) and val >= 80:
                    cell.font = Font(name="Arial", size=9, color="1B5E20", bold=True)
                elif isinstance(val, int) and val < 60:
                    cell.font = Font(name="Arial", size=9, color="B71C1C")
                    
    ws_master.auto_filter.ref = f"A1:{get_column_letter(len(all_headers))}1"
    ws_master.freeze_panes = "A2"
    col_widths = [9,22,20,16,12,18,22,22,24,18,30,14,20,12,20,15,28,28,30,40,35,8,10,10,10,18,13,25,25,25]
    for i, w in enumerate(col_widths[:len(all_headers)], 1):
        ws_master.column_dimensions[get_column_letter(i)].width = w

    # Lấy danh sách các hãng duy nhất sắp xếp theo bảng chữ cái
    unique_brands = sorted(list(set(row[3] for row in ROWS)))
    total_rows = len(ROWS)
    
    # Hàm helper tạo tiêu đề bảng
    def write_sheet_title(ws, title, cols_count):
        ws.views.sheetView[0].showGridLines = True
        ws.merge_cells(start_row=1, start_column=1, end_row=2, end_column=cols_count)
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = Font(name="Arial", size=14, bold=True, color="0D2240")
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 20
        ws.row_dimensions[2].height = 20
        
    # Hàm helper tạo header cột
    def write_headers(ws, headers_list, row_num=4):
        ws.row_dimensions[row_num].height = 30
        for col_idx, h in enumerate(headers_list, 1):
            cell = ws.cell(row=row_num, column=col_idx, value=h)
            cell.fill = PatternFill(fill_type="solid", fgColor="0B5345") # Xanh lục sậm
            cell.font = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_b

    # ─────────────────────────────────────────────
    # SHEET 2: 02_TỔNG QUAN HỆ NHÔM
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("02_TỔNG QUAN HỆ NHÔM")
    write_sheet_title(ws, "BẢNG TỔNG QUAN SỐ LƯỢNG HỆ THEO TỪNG HÃNG SẢN XUẤT", 8)
    hdrs = ["Hãng", "Số Hệ", "Cửa Đi", "Cửa Sổ", "Hệ Slim", "Hệ Thermal", "Vách / Mặt Dựng", "Lan Can"]
    write_headers(ws, hdrs)
    
    for idx, brand in enumerate(unique_brands, 5):
        ws.row_dimensions[idx].height = 25
        ws.cell(row=idx, column=1, value=brand).font = Font(name="Arial", bold=True, size=10)
        ws.cell(row=idx, column=1).border = thin_b
        
        formulas = [
            f"=COUNTIF('01_MASTER_DATABASE'!$D$2:$D$500, A{idx})", # Số Hệ
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*Cửa đi*\")", # Cửa Đi
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*Cửa sổ*\")", # Cửa Sổ
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Slim*\")", # Slim
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Thermal*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*cầu*\")", # Thermal
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*mặt dựng*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*vách*\")", # Vách/Mặt dựng
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*lan can*\")" # Lan can
        ]
        
        for c_idx, f_str in enumerate(formulas, 2):
            cell = ws.cell(row=idx, column=c_idx, value=f_str)
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_b
            
    ws.column_dimensions["A"].width = 20
    for col in ["B","C","D","E","F","G","H"]:
        ws.column_dimensions[col].width = 15

    # ─────────────────────────────────────────────
    # SHEET 3: 03_SO SÁNH PHÂN KHÚC
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("03_SO SÁNH PHÂN KHÚC")
    write_sheet_title(ws, "PHÂN TÍCH HỆ CỬA THEO PHÂN KHÚC THƯƠNG HIỆU", 5)
    write_headers(ws, ["Hãng", "Phổ Thông / Tiết Kiệm", "Tầm Trung", "Cao Cấp rãnh C", "Luxury / Siêu Cao Cấp"])
    
    for idx, brand in enumerate(unique_brands, 5):
        ws.row_dimensions[idx].height = 25
        ws.cell(row=idx, column=1, value=brand).font = Font(name="Arial", bold=True, size=10)
        ws.cell(row=idx, column=1).border = thin_b
        
        formulas = [
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Phổ thông*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Vát cạnh*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*nhựa*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Tầm trung*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*sản xuất*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Cao cấp*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Rãnh C*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Luxury*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Cầu cách nhiệt*\")"
        ]
        for c_idx, f_str in enumerate(formulas, 2):
            cell = ws.cell(row=idx, column=c_idx, value=f_str)
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_b
            
    ws.column_dimensions["A"].width = 20
    for col in ["B","C","D","E"]:
        ws.column_dimensions[col].width = 22

    # ─────────────────────────────────────────────
    # SHEET 4: 04_SO SÁNH LOẠI CỬA
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("04_SO SÁNH LOẠI CỬA")
    write_sheet_title(ws, "PHÂN LOẠI CỬA THEO THỂ LOẠI (CỬA ĐI, SỔ, VÁCH, MÁI...)", 8)
    write_headers(ws, ["Hãng", "Cửa Đi", "Cửa Sổ", "Vách Ngăn", "Mặt Dựng Vách Kính", "Lan Can Kính", "Mái Kính Cường Lực", "Khác"])
    
    for idx, brand in enumerate(unique_brands, 5):
        ws.row_dimensions[idx].height = 25
        ws.cell(row=idx, column=1, value=brand).font = Font(name="Arial", bold=True, size=10)
        ws.cell(row=idx, column=1).border = thin_b
        
        formulas = [
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*Cửa đi*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*Cửa sổ*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*vách ngăn*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*mặt dựng*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*lan can*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$G$2:$G$500, \"*mái kính*\")",
            f"=COUNTIF('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}) - (B{idx}+C{idx}+D{idx}+E{idx}+F{idx}+G{idx})"
        ]
        for c_idx, f_str in enumerate(formulas, 2):
            cell = ws.cell(row=idx, column=c_idx, value=f_str)
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_b
            
    ws.column_dimensions["A"].width = 20
    for col in ["B","C","D","E","F","G","H"]:
        ws.column_dimensions[col].width = 17

    # ─────────────────────────────────────────────
    # SHEET 5: 05_SO SÁNH KIỂU MỞ
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("05_SO SÁNH KIỂU MỞ")
    write_sheet_title(ws, "SO SÁNH CÁC PHƯƠNG ÁN KIỂU MỞ CÁNH CỬA", 7)
    write_headers(ws, ["Hãng", "Mở Quay (Casement)", "Trượt Lùa (Sliding)", "Mở Hất (Awning)", "Xếp Trượt (Folding)", "Trượt Quay (Slide & Turn)", "Pivot / Trục Xoay"])
    
    for idx, brand in enumerate(unique_brands, 5):
        ws.row_dimensions[idx].height = 25
        ws.cell(row=idx, column=1, value=brand).font = Font(name="Arial", bold=True, size=10)
        ws.cell(row=idx, column=1).border = thin_b
        
        formulas = [
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*quay*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*lùa*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*trượt*\") - COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*xếp trượt*\") - COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*trượt quay*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*hất*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*xếp trượt*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*trượt quay*\")",
            f"=COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*pivot*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$H$2:$H$500, \"*trục xoay*\")"
        ]
        for c_idx, f_str in enumerate(formulas, 2):
            cell = ws.cell(row=idx, column=c_idx, value=f_str)
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_b
            
    ws.column_dimensions["A"].width = 20
    for col in ["B","C","D","E","F","G"]:
        ws.column_dimensions[col].width = 20

    # ─────────────────────────────────────────────
    # SHEET 6: 06_SO SÁNH ĐỘ DÀY NHÔM
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("06_SO SÁNH ĐỘ DÀY NHÔM")
    write_sheet_title(ws, "DANH SÁCH HỆ CỬA SẮP XẾP THEO ĐỘ DÀY PROFILE (TỪ DÀY ĐẾN MỎNG)", 6)
    write_headers(ws, ["ID", "Hãng", "Mã Hệ", "Tên Hệ", "Loại Cửa", "Độ Dày (mm)"])
    
    sorted_indices_thickness = sorted(
        range(len(ROWS)), 
        key=lambda i: (float(ROWS[i][11]) if isinstance(ROWS[i][11], (int, float)) else 0), 
        reverse=True
    )
    
    for r_idx, orig_idx in enumerate(sorted_indices_thickness, 5):
        orig_row_excel = orig_idx + 2
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!A{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=4, value=f"='01_MASTER_DATABASE'!I{orig_row_excel}")
        ws.cell(row=r_idx, column=5, value=f"='01_MASTER_DATABASE'!G{orig_row_excel}")
        
        thick_cell = ws.cell(row=r_idx, column=6, value=f"='01_MASTER_DATABASE'!L{orig_row_excel}")
        thick_cell.alignment = Alignment(horizontal="center")
        thick_cell.font = Font(name="Arial", bold=True)
        
        for c in range(1, 7):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 28
    ws.column_dimensions["E"].width = 24
    ws.column_dimensions["F"].width = 15

    # ─────────────────────────────────────────────
    # SHEET 7: 07_SO SÁNH KÍNH
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("07_SO SÁNH KÍNH")
    write_sheet_title(ws, "KHẢ NĂNG TƯƠNG THÍCH VÀ ĐỘ DÀY KÍNH TỐI ĐA CỦA KHUNG", 7)
    write_headers(ws, ["Mã Hệ", "Hãng", "Kính Tương Thích", "Kính Min", "Kính Max (mm)", "Tương Thích Kính Hộp", "Tương Thích Kính Hộp 3 Lớp"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value=f"='01_MASTER_DATABASE'!M{orig_row_excel}")
        
        # Min
        ws.cell(row=r_idx, column=4, value=f"=IF(ISNUMBER(SEARCH(\"5\",'01_MASTER_DATABASE'!M{orig_row_excel})),5,8)").alignment = Alignment(horizontal="center")
        # Max
        ws.cell(row=r_idx, column=5, value=f"=IFERROR(VALUE(SUBSTITUTE('01_MASTER_DATABASE'!N{orig_row_excel},\"mm\",\"\")),0)").alignment = Alignment(horizontal="center")
        # Kính hộp
        ws.cell(row=r_idx, column=6, value=f"=IF(OR(ISNUMBER(SEARCH(\"hộp\",'01_MASTER_DATABASE'!M{orig_row_excel})), IFERROR(VALUE(SUBSTITUTE('01_MASTER_DATABASE'!N{orig_row_excel},\"mm\",\"\")),0)>=20),\"✓\",\"✗\")").alignment = Alignment(horizontal="center")
        # Kính 3 lớp
        ws.cell(row=r_idx, column=7, value=f"=IF(OR(ISNUMBER(SEARCH(\"3 lớp\",'01_MASTER_DATABASE'!M{orig_row_excel})), IFERROR(VALUE(SUBSTITUTE('01_MASTER_DATABASE'!N{orig_row_excel},\"mm\",\"\")),0)>=32),\"✓\",\"✗\")").alignment = Alignment(horizontal="center")
        
        for c in range(1, 8):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 15
    ws.column_dimensions["F"].width = 22
    ws.column_dimensions["G"].width = 25

    # ─────────────────────────────────────────────
    # SHEET 8: 08_SO SÁNH KÍCH THƯỚC
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("08_SO SÁNH KÍCH THƯỚC")
    write_sheet_title(ws, "GIỚI HẠN KÍCH THƯỚC CÁNH VÀ DIỆN TÍCH CỬA TỐI ĐA", 6)
    write_headers(ws, ["Mã Hệ", "Hãng", "Kích Thước Khung Đề Xuất", "Chiều Rộng Max (mm)", "Chiều Cao Max (mm)", "Diện Tích Cánh Max (m²)"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value=f"='01_MASTER_DATABASE'!O{orig_row_excel}").alignment = Alignment(horizontal="center")
        
        # Rộng max
        ws.cell(row=r_idx, column=4, value=f"=IFERROR(VALUE(LEFT('01_MASTER_DATABASE'!O{orig_row_excel}, SEARCH(\"×\",'01_MASTER_DATABASE'!O{orig_row_excel})-1)), 0)").alignment = Alignment(horizontal="center")
        # Cao max
        ws.cell(row=r_idx, column=5, value=f"=IFERROR(VALUE(RIGHT('01_MASTER_DATABASE'!O{orig_row_excel}, LEN('01_MASTER_DATABASE'!O{orig_row_excel}) - SEARCH(\"×\",'01_MASTER_DATABASE'!O{orig_row_excel}))), 0)").alignment = Alignment(horizontal="center")
        # Diện tích max
        ws.cell(row=r_idx, column=6, value=f"=IFERROR((D{r_idx}*E{r_idx})/1000000, 0)").alignment = Alignment(horizontal="center")
        
        for c in range(1, 7):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 20
    ws.column_dimensions["F"].width = 22

    # ─────────────────────────────────────────────
    # SHEET 9: 09_SO SÁNH TẢI TRỌNG
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("09_SO SÁNH TẢI TRỌNG")
    write_sheet_title(ws, "TẢI TRỌNG CÁNH CỬA TỐI ĐA (CHỊU LỰC BẢN LỀ / BÁNH XE)", 5)
    write_headers(ws, ["Mã Hệ", "Hãng", "Loại Chịu Lực Bản Lề", "Loại Chịu Lực Bánh Xe", "Tải Trọng Max (kg/cánh)"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        
        # Bản lề
        ws.cell(row=r_idx, column=3, value=f"=IF(ISNUMBER(SEARCH(\"quay\",'01_MASTER_DATABASE'!G{orig_row_excel})), \"Bản lề 3D/4D rãnh C\", \"N/A\")")
        # Bánh xe
        ws.cell(row=r_idx, column=4, value=f"=IF(ISNUMBER(SEARCH(\"lùa\",'01_MASTER_DATABASE'!G{orig_row_excel})), \"Bánh xe trượt chịu lực\", \"N/A\")")
        # Tải trọng
        ws.cell(row=r_idx, column=5, value=f"='01_MASTER_DATABASE'!P{orig_row_excel}").alignment = Alignment(horizontal="center")
        
        for c in range(1, 6):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 25
    ws.column_dimensions["E"].width = 24

    # ─────────────────────────────────────────────
    # SHEET 10: 10_SO SÁNH HIỆU NĂNG
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("10_SO SÁNH HIỆU NĂNG")
    write_sheet_title(ws, "ĐÁNH GIÁ CHỈ SỐ HIỆU NĂNG (CÁCH ÂM, CÁCH NHIỆT, CHỐNG GIÓ, NƯỚC)", 7)
    write_headers(ws, ["Mã Hệ", "Hãng", "Khả Năng Cách Âm", "Khả Năng Cách Nhiệt", "Khả Năng Chống Gió", "Kín Nước", "Kín Khí"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        
        # Cách âm
        ws.cell(row=r_idx, column=3, value=f"=IF(OR(ISNUMBER(SEARCH(\"Thermal\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"cầu\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"uPVC\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"Xuất sắc (>=40dB)\",\"Tốt (>=30dB)\")")
        # Cách nhiệt
        ws.cell(row=r_idx, column=4, value=f"=IF(OR(ISNUMBER(SEARCH(\"Thermal\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"cầu\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"uPVC\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"Xuất sắc (Uf <= 1.6)\",\"Trung bình (Không cầu)\")")
        # Chống gió
        ws.cell(row=r_idx, column=5, value=f"=IF(OR(ISNUMBER(SEARCH(\"Thermal\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"cầu\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"uPVC\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"Class 4 (Đến 2000 Pa)\",\"Class 3 (Đến 1500 Pa)\")")
        # Kín nước
        ws.cell(row=r_idx, column=6, value=f"=IF(OR(ISNUMBER(SEARCH(\"Thermal\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"cầu\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"uPVC\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"Tuyệt đối (Class 9A)\",\"Tốt (Class 7A)\")")
        # Kín khí
        ws.cell(row=r_idx, column=7, value=f"=IF(OR(ISNUMBER(SEARCH(\"Thermal\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"cầu\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"uPVC\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"Tuyệt đối (Class 4)\",\"Tốt (Class 3)\")")
        
        for c in range(1, 8):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    for col in ["C","D","E","F","G"]:
        ws.column_dimensions[col].width = 22

    # ─────────────────────────────────────────────
    # SHEET 11: 11_SO SÁNH PHỤ KIỆN
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("11_SO SÁNH PHỤ KIỆN")
    write_sheet_title(ws, "MẪU RÃNH PROFILE VÀ PHỤ KIỆN TƯƠNG THÍCH", 7)
    write_headers(ws, ["Mã Hệ", "Hãng", "Chuẩn Rãnh Profile", "Thương Hiệu Phụ Kiện Tương Thích", "Loại Khóa Đề Xuất", "Tay Nắm", "Bản Lề Đi Kèm"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        
        # Rãnh
        ws.cell(row=r_idx, column=3, value=f"=IF(ISNUMBER(SEARCH(\"rãnh C\",'01_MASTER_DATABASE'!I{orig_row_excel})), \"Rãnh C Tiêu Chuẩn Châu Âu\", \"Rãnh 22 Truyền Thống\")")
        # Phụ kiện
        ws.cell(row=r_idx, column=4, value=f"='01_MASTER_DATABASE'!Q{orig_row_excel}")
        # Khóa
        ws.cell(row=r_idx, column=5, value=f"=IF(ISNUMBER(SEARCH(\"toilet\",'01_MASTER_DATABASE'!C{orig_row_excel})), \"Khóa đơn điểm\", \"Khóa đa điểm an toàn\")")
        # Tay nắm
        ws.cell(row=r_idx, column=6, value="Đồng bộ theo hãng phụ kiện")
        # Bản lề
        ws.cell(row=r_idx, column=7, value=f"=IF(ISNUMBER(SEARCH(\"bản lề ẩn\",'01_MASTER_DATABASE'!Q{orig_row_excel})), \"Bản lề ẩn cao cấp\", \"Bản lề nổi tiêu chuẩn\")")
        
        for c in range(1, 8):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["D"].width = 30
    ws.column_dimensions["E"].width = 20
    ws.column_dimensions["F"].width = 22
    ws.column_dimensions["G"].width = 22

    # ─────────────────────────────────────────────
    # SHEET 12: 12_SO SÁNH GIOĂNG
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("12_SO SÁNH GIOĂNG")
    write_sheet_title(ws, "HỆ THỐNG VẬT TƯ GIOĂNG VÀ KEO ĐỒNG BỘ", 6)
    write_headers(ws, ["Mã Hệ", "Hãng", "Gioăng Cao Su EPDM", "Gioăng Nhựa Dẻo TPE", "Gioăng Lông (Ray Lùa)", "Keo Ép Góc / Silicon"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        
        # EPDM
        ws.cell(row=r_idx, column=3, value=f"=IF(ISNUMBER(SEARCH(\"EPDM\",'01_MASTER_DATABASE'!R{orig_row_excel})), \"✓\", \"✗\")").alignment = Alignment(horizontal="center")
        # TPE
        ws.cell(row=r_idx, column=4, value=f"=IF(ISNUMBER(SEARCH(\"TPE\",'01_MASTER_DATABASE'!R{orig_row_excel})), \"✓\", \"✗\")").alignment = Alignment(horizontal="center")
        # Gioăng lông
        ws.cell(row=r_idx, column=5, value=f"=IF(ISNUMBER(SEARCH(\"nỉ\",'01_MASTER_DATABASE'!R{orig_row_excel})), \"✓\", \"✗\")").alignment = Alignment(horizontal="center")
        # Keo
        ws.cell(row=r_idx, column=6, value=f"='01_MASTER_DATABASE'!R{orig_row_excel}")
        
        for c in range(1, 7):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 20
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 20
    ws.column_dimensions["F"].width = 32

    # ─────────────────────────────────────────────
    # SHEET 13: 13_SO SÁNH BỀ MẶT
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("13_SO SÁNH BỀ MẶT")
    write_sheet_title(ws, "CÔNG NGHỆ XỬ LÝ BỀ MẶT VÀ HOÀN THIỆN", 6)
    write_headers(ws, ["Mã Hệ", "Hãng", "Sơn Tĩnh Điện", "Mạ Điện Di Anodize ED", "Phủ Hạt Mịn (PVDF)", "Sơn Chống Vân Tay (MED)"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        
        # Sơn
        ws.cell(row=r_idx, column=3, value=f"=IF(ISNUMBER(SEARCH(\"Sơn\",'01_MASTER_DATABASE'!S{orig_row_excel})), \"✓\", \"✗\")").alignment = Alignment(horizontal="center")
        # Anodize
        ws.cell(row=r_idx, column=4, value=f"=IF(ISNUMBER(SEARCH(\"Anodize\",'01_MASTER_DATABASE'!S{orig_row_excel})), \"✓\", \"✗\")").alignment = Alignment(horizontal="center")
        # PVDF
        ws.cell(row=r_idx, column=5, value=f"=IF(ISNUMBER(SEARCH(\"PVDF\",'01_MASTER_DATABASE'!S{orig_row_excel})), \"✓\", \"✗\")").alignment = Alignment(horizontal="center")
        # MED
        ws.cell(row=r_idx, column=6, value=f"=IF(ISNUMBER(SEARCH(\"MED\",'01_MASTER_DATABASE'!S{orig_row_excel})), \"✓\", \"✗\")").alignment = Alignment(horizontal="center")
        
        for c in range(1, 7):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 22
    ws.column_dimensions["E"].width = 22
    ws.column_dimensions["F"].width = 24

    # ─────────────────────────────────────────────
    # SHEET 14: 14_SO SÁNH ỨNG DỤNG
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("14_SO SÁNH ỨNG DỤNG")
    write_sheet_title(ws, "MỨC ĐỘ THÍCH HỢP CHO TỪNG LOẠI CÔNG TRÌNH XÂY DỰNG", 7)
    write_headers(ws, ["Mã Hệ", "Hãng", "Nhà Phố Dân Dụng", "Biệt Thự / Villa", "Resort / Nghỉ Dưỡng", "Chung Cư Cao Tầng", "Văn Phòng / Showroom"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        
        # Nhà phố
        ws.cell(row=r_idx, column=3, value=f"=IF(ISNUMBER(SEARCH(\"nhà phố\",'01_MASTER_DATABASE'!K{orig_row_excel})), \"Tối Ưu\", \"Có Thể Dùng\")").alignment = Alignment(horizontal="center")
        # Biệt thự
        ws.cell(row=r_idx, column=4, value=f"=IF(ISNUMBER(SEARCH(\"biệt thự\",'01_MASTER_DATABASE'!K{orig_row_excel})), \"Tối Ưu\", \"Có Thể Dùng\")").alignment = Alignment(horizontal="center")
        # Resort
        ws.cell(row=r_idx, column=5, value=f"=IF(ISNUMBER(SEARCH(\"nghỉ dưỡng\",'01_MASTER_DATABASE'!K{orig_row_excel})), \"Tối Ưu\", \"Có Thể Dùng\")").alignment = Alignment(horizontal="center")
        # Chung cư
        ws.cell(row=r_idx, column=6, value=f"=IF(ISNUMBER(SEARCH(\"chung cư\",'01_MASTER_DATABASE'!K{orig_row_excel})), \"Tối Ưu\", \"Có Thể Dùng\")").alignment = Alignment(horizontal="center")
        # Văn phòng
        ws.cell(row=r_idx, column=7, value=f"=IF(ISNUMBER(SEARCH(\"văn phòng\",'01_MASTER_DATABASE'!K{orig_row_excel})), \"Tối Ưu\", \"Có Thể Dùng\")").alignment = Alignment(horizontal="center")
        
        for c in range(1, 8):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    for col in ["C","D","E","F","G"]:
        ws.column_dimensions[col].width = 20

    # ─────────────────────────────────────────────
    # SHEET 15: 15_SO SÁNH GIÁ
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("15_SO SÁNH GIÁ")
    write_sheet_title(ws, "SO SÁNH PHÂN KHÚC VÀ ĐƠN GIÁ THÀNH PHẨM ƯỚC TÍNH", 5)
    write_headers(ws, ["Mã Hệ", "Hãng", "Phân Khúc Định Vị", "Giá Vật Tư Ước Lượng (đ/kg)", "Giá Cửa Thành Phẩm Dự Kiến (đ/m²)"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value=f"='01_MASTER_DATABASE'!B{orig_row_excel}")
        
        # Vật tư
        ws.cell(row=r_idx, column=4, value=f"=IF(OR(ISNUMBER(SEARCH(\"Luxury\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"Cầu\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"Cao (>= 120k/kg)\",IF(ISNUMBER(SEARCH(\"Tầm trung\",'01_MASTER_DATABASE'!B{orig_row_excel})),\"Trung bình (85k - 100k/kg)\",\"Kinh tế (<= 85k/kg)\"))").alignment = Alignment(horizontal="center")
        # Thành phẩm
        ws.cell(row=r_idx, column=5, value=f"=IF(OR(ISNUMBER(SEARCH(\"Luxury\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"Cầu\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\">= 5,000,000\",IF(ISNUMBER(SEARCH(\"Tầm trung\",'01_MASTER_DATABASE'!B{orig_row_excel})),\"2,500,000 - 4,000,000\",\"1,500,000 - 2,200,000\"))").alignment = Alignment(horizontal="center")
        
        for c in range(1, 6):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["D"].width = 28
    ws.column_dimensions["E"].width = 30

    # ─────────────────────────────────────────────
    # SHEET 16: 16_SO SÁNH ƯU ĐIỂM
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("16_SO SÁNH ƯU ĐIỂM")
    write_sheet_title(ws, "DANH SÁCH CÁC ƯU ĐIỂM NỔI BẬT CỦA TỪNG HỆ", 3)
    write_headers(ws, ["Mã Hệ", "Hãng", "Ưu Điểm Nổi Bật"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value=f"='01_MASTER_DATABASE'!T{orig_row_excel}")
        
        for c in range(1, 4):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 60

    # ─────────────────────────────────────────────
    # SHEET 17: 17_SO SÁNH NHƯỢC ĐIỂM
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("17_SO SÁNH NHƯỢC ĐIỂM")
    write_sheet_title(ws, "DANH SÁCH CÁC NHƯỢC ĐIỂM / HẠN CHẾ R&D CỦA TỪNG HỆ", 3)
    write_headers(ws, ["Mã Hệ", "Hãng", "Hạn Chế Cần Lưu Ý"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value=f"='01_MASTER_DATABASE'!U{orig_row_excel}")
        
        for c in range(1, 4):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 60

    # ─────────────────────────────────────────────
    # SHEET 18: 18_SO SÁNH TÀI LIỆU
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("18_SO SÁNH TÀI LIỆU")
    write_sheet_title(ws, "TÌNH TRẠNG CÓ SẴN TÀI LIỆU KỸ THUẬT VÀ BẢN VẼ", 6)
    write_headers(ws, ["Mã Hệ", "Hãng", "Bản Vẽ CAD (.dxf)", "Shopdrawing (.dwg)", "Bản Vẽ BIM / Revit", "Catalogue / Manual"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value=f"='01_MASTER_DATABASE'!V{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=4, value=f"='01_MASTER_DATABASE'!W{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=5, value="✗").alignment = Alignment(horizontal="center") # BIM chưa số hóa
        ws.cell(row=r_idx, column=6, value=f"='01_MASTER_DATABASE'!X{orig_row_excel}").alignment = Alignment(horizontal="center")
        
        for c in range(1, 7):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    for col in ["C","D","E","F"]:
        ws.column_dimensions[col].width = 20

    # ─────────────────────────────────────────────
    # SHEET 19: 19_SO SÁNH HÌNH ẢNH
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("19_SO SÁNH HÌNH ẢNH")
    write_sheet_title(ws, "THƯ VIỆN HÌNH ẢNH MINH HỌA MẶT CẮT HỆ NHÔM", 4)
    write_headers(ws, ["Mã Hệ", "Hãng", "Vị Trí Lưu Trữ Ảnh Thumbnail", "Đường Dẫn File Thực Tế"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        ws.cell(row=r_idx, column=3, value="Thư mục 12_Hình ảnh & Video").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=4, value=f"BaoCao_ThuVien_CuaNhom/THƯ VIỆN HỆ NHÔM SAO VÀNG/12_Hình ảnh & Video/Section_{orig_row_excel-1}.png")
        
        for c in range(1, 5):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 50

    # ─────────────────────────────────────────────
    # SHEET 20: 20_ĐÁNH GIÁ SAO VÀNG
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("20_ĐÁNH GIÁ SAO VÀNG")
    write_sheet_title(ws, "BẢNG ĐÁNH GIÁ ĐẶC QUYỀN NỘI BỘ SAO VÀNG GROUP", 7)
    write_headers(ws, ["Mã Hệ", "Hãng", "Chất Lượng Cấu Trúc", "Độ Dễ Gia Công", "Độ Dễ Lắp Dựng", "Chi Phí Bảo Trì", "Khuyến Nghị Tư Vấn"])
    
    for r_idx in range(5, total_rows + 5):
        orig_row_excel = r_idx - 3
        ws.row_dimensions[r_idx].height = 25
        
        ws.cell(row=r_idx, column=1, value=f"='01_MASTER_DATABASE'!C{orig_row_excel}").alignment = Alignment(horizontal="center")
        ws.cell(row=r_idx, column=2, value=f"='01_MASTER_DATABASE'!D{orig_row_excel}")
        
        # Chất lượng
        ws.cell(row=r_idx, column=3, value=f"=IF(OR(ISNUMBER(SEARCH(\"Luxury\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"Cầu\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"⭐⭐⭐⭐⭐\",\"⭐⭐⭐⭐\")").alignment = Alignment(horizontal="center")
        # Gia công
        ws.cell(row=r_idx, column=4, value=f"=IF(ISNUMBER(SEARCH(\"Slim\",'01_MASTER_DATABASE'!B{orig_row_excel})),\"⭐⭐⭐\",\"⭐⭐⭐⭐\")").alignment = Alignment(horizontal="center")
        # Thi công
        ws.cell(row=r_idx, column=5, value="⭐⭐⭐⭐").alignment = Alignment(horizontal="center")
        # Bảo trì
        ws.cell(row=r_idx, column=6, value="⭐⭐⭐⭐").alignment = Alignment(horizontal="center")
        # Khuyến nghị
        ws.cell(row=r_idx, column=7, value=f"=IF(OR(ISNUMBER(SEARCH(\"Luxury\",'01_MASTER_DATABASE'!B{orig_row_excel})),ISNUMBER(SEARCH(\"Cầu\",'01_MASTER_DATABASE'!B{orig_row_excel}))),\"Khuyên Dùng Cho Biệt Thự / Penhouse\",\"Khuyên Dùng Cho Nhà Phố / Căn Hộ\")")
        
        for c in range(1, 8):
            ws.cell(row=r_idx, column=c).border = thin_b
            ws.cell(row=r_idx, column=c).font = Font(name="Arial", size=9.5)
            
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 15
    for col in ["C","D","E","F"]:
        ws.column_dimensions[col].width = 18
    ws.column_dimensions["G"].width = 35

    # ─────────────────────────────────────────────
    # SHEET 21: 21_DASHBOARD
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("21_DASHBOARD")
    ws.views.sheetView[0].showGridLines = True
    
    ws["A1"] = "BẢNG ĐIỀU KHIỂN & BIỂU ĐỒ THỐNG KÊ"
    ws["A1"].font = Font(name="Arial Black", size=16, color="0D2240")
    ws.merge_cells("A1:D2")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    
    stats_list = [
        ("Tổng số hệ trong thư viện", "=COUNTA('01_MASTER_DATABASE'!A2:A500)"),
        ("Số hệ Cầu Cách Nhiệt (Thermal)", "=COUNTIF('01_MASTER_DATABASE'!B2:B500, \"*Cầu cách nhiệt*\") + COUNTIF('01_MASTER_DATABASE'!B2:B500, \"*Thermal*\")"),
        ("Số hệ cửa Slim nội/ngoại thất", "=COUNTIF('01_MASTER_DATABASE'!B2:B500, \"*Slim*\")"),
        ("Số hệ đạt tiêu chuẩn rãnh C châu Âu", "=COUNTIF('01_MASTER_DATABASE'!Q2:Q500, \"*rãnh C*\")"),
        ("Số hệ tương thích kính hộp dày (>=20mm)", "=COUNTIFS('01_MASTER_DATABASE'!N2:N500, \"*20mm*\") + COUNTIFS('01_MASTER_DATABASE'!N2:N500, \"*24mm*\") + COUNTIFS('01_MASTER_DATABASE'!N2:N500, \"*28mm*\") + COUNTIFS('01_MASTER_DATABASE'!N2:N500, \"*32mm*\") + COUNTIFS('01_MASTER_DATABASE'!N2:N500, \"*36mm*\")"),
        ("Số hệ có sẵn bản vẽ CAD", "=COUNTIF('01_MASTER_DATABASE'!V2:V500, \"*✓*\")")
    ]
    
    ws.row_dimensions[1].height = 20
    ws.row_dimensions[2].height = 20
    for idx, (lbl, f_str) in enumerate(stats_list, 4):
        ws.row_dimensions[idx].height = 30
        cell_lbl = ws.cell(row=idx, column=1, value=lbl)
        cell_lbl.font = Font(name="Arial", bold=True, size=11)
        cell_lbl.border = thin_b
        
        cell_val = ws.cell(row=idx, column=2, value=f_str)
        cell_val.font = Font(name="Arial Black", size=11, color="1B5E20")
        cell_val.alignment = Alignment(horizontal="center")
        cell_val.border = thin_b
        
    ws.column_dimensions["A"].width = 40
    ws.column_dimensions["B"].width = 20

    # ─────────────────────────────────────────────
    # SHEET 22: 22_SO SÁNH TRỰC TIẾP THEO MÃ HỆ
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("22_SO SÁNH TRỰC TIẾP THEO MÃ HỆ")
    write_sheet_title(ws, "BẢNG SO SÁNH TRỰC TIẾP side-by-side PHÂN HỆ 55 PHỔ BIẾN", 6)
    
    def find_row_for_code(code_substr, brand_name):
        for idx, row in enumerate(ROWS, 2):
            if brand_name.lower() in row[3].lower() and code_substr.lower() in row[2].lower():
                return idx
        return 2
        
    row_pma = find_row_for_code("55", "PMA")
    row_xingfa = find_row_for_code("55", "Xingfa")
    row_topal = find_row_for_code("Prima", "Topal")
    row_vp = find_row_for_code("55", "Việt Pháp")
    row_civro = find_row_for_code("AW55", "Civro")
    
    hdrs_comp = ["Tiêu Chí So Sánh", "Xingfa Quảng Đông (Hệ 55)", "PMA Vina (Hệ 55)", "Topal Prima (Hệ 55)", "Việt Pháp SHAL (Hệ 55)", "Civro Đức (Hệ AW55)"]
    write_headers(ws, hdrs_comp)
    
    ws.row_dimensions[5].height = 25
    ws.cell(row=5, column=1, value="Độ dày Profile (mm)").font = Font(name="Arial", bold=True)
    ws.cell(row=5, column=2, value=f"='01_MASTER_DATABASE'!L{row_xingfa}")
    ws.cell(row=5, column=3, value=f"='01_MASTER_DATABASE'!L{row_pma}")
    ws.cell(row=5, column=4, value=f"='01_MASTER_DATABASE'!L{row_topal}")
    ws.cell(row=5, column=5, value=f"='01_MASTER_DATABASE'!L{row_vp}")
    ws.cell(row=5, column=6, value=f"='01_MASTER_DATABASE'!L{row_civro}")
    
    ws.row_dimensions[6].height = 25
    ws.cell(row=6, column=1, value="Kính tối đa").font = Font(name="Arial", bold=True)
    ws.cell(row=6, column=2, value=f"='01_MASTER_DATABASE'!N{row_xingfa}")
    ws.cell(row=6, column=3, value=f"='01_MASTER_DATABASE'!N{row_pma}")
    ws.cell(row=6, column=4, value=f"='01_MASTER_DATABASE'!N{row_topal}")
    ws.cell(row=6, column=5, value=f"='01_MASTER_DATABASE'!N{row_vp}")
    ws.cell(row=6, column=6, value=f"='01_MASTER_DATABASE'!N{row_civro}")

    ws.row_dimensions[7].height = 25
    ws.cell(row=7, column=1, value="Hệ Phụ Kiện").font = Font(name="Arial", bold=True)
    ws.cell(row=7, column=2, value=f"='01_MASTER_DATABASE'!Q{row_xingfa}")
    ws.cell(row=7, column=3, value=f"='01_MASTER_DATABASE'!Q{row_pma}")
    ws.cell(row=7, column=4, value=f"='01_MASTER_DATABASE'!Q{row_topal}")
    ws.cell(row=7, column=5, value=f"='01_MASTER_DATABASE'!Q{row_vp}")
    ws.cell(row=7, column=6, value=f"='01_MASTER_DATABASE'!Q{row_civro}")

    ws.row_dimensions[8].height = 25
    ws.cell(row=8, column=1, value="Chỉ Số Cách Âm").font = Font(name="Arial", bold=True)
    ws.cell(row=8, column=2, value="Tốt (32dB)")
    ws.cell(row=8, column=3, value="Trung bình (28dB)")
    ws.cell(row=8, column=4, value="Khá (30dB)")
    ws.cell(row=8, column=5, value="Trung bình (28dB)")
    ws.cell(row=8, column=6, value="Xuất sắc (>=40dB)")

    for r in range(5, 9):
        for c in range(1, 7):
            ws.cell(row=r, column=c).border = thin_b
            ws.cell(row=r, column=c).font = Font(name="Arial", size=9.5)
            if c > 1:
                ws.cell(row=r, column=c).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    ws.column_dimensions["A"].width = 22
    for col in ["B","C","D","E","F"]:
        ws.column_dimensions[col].width = 24

    # ─────────────────────────────────────────────
    # SHEET 23: 23_MA TRẬN HỆ NHÔM
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("23_MA TRẬN HỆ NHÔM")
    write_sheet_title(ws, "MA TRẬN HỆ CỬA PHÂN BỔ THEO THƯƠNG HIỆU VÀ KÍCH THƯỚC BẢN RỘNG (55, 60, 65...)", 15)
    
    matrix_hdrs = ["Hãng", "Hệ 55", "Hệ 58", "Hệ 60", "Hệ 65", "Hệ 70", "Hệ 75", "Hệ 80", "Hệ 83", "Hệ 93", "Hệ 95", "Hệ 115", "Hệ 120", "Phân Hệ Slim", "Phân Hệ Thermal"]
    write_headers(ws, matrix_hdrs)
    
    for idx, brand in enumerate(unique_brands, 5):
        ws.row_dimensions[idx].height = 25
        ws.cell(row=idx, column=1, value=brand).font = Font(name="Arial", bold=True, size=10)
        ws.cell(row=idx, column=1).border = thin_b
        
        formulas = [
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*55*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*58*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*60*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*65*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*70*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*75*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*80*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*83*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*93*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*95*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*115*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$C$2:$C$500, \"*120*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Slim*\")>0, \"✓\", \"✗\")",
            f"=IF(COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*Thermal*\") + COUNTIFS('01_MASTER_DATABASE'!$D$2:$D$500, A{idx}, '01_MASTER_DATABASE'!$B$2:$B$500, \"*cầu*\")>0, \"✓\", \"✗\")",
        ]
        
        for c_idx, f_str in enumerate(formulas, 2):
            cell = ws.cell(row=idx, column=c_idx, value=f_str)
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_b
            
    ws.column_dimensions["A"].width = 20
    for col in [get_column_letter(c) for c in range(2, 16)]:
        ws.column_dimensions[col].width = 12

    # ─────────────────────────────────────────────
    # SHEET 24: 24_MA TRẬN HÃNG
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("24_MA TRẬN HÃNG")
    write_sheet_title(ws, "MA TRẬN SO SÁNH ĐẶC TÍNH CỐT LÕI GIỮA CÁC THƯƠNG HIỆU HÃNG", 7)
    
    hdrs_matrix = ["Tiêu Chí Thương Hiệu", "Xingfa (Quảng Đông)", "PMA (Việt Nam)", "Maxpro JP (Nhật Bản)", "Civro (Đức)", "Topal (Việt Nam)", "Eurowindow (Việt Nam)"]
    write_headers(ws, hdrs_matrix)
    
    matrix_rows = [
        ("Nguồn gốc xuất xứ chính", "Quảng Đông (Trung Quốc)", "Việt Nam sản xuất", "Công nghệ Nhật Bản/VN", "Đức nhập khẩu", "Việt Nam (Khang Minh/Topal)", "Đức/Việt Nam (Eurowindow)"),
        ("Công nghệ sơn phủ", "Sơn tĩnh điện AkzoNobel", "Sơn tĩnh điện AkzoNobel", "Mạ điện di Anodize ED", "Sơn tĩnh điện phủ MED Đức", "Sơn tĩnh điện AkzoNobel", "Sơn tĩnh điện bảo hành 20 năm"),
        ("Chuẩn rãnh phụ kiện", "Rãnh 22 truyền thống", "Rãnh 22 truyền thống", "Rãnh C Châu Âu / Rãnh 22", "Rãnh C Châu Âu cao cấp", "Rãnh 22 truyền thống", "Rãnh C Châu Âu (nhôm mới)"),
        ("Phân khúc giá hoàn thiện", "Từ 2.2tr - 3.2tr/m²", "Từ 1.6tr - 2.5tr/m²", "Từ 3.5tr - 5.0tr/m²", "Từ 6.0tr - 12.0tr/m²", "Từ 1.8tr - 2.8tr/m²", "Từ 3.0tr - 6.5tr/m²"),
        ("Thế mạnh lớn nhất", "Thương hiệu quốc dân, uy tín cao", "Giá bình dân, đại chúng dễ tiếp cận", "Mạ điện di chống muối biển 40 năm", "Cách âm nhiệt siêu việt cao cấp nhất", "Hệ thống đại lý phân phối lớn", "Cách âm cách nhiệt hàng đầu, bảo hành tốt"),
        ("Điểm yếu R&D", "Nhiều hàng giả hàng nhái trên thị trường", "Cấu trúc định hình mỏng không chịu bão lớn", "Giá phôi anodize khá đắt đỏ", "Giá cực kỳ cao, kén khách hàng đại chúng", "Độ nhận diện thương hiệu tầm trung", "Khung bao thô dày (nhựa), thi công đắt")
    ]
    
    for r_idx, row_vals in enumerate(matrix_rows, 5):
        ws.row_dimensions[r_idx].height = 35
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx == 1:
                cell.font = Font(name="Arial", bold=True, size=9.5)
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                
    ws.column_dimensions["A"].width = 24
    for col in ["B","C","D","E","F","G"]:
        ws.column_dimensions[col].width = 24

    # ─────────────────────────────────────────────
    # SHEET 25: 25_DANH MỤC PHÂN HỆ NỔI BẬT
    # ─────────────────────────────────────────────
    ws = wb.create_sheet("25_DANH MỤC PHÂN HỆ NỔI BẬT")
    ws.views.sheetView[0].showGridLines = True
    
    # Tiêu đề lớn của sheet
    ws.merge_cells("A1:I2")
    title_cell = ws.cell(row=1, column=1, value="DANH MỤC PHÂN LOẠI CÁC PHÂN HỆ NỔI BẬT (R&D SAO VÀNG)")
    title_cell.font = Font(name="Arial", size=14, bold=True, color="FFFFFF")
    title_cell.fill = PatternFill(fill_type="solid", fgColor="0F172A") # Màu Slate sậm
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 20
    ws.row_dimensions[2].height = 20
    
    # Định nghĩa 18 phân hệ nổi bật
    categories_def = [
        {
            "title": "I. DANH MỤC HỆ CỬA NHÔM CẦU CÁCH NHIỆT (THERMAL BREAK & PASSIVE HOUSE)",
            "color": "1E4620", # Xanh lá đậm
            "filter_fn": lambda r: "cầu" in r[1].lower() or "thermal" in r[1].lower() or "passive" in r[1].lower()
        },
        {
            "title": "II. DANH MỤC HỆ CỬA SLIM SIÊU MỎNG (NỘI THẤT & NGOẠI THẤT)",
            "color": "9D174D", # Hồng đỏ đậm
            "filter_fn": lambda r: "slim" in r[1].lower()
        },
        {
            "title": "III. DANH MỤC CỬA CHÍNH BIỆT THỰ / LUXURY / HỆ THỦY LỰC / BẢN LỀ ẨN",
            "color": "9A3412", # Cam nâu đậm
            "filter_fn": lambda r: any(k in r[1].lower() for k in ["thủy lực", "luxury", "tân cổ điển"]) or any(k in r[2].lower() for k in ["ad75", "vasona", "bản lề ẩn"])
        },
        {
            "title": "IV. DANH MỤC CỬA SIÊU TRƯỜNG SIÊU TRỌNG & NÂNG TRƯỢT (LIFT & SLIDE / HEAVY DUTY)",
            "color": "5B21B6", # Tím đậm
            "filter_fn": lambda r: any(k in r[1].lower() for k in ["lift", "lùa siêu trường", "nâng trượt", "xếp trượt", "bi-fold", "bifold"]) or any(k in r[2].lower() for k in ["ls100", "ls180"])
        },
        {
            "title": "V. DANH MỤC HỆ CỬA MỞ QUAY PHỔ THÔNG (CASEMENT)",
            "color": "0F766E", # Xanh Teal
            "filter_fn": lambda r: not ("cầu" in r[1].lower() or "thermal" in r[1].lower() or "slim" in r[1].lower() or "thủy lực" in r[1].lower() or "luxury" in r[1].lower() or "lift" in r[1].lower() or "xếp trượt" in r[1].lower()) and ("quay" in r[7].lower() or "quay" in r[6].lower() or "casement" in r[1].lower() or "quay" in r[2].lower())
        },
        {
            "title": "VI. DANH MỤC HỆ CỬA TRƯỢT PHỔ THÔNG (SLIDING)",
            "color": "1D4ED8", # Xanh dương
            "filter_fn": lambda r: not ("cầu" in r[1].lower() or "thermal" in r[1].lower() or "slim" in r[1].lower() or "thủy lực" in r[1].lower() or "luxury" in r[1].lower() or "lift" in r[1].lower() or "xếp trượt" in r[1].lower()) and ("lùa" in r[7].lower() or "trượt" in r[7].lower() or "lùa" in r[6].lower() or "trượt" in r[6].lower() or "sliding" in r[1].lower() or "lùa" in r[2].lower())
        },
        {
            "title": "VII. DANH MỤC HỆ CỬA SỔ",
            "color": "4338CA", # Chàm (Indigo)
            "filter_fn": lambda r: "cửa sổ" in r[6].lower() or "cửa sổ" in r[8].lower() or "sổ" in r[6].lower()
        },
        {
            "title": "VIII. DANH MỤC HỆ PIVOT (CỬA TRỤC XOAY TRUNG TÂM)",
            "color": "78350F", # Nâu đất
            "filter_fn": lambda r: "pivot" in r[1].lower() or "pivot" in r[7].lower() or "pivot" in r[2].lower() or "trục xoay" in r[1].lower() or "trục xoay" in r[7].lower()
        },
        {
            "title": "IX. DANH MỤC HỆ TRƯỢT QUAY (SLIDE & TURN)",
            "color": "B45309", # Cam đất
            "filter_fn": lambda r: "trượt quay" in r[1].lower() or "trượt quay" in r[7].lower() or "trượt quay" in r[2].lower()
        },
        {
            "title": "X. DANH MỤC VÁCH KÍNH - CURTAIN WALL (MẶT DỰNG)",
            "color": "047857", # Xanh lá đậm
            "filter_fn": lambda r: "mặt dựng" in r[1].lower() or "mặt dựng" in r[6].lower() or "curtain wall" in r[1].lower() or "mặt dựng" in r[8].lower()
        },
        {
            "title": "XI. DANH MỤC VÁCH NGĂN NỘI THẤT",
            "color": "475569", # Slate xám
            "filter_fn": lambda r: "vách ngăn" in r[1].lower() or "vách ngăn" in r[6].lower() or "vách ngăn" in r[8].lower() or ("nội thất" in r[1].lower() and "vách" in r[6].lower())
        },
        {
            "title": "XII. DANH MỤC MÁI KÍNH",
            "color": "0369A1", # Xanh biển
            "filter_fn": lambda r: "mái kính" in r[1].lower() or "mái kính" in r[6].lower() or "mái kính" in r[8].lower()
        },
        {
            "title": "XIII. DANH MỤC LAN CAN",
            "color": "BE123C", # Hồng sậm
            "filter_fn": lambda r: "lan can" in r[1].lower() or "lan can" in r[6].lower() or "lan can" in r[8].lower()
        },
        {
            "title": "XIV. DANH MỤC LAM NHÔM",
            "color": "4D7C0F", # Olive
            "filter_fn": lambda r: "lam" in r[1].lower() or "lam" in r[6].lower() or "lam" in r[8].lower()
        },
        {
            "title": "XV. DANH MỤC CỬA CHUYÊN DỤNG (CHỐNG CHÁY, PHÒNG SẠCH, NHỰA ABS...)",
            "color": "991B1B", # Crimson đỏ
            "filter_fn": lambda r: "chuyên dụng" in r[1].lower() or "chống cháy" in r[1].lower() or "chống cháy" in r[6].lower() or "chống cháy" in r[8].lower() or "chuyên dụng" in r[8].lower() or "abs" in r[1].lower() or "abs" in r[2].lower()
        },
        {
            "title": "XVI. DANH MỤC HỆ NỘI THẤT NHÔM (TỦ BẾP, CÁNH KÍNH...)",
            "color": "6D28D9", # Tím Violet
            "filter_fn": lambda r: "tủ" in r[1].lower() or "tủ" in r[6].lower() or "tủ" in r[8].lower() or "nội thất nhôm" in r[1].lower() or "nội thất nhôm" in r[8].lower() or "cánh kính" in r[1].lower() or "cánh kính" in r[8].lower()
        },
        {
            "title": "XVII. DANH MỤC PROFILE TRANG TRÍ (NẸP NHÔM, PHÀO CHỈ...)",
            "color": "3F3F46", # Zinc xám
            "filter_fn": lambda r: "nẹp" in r[1].lower() or "nẹp" in r[8].lower() or "trang trí" in r[1].lower() or "trang trí" in r[8].lower() or "phào" in r[1].lower() or "phào" in r[8].lower()
        },
        {
            "title": "XVIII. DANH MỤC HỆ CỬA TỰ ĐỘNG (AUTOMATIC DOORS)",
            "color": "A21CAF", # Magenta
            "filter_fn": lambda r: "tự động" in r[1].lower() or "tự động" in r[7].lower() or "tự động" in r[8].lower() or "automatic" in r[1].lower() or "automatic" in r[7].lower() or "automatic" in r[8].lower()
        }
    ]
    
    current_row = 4
    for cat in categories_def:
        # Dòng tiêu đề phân mục
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row+1, end_column=9)
        sec_cell = ws.cell(row=current_row, column=1, value=cat["title"])
        sec_cell.font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
        sec_cell.fill = PatternFill(fill_type="solid", fgColor=cat["color"])
        sec_cell.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[current_row].height = 20
        ws.row_dimensions[current_row+1].height = 20
        current_row += 2
        
        # Dòng Header của bảng phân mục
        hdrs_sub = ["STT", "Hãng", "Mã Hệ", "Tên Hệ", "Phân Nhóm Hệ", "Độ Dày (mm)", "Loại Sản Phẩm", "Kiểu Mở", "Đặc Điểm Kỹ Thuật"]
        ws.row_dimensions[current_row].height = 25
        for col_idx, h in enumerate(hdrs_sub, 1):
            cell = ws.cell(row=current_row, column=col_idx, value=h)
            cell.font = Font(name="Arial", size=9, bold=True, color="333333")
            cell.fill = PatternFill(fill_type="solid", fgColor="E2E8F0") # Xám sáng nhạt
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_b
        current_row += 1
        
        # Lọc và điền các hệ nhôm phù hợp
        matching_rows = [r for r in ROWS if cat["filter_fn"](r)]
        if not matching_rows:
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=9)
            cell = ws.cell(row=current_row, column=1, value="Chưa có dữ liệu phân loại cho mục này")
            cell.font = Font(name="Arial", italic=True, size=9.5, color="888888")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_b
            ws.row_dimensions[current_row].height = 25
            current_row += 1
        else:
            for s_idx, r in enumerate(matching_rows, 1):
                ws.row_dimensions[current_row].height = 25
                
                # STT
                cell = ws.cell(row=current_row, column=1, value=s_idx)
                cell.alignment = Alignment(horizontal="center")
                # Hãng
                ws.cell(row=current_row, column=2, value=r[3])
                # Mã Hệ
                ws.cell(row=current_row, column=3, value=r[2]).alignment = Alignment(horizontal="center")
                # Tên Hệ
                ws.cell(row=current_row, column=4, value=r[8])
                # Nhóm Hệ
                ws.cell(row=current_row, column=5, value=r[1])
                # Độ Dày
                cell_thick = ws.cell(row=current_row, column=6, value=r[11])
                cell_thick.alignment = Alignment(horizontal="center")
                cell_thick.font = Font(name="Arial", bold=True)
                # Loại Cửa
                ws.cell(row=current_row, column=7, value=r[6])
                # Kiểu Mở
                ws.cell(row=current_row, column=8, value=r[7])
                # Đặc Điểm Kỹ Thuật
                ws.cell(row=current_row, column=9, value=r[10])
                
                for col_idx in range(1, 10):
                    c_cell = ws.cell(row=current_row, column=col_idx)
                    c_cell.font = Font(name="Arial", size=9)
                    c_cell.border = thin_b
                    if col_idx not in [1, 3, 6]:
                        c_cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
                        
                current_row += 1
                
        # Dòng trống ngăn cách
        current_row += 1
        
    # Cấu hình chiều rộng cột
    col_widths_sub = [8, 16, 16, 25, 20, 14, 25, 25, 45]
    for i, w in enumerate(col_widths_sub, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # ─────────────────────────────────────────────
    # LƯU FILE KẾT QUẢ
    # ─────────────────────────────────────────────
    path = os.path.join(ROOT, "02_Cơ sở Dữ liệu Gốc.xlsx")
    safe_save_excel(wb, path)
    print(f"  >> Saved Master Excel with 25 Sheets to: {path}")

# ─────────────────────────────────────────────
# TẠO FILE 3: 03_Danh mục Catalogue theo Hãng.xlsx
# ─────────────────────────────────────────────
def create_catalogue_index():
    wb = Workbook()
    ws = wb.active
    ws.title = "Danh mục Catalogue"

    headers = ["STT","Hãng","Tên Catalogue / Tài liệu","Loại tài liệu",
               "Phiên bản / Năm","Ngôn ngữ","Tình trạng",
               "Link PDF (Trực tuyến)","Link File Nội bộ","Ghi chú"]
    hdr_fill = PatternFill(fill_type="solid", fgColor="705D30")
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")

    ws.row_dimensions[1].height = 40
    for i,h in enumerate(headers,1):
        c = ws.cell(row=1, column=i, value=h)
        c.fill = hdr_fill
        c.font = hdr_font
        c.alignment = Alignment(horizontal="center",vertical="center",wrap_text=True)
        c.border = thin_b

    cat_rows = [
        (1,"Viralwindow","Catalogue VRE65 Hệ Châu Âu","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://viralwindow.com","",""),
        (2,"Viralwindow","Catalogue VRA55/94 Hệ Châu Á","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://viralwindow.com","",""),
        (3,"Civro","Catalogue AW/AD 55/65/75","Catalogue kỹ thuật + Profile","2025","Tiếng Việt","Có","https://civrowindow.com","",""),
        (4,"Civro","Hướng dẫn lắp đặt Civro","Tài liệu lắp đặt","2025","Tiếng Việt","Có","https://civrowindow.com","",""),
        (5,"Civro","Catalogue CSD Lift Slide","Catalogue kỹ thuật","2025","Tiếng Việt","Có","https://civrowindow.com","",""),
        (6,"Maxpro JP","Catalogue R55/65/70/83","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-maxpro-7.html","",""),
        (7,"Maxpro JP","Catalogue SD83/SD115/SFD80","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-maxpro-7.html","",""),
        (8,"Maxpro JP","Hướng dẫn gia công Maxpro JP","Tài liệu gia công","2024","Tiếng Việt","Có","https://maxprojp.com.vn","",""),
        (9,"Soco","Catalogue Soco 65/94/120/180","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomsocowindow.com","",""),
        (10,"Soco","Catalogue Soco LS100 Lift Slide","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomsocowindow.com","",""),
        (11,"PAG","Catalogue PAG 60/65/80/83/120/200","Catalogue kỹ thuật","2025","Tiếng Việt","Có","https://nhompag.com.vn","",""),
        (12,"PAG","Catalogue PAG Seamless Welded","Catalogue kỹ thuật","2025","Tiếng Việt","Có","https://nhompag.com.vn","",""),
        (13,"Owin","Catalogue Owin HL180 Thủy lực","Catalogue kỹ thuật","2024","Tiếng Việt","Có","","",""),
        (14,"Owin","Catalogue Owin Trượt quay","Catalogue kỹ thuật","2024","Tiếng Việt","Có","","",""),
        (15,"PMA","Catalogue PMA 55/Classic/Platinum","Catalogue kỹ thuật","2025","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma.html","",""),
        (16,"Topal","Catalogue Topal Prima / Slima","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://topal.vn","",""),
        (17,"Việt Pháp","Catalogue Việt Pháp SHAL","Catalogue kỹ thuật","2021","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-viet-phap-shal-18.html","",""),
        (18,"EuroVN","Catalogue EuroVN Gold/VIP","Catalogue kỹ thuật","2024","Tiếng Việt","Có","","",""),
        (19,"EuroVN","Catalogue EuroVN Thủy lực","Catalogue kỹ thuật","2024","Tiếng Việt","Có","","",""),
        (20,"Yangli","Catalogue Yangli 55/65/Lùa/Slim","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://yangli.vn","",""),
        (21,"Việt Ý","Catalogue Nhôm Việt Ý Italumi","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pma-9.html","",""),
        (22,"Xingfa","Catalogue Xingfa Quảng Đông Nhập Khẩu","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-nhap-khau-quang-dong-1.html","",""),
        (23,"Xingfa Việt Nam","Catalogue Nhôm Xingfa Việt Nam","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-xingfa-viet-nam-10.html","",""),
        (24,"PMI","Catalogue Nhôm PMI nhập khẩu Malaysia","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-pmi-16.html","",""),
        (25,"YKK AP","Catalogue YKK WS55/SL90","Catalogue kỹ thuật","2024","Tiếng Anh/Nhật","Có","https://www.ykkap.co.jp","",""),
        (26,"Hondalex","Catalogue Nhôm Hondalex LongVanGroup","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-hondalex-longvangroup-50.html","",""),
        (27,"KOSO","Catalogue Cửa Nhựa ABS KOS Hàn Quốc","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-cua-nhua-abs-kos-48.html","",""),
        (28,"Eurowindow","Catalogue Cửa Nhựa uPVC Eurowindow","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-cua-nhua-loi-thep-eurowindow-36.html","",""),
        (29,"Eurowindow","Catalogue Cửa Nhôm Eurowindow","Catalogue kỹ thuật","2024","Tiếng Việt","Có","https://nhomkinhdaiphuc.com/catalogue/catalogue-nhom-eurowindow-28.html","",""),
        (30,"Schüco","Technical Manual AWS 65 / AWS 75.SI","Technical manual Đức","2025","Tiếng Anh/Đức","Cần mua","https://www.schueco.com","","Cần đăng ký tài khoản hệ thống Schüco"),
        (31,"Schüco","Technical Manual FW 50+ Curtain Wall","Technical manual Đức","2025","Tiếng Anh","Cần mua","https://www.schueco.com","",""),
        (32,"Reynaers","Technical Manual CS 60 / CS 77","Technical manual Bỉ","2025","Tiếng Anh","Cần mua","https://www.reynaers.com","",""),
        (33,"Reynaers","Technical Manual CP 155-LS","Technical manual Bỉ","2025","Tiếng Anh","Cần mua","https://www.reynaers.com","",""),
        (34,"Roto (Phụ kiện)","Catalogue Roto Roto NT / Patio Life / Fold","Catalogue phụ kiện","2025","Tiếng Anh","Có","https://www.roto-frank.com","",""),
        (35,"Siegenia (Phụ kiện)","Catalogue Siegenia Titan LSS / Aeromat","Catalogue phụ kiện","2025","Tiếng Anh","Có","https://www.siegenia.com","",""),
        (36,"Dorma (Phụ kiện)","Catalogue Dorma TS98 Floor Spring / Slido","Catalogue phụ kiện","2025","Tiếng Anh","Có","https://www.dormakaba.com","",""),
        (37,"CMECH (Phụ kiện)","Catalogue CMECH 3D Hinges","Catalogue phụ kiện","2024","Tiếng Anh","Có","https://www.cmech.com","",""),
        (38,"Häfele (Phụ kiện)","Catalogue Häfele Sliding Pocket Slido","Catalogue phụ kiện","2025","Tiếng Anh/Việt","Có","https://www.hafele.com.vn","",""),
        (39,"Hettich (Phụ kiện)","Catalogue Hettich Sliding Systems","Catalogue phụ kiện","2025","Tiếng Anh","Có","https://www.hettich.com","",""),
    ]

    row_colors = ["F5F5F5","FFFFFF"]
    for r_idx, row in enumerate(cat_rows, 2):
        fill_c = PatternFill(fill_type="solid", fgColor=row_colors[r_idx%2])
        ws.row_dimensions[r_idx].height = 30
        for c_idx, val in enumerate(row, 1):
            c = ws.cell(row=r_idx, column=c_idx, value=val)
            c.fill = fill_c
            c.font = Font(name="Arial", size=9)
            c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            c.border = thin_b

    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"
    ws.freeze_panes = "A2"
    col_ws = [6,20,40,22,14,16,16,35,35,30]
    for i,w in enumerate(col_ws,1):
        ws.column_dimensions[get_column_letter(i)].width = w

    path = os.path.join(ROOT, "03_Danh mục Catalogue theo Hãng.xlsx")
    safe_save_excel(wb, path)

# ─────────────────────────────────────────────
# TẠO FILE 4: 04_Bảng theo dõi Tiến độ.xlsx
# ─────────────────────────────────────────────
def create_tracker():
    wb = Workbook()
    ws = wb.active
    ws.title = "Bảng theo dõi Tiến độ"

    tracker_hdrs = [
        "STT","Mã Hệ","Hãng","Nhóm phân khúc",
        "Thu thập Catalogue","Thu thập CAD","Thu thập Shopdrawing",
        "Thu thập HD Lắp đặt","Phân tích kỹ thuật","Nhập Database",
        "Đánh giá Ưu/Nhược","Viết báo cáo kỹ thuật",
        "Mức độ hoàn thiện (%)","Người phụ trách","Hạn hoàn thành","Ghi chú / Vướng mắc"
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="4B5563")
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    ws.row_dimensions[1].height = 42

    for i,h in enumerate(tracker_hdrs, 1):
        c = ws.cell(row=1, column=i, value=h)
        c.fill = hdr_fill
        c.font = hdr_font
        c.alignment = Alignment(horizontal="center",vertical="center",wrap_text=True)
        c.border = thin_b

    STATUS_FILL = {
        "✓ Hoàn thành": PatternFill(fill_type="solid", fgColor="C8E6C9"),
        "⚡ Đang làm":   PatternFill(fill_type="solid", fgColor="FFF9C4"),
        "✗ Chưa làm":   PatternFill(fill_type="solid", fgColor="FFEBEE"),
    }

    for r_idx, row in enumerate(ROWS, 2):
        stt = r_idx - 1
        ma_he = row[2]
        hang  = row[3]
        nhom  = row[1]
        cat   = "✓ Hoàn thành" if row[23]=="✓" else "✗ Chưa làm"
        cad   = "✓ Hoàn thành" if row[21]=="✓" else "✗ Chưa làm"
        sd    = "✓ Hoàn thành" if row[22]=="✓" else "✗ Chưa làm"
        hd    = "✓ Hoàn thành" if row[24]=="✓" else "✗ Chưa làm"
        pct   = row[26]
        tinh_trang = row[25]

        pta  = "✓ Hoàn thành" if tinh_trang=="Hoàn thành" else ("⚡ Đang làm" if tinh_trang=="Đang nghiên cứu" else "✗ Chưa làm")
        nhap = "✓ Hoàn thành" if pct >= 70 else ("⚡ Đang làm" if pct >= 50 else "✗ Chưa làm")
        dg   = "✓ Hoàn thành" if pct >= 80 else "✗ Chưa làm"
        viet = "✓ Hoàn thành" if pct >= 90 else "✗ Chưa làm"

        tracker_row = [
            stt, ma_he, hang, nhom,
            cat, cad, sd, hd, pta, nhap, dg, viet,
            f"{pct}%", "Phòng R&D Sao Vàng", "2026-Q4",
            "" if tinh_trang=="Hoàn thành" else f"Cần thu thập thêm ({tinh_trang})"
        ]

        row_fill = PatternFill(fill_type="solid",
            fgColor="EDF7ED" if tinh_trang=="Hoàn thành" else "FFFDE7" if tinh_trang=="Đang nghiên cứu" else "FDE8E8")
        ws.row_dimensions[r_idx].height = 28

        for c_idx, val in enumerate(tracker_row, 1):
            c = ws.cell(row=r_idx, column=c_idx, value=val)
            c.font   = Font(name="Arial", size=9)
            c.border = thin_b
            # Status cells color
            if c_idx in (5,6,7,8,9,10,11,12) and val in STATUS_FILL:
                c.fill = STATUS_FILL[val]
                c.alignment = Alignment(horizontal="center", vertical="center")
            else:
                c.fill = row_fill
                c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws.auto_filter.ref = f"A1:{get_column_letter(len(tracker_hdrs))}1"
    ws.freeze_panes = "A2"
    tw = [6,20,18,22,16,16,18,18,18,16,16,20,16,22,16,35]
    for i,w in enumerate(tw[:len(tracker_hdrs)], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    path = os.path.join(ROOT, "04_Bảng theo dõi Tiến độ.xlsx")
    safe_save_excel(wb, path)

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("THU VIEN HE NHOM SAO VANG v3.0 – Building 4 Master Docs")
    print("=" * 60)
    print(f"  Total systems in database: {len(ROWS)}")
    print()

    print("[1/4] Creating 01_So tay Ky thuat Tong hop.docx ...")
    create_handbook()

    print("[2/4] Creating 02_Co so Du lieu Goc.xlsx ...")
    create_database()

    print("[3/4] Creating 03_Danh muc Catalogue theo Hang.xlsx ...")
    create_catalogue_index()

    print("[4/4] Creating 04_Bang theo doi Tien do.xlsx ...")
    create_tracker()

    print()
    print("=" * 60)
    print("ALL 4 MASTER DOCUMENTS GENERATED SUCCESSFULLY!")
    print(f"Output folder: {ROOT}")
    print("=" * 60)
