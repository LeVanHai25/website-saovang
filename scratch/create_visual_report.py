# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo Báo cáo Tóm tắt Trực quan (BaoCao_TomTat_TrucQuan.docx) kèm hình ảnh minh họa
Phiên bản 1.0
"""
import os, sys, shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

def safe_save_docx(doc, path):
    try:
        doc.save(path)
        print(f"  >> Saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_DANHGIA_MOI.docx")
        doc.save(alt_path)
        print(f"  [!] CANH BAO: Tep '{os.path.basename(path)}' dang duoc mo trong Word.")
        print(f"  >> Da luu tam sang: {alt_path}")

# Đường dẫn
ARTIFACT_DIR = r"C:\Users\Admin\.gemini\antigravity\brain\5a84d0ab-58a2-4873-b63f-009ddcae0916"
BASE  = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"
ROOT  = os.path.join(BASE, "THƯ VIỆN HỆ NHÔM SAO VÀNG")
MEDIA_DIR = os.path.join(ROOT, "12_Hình ảnh & Video")

# Copy ảnh từ artifact dir vào thư mục media của thư viện
images_src = {
    "cua_slim.png": "cua_slim_mockup_1782967355773.png",
    "cau_cach_nhiet.png": "cau_cach_nhiet_profile_1782967368909.png",
    "cua_thuy_luc.png": "cua_thuy_luc_luxury_1782967386097.png",
    "cua_xep_truot.png": "cua_xep_truot_bifold_1782967396462.png"
}

images_dest = {}
for alias, name in images_src.items():
    src_path = os.path.join(ARTIFACT_DIR, name)
    dest_path = os.path.join(MEDIA_DIR, alias)
    if os.path.exists(src_path):
        shutil.copy2(src_path, dest_path)
        images_dest[alias] = dest_path
        print(f"  [+] Copied image to: {dest_path}")
    else:
        print(f"  [!] Warning: Source image not found at {src_path}")

# Khởi tạo Document
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Định nghĩa các hàm helper tạo kiểu chữ
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x0D,0x22,0x40) # Navy
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.name = "Arial"
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x70,0x5D,0x30) # Vàng Antique
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.name = "Arial"
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def para(text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Arial"
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p

# ── TRANG BÌA ──────────────────────────────────────────
for _ in range(3): doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BÁO CÁO TÓM TẮT TRỰC QUAN")
run.font.name = "Arial Black"
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x0D,0x22,0x40)
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("HỆ THỐNG HỆ CỬA NHÔM KÍNH TRÊN THỊ TRƯỜNG\nPhục vụ Tư vấn, Bán hàng & Thiết kế nhanh")
run2.font.name = "Arial"
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x70,0x5D,0x30)
run2.bold = True

for _ in range(4): doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_info = info.add_run(
    "THƯ VIỆN HỆ NHÔM SAO VÀNG\n"
    "Phòng Nghiên cứu & Phát triển sản phẩm (R&D) – Sao Vàng Group\n"
    "Phiên bản 1.0  |  Tháng 7 năm 2026"
)
run_info.font.name = "Arial"
run_info.font.size = Pt(11)
run_info.italic = True

doc.add_page_break()

# ── LỜI MỞ ĐẦU ──────────────────────────────────────────
h1("LỜI MỞ ĐẦU")
para(
    "Báo cáo này được Phòng R&D Sao Vàng biên soạn nhằm cung cấp một tài liệu trực quan, dễ đọc, "
    "giúp Ban Giám đốc, bộ phận Kinh doanh và khách hàng có thể nhanh chóng nắm bắt bản chất kỹ thuật "
    "của các hệ cửa nhôm kính cốt lõi đang thịnh hành trên thị trường. Thay vì tra cứu hàng chục cột thông tin "
    "trong bảng Excel phẳng phức tạp, tài liệu này chắt lọc cấu trúc cốt lõi nhất kèm hình ảnh minh họa sinh động."
)

# ── BẢN ĐỒ PHÂN KHÚC HỆ NHÔM ──────────────────────────────
h1("PHẦN I: BẢN ĐỒ PHÂN KHÚC HỆ NHÔM THỊ TRƯỜNG")
para(
    "Thị trường cửa nhôm kính tại Việt Nam hiện nay có thể chia thành 4 phân khúc chính dựa trên "
    "nguồn gốc xuất xứ, độ dày profile nhôm, loại phụ kiện và công nghệ sơn phủ bề mặt:"
)

# Tạo bảng phân khúc
tbl_seg = doc.add_table(rows=1, cols=4)
tbl_seg.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_seg.style = "Table Grid"
hdr_cells = tbl_seg.rows[0].cells
hdr_cells[0].text = "Phân khúc"
hdr_cells[1].text = "Thương hiệu đại diện"
hdr_cells[2].text = "Đặc tính kỹ thuật chính"
hdr_cells[3].text = "Giá trị sử dụng tối ưu"
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

segments_data = [
    ("Phổ thông\n(Economy)", "Yangli, Xingfa cỏ, Topal XFEC, EuroVN", "Nhôm dày 1.0-1.4mm, sơn bột thường, phụ kiện rãnh 22 đại trà.", "Nhà cấp 4, nhà trọ, công trình ngân sách hạn chế."),
    ("Trung cấp\n(Mid-range)", "Viralwindow VRA, Topal XFAD, Soco 65, Maxpro R55", "Nhôm dày 1.4-1.6mm, sơn tĩnh điện bảo hành 5-10 năm, rãnh C Châu Âu.", "Nhà phố phổ thông, chung cư thương mại, văn phòng."),
    ("Cao cấp\n(Premium)", "Civro AW/AD, Maxpro R65/83, PAG 60/80, VRE65, Kogen", "Nhôm dày 1.6-2.0mm, Anodized ED bền màu 40 năm, có cầu cách nhiệt PA66 hoặc ẩn bản lề độc đáo.", "Biệt thự cao cấp, villa biển, căn hộ hạng sang, penthouse."),
    ("Siêu cao cấp\n(Luxury)", "Schüco, Reynaers, Technal, Owin HL180, Soco LS100", "Nhôm dày 2.0-3.0mm, cầu cách nhiệt Technoform đa khoang, Passive House.", "Siêu biệt thự lâu đài, sảnh resort 5-6 sao, công trình biểu tượng.")
]

for seg, brand, spec, val in segments_data:
    row_cells = tbl_seg.add_row().cells
    row_cells[0].text = seg
    row_cells[1].text = brand
    row_cells[2].text = spec
    row_cells[3].text = val
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_page_break()

# ── PHẦN II: 4 HỆ CỬA CHỦ LỰC KÈM HÌNH ẢNH MINH HỌA ────────────────
h1("PHẦN II: 4 HỆ CỬA CHỦ LỰC VÀ BẢN VẼ PHÂN TÍCH")

# 1. Hệ Cửa Slim (Minimalist Sliding)
h2("1. Hệ Cửa Slim (Cửa trượt mỏng tối giản)")
para(
    "Hệ cửa Slim là bước đột phá về thiết kế nội ngoại thất hiện đại, tối ưu hóa diện tích bề mặt kính "
    "bằng cách thu hẹp tối đa bản rộng khung nhôm (chỉ từ 40mm - 45mm)."
)

bullet("Ứng dụng cốt lõi: Làm vách ngăn phòng khách, cửa thông phòng WC, phòng thay đồ hoặc cửa ra ban công có mái che.")
bullet("Các hãng tiêu biểu: Kogen Slim, Yangli Slim, Vitrocsa (Thụy Sỹ), các hệ Slim sản xuất lắp ráp trong nước.")
bullet("Ưu điểm: Thiết kế siêu mỏng sang trọng, mở rộng tầm nhìn tối đa (nhìn thấy 95% kính), trượt êm ái nhờ ray treo giảm chấn tự hãm không ray dưới sàn.")
bullet("Nhược điểm: Profile mỏng nên khả năng cách âm, chống nước mưa và chịu gió bão ngoài trời hạn chế.")

# Chèn ảnh Slim
if "cua_slim.png" in images_dest:
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(images_dest["cua_slim.png"], width=Inches(5.0))
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_caption.add_run("Hình 1: Cửa lùa Slim mỏng nội thất phong cách tối giản hiện đại")
    run_cap.font.size = Pt(9.5)
    run_cap.italic = True
else:
    para("[Hình ảnh: Cửa lùa Slim mỏng tối giản]", italic=True)

doc.add_page_break()

# 2. Hệ Cửa Cầu Cách Nhiệt (Thermal Break)
h2("2. Hệ Nhôm Cầu Cách Nhiệt (Thermal Break Aluminum)")
para(
    "Hệ nhôm cầu cách nhiệt được cấu tạo gồm 2 thanh nhôm độc lập ghép lại với nhau bởi dải cầu Polyamide "
    "PA66 GF25 ở giữa. Dải cầu này có tác dụng ngăn cản sự truyền nhiệt từ môi trường ngoài vào phòng."
)

bullet("Ứng dụng cốt lõi: Mặt dựng ngoại thất biệt thự, văn phòng, phòng ngủ hướng Tây, vùng khí hậu nhiệt đới nắng nóng gay gắt hoặc vùng núi cao lạnh giá.")
bullet("Các hãng tiêu biểu: Civro (cầu Technoform), PAG (cầu PA66 AkzoNobel), Schüco AWS 65/75, Reynaers CS 77.")
bullet("Ưu điểm: Khả năng cách nhiệt vượt trội (giảm tiền điện điều hòa đến 35%), cách âm cực cao (lên tới 42-45dB), loại bỏ đọng sương bề mặt kính trong phòng máy lạnh.")
bullet("Nhược điểm: Giá thành cao hơn 40% - 60% so với nhôm không cầu cách nhiệt thông thường, yêu cầu máy móc gia công ép góc CNC và keo cấu trúc chuyên dụng.")

# Chèn ảnh cầu cách nhiệt
if "cau_cach_nhiet.png" in images_dest:
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(images_dest["cau_cach_nhiet.png"], width=Inches(5.0))
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_caption.add_run("Hình 2: Sơ đồ mặt cắt thanh profile nhôm cầu cách nhiệt đa khoang với cầu PA66")
    run_cap.font.size = Pt(9.5)
    run_cap.italic = True
else:
    para("[Hình ảnh: Mặt cắt thanh nhôm cầu cách nhiệt]", italic=True)

doc.add_page_break()

# 3. Hệ Cửa Thủy Lực (Hydraulic Door)
h2("3. Hệ Cửa Thủy Lực Bản Lớn (Grand Entrance Door)")
para(
    "Là hệ cửa đi mở quay bản cánh siêu lớn (từ 140mm - 200mm) sử dụng bản lề sàn thủy lực "
    "hoặc piston thủy lực tự đóng êm."
)

bullet("Ứng dụng cốt lõi: Làm cửa đi chính đại sảnh biệt thự lâu đài, sảnh đón khách sạn, showroom, cửa chính mặt tiền tòa nhà văn phòng.")
bullet("Các hãng tiêu biểu: Owin HL180, Maxpro R200 Thủy lực, EuroVN Thủy lực, Yangli Thủy lực.")
bullet("Ưu điểm: Thẩm mỹ uy nghi hoành tráng, chịu được cánh siêu rộng và nặng (tải trọng lên đến 800 - 1500kg), tự động đóng mở bằng piston rất nhẹ nhàng.")
bullet("Nhược điểm: Cần gia cố kết cấu bê tông sàn kỹ lưỡng trước khi lắp bản lề sàn, chi phí phụ kiện thủy lực nhập khẩu rất cao.")

# Chèn ảnh cửa thủy lực
if "cua_thuy_luc.png" in images_dest:
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(images_dest["cua_thuy_luc.png"], width=Inches(5.0))
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_caption.add_run("Hình 3: Cửa đi chính thủy lực bản lớn biệt thự phong cách Luxury")
    run_cap.font.size = Pt(9.5)
    run_cap.italic = True
else:
    para("[Hình ảnh: Cửa thủy lực bản lớn]", italic=True)

doc.add_page_break()

# 4. Hệ Cửa Xếp Trượt (Bi-fold / Folding Door)
h2("4. Hệ Cửa Xếp Trượt (Cửa xếp gấp đa cánh)")
para(
    "Hệ cửa cho phép gấp các cánh cửa xếp gọn lại về một hoặc hai phía, giải phóng hoàn toàn "
    "không gian lối đi khi mở cửa."
)

bullet("Ứng dụng cốt lõi: Lối ra ban công sân vườn biệt thự, ngăn không gian phòng tiệc nhà hàng, sảnh quán cà phê vườn, lối ra hồ bơi.")
bullet("Các hãng tiêu biểu: Kogen Slim xếp gấp 50/68 (ẩn bản lề), Civro Folding, Maxpro SFD80, PAG 83, Soco Folding, Yangli Folding 80.")
bullet("Ưu điểm: Mở rộng 100% diện tích ô cửa, kết nối không gian bên trong và ngoài trời hoàn hảo.")
bullet("Nhược điểm: Độ kín khít cách âm cách nước kém hơn hệ cửa quay/trượt thường do số lượng gioăng giáp ranh cánh nhiều, bánh xe và bản lề xếp dễ mài mòn cần bảo trì định kỳ.")

# Chèn ảnh cửa xếp trượt
if "cua_xep_truot.png" in images_dest:
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(images_dest["cua_xep_truot.png"], width=Inches(5.0))
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_caption.add_run("Hình 4: Hệ cửa đi xếp trượt nhiều cánh mở rộng lối đi ra sân vườn biệt thự")
    run_cap.font.size = Pt(9.5)
    run_cap.italic = True
else:
    para("[Hình ảnh: Cửa xếp trượt đa cánh]", italic=True)

doc.add_page_break()

# ── PHẦN III: HƯỚNG DẪN LỰA CHỌN THEO CÔNG TRÌNH ───────────────────
h1("PHẦN III: KHUYẾN NGHỊ TRA CỨU THEO LOẠI CÔNG TRÌNH")
para(
    "Để phục vụ công tác bán hàng và tư vấn khách hàng nhanh, Phòng R&D đề xuất cấu hình hệ nhôm "
    "cho 5 nhóm công trình phổ biến nhất tại Việt Nam:"
)

tbl_rec = doc.add_table(rows=1, cols=3)
tbl_rec.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_rec.style = "Table Grid"
hdr_rec = tbl_rec.rows[0].cells
hdr_rec[0].text = "Loại công trình"
hdr_rec[1].text = "Cấu hình đề xuất R&D"
hdr_rec[2].text = "Phân tích công năng và Giá"
for cell in hdr_rec:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

recs_data = [
    ("Nhà phố thương mại\n(Ngân sách trung bình)", 
     "Khung bao và cánh: Yangli 55 rãnh C hoặc Topal XFAD 55.\nPhụ kiện: Bogo / Kinlong đồng bộ.", 
     "Tối ưu chi phí đầu tư ban đầu, vận hành ổn định, dễ thay thế phụ kiện đại trà."),
    
    ("Biệt thự phố, liền kề\n(Ngân sách khá)", 
     "Cửa đi/sổ chính: Viralwindow VRE65 hoặc Maxpro R65.\nCửa đi chính: Owin Thủy lực.\nPhụ kiện: CMECH / Sigico Châu Âu.", 
     "Tối ưu tính thẩm mỹ, bền màu 20-40 năm, chống rung gió tốt, phụ kiện êm ái sang trọng."),
    
    ("Biệt thự vườn luxury\n(Cao cấp & Siêu cao cấp)", 
     "Cửa sổ/đi phòng ngủ: Civro AW65/AD65 cầu cách nhiệt.\nLối ra hồ bơi/sân vườn: PAG 125 Lift Slide có cầu cách nhiệt hoặc Civro Folding.\nPhụ kiện: Sobinco (Bỉ) / Roto (Đức).", 
     "Cách âm phòng ngủ tuyệt đối (40dB+), cản nhiệt hướng Tây 85%, trượt nâng chịu tải cánh lớn cực tốt."),
    
    ("Villa, Resort sát biển\n(Yêu cầu chống muối mặn)", 
     "Khung nhôm: Maxpro SD115 chống bão hoặc Soco 120/180.\nPhụ kiện: Inox SUS316L chống ăn mòn muối mặn.\nBề mặt: Anodize ED phủ mạ điện di chịu mặn.", 
     "Chống rỉ sét bề mặt do gió biển mang muối mặn, chịu được sức gió bão cấp 14 ngoài khơi."),
    
    ("Văn phòng, Tòa nhà cao tầng\n(Mặt dựng vách kính)", 
     "Vách dựng bao quanh: Civro CCW50 Stick (tòa nhà vừa) hoặc Soco UCW Unitized (tòa nhà cao tầng 30+ tầng).\nSilicon: Dow Corning 895 kết cấu chịu lực.", 
     "Tốc độ thi công Unitized cực nhanh (3 ngày/tầng), chất lượng đồng đều đúc xưởng, chịu rung chấn cao động đất.")
]

for build, config, logic in recs_data:
    row_cells = tbl_rec.add_row().cells
    row_cells[0].text = build
    row_cells[1].text = config
    row_cells[2].text = logic
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_page_break()

# ── PHẦN IV: THÔNG TIN LIÊN HỆ ─────────────────────────────────────
h1("PHẦN IV: THÔNG TIN LIÊN HỆ VÀ LIÊN KẾT TRA CỨU")
para(
    "Để có thêm thông tin chi tiết về các bản vẽ CAD, bản vẽ Shopdrawing và Catalogue gốc của từng hệ cửa nhôm, "
    "kính mời quý bộ phận tra cứu các tệp tin lưu trữ trực tiếp trên hệ thống máy chủ R&D Sao Vàng Group:"
)

bullet("Cơ sở dữ liệu phẳng chi tiết 110 hệ: 02_Cơ sở Dữ liệu Gốc.xlsx")
bullet("Danh mục tài liệu PDF chính hãng: 03_Danh mục Catalogue theo Hãng.xlsx")
bullet("Theo dõi tiến trình khảo sát số hóa hệ mới: 04_Bảng theo dõi Tiến độ.xlsx")
bullet("Thư mục bản vẽ CAD mẫu chi tiết: 05_Thư viện CAD")
bullet("Thư mục mặt cắt nhôm thực tế: 06_Thư viện Profile Nhôm")

para("\nTài liệu được lưu hành nội bộ và bảo mật bởi Sao Vàng Group.", italic=True, size=10, space_after=12)

# Lưu document
path_doc = os.path.join(ROOT, "BaoCao_TomTat_TrucQuan.docx")
safe_save_docx(doc, path_doc)
