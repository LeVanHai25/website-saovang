# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo 2 tài liệu riêng biệt cho hãng Việt Pháp SHAL:
1. Catalogue_HeCuaNhom_VietPhap.xlsx
2. ThuVien_HeCuaNhom_VietPhap.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

VIETPHAP_SYSTEMS = [
    ('1', 'Phổ thông', 'Việt Pháp 450', 'Cửa đi mở quay hệ 450', 
     'Cửa đi mở quay 1 cánh, 2 cánh hoặc 4 cánh, nẹp liền sập dễ sản xuất, lắp đặt thông dụng.', 
     '1.2mm - 1.4mm', 'Phụ kiện đồng bộ Việt Pháp / Kinlong rãnh 22', 'Gioăng EPDM đơn / Silicon Apollo A500', 
     'Giá thành rẻ, phổ biến, dễ thi công, tuy nhiên độ dày mỏng và cách âm trung bình.'),
     
    ('2', 'Phổ thông', 'Việt Pháp 4400', 'Cửa sổ mở quay/hất hệ 4400', 
     'Cửa sổ mở quay, mở hất hoặc vách cố định bản nhỏ gọn nhẹ, dễ sản xuất nhanh.', 
     '1.1mm - 1.2mm', 'Khớp chốt Việt Pháp / Kinlong', 'Gioăng cao su chèn sập / Silicon Apollo', 
     'Nhẹ, chi phí cực thấp, phổ biến rộng rãi cho công trình dân dụng bình dân.'),
     
    ('3', 'Phổ thông', 'Việt Pháp 2600', 'Cửa sổ & đi trượt lùa hệ 2600', 
     'Cửa đi trượt lùa hoặc cửa sổ trượt lùa 2 ray tiết kiệm không gian đóng mở tối đa.', 
     '1.1mm - 1.2mm', 'Bánh xe đơn trượt & khóa sập Việt Pháp', 'Gioăng chèn nỉ mỏng / Apollo', 
     'Giá thành rẻ, tối ưu diện tích. Tuy nhiên kết cấu mỏng, trượt không đầm tay.'),
     
    ('4', 'Tầm trung', 'Shal 55', 'Hệ nhôm Việt Pháp Shal 55 mở quay', 
     'Cửa đi, cửa sổ mở quay kiểu dáng Xingfa định hình bản 55mm thiết kế phẳng mượt góc.', 
     '1.2mm - 1.4mm', 'Phụ kiện PMA / Draho / Kinlong đồng bộ', 'Gioăng EPDM đơn / Silicon Apollo', 
     'Kiểu dáng hiện đại giống Xingfa, giá hợp lý cho nhà phố cận trung cấp.'),
     
    ('5', 'Mặt dựng', 'Việt Pháp 1100', 'Hệ vách mặt dựng Stick 1100', 
     'Vách mặt dựng Stick chìm hoặc nổi liên kết kính cho các công trình showroom, văn phòng tầm trung.', 
     '1.5mm - 2.0mm', 'Ke chịu lực & bulong liên kết inox', 'Silicon kết cấu chống thấm Dow Corning', 
     'Thi công đơn giản ngoài công trường, chi phí hợp lý. Chỉ phù hợp nhà trung và thấp tầng.'),
     
    ('6', 'Cao cấp Premium', 'Shal Premium', 'Hệ nhôm rãnh C cao cấp Shal Premium', 
     'Cửa đi, cửa sổ rãnh C chuẩn Châu Âu, tích hợp phụ kiện cao cấp, nước sơn cao cấp liên doanh Sepalumic (Pháp).', 
     '1.6mm - 2.0mm', 'Phụ kiện CMECH / Roto rãnh C Châu Âu', 'Gioăng EPDM đa khoang / Silicon Dow Corning', 
     'Chất lượng sơn cao cấp bảo hành 30 năm, rãnh C cách âm và chống muối mặn tuyệt vời.')
]

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_VietPhap.xlsx
# ─────────────────────────────────────────────
def create_vietphap_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue VietPhap"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM LIÊN DOANH VIỆT PHÁP SHAL"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="003E3E")
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    fill_row = PatternFill(fill_type="solid", fgColor="F0F7F7")
    for r_idx, row_vals in enumerate(VIETPHAP_SYSTEMS, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_VietPhap.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel VietPhap saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel VietPhap to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_VietPhap.docx
# ─────────────────────────────────────────────
def create_vietphap_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x00,0x3E,0x3E))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM LIÊN DOANH VIỆT PHÁP SHAL", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật tổng hợp hệ cửa nhôm liên doanh công nghệ Sepalumic (Pháp)\nTích hợp phân tích ưu nhược điểm R&D và ma trận so sánh phân khúc thị trường", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM VIỆT PHÁP SHAL", bold=True, size=14, color=RGBColor(0x00,0x3E,0x3E), space_before=12)
    add_p(
        "Nhôm Việt Pháp SHAL là sản phẩm liên doanh giữa Tập đoàn Sepalumic (Pháp) và Công ty HACIPCO thuộc UBND thành phố Hà Nội. "
        "Với hơn 2 thập kỷ hoạt động phát triển tại Việt Nam, Việt Pháp SHAL là một trong những thương hiệu nhôm định hình đầu tiên "
        "đặt nền móng cho nhôm hệ định hình sơn tĩnh điện trong nước. "
        "Năm 2015, nhà máy hiện đại quy mô lớn được khánh thành tại KCN Phúc Sơn, Ninh Bình với tổng mức đầu tư 1.180 tỷ đồng, "
        "nhập khẩu toàn bộ hệ thống máy đùn ép và sơn tĩnh điện từ Tây Ban Nha, Ý, Đức. "
        "Nhôm Việt Pháp SHAL đạt đầy đủ các tiêu chuẩn quốc tế khắc khe như AAMA, GSB, Qualicoat, JIS H4100, ASTM B221M và TCVN. "
        "Đặc biệt, hãng cam kết bảo hành bề mặt lớp sơn phủ lên tới 30 năm cho các dòng sản phẩm cao cấp rãnh C Châu Âu."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ, ƯU NHƯỢC ĐIỂM R&D CÁC HỆ NHÔM VIỆT PHÁP SHAL", bold=True, size=14, color=RGBColor(0x00,0x3E,0x3E), space_before=12)
    
    # 1. Việt Pháp 450
    add_p("1. Việt Pháp Hệ 450 (Cửa đi mở quay)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.4mm.")
    add_bullet("Ưu điểm R&D: Giá thành rất rẻ, dễ gia công bằng máy đột dập thông dụng, phụ kiện Kinlong hoặc Việt Pháp nội địa rất sẵn, phù hợp cho nhà cấp 4, nhà phố giá rẻ, kho xưởng.")
    add_bullet("Nhược điểm R&D: Bản cánh hẹp, độ dày mỏng nên chống rung gió yếu, không tương thích phụ kiện rãnh C cao cấp, cách âm ở mức trung bình.")
    
    # 2. Việt Pháp 4400
    add_p("2. Việt Pháp Hệ 4400 (Cửa sổ mở quay/hất)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.1mm - 1.2mm.")
    add_bullet("Ưu điểm R&D: Bản nhôm mỏng nhẹ, chi phí nguyên vật liệu siêu thấp, lắp dựng cực nhanh, đa năng làm được cả cửa sổ quay và hất.")
    add_bullet("Nhược điểm R&D: Độ kín khít nước mưa trung bình, bản cánh mỏng dễ bị biến dạng nếu va đập gió lớn.")

    # 3. Việt Pháp 2600
    add_p("3. Việt Pháp Hệ 2600 (Cửa lùa trượt)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.1mm - 1.2mm.")
    add_bullet("Ưu điểm R&D: Hệ lùa trượt 2 ray kinh điển, giá thành cực rẻ tiết kiệm diện tích tối đa.")
    add_bullet("Nhược điểm R&D: Ray nhôm dễ mòn, độ trượt lùa nhẹ nhàng kém hơn các hệ lùa bản lớn có ray inox chặn.")

    # 4. Shal 55
    add_p("4. Shal 55 (Nhôm hệ trung cấp)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.2mm - 1.4mm.")
    add_bullet("Ưu điểm R&D: Thiết kế phẳng phiu kiểu dáng rãnh C tương tự Xingfa hiện đại, khắc phục nhược điểm thẩm mỹ của hệ 450 truyền thống.")
    add_bullet("Nhược điểm R&D: Phân khúc cạnh tranh khốc liệt với nhôm PMA và Xingfa mỏng trong nước.")

    # 5. Việt Pháp 1100
    add_p("5. Việt Pháp 1100 (Vách mặt dựng Stick)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.5mm - 2.0mm.")
    add_bullet("Ưu điểm R&D: Giá nhôm mặt dựng Stick nội địa rẻ, lắp đặt không cần cầu trục chuyên dụng lớn, giảm thiểu chi phí đầu tư.")
    add_bullet("Nhược điểm R&D: Khả năng chống thấm dột dài hạn phụ thuộc hoàn toàn vào tay nghề đi keo ngoài công trường.")

    # 6. Shal Premium
    add_p("6. Shal Premium (Hệ rãnh C cao cấp bảo hành 30 năm)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.6mm - 2.0mm chuẩn rãnh C.")
    add_bullet("Ưu điểm R&D: Lớp sơn phủ cao cấp nhập khẩu bảo hành bề mặt lên tới 30 năm chống phai màu, kết cấu rãnh C Châu Âu tích hợp CMECH/Roto cách âm cách nhiệt hoàn hảo.")
    add_bullet("Nhược điểm R&D: Chi phí sản xuất cao, định vị phân khúc khách hàng cao cấp khó bán hơn các dòng nhôm bình dân.")

    # ── PHẦN SO SÁNH MA TRẬN ─────────────────────
    add_p("PHẦN SO SÁNH MA TRẬN PHÂN KHÚC HỆ VIỆT PHÁP 450 VỚI ĐỐI THỦ CẠNH TRANH", bold=True, size=14, color=RGBColor(0x00,0x3E,0x3E), space_before=12)
    add_p(
        "Bảng ma trận so sánh chéo đặc tính kỹ thuật hệ nhôm Việt Pháp 450 truyền thống "
        "với các dòng nhôm phân khúc phổ thông và cận trung cấp trên thị trường:"
    )
    
    tbl_comp = doc.add_table(rows=1, cols=5)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_comp.style = "Table Grid"
    
    hdr_c = tbl_comp.rows[0].cells
    hdr_c[0].text = "Đặc tính so sánh"
    hdr_c[1].text = "Việt Pháp 450"
    hdr_c[2].text = "PMA 55"
    hdr_c[3].text = "Topal Slida"
    hdr_c[4].text = "Xingfa Việt Nam"
    for cell in hdr_c:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    comp_data = [
        ("Độ dày cánh cửa", "1.2mm - 1.4mm", "1.2mm - 1.4mm", "1.2mm", "1.4mm"),
        ("Kiểu dáng nẹp kính", "Nẹp sập rời truyền thống", "Nẹp liền sập vát cạnh", "Nẹp sập rời vuông", "Nẹp sập bo tròn"),
        ("Độ bền màu sơn phủ", "Bảo hành 5-10 năm", "Bảo hành 10-15 năm", "Bảo hành 10 năm", "Bảo hành 10 năm"),
        ("Hệ phụ kiện phổ biến", "Chốt gạt Việt Pháp / Kinlong", "PMA / Kinlong rãnh 22", "Topal đồng bộ", "Kinlong / Draho rãnh 22"),
        ("Phân khúc giá", "$ (Siêu tiết kiệm)", "$$ (Bình dân giá tốt)", "$$ (Tầm trung)", "$$ (Tầm trung)")
    ]
    
    for row in comp_data:
        cells = tbl_comp.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True
                
    doc.add_paragraph()
    add_p("Tài liệu lưu hành nội bộ và thuộc bản quyền R&D Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_VietPhap.docx")
    try:
        doc.save(path)
        print(f"  >> Docx VietPhap saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx VietPhap to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("CREATING standalone VietPhap docx & xlsx in BASE directory...")
    print("=" * 60)
    create_vietphap_xlsx()
    create_vietphap_docx()
    print("=" * 60)
