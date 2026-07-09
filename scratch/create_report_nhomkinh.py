# -*- coding: utf-8 -*-
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Sets light borders for the table."""
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def main():
    doc = Document()
    
    # 1. Page Setup & Margins (1 inch / 2.54 cm all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Nhôm Kính Sao Vàng (NKSV) | Báo cáo Định hướng & Chiến lược Phát triển Thương hiệu"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = 'Arial'
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Công ty Cổ phần Sản xuất Cơ khí Sao Vàng - Phân khúc Nhôm Kính - Lưu hành nội bộ"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.name = 'Arial'
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Styling Helpers
    PRIMARY_COLOR = RGBColor(0, 102, 102)      # Deep Teal #006666 (Fresh color for Glass/Alum)
    SECONDARY_COLOR = RGBColor(176, 126, 40)  # Muted Gold #B07E28
    TEXT_COLOR = RGBColor(38, 38, 38)          # Off-Black #262626
    
    def add_heading_1(text, space_before=18, space_after=6):
        h = doc.add_heading('', level=1)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return h

    def add_heading_2(text, space_before=12, space_after=4):
        h = doc.add_heading('', level=2)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return h

    def add_heading_3(text, space_before=8, space_after=2):
        h = doc.add_heading('', level=3)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        return h

    def add_para(text, bold_prefix="", space_after=6, is_bullet=False, indent_level=0):
        style = 'List Bullet' if is_bullet else 'Normal'
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.2
        
        if indent_level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = TEXT_COLOR
            
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = TEXT_COLOR
        return p

    def add_callout(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(10.5)
            r_bold.font.bold = True
            r_bold.font.italic = True
            r_bold.font.color.rgb = PRIMARY_COLOR
            
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(10.5)
        r_text.font.italic = True
        r_text.font.color.rgb = RGBColor(80, 80, 80)
        return p

    # --- Title Page ---
    p_comp = doc.add_paragraph()
    p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_comp.paragraph_format.space_before = Pt(36)
    p_comp.paragraph_format.space_after = Pt(18)
    run_comp = p_comp.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_comp.font.name = 'Arial'
    run_comp.font.size = Pt(12)
    run_comp.font.bold = True
    run_comp.font.color.rgb = PRIMARY_COLOR

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("BÁO CÁO NỘI DUNG & ĐỊNH HƯỚNG CHIẾN LƯỢC PHÁT TRIỂN")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_before = Pt(6)
    p_subtitle.paragraph_format.space_after = Pt(72)
    run_sub = p_subtitle.add_run("PHÂN KHÚC: NHÔM KÍNH SAO VÀNG (NKSV)\nTái Định Vị thành Nhà Cung Cấp Giải Pháp Cửa & Mặt Dựng Nhôm Kính Kiến Trúc Tổng Thể")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.left_indent = Inches(1.5)
    p_meta.paragraph_format.space_after = Pt(4)
    r_meta_rec = p_meta.add_run("Kính gửi: ")
    r_meta_rec.font.bold = True
    r_meta_rec.font.name = 'Arial'
    r_meta_rec_val = p_meta.add_run("Ban Giám đốc Công ty Cổ phần Sản xuất Cơ khí Sao Vàng")
    r_meta_rec_val.font.name = 'Arial'

    p_meta2 = doc.add_paragraph()
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta2.paragraph_format.left_indent = Inches(1.5)
    p_meta2.paragraph_format.space_after = Pt(4)
    r_meta_pre = p_meta2.add_run("Người trình bày: ")
    r_meta_pre.font.bold = True
    r_meta_pre.font.name = 'Arial'
    r_meta_pre_val = p_meta2.add_run("Bộ phận Marketing & Kỹ thuật")
    r_meta_pre_val.font.name = 'Arial'

    p_meta3 = doc.add_paragraph()
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta3.paragraph_format.left_indent = Inches(1.5)
    p_meta3.paragraph_format.space_after = Pt(12)
    r_meta_date = p_meta3.add_run("Thời gian: ")
    r_meta_date.font.bold = True
    r_meta_date.font.name = 'Arial'
    r_meta_date_val = p_meta3.add_run("Tháng 07 năm 2026")
    r_meta_date_val.font.name = 'Arial'

    doc.add_page_break()

    # --- LỜI MỞ ĐẦU ---
    add_heading_1("LỜI MỞ ĐẦU: LÝ DO TÁI ĐỊNH VỊ THƯƠNG HIỆU NHÔM KÍNH SAO VÀNG")
    add_para("Kính gửi Sếp và Ban Giám đốc,", space_after=10)
    
    add_para("Lĩnh vực nhôm kính kiến trúc tại Việt Nam hiện đang bị bão hòa cực kỳ nghiêm trọng. Hơn 95% đơn vị thi công trên thị trường là các xưởng gia công quy mô nhỏ lẻ, hoạt động theo mô hình cơ học: nhận kích thước sổ đo -> mua nhôm cây (chủ yếu là Xingfa Quảng Đông phổ thông hoặc hàng nhái loại 2, 3) -> mua phụ kiện giá rẻ -> lắp đặt tính tiền theo m2 thô. Cách làm này kéo doanh nghiệp vào cuộc đua phá giá khốc liệt, làm mất đi uy tín kỹ thuật và khả năng đảm nhận các công trình lớn.")
    
    add_para("Chủ đầu tư biệt thự cao cấp, resort ven biển, và các dự án tòa nhà hiện đại đòi hỏi khắt khe về tính chống ăn mòn (muối mặn biển), khả năng cản gió bão cực đoan, chống thấm ngấm tuyệt đối, tính cách âm cách nhiệt và tính thẩm mỹ nghệ thuật cao như uốn vòm, cửa lùa slim cánh mỏng chịu lực.")
    
    add_para("Do đó, chúng tôi đề xuất tái định vị phân khúc Nhôm Kính Sao Vàng (NKSV) từ một đơn vị làm cửa nhôm kính thông thường thành ")
    p_last = doc.paragraphs[-1]
    r_em = p_last.add_run("Nhà Cung Cấp Giải Pháp Cửa & Mặt Dựng Nhôm Kính Kiến Trúc Tổng Thể (Architectural Aluminum & Glass Solutions Provider)")
    r_em.font.bold = True
    r_em.font.name = 'Arial'
    r_last = p_last.add_run(" theo mô hình chuỗi giá trị khép kín EPCM (Engineering – Procurement – Construction – Maintenance).")
    r_last.font.name = 'Arial'
    
    add_para("Tài liệu này tổng hợp toàn bộ nội dung định hướng, phân tách 19 hệ thống nhôm chủ lực mà bộ phận kỹ thuật đã thu thập và chuẩn hóa dữ liệu ngày hôm nay để Sếp duyệt trước khi đưa lên hệ thống website và hồ sơ năng lực chính thức.")

    # --- PHẦN A: ĐỊNH HƯỚNG TẦM NHÌN, SỨ MỆNH & GIÁ TRỊ CỐT LÕI ---
    add_heading_1("PHẦN A: ĐỊNH HƯỚNG TẦM NHÌN, SỨ MỆNH & GIÁ TRỊ CỐT LÕI (NKSV)")
    
    add_heading_2("1. TẦM NHÌN (Vision)")
    add_para("Trở thành đơn vị hàng đầu tại Việt Nam trong việc cung cấp giải pháp cửa và mặt dựng nhôm kính kiến trúc phân khúc cao cấp và luxury, tiên phong ứng dụng các công nghệ xử lý bề mặt thế hệ mới (Xi mạ điện di Anodized ED) và các hệ nhôm rãnh C chuẩn Châu Âu đạt độ bền trên 40 năm.")

    add_heading_2("2. SỨ MỆNH (Mission)")
    add_para("Kiến tạo những khoảng mở hoàn mỹ, kết nối không gian thiên nhiên và con người bằng các sản phẩm cửa nhôm kính đạt độ thẩm mỹ tinh xảo nhất và hiệu năng kỹ thuật an toàn cao nhất.", is_bullet=True)
    add_para("Đồng hành cùng các kiến trúc sư, chủ đầu tư để hiện thực hóa những thiết kế vòm cong mỹ thuật, vách kính mặt dựng nhịp lớn khó thi công mà các đơn vị khác từ chối thực hiện.", is_bullet=True)
    add_para("Cam kết loại bỏ vấn nạn hàng nhái, hàng giả linh phụ kiện trên thị trường bằng sự trung thực tuyệt đối về nguồn gốc vật tư xuất xứ chính hãng.", is_bullet=True)

    add_heading_2("3. GIÁ TRỊ CỐT LÕI (Core Values)")
    add_para("Đường keo silicon láng mịn, khe ép góc kín khít như sợi chỉ, kết cấu cửa cân đối hoàn hảo và vận hành êm ái.", bold_prefix="Thẩm Mỹ Tinh Xảo: ", is_bullet=True)
    add_para("Tính toán kỹ lưỡng tải trọng gió, hệ số cản nhiệt U-value, khả năng cách âm để tư vấn giải pháp nhôm kính đúng kỹ thuật cho từng công trình.", bold_prefix="Hiệu Năng Kỹ Thuật: ", is_bullet=True)
    add_para("100% nhôm cây, kính cường lực/kính hộp, phụ kiện, gioăng, keo đều chính hãng rõ ràng nguồn gốc chứng chỉ (CO/CQ). Nói không với hàng giả nhái.", bold_prefix="Minh Bạch Tuyệt Đối: ", is_bullet=True)
    add_para("Bảo hành bề mặt màu sắc nhôm lên đến 40 năm. Đồng hành chăm sóc bảo dưỡng cửa trọn đời dự án.", bold_prefix="Cam Kết Dài Hạn: ", is_bullet=True)

    # --- PHẦN B: NKSV GIẢI QUYẾT NHỮNG BÀI TOÁN GÌ CHO KHÁCH HÀNG? ---
    add_heading_1("PHẦN B: NKSV GIẢI QUYẾT NHỮNG BÀI TOÁN GÌ CHO KHÁCH HÀNG?")
    add_para("Thị trường thi công nhôm kính tại Việt Nam có rất nhiều bất cập lớn, NKSV tập trung giải quyết triệt để 5 bài toán đau đầu của khách hàng:")

    add_heading_2("Bài toán 1: Hàng nhái, hàng giả tràn lan trên thị trường nhôm kính")
    add_callout("Thanh nhôm Xingfa bị làm giả tem nhãn, đặc biệt là phụ kiện khóa bản lề Kinlong bị nhái loại 2, loại 3 tinh vi đến mức người thường không phân biệt được. Công trình sử dụng vài tháng đã rỉ sét bản lề, kẹt khóa, xệ cánh.", "Nỗi đau: ")
    add_para("NKSV thiết lập quy trình kiểm soát chuỗi cung ứng khắt khe. Chúng tôi là đối tác trực tiếp của các nhà phân phối chính hãng lớn (như Maxpro Nhật Bản, Civro Đức, Xingfa Quảng Đông tem đỏ, phụ kiện CMECH Mỹ, Roto Đức). Mọi lô hàng nhập xưởng đều đi kèm hồ sơ kiểm định, CO/CQ đầy đủ. Chúng tôi cam kết đền bù 200% nếu phát hiện hàng giả, hàng nhái.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 2: Ngấm nước mưa qua khe cửa & cách âm kém")
    add_callout("Sau các trận bão lớn, nước mưa tràn qua các góc ép cánh cửa, thấm đẫm vào sàn gỗ đắt tiền của biệt thự. Hệ gioăng lỏng lẻo khiến tiếng ồn còi xe, tiếng gió rít vẫn lọt vào phòng ngủ dù đã đóng kín cửa.", "Nỗi đau: ")
    add_para("Ứng dụng biện pháp đúc ép góc thủy lực kết hợp bơm keo chuyên dụng 2 thành phần (như keo đặc chủng Xylotex/Dow Corning) tại xưởng. Đồng thời, NKSV sử dụng hệ gioăng EPDM 3 lớp khép kín, kết hợp phụ kiện rãnh C tiêu chuẩn Châu Âu đồng bộ ép chặt cánh cửa vào khung bao, giải quyết dứt điểm hiện tượng thấm dột và đạt hiệu số cách âm lên đến 40dB.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 3: Cửa bị xệ cánh và kẹt khóa sau một thời gian ngắn sử dụng")
    add_callout("Cửa nhôm kính có diện tích cánh lớn (đặc biệt là cửa đi mở quay hoặc xếp lùa) chịu tải trọng kính hộp rất nặng. Nếu xưởng sử dụng nhôm mỏng, liên kết góc lỏng lẻo và phụ kiện chịu lực kém, cửa sẽ nhanh chóng bị xệ cánh, kẹt cạ nền nhà và không thể khóa được.", "Nỗi đau: ")
    add_para("Sử dụng các thanh nhôm profile có độ dày tiêu chuẩn từ 1.8mm - 2.5mm cho các hệ cửa lớn, kết hợp với góc ép cơ học máy ép góc thủy lực lực ép 15 tấn, ke góc dày 4mm. Chúng tôi chỉ phối hợp các thương hiệu phụ kiện chịu tải cao cấp như CMECH (Mỹ), Sobinco (Bỉ), Roto (Đức) để đảm bảo độ bền đóng mở trên 100,000 lần không bị xệ cạ.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 4: Mất thẩm mỹ ở các chi tiết hoàn thiện bề mặt và góc liên kết")
    add_callout("Các khe ghép góc 45 độ bị hở lớn, silicon bắn nham nhở, các đường uốn vòm tròn bị móp méo gãy khúc, bề mặt nhôm dễ bong tróc sơn do thời tiết nắng nóng mặn biển khắc nghiệt.", "Nỗi đau: ")
    add_para("Sản xuất trên dây chuyền máy cắt CNC 2 đầu tự động nhập khẩu, máy uốn vòm CNC kỹ thuật số chính xác. Thợ bắn silicone của NKSV có tay nghề cao đảm bảo đường keo phẳng mịn, đồng đều. Chúng tôi cung cấp các dòng nhôm công nghệ xi mạ Anodized ED chống phai màu do tia UV và chống muối biển tuyệt đối.", "Giải pháp của NKSV: ")

    add_heading_2("Bài toán 5: Rủi ro khi kết hợp mặt dựng nhôm kính với khung thép chịu lực")
    add_callout("Khi làm vách kính lớn hoặc mái đón kính trung tâm, phần kết cấu thép đỡ do một bên cơ khí làm, phần nhôm kính do bên cửa làm. Khi hai bên không khớp bản vẽ, kính lắp vào bị nứt vỡ hoặc không thể lắp vừa, hai nhà thầu đổ lỗi lẫn nhau.", "Nỗi đau: ")
    add_para("NKSV kế thừa nền tảng cơ khí chịu lực nặng từ CKSV. Chúng tôi có khả năng tự tính toán kết cấu thép đỡ kết hợp hệ nhôm kính mặt dựng Spider/Curtain Wall một cách đồng bộ. Thiết kế 3D tích hợp giúp loại bỏ hoàn toàn các sai lệch giao diện cơ khí - nhôm kính, chịu trách nhiệm trọn gói trước chủ đầu tư.", "Giải pháp của NKSV: ")

    # --- PHẦN I: HỆ THỐNG NHẬN DIỆN THƯƠNG HIỆU & THÔNG ĐIỆP ĐỊNH VỊ ---
    add_heading_1("PHẦN I: HỆ THỐNG NHẬN DIỆN THƯƠNG HIỆU & THÔNG ĐIỆP ĐỊNH VỊ")
    
    add_heading_2("1. Tuyên ngôn định vị thương hiệu (Brand Positioning)")
    add_para('"Nhôm Kính Sao Vàng không chỉ cung cấp cửa nhôm thông thường — Chúng tôi mang đến giải pháp cửa và mặt dựng kiến trúc tổng thể, kết hợp công nghệ vật liệu hiện đại và kỹ nghệ lắp ráp đỉnh cao để kiến tạo nên những không gian sống đẳng cấp và bền vững cùng thời gian."')
    
    add_heading_2("2. Thông điệp truyền thông cốt lõi (Slogan & Key Message)")
    add_para("MỞ RỘNG KHÔNG GIAN — NÂNG TẦM KIẾN TRÚC", bold_prefix="Slogan chính: ")
    add_para("Một Điểm Chạm – Trọn Vẹn Giải Pháp – Bền Vững Tương Lai.", bold_prefix="Triết lý vận hành: ")
    add_para("Khách hàng từ chủ đầu tư, kiến trúc sư đến chủ nhà chỉ cần làm việc với NKSV để được đáp ứng trọn gói từ khâu khảo sát, thiết kế hệ cửa phù hợp, nhập khẩu vật tư, sản xuất CNC, thi công hiện trường và bảo dưỡng định kỳ.", bold_prefix="Một Điểm Chạm: ", indent_level=1)
    add_para("Chúng tôi giải quyết toàn diện bài toán kiến trúc của công trình: từ hệ cửa đi quay/lùa, mặt dựng kính nhịp lớn, lan can kính cong nghệ thuật đến mái kính lấy sáng, đảm bảo tính đồng bộ thẩm mỹ.", bold_prefix="Trọn Vẹn Giải Pháp: ", indent_level=1)
    add_para("Sự yên tâm tuyệt đối với tuổi thọ lớp sơn màu nhôm lên đến 40 năm, cam kết không thấm nước, cách âm cách nhiệt giúp tiết kiệm năng lượng điện điều hòa.", bold_prefix="Bền Vững Tương Lai: ", indent_level=1)

    # --- PHẦN II: 5 LỢI THẾ CẠNH TRANH CỐT LÕI (USPs) ---
    add_heading_1("PHẦN II: 5 LỢI THẾ CẠNH TRANH CỐT LÕI (USPs)")
    
    add_para("NKSV làm chủ hệ thống dữ liệu kỹ thuật và phân phối của 19 thương hiệu nhôm lớn nhất thị trường, từ các dòng nhôm siêu cao cấp nhập khẩu Châu Âu/Nhật Bản (Civro, Maxpro, Hondalex, Soco, Kogen, KOSO) đến phân khúc phổ thông (Xingfa, PMA, Topal, VietPhap). Chúng tôi sẵn sàng tư vấn khách quan nhất để tối ưu ngân sách cho từng dự án.", bold_prefix="1. Hệ sinh thái 19 dòng nhôm đa dạng bậc nhất: ", is_bullet=True)
    
    add_para("Chúng tôi đi đầu trong thi công các hệ nhôm xử lý bề mặt bằng công nghệ mạ điện di Anodise ED (Maxpro, Hondalex, Soco) chống oxy hóa hoàn hảo dưới tác động của gió biển chứa muối mặn, là lựa chọn số một cho các resort ven biển và biệt thự cao cấp.", bold_prefix="2. Làm chủ công nghệ mạ điện di Anodise ED bền màu 40 năm: ", is_bullet=True)
    
    add_para("Kế thừa nhà xưởng cơ khí hiện đại, NKSV uốn vòm CNC nhôm kính với bán kính cong cực nhỏ mà không bị móp méo mép nhôm hay rạn nứt bề mặt sơn. Chúng tôi uốn được các hệ vòm tân cổ điển của nhôm Xingfa, Maxpro, Civro đạt độ thẩm mỹ tinh xảo đỉnh cao.", bold_prefix="3. Kỹ nghệ uốn vòm nhôm kính nghệ thuật độc quyền: ", is_bullet=True)
    
    add_para("Liên kết khung cửa góc 45 độ được xử lý bằng ke góc đúc dày, ép góc thủy lực lực ép lớn kết hợp keo góc 2 thành phần đặc chủng. Đường keo silicone ngoài trời và trong nhà được bắn thẳng tắp, bo tròn mép mịn màng không lem nhem.", bold_prefix="4. Tiêu chuẩn lắp ghép kín khít và bắn keo silicone láng mịn: ", is_bullet=True)
    
    add_para("Quy trình khép kín giúp NKSV kiểm soát chặt chẽ tiến độ sản xuất tại xưởng và thi công lắp đặt tại công trình, cam kết giao hàng đúng tiến độ thầu và có chính sách bảo trì cửa định kỳ 1 năm/lần cho đối tác.", bold_prefix="5. Dịch vụ bảo trì định kỳ trọn vòng đời cửa độc nhất: ", is_bullet=True)

    # --- PHẦN III: 4 TRỤ CỘT NĂNG LỰC DỊCH VỤ NHÔM KÍNH ---
    add_heading_1("PHẦN III: 4 TRỤ CỘT NĂNG LỰC DỊCH VỤ NHÔM KÍNH")
    add_para("Năng lực dịch vụ của NKSV được cấu trúc đồng bộ theo 4 giai đoạn khép kín:")

    # Table Summary
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table, color="A0C4C4")
    
    headers = [
        "01. ARCHITECTURAL\nENGINEERING",
        "02. ADVANCED\nMANUFACTURING",
        "03. PRECISION\nINSTALLATION",
        "04. LIFECYCLE\nMAINTENANCE"
    ]
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_background(cell, "006666")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    sub_headers = [
        "Tư vấn giải pháp & Thiết kế hệ cửa tối ưu",
        "Sản xuất cửa CNC & Uốn vòm nghệ thuật tại xưởng",
        "Thi công lắp dựng hiện trường chống ngấm thấm",
        "Chăm sóc định kỳ & Sửa chữa sự cố trong 24 giờ"
    ]
    for col_idx, sub_text in enumerate(sub_headers):
        cell = table.cell(1, col_idx)
        set_cell_background(cell, "EBF2F2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(sub_text)
        run.font.name = 'Arial'
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = SECONDARY_COLOR

    descriptions = [
        "Khảo sát đo laser hiện trạng, vẽ thiết kế 3D hệ cửa/mặt dựng, tính toán U-value cản nhiệt & sức cản gió, Shop Drawing, lập BOM vật tư kính/nhôm thanh.",
        "Cắt góc CNC tự động 2 đầu, phay lỗ khóa/phụ kiện trên máy CNC rãnh C, ép góc thủy lực bơm keo đặc chủng, uốn vòm CNC định hình profile.",
        "Cân chỉnh khung bao bằng máy laser, cố định vít nở nở inox, chèn keo bọt nở cách âm chống thấm, lắp kính hộp, bắn keo silicone hoàn thiện.",
        "Bảo dưỡng định kỳ hàng năm (kiểm tra gioăng, keo silicone, bôi trơn bản lề khóa), thay thế kính vỡ khẩn cấp, sửa cửa kẹt xệ cánh trong vòng 24 giờ."
    ]
    for col_idx, desc_text in enumerate(descriptions):
        cell = table.cell(2, col_idx)
        set_cell_background(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(desc_text)
        run.font.name = 'Arial'
        run.font.size = Pt(8.5)
        run.font.color.rgb = TEXT_COLOR

    add_para("", space_after=12) # Spacer

    add_heading_2("TRỤ CỘT 01: ARCHITECTURAL ENGINEERING (Kỹ Thuật & Thiết Kế)")
    add_para("Đo đạc hiện trạng bằng máy đo khoảng cách laser độ chính xác cao. Khảo sát kỹ lưỡng hướng nắng, hướng gió và độ rung chấn của công trình.", bold_prefix="Khảo sát kỹ thuật chuyên sâu: ", is_bullet=True)
    add_para("Tư vấn chọn hệ nhôm phù hợp trong 19 dòng nhôm dựa trên ngân sách và yêu cầu kỹ thuật. Đề xuất phương án kính (kính cường lực, kính dán an toàn, kính hộp Low-E cản nhiệt, kính phản quang).", bold_prefix="Tư vấn giải pháp tối ưu: ", is_bullet=True)
    add_para("Thiết kế bản vẽ Shop Drawing chi tiết từng mặt cắt thanh nhôm, điểm liên kết gioăng, điểm bắn keo silicon và phương án lắp ghép phụ kiện.", bold_prefix="Triển khai bản vẽ Shop Drawing: ", is_bullet=True)
    add_para("Tính toán tải trọng gió chịu được của các vách kính lớn (Wind Load Analysis) và tính toán khả năng truyền nhiệt của hệ cửa để đảm bảo tiêu chuẩn tiết kiệm năng lượng công trình xanh.", bold_prefix="Mô phỏng ứng suất & Nhiệt học cửa: ", is_bullet=True)

    add_heading_2("TRỤ CỘT 02: ADVANCED MANUFACTURING (Sản Xuất & Chế Tạo Cửa CNC)")
    add_para("Sử dụng máy cắt hai đầu CNC nhập khẩu cắt góc 45 độ chuẩn xác tuyệt đối, tránh hiện tượng hở khe góc khi ghép.", bold_prefix="Cắt góc CNC chính xác cao: ", is_bullet=True)
    add_para("Phay lỗ khóa, lỗ tay nắm, lỗ thoát nước mưa tự động trên máy phay trung tâm CNC rãnh C Châu Âu đảm bảo khớp phụ kiện 100%.", bold_prefix="Phay lỗ phụ kiện tự động: ", is_bullet=True)
    add_para("Liên kết góc cánh bằng máy ép góc thủy lực ép lực lớn 15 tấn kết hợp bơm keo góc 2 thành phần chuyên dụng Xylotex, giúp góc cửa cứng vững và kín khít nước mưa.", bold_prefix="Ép góc thủy lực bơm keo đặc chủng: ", is_bullet=True)
    add_para("Sở hữu máy uốn vòm chuyên dụng uốn cong thanh nhôm profile bản lớn theo các hình cung tròn, hình elip nghệ thuật mà không làm biến dạng mặt cắt nhôm.", bold_prefix="Uốn vòm nghệ thuật CNC: ", is_bullet=True)

    add_heading_2("TRỤ CỘT 03: PRECISION INSTALLATION (Thi Công & Lắp Dựng Hiện Trường)")
    add_para("Định vị khung bao cửa vào tường bằng máy cân mực laser 5 tia siêu sáng. Đảm bảo độ dung sai thẳng đứng và nằm ngang dưới 1mm.", bold_prefix="Cân chỉnh laser chuẩn xác: ", is_bullet=True)
    add_para("Liên kết khung bao vào bê tông bằng vít nở inox chịu lực chống rỉ sét. Bơm keo bọt nở cách âm chịu nước vào khe hở tường và khung bao để triệt tiêu tiếng ồn.", bold_prefix="Liên kết chịu lực & Bơm keo bọt nở: ", is_bullet=True)
    add_para("Đội ngũ thợ lắp ráp lành nghề trực tiếp thực hiện việc đi keo silicone ngoài trời đạt đường nét láng mịn, phẳng lỳ như robot bắn, ngăn nước mưa tuyệt đối.", bold_prefix="Kỹ nghệ bắn silicone chuyên nghiệp: ", is_bullet=True)
    add_para("Hiệu chỉnh tay nắm, bản lề, chốt đa điểm ôm khít gioăng cao su EPDM. Nghiệm thu cửa đóng mở êm ái, nhẹ nhàng, không phát ra tiếng cọ kẹt âm thanh.", bold_prefix="Tinh chỉnh vận hành & Nghiệm thu: ", is_bullet=True)

    add_heading_2("TRỤ CỘT 04: LIFECYCLE MAINTENANCE (Bảo Trì & Vận Hành Trọn Đời)")
    add_para("Định kỳ 1 năm/lần cử kỹ sư đến kiểm tra độ mỏi bản lề, gioăng cao su, đường keo silicone và bảo dưỡng tra dầu mỡ phụ kiện cửa miễn phí trong thời gian bảo hành.", bold_prefix="Dịch vụ bảo trì định kỳ hàng năm: ", is_bullet=True)
    add_para("Cung cấp đội phản ứng nhanh xử lý các sự cố cửa kẹt khóa, xệ cánh, rò rỉ nước hoặc vỡ kính do va đập thiên tai trong vòng 24 giờ kể từ khi nhận thông tin.", bold_prefix="Sửa chữa sự cố khẩn cấp 24/7: ", is_bullet=True)
    add_para("Tính toán phương án gia cường cửa, thay thế kính thường sang kính hộp cách âm cách nhiệt tốt hơn, hoặc thi công mở rộng thêm các hệ cửa mới mà không ảnh hưởng kết cấu cũ.", bold_prefix="Cải tạo, Nâng cấp hệ thống cửa: ", is_bullet=True)

    # --- PHẦN IV: CHIẾN LƯỢC PHÂN TÁCH ĐỐI TƯỢNG KHÁCH HÀNG ---
    add_heading_1("PHẦN IV: CHIẾN LƯỢC PHÂN TÁCH ĐỐI TƯỢNG KHÁCH HÀNG")
    add_para("Để tối ưu hóa chiến dịch truyền thông và bán hàng, NKSV phân chia khách hàng nhôm kính thành 2 phân khúc chiến lược:")
    
    add_heading_2("1. Phân khúc Khách hàng Dự án & Công nghiệp (B2B)")
    add_para("Chủ đầu tư tòa nhà cao ốc, khách sạn nghỉ dưỡng, Tổng thầu xây dựng cần hoàn thiện gói thầu mặt dựng kính curtain wall, vách kính spider sân bay, cửa nhôm kính dự án căn hộ.", bold_prefix="Đối tượng: ")
    add_para("Năng lực hồ sơ thầu, giá thành thầu cạnh tranh, giấy chứng nhận CO/CQ thanh nhôm và phụ kiện chính hãng, chứng nhận thí nghiệm kín nước kín khí cản gió của hệ cửa, năng lực tài chính và tiến độ sản xuất chế tạo hàng loạt tại xưởng.", bold_prefix="Tiêu chí quyết định: ")
    add_para("Sử dụng văn phong kỹ thuật, chuyên nghiệp. Tập trung vào năng lực sản xuất máy móc xưởng, tiến độ thi công thực tế tại công trường, và các chứng chỉ thử nghiệm mẫu cửa đạt chuẩn.", bold_prefix="Cách trình bày nội dung: ")

    add_heading_2("2. Phân khúc Khách hàng Kiến trúc & Dân dụng cao cấp (B2C & B2B2C)")
    add_para("Chủ biệt thự lâu đài, siêu biệt thự phố hiện đại, penthouse, resort nghỉ dưỡng cao cấp; các văn phòng kiến trúc sư thiết kế nội ngoại thất sang trọng.", bold_prefix="Đối tượng: ")
    add_para("Độ hoàn thiện thẩm mỹ tinh xảo từng góc ghép nhôm, công nghệ màu nhôm độc lạ sang trọng (Anodized ED sâm panh, ghi metallic), phụ kiện tay nắm CMECH kiểu dáng Mỹ mạ vàng, khả năng cách âm cách nhiệt tuyệt đối, và sự uy tín tận tâm của dịch vụ chăm sóc khách hàng.", bold_prefix="Tiêu chí quyết định: ")
    add_para("Tập trung vào hình ảnh chụp thực tế sắc nét cận cảnh các góc cửa kín khít, đường keo silicone láng mịn, video vận hành cửa trượt siêu mỏng nhẹ nhàng êm ái, các công trình biệt thự sang trọng đã thi công thực tế.", bold_prefix="Cách trình bày nội dung: ")

    # --- PHẦN V: DANH MỤC PHÂN PHỐI & THƯ VIỆN 19 HỆ NHÔM CHI TIẾT ---
    add_heading_1("PHẦN V: DANH MỤC PHÂN PHỐI & THƯ VIỆN 19 HỆ NHÔM CHI TIẾT")
    add_para("NKSV làm chủ công nghệ gia công, lắp đặt và là đối tác cung cấp của 19 hệ nhôm lớn trên thị trường hiện nay. Dưới đây là bảng phân tách chi tiết hệ thống nhôm phân chia theo phân khúc và đặc tính kỹ thuật tiêu biểu:")

    # 19 Systems Table
    table_nhom = doc.add_table(rows=20, cols=5)
    set_table_borders(table_nhom, color="CCCCCC")
    
    # Table widths setup
    widths = [Inches(0.4), Inches(1.2), Inches(1.0), Inches(2.0), Inches(2.2)]
    
    headers_nhom = ["STT", "Hãng Nhôm", "Phân Khúc", "Hệ Nhôm Chủ Lực & Độ Dày", "Đặc Trưng Công Nghệ & Ứng Dụng Tiêu Biểu"]
    for col_idx, h_text in enumerate(headers_nhom):
        cell = table_nhom.cell(0, col_idx)
        set_cell_background(cell, "006666")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    data_nhom = [
        ("1", "Civro (Đức)", "Siêu Cao Cấp\n(Luxury)", "Hệ mở quay C70, Hệ lùa R120, Hệ xếp trượt SS80\n(Độ dày: 2.0mm - 3.0mm)", "Rãnh C chuẩn Châu Âu, tích hợp phụ kiện CMECH/Sobinco. Thiết kế cầu cách nhiệt, khóa vân tay. Ứng dụng: Biệt thự siêu sang, penthouse hiện đại."),
        ("2", "Xingfa (QĐ)", "Trung Cấp", "Hệ 55 (1.4-2.0mm), Hệ 93 lùa (2.0mm), Hệ 65 mặt dựng (2.5mm)", "Nhôm Xingfa Quảng Đông nhập khẩu tem đỏ. Sơn tĩnh điện bền bỉ. Độ nhận diện cao nhất thị trường. Phù hợp nhà dân, văn phòng, dự án chung cư."),
        ("3", "Maxpro.JP", "Cao Cấp / Luxury", "Hệ R55, R65 (1.4-1.6mm), Hệ R70 vòm (1.8-2.0mm), Hệ lùa 115 (2.0mm)", "Công nghệ mạ điện di Anodise ED Nhật Bản (JIS H8602). Chống bám vân tay, chống trầy xước, bền màu 40 năm. Phù hợp biệt thự biển, resort duyên hải."),
        ("4", "Hondalex", "Cao Cấp / Luxury", "Hệ LV56, LV60, LV120, LV150\n(Độ dày: 1.5mm - 2.0mm)", "Nhôm Nhật Bản do nhà máy Long Vân sản xuất. Công nghệ phủ màu kim loại anode cao cấp. Chống ăn mòn muối biển tuyệt đối. Phù hợp biệt thự cao cấp."),
        ("5", "Kogen", "Cao Cấp", "Hệ lùa Slim 138, Hệ xếp trượt rãnh C, Cửa sổ đồng bộ", "Thiết kế Đức phong cách tối giản (Minimalist) cánh mỏng chịu lực. Rãnh C đồng bộ phụ kiện Châu Âu. Tạo tầm nhìn Panorama tối đa cho biệt thự."),
        ("6", "KOSO", "Cao Cấp / Luxury", "Hệ K55, K65, K120, K150\n(Độ dày: 1.6mm - 2.2mm)", "Thương hiệu nhôm Nhật Bản cao cấp. Profile chắc khỏe, chất lượng mạ Anodized xuất sắc, thiết kế thẩm mỹ cao, cách âm vượt trội."),
        ("7", "Soco", "Cao Cấp / Luxury", "Hệ S55, S65, S120, S180 xếp trượt\n(Độ dày: 1.8mm - 2.5mm)", "Nhôm Soco Yongli xi mạ Anodized ED chất lượng cao rãnh C Châu Âu. Vận hành cực kỳ êm ái, đóng mở đầm tay. Thiết kế sang trọng cho biệt thự."),
        ("8", "Slim (Hệ nhôm)", "Cao Cấp / Luxury", "Hệ lùa siêu mỏng nội thất, Hệ lùa siêu mỏng ngoại thất (1.6-2.5mm)", "Profile nhôm siêu mỏng tối giản. Phụ kiện bánh xe chịu tải trọng lớn treo trên hoặc ray dưới chìm đất. Ứng dụng: Cửa thông phòng, vách ngăn phòng."),
        ("9", "PMI (Malaysia)", "Cao Cấp", "Hệ PE55 mở quay, Hệ PG70 lùa bản dày\n(Độ dày: 1.4mm - 2.0mm)", "Thương hiệu cao cấp của Tập đoàn Press Metal (Malaysia). Thanh nhôm dày cứng vững, xử lý góc ép khít kín, sơn phủ bền bỉ chịu thời tiết nóng ẩm."),
        ("10", "PAG", "Cao Cấp", "Hệ P55 rãnh C, Hệ P90 lùa\n(Độ dày: 1.6mm - 2.0mm)", "Nhôm rãnh C Châu Âu thiết kế thông minh, tương thích hoàn hảo các dòng phụ kiện CMECH/Roto chốt đa điểm an ninh. Thẩm mỹ tinh tế, sắc nét."),
        ("11", "Eurowindow", "Cao Cấp", "Hệ nhôm cầu cách nhiệt EA55, EA90, Mặt dựng kính", "Thương hiệu hàng đầu Việt Nam lâu năm. Cửa nhôm cầu cách nhiệt tích hợp kính hộp cách âm cách nhiệt tối đa. Phù hợp tòa nhà văn phòng, biệt thự cao cấp."),
        ("12", "Topal (Austdoor)", "Trung - Cao Cấp", "Hệ Topal Prima 55, Topal Slim 55, Topal Prima lùa 120", "Thương hiệu của tập đoàn Austdoor. Thiết kế đồng bộ gioăng EPDM kép kín khít cản nước tốt, khóa chốt thiết kế riêng của Austdoor. Phù hợp nhà dân dụng."),
        ("13", "Viralwindow", "Trung - Cao Cấp", "Hệ V55 mở quay, Hệ V93 lùa, Hệ xếp trượt rãnh C", "Thương hiệu Việt Nam chất lượng cao. Bề mặt sơn bảo hành lên đến 20 năm của AkzoNobel. Tích hợp đồng bộ phụ kiện đồng bộ cao cấp. Phù hợp biệt thự phố."),
        ("14", "Owin", "Trung - Cao Cấp", "Hệ mở quay O55, Cửa lùa O90, Hệ cửa xếp trượt O80", "Thiết kế nhôm chắc khỏe, hệ gioăng cao su đàn hồi cao, đóng mở êm ái, chống ồn tốt. Phù hợp nhà phố biệt thự phân khúc trung-cao."),
        ("15", "PMA", "Phổ Thông / Trung", "Hệ PMA 55 vát cạnh, PMA Platinum, PMA hệ lùa", "Nhôm Việt Nam thiết kế vát cạnh nghệ thuật giúp tiết kiệm phôi nhưng vẫn chịu lực tốt, kinh tế. Phù hợp nhà phố bình dân, công trình dân dụng."),
        ("16", "Yangli", "Phổ Thông / Trung", "Hệ Y55 vát cạnh, Y93 lùa\n(Độ dày: 1.2mm - 1.4mm)", "Hệ nhôm vát cạnh phân khúc trung cấp giá cả cạnh tranh cao, mẫu mã bắt mắt, sơn tĩnh điện ngoài trời bền bỉ. Phù hợp nhà dân, nhà cho thuê."),
        ("17", "EuroVN", "Phổ Thông / Trung", "Hệ EuroVN Gold 55, EuroVN vát cạnh", "Sản phẩm của nhà máy nhôm Euroha. Hệ nhôm cải tiến độ dày vừa phải, giá thành hợp lý, cạnh tranh tốt với Xingfa giá rẻ. Phù hợp nhà phố phổ thông."),
        ("18", "VietPhap", "Phổ Thông", "Hệ 4400 sổ, Hệ 4500 đi quay, Hệ 2600 lùa", "Dòng nhôm liên doanh Việt Pháp đời đầu rất phổ biến. Tiết kiệm chi phí tối đa, chất lượng ổn định. Phù hợp nhà cấp 4, phòng trọ, cửa sổ phụ."),
        ("19", "VietY", "Phổ Thông", "Hệ nhôm Việt Ý mở quay, hệ cửa lùa", "Hệ nhôm nội địa Việt Ý thiết kế hiện đại, giá thành rẻ, phù hợp lắp dựng các hạng mục cửa trong nhà hoặc dự án chi phí thấp.")
    ]

    for row_idx, row_data in enumerate(data_nhom, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table_nhom.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            
            # Formatting cells based on data type
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(9)
            elif col_idx == 1:
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.bold = True
            elif col_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(8.5)
                # Highlight luxury systems
                if "Luxury" in text:
                    run.font.color.rgb = SECONDARY_COLOR
                    run.font.bold = True
            else:
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(8.5)
                
            # Zebra striping for table rows
            if row_idx % 2 == 0:
                set_cell_background(cell, "F9FBFB")
            else:
                set_cell_background(cell, "FFFFFF")

    # Set cell widths
    for row in table_nhom.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    add_para("", space_after=12) # Spacer

    # --- PHẦN VI: CHÍNH SÁCH BẢO MẬT & KIỂM SOÁT CHẤT LƯỢNG ---
    add_heading_1("PHẦN VI: CHÍNH SÁCH BẢO MẬT & KIỂM SOÁT CHẤT LƯỢNG (TIÊU CHUẨN TRÌNH BÀY WEB)")
    
    add_heading_2("1. Cam kết bảo mật thông tin dự án đầu tư (NDA Commitment)")
    add_callout('"Nhôm Kính Sao Vàng cam kết bảo vệ tuyệt đối tất cả các tài liệu kỹ thuật, bản vẽ biện pháp thi công mặt dựng kính và hồ sơ dự án của khách hàng. Chúng tôi hiểu rằng bản vẽ kiến trúc và kết cấu mặt dựng là tài sản trí tuệ quan trọng của chủ đầu tư. NKSV sẵn sàng ký kết Thỏa thuận bảo mật thông tin (NDA) trước khi tiếp nhận bản vẽ AutoCAD chi tiết nhằm loại bỏ hoàn toàn rủi ro rò rỉ thông tin dự án."')
    
    add_heading_2("2. Chính sách quản lý chất lượng nhôm kính (Quality Assurance Policy)")
    add_para("Chỉ nhập khẩu thanh nhôm và phụ kiện từ nhà sản xuất gốc có đầy đủ chứng chỉ CO/CQ. Nói không với hàng nhái cỏ, hàng kém chất lượng trôi nổi trên thị trường.", bold_prefix="Cam kết chính hãng 100%: ", is_bullet=True)
    add_para("Khung cửa sau khi cắt và ép góc đều được đo kiểm bằng thước điện tử đảm bảo sai số đường chéo góc dưới 1mm. Mối ép góc đúc ép thủy lực bơm keo kín khít tuyệt đối, không có khe hở ánh sáng.", bold_prefix="Kiểm soát dung sai liên kết chặt chẽ: ", is_bullet=True)
    add_para("Cam kết bảo hành lớp sơn tĩnh điện/xi mạ màu sắc bề mặt nhôm lên tới 40 năm đối với hệ nhôm Maxpro và 10-20 năm đối với các dòng nhôm cao cấp khác.", bold_prefix="Bảo hành màu sắc bề mặt dài hạn: ", is_bullet=True)
    add_para("Bảo hành 3 năm nước mưa không rò rỉ, thấm ngấm qua các khe cửa và đường silicon hoàn thiện.", bold_prefix="Bảo hành chống thấm dột 3 năm: ", is_bullet=True)

    # --- PHẦN VII: KẾ HOẠCH BỐ CỤC NỘI DUNG TRÊN WEBSITE MỚI ---
    add_heading_1("PHẦN VII: KẾ HOẠCH BỐ CỤC NỘI DUNG TRÊN WEBSITE MỚI")
    add_para("Bố cục website mới của NKSV được thiết kế thân thiện với người dùng B2B và B2C cao cấp:")
    
    add_para("Banner giới thiệu thông điệp 'Giải pháp cửa & mặt dựng nhôm kính kiến trúc tổng thể', slide hình ảnh công trình biệt thự mặt biển thực tế sang trọng, tóm tắt 5 USPs nổi bật, các biểu tượng 4 trụ cột dịch vụ, đánh giá từ kiến trúc sư đối tác và nút nhận báo giá nhanh.", bold_prefix="Trang chủ Nhôm Kính (Homepage nhom-kinh): ", is_bullet=True)
    
    add_para("Nơi số hóa toàn bộ dữ liệu của 19 dòng nhôm. Mỗi hệ nhôm sẽ có 1 trang con giới thiệu chi tiết về thông số kỹ thuật độ dày nhôm, các hệ profile chủ lực, bảng màu xi mạ/sơn tĩnh điện tiêu biểu, nút tải Catalogue PDF chính hãng và nút tải File AutoCAD mặt cắt nhôm phục vụ kiến trúc sư vẽ thiết kế.", bold_prefix="Trang Thư viện hệ nhôm (thu-vien-he-nhom.html): ", is_bullet=True)
    
    add_para("Diễn giải sâu sắc về quy trình EPCM khép kín: Architectural Engineering -> Advanced Manufacturing -> Precision Installation -> Lifecycle Maintenance kèm video thực tế gia công ép góc CNC và thợ đi silicone tại công trình.", bold_prefix="Trang Lĩnh vực hoạt động (linh-vuc-hoat-dong.html): ", is_bullet=True)
    
    add_para("Thư viện ảnh thực tế phân chia rõ theo các hạng mục: Siêu biệt thự hiện đại, Lâu đài tân cổ điển uốn vòm, Resort ven biển chống mặn, Hệ mặt dựng kính lớn tòa nhà, Mái đón kính cường lực kết hợp giàn thép đỡ.", bold_prefix="Trang Dự án tiêu biểu (du-an.html): ", is_bullet=True)
    
    add_para("Tích hợp công cụ tính toán nhanh sơ bộ chi phí dựa trên diện tích cửa và phân khúc nhôm (Phổ thông - Trung - Cao cấp), kết hợp nút upload bản vẽ AutoCAD/PDF để nhận báo giá chi tiết, cam kết NDA bảo mật hiển thị rõ bên cạnh nút gửi.", bold_prefix="Form Báo giá thông minh (Get a Quote Form): ", is_bullet=True)

    # --- PHẦN VIII: TỔNG KẾT CHIẾN LƯỢC ---
    add_heading_1("PHẦN VIII: TỔNG KẾT CHIẾN LƯỢC PHÁT TRIỂN")
    
    add_heading_2("1. Nhôm Kính Sao Vàng giỏi nhất ở đâu? (Core Expertise)")
    add_para("NKSV tự hào dẫn đầu thị trường và thực hiện tốt nhất ở 3 mảng kỹ thuật cốt lõi sau:")
    add_para("Chúng tôi sở hữu máy móc chuyên dụng và kỹ nghệ uốn cong profile nhôm bản lớn (nhôm dày từ 1.8mm - 2.5mm của Xingfa, Maxpro, Civro) tạo thành các hệ vách kính vòm cong cổ điển, elip nghệ thuật với sai số dưới 1mm, không hề có hiện tượng nhăn nheo bề mặt xi mạ sơn nhôm hoặc móp méo kết cấu cánh cửa đóng mở.", bold_prefix="Kỹ nghệ uốn vòm nhôm kính nghệ thuật độ khó cao: ", is_bullet=True)
    add_para("Là đơn vị hàng đầu làm chủ công nghệ gia công nhôm xi mạ điện di Anodize ED thế hệ mới (độ bền 40 năm) kết hợp phụ kiện rãnh C đồng bộ Châu Âu. Chúng tôi thiết kế phương án cửa trượt Slim cánh siêu mỏng chịu tải lớn và hệ xếp trượt 90 độ không cần cột góc nâng tầm không gian panorama.", bold_prefix="Gia công lắp đặt hệ nhôm xi mạ Anodized ED & Cửa Slim thế hệ mới: ", is_bullet=True)
    add_para("Nhờ kế thừa thế mạnh kết cấu cơ khí thép của CKSV, NKSV là đơn vị hiếm hoi thi công hoàn mỹ các hạng mục kết hợp đa nguyên vật liệu: Mặt dựng kính lớn (Curtain Wall) hoặc vách kính chân nhện Spider nhịp lớn liên kết trực tiếp vào khung dầm thép đỡ chịu lực phức tạp của sảnh đón tòa nhà.", bold_prefix="Tích hợp đồng bộ Hệ kết cấu thép chịu lực và Vách kính lớn: ", is_bullet=True)

    add_heading_2("2. Tại sao khách hàng phải chọn Nhôm Kính Sao Vàng? (Competitive Advantages)")
    add_para("Khách hàng chủ đầu tư, tổng thầu và kiến trúc sư lựa chọn hợp tác với NKSV vì những điểm khác biệt duy nhất sau:")
    add_para("Danh mục 19 hãng nhôm đầy đủ từ bình dân đến luxury giúp khách hàng thoải mái so sánh lựa chọn tại một đối tác duy nhất. Kỹ sư NKSV tư vấn trung thực theo công năng thực tế và ngân sách đầu tư chứ không ép khách hàng mua một loại nhôm cố định.", bold_prefix="Hệ sinh thái sản phẩm phong phú nhất thị trường (19 hệ nhôm): ", is_bullet=True)
    add_para("Mọi góc cửa đều được ép góc thủy lực lực ép 15 tấn kết hợp ke góc đúc cơ học dày 4mm và bơm keo 2 thành phần đặc chủng. Kết hợp gioăng EPDM 3 lớp khép kín giúp giải quyết triệt để 2 vấn nạn phổ biến của cửa nhôm kính là ngấm nước mưa qua khe góc và xệ cánh cửa sau vài năm sử dụng.", bold_prefix="Triết lý thiết kế DFM chống xệ cánh và chống thấm dột tuyệt đối: ", is_bullet=True)
    add_para("Chúng tôi công khai xuất xứ, cung cấp chứng chỉ CO/CQ chính hãng của thanh nhôm và hệ phụ kiện khóa bản lề cao cấp. NKSV cam kết hoàn tiền gấp đôi nếu khách hàng phát hiện hàng nhái Kinlong hay Xingfa giả tem trên công trình.", bold_prefix="Minh bạch nguồn gốc xuất xứ vật tư và nói không với hàng giả: ", is_bullet=True)
    add_para("Cam kết bảo hành nước mưa không thấm qua khe cửa lên đến 3 năm, bảo hành lớp xi mạ màu sắc bề mặt lên đến 40 năm và đặc biệt cung cấp dịch vụ bảo dưỡng cửa định kỳ hàng năm giúp cửa luôn hoạt động trơn tru êm ái.", bold_prefix="Chính sách bảo hành bề mặt lên đến 40 năm & Bảo trì định kỳ trọn đời: ", is_bullet=True)
    add_para("Khả năng uốn vòm nghệ thuật, làm mặt dựng kính nhịp lớn kết hợp khung thép đỡ giúp NKSV đáp ứng trọn gói mọi yêu cầu kiến trúc sáng tạo nhất của KTS mà các xưởng cửa thông thường không thể nhận thầu.", bold_prefix="Năng lực thi công các hạng mục nhôm kính phi tiêu chuẩn siêu khó: ", is_bullet=True)

    # Save document
    output_filename = "BaoCao_ChienLuoc_TaiDinhVi_NKSV.docx"
    doc.save(output_filename)
    print(f"Document successfully created and saved as: {output_filename}")

if __name__ == "__main__":
    main()
