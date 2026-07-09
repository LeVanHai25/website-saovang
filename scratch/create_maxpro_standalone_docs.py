# -*- coding: utf-8 -*-
"""
THƯ VIỆN HỆ NHÔM SAO VÀNG
Script tạo/ghi đè 2 tài liệu độc lập cho hãng Maxpro JP dựa trên catalogue mới:
1. Catalogue_HeCuaNhom_Maxpro.xlsx
2. ThuVien_HeCuaNhom_Maxpro.docx
Phiên bản 1.0
"""
import os, sys
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

BASE = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"

thin_b = Border(
    left=Side(style="thin", color="DDDDDD"),
    right=Side(style="thin", color="DDDDDD"),
    top=Side(style="thin", color="DDDDDD"),
    bottom=Side(style="thin", color="DDDDDD"),
)

# ─────────────────────────────────────────────
# 1. TẠO Catalogue_HeCuaNhom_Maxpro.xlsx
# ─────────────────────────────────────────────
def create_maxpro_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Catalogue Maxpro JP"
    
    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "DANH MỤC HỆ THỐNG CỬA NHÔM ANODISE ED NHẬT BẢN MAXPRO.JP"
    title_cell.font = Font(name="Arial", bold=True, size=14, color="0D2240")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 40
    
    # Header row
    headers = [
        'STT', 'Nhóm Hệ', 'Mã Hệ', 'Tên Hệ', 
        'Loại Sản Phẩm', 'Độ Dày (mm)', 
        'Phụ Kiện', 'Gioăng & Keo', 'Đặc Điểm Kỹ Thuật'
    ]
    hdr_fill = PatternFill(fill_type="solid", fgColor="705D30") # Tone màu Maxpro vàng đồng
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[4].height = 35
    
    for i, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=i, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = hdr_align
        cell.border = thin_b
        
    # Maxpro Data
    maxpro_rows = [
        ('1', 'Anodized Phổ thông', 'R55', 'Cửa đi + Cửa sổ mở quay R55', 
         'Cửa đi, cửa sổ mở quay bản khung bao 55mm phổ thông, cánh trơn phẳng tối giản kiểu Nhật Bản.', 
         '1.4mm', 'CMECH (Mỹ) / BOGO rãnh C đồng bộ', 'Gioăng EPDM 2 khoang / Silicon trung tính', 
         'Anodized ED mạ điện di chuẩn Nhật JIS H8602 bền màu 40 năm.'),
         
        ('2', 'Anodized Cao cấp', 'R65', 'Cửa đi + Cửa sổ mở quay R65', 
         'Hệ cửa sổ mở quay hất hoặc cửa đi bản rộng 65mm dày dặn, tương thích nhiều loại kính hộp.', 
         '1.6mm', 'Phụ kiện ROTO (Đức) / CMECH chốt đa điểm', 'Gioăng EPDM đa khoang / Silicon Sika', 
         'Màu sắc Champagne/Vàng Gold/Nâu Cafe cực kỳ sang trọng quý phái.'),
         
        ('3', 'Anodized Siêu cao cấp', 'R65 Plus', 'Cửa sổ mở quay R65 Plus cải tiến', 
         'Cửa sổ mở quay cải tiến thêm lớp gioăng trung tâm cách âm cách nước vượt trội hoàn toàn.', 
         '1.6mm', 'CMECH / Roto rãnh C Châu Âu cao cấp', 'Gioăng EPDM 3 lớp cải tiến', 
         'Công nghệ xi mạ ED ngăn bám bụi vân tay bề mặt hoàn mỹ.'),
         
        ('4', 'Tân cổ điển Cao cấp', 'R70', 'Cửa đi tân cổ điển vát cạnh R70', 
         'Cửa đi mở quay bản khung 70mm với thiết kế vát cạnh đặc trưng tân cổ điển tạo hiệu ứng 3D.', 
         '2.0mm', 'Bản lề chịu tải lớn CMECH / khóa đa điểm', 'Gioăng EPDM kép chèn ép lực chìm', 
         'Tạo điểm nhấn cổ điển sang trọng cho biệt thự, lâu đài Pháp.'),
         
        ('5', 'Tân cổ điển Siêu cao cấp', 'R83', 'Cửa đi mở quay vòm cong R83', 
         'Cửa đi vòm cung uốn nghệ thuật bản khung 83mm siêu dày, chịu lực nén ép góc thủy lực cao.', 
         '1.8mm', 'Bản lề CMECH & khóa tay nắm đồng bộ cổ điển', 'Gioăng EPDM đa khoang / Dow Corning', 
         'Chuyên biệt cho các hệ cửa vòm lớn lâu đài bán cổ điển.'),
         
        ('6', 'Thủy lực Siêu cao cấp', 'R200 Thủy lực', 'Cửa đi thủy lực đại sảnh R200', 
         'Cửa đi thủy lực bản cánh rộng 200mm siêu nặng, kính cường lực lớn, đóng mở piston thủy lực sàn.', 
         '2.0mm', 'Piston thủy lực Dorma / DICTATOR chịu tải 800kg', 'Silicon kết cấu chuyên dụng + Gioăng EPDM', 
         'Lối đi đại sảnh uy nghi tráng lệ biệt thự, showroom xe sang.'),
         
        ('7', 'Trượt lùa Trung cấp', 'SW55', 'Cửa sổ trượt lùa SW55', 
         'Cửa sổ trượt lùa nhẹ nhàng tối ưu không gian cho nhà phố chung cư phổ thông.', 
         '1.4mm', 'Bánh xe chịu lực & tay khóa Bogo/Kinlong', 'Gioăng EPDM chèn nỉ chống ồn', 
         'Mạ màu Anodized ED chống muối mặn han rỉ ven biển.'),
         
        ('8', 'Trượt lùa Cao cấp', 'SD83', 'Cửa đi trượt phẳng SD83', 
         'Cửa đi trượt phẳng lướt nhẹ êm trên ray inox SUS304 chống mài mòn chịu lực tốt.', 
         '1.6mm', 'Ray inox SUS304 & bánh xe chịu tải CMECH', 'Gioăng EPDM kép / Keo dán Apollo A500', 
         'Thẩm mỹ phẳng tinh tế kiểu Nhật Bản hiện đại gọn gàng.'),
         
        ('9', 'Trượt lùa Siêu cao cấp', 'SD115 Chống bão', 'Cửa đi trượt lùa chịu gió bão SD115', 
         'Cửa đi trượt lùa bản dày 115mm chịu bão gió cấp 14, gioăng nén kín nước tiêu chuẩn khắt khe.', 
         '2.0mm - 2.5mm', 'Ray inox SUS316L & bánh xe đặc chủng chịu bão', 'Gioăng EPDM cầu cách nhiệt đa khoang', 
         'Chuyên biệt biệt thự mặt biển, resort duyên hải chịu bão gió.'),
         
        ('10', 'Xếp trượt Cao cấp', 'SFD80', 'Cửa đi xếp trượt nhiều cánh SFD80', 
         'Cửa đi xếp gấp trượt dồn cánh mở thoáng 100% lối ra ban công sân vườn rộng biệt thự.', 
         '1.8mm', 'Hệ bản lề xếp & bánh xe treo OPK/Roto', 'Gioăng EPDM viền mềm đàn hồi', 
         'Bản lề treo chịu lực xoay đóng mở nhẹ nhàng bền bỉ.')
    ]
    
    fill_row = PatternFill(fill_type="solid", fgColor="FFF9C4") # Vàng chanh nhạt tinh tế
    for r_idx, row_vals in enumerate(maxpro_rows, 5):
        ws.row_dimensions[r_idx].height = 40
        for c_idx, val in enumerate(row_vals, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill = fill_row
            cell.font = Font(name="Arial", size=9.5)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_b
            if c_idx in (1, 3, 6):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
    # Column widths
    col_widths = [6, 18, 20, 24, 38, 14, 28, 25, 30]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=4, column=i).column_letter].width = w
        
    path = os.path.join(BASE, "Catalogue_HeCuaNhom_Maxpro.xlsx")
    try:
        wb.save(path)
        print(f"  >> Excel Maxpro saved: {path}")
    except PermissionError:
        alt_path = path.replace(".xlsx", "_temp.xlsx")
        wb.save(alt_path)
        print(f"  [!] Locked. Saved Excel Maxpro to: {alt_path}")

# ─────────────────────────────────────────────
# 2. TẠO ThuVien_HeCuaNhom_Maxpro.docx
# ─────────────────────────────────────────────
def create_maxpro_docx():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)
        
    def add_p(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return p

    # Bìa tài liệu
    add_p("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x0D,0x22,0x40))
    add_p("THƯ VIỆN KỸ THUẬT & CATALOGUE\nHỆ CỬA NHÔM ANODISE ED NHẬT BẢN MAXPRO.JP", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x70,0x5D,0x30), space_before=15)
    add_p("Tài liệu kỹ thuật và hướng dẫn tra cứu các hệ thống cửa nhôm công nghệ Anodise ED bền màu 40 năm\nTham chiếu catalogue Đại Phúc mới nhất năm 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_p("Năm 2026 - Lưu hành nội bộ", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    
    # Giới thiệu
    add_p("GIỚI THIỆU THƯƠNG HIỆU NHÔM CAO CẤP MAXPRO.JP (NHẬT BẢN)", bold=True, size=14, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    add_p(
        "Maxpro.JP là dòng cửa nhôm hệ cao cấp được sản xuất trên dây chuyền công nghệ mạ điện di Anodise ED hiện đại bậc nhất theo tiêu chuẩn công nghiệp Nhật Bản JIS H8602 (Class A1). "
        "Với công nghệ này, thanh nhôm được bao phủ một lớp bảo vệ siêu mịn chống bám vân tay, chống trầy xước, ngăn ngừa hoàn toàn sự ăn mòn của axit và muối mặn biển. "
        "Maxpro tự hào cam kết độ bền màu lên đến 40 năm, là sự lựa chọn tối ưu cho các công trình biệt thự mặt biển, resort cao cấp duyên hải hoặc siêu biệt thự phố thị hiện đại. "
        "Hệ thống nhôm Maxpro rãnh C Châu Âu tương thích hoàn hảo với các dòng phụ kiện đỉnh cao thế giới như CMECH (Mỹ), Roto (Đức), Bogo, Sigico đem lại trải nghiệm vận hành êm ái hoàn mỹ."
    )
    
    # Chi tiết hệ
    add_p("CHI TIẾT THÔNG SỐ & ỨNG DỤNG CÁC HỆ NHÔM CHỦ LỰC MAXPRO", bold=True, size=14, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    
    add_p("1. Hệ cửa mở quay R55 / R65 / R65 Plus", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.4mm (hệ R55) đến 1.6mm (hệ R65 và R65 Plus).")
    add_bullet("Hệ R65 Plus cải tiến vượt bậc bằng cách bố trí thêm gioăng trung tâm cách âm, cách nước vượt trội.")
    add_bullet("Rãnh C chuẩn Châu Âu tương thích phụ kiện CMECH/Roto chốt đa điểm an ninh cao chống cạy khóa.")
    
    add_p("2. Hệ cửa tân cổ điển uốn vòm R70 / R83", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 1.8mm - 2.0mm cứng vững.")
    add_bullet("Đặc trưng: Profile vát cạnh nghệ thuật tạo hiệu ứng 3D đậm chất tân cổ điển Pháp sang trọng.")
    add_bullet("Thiết kế vòm: Thích hợp uốn các vách kính vòm đỉnh đầu sang trọng biệt thự lâu đài lớn.")
    
    add_p("3. Hệ cửa thủy lực đại sảnh R200", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Độ dày nhôm: 2.0mm siêu dày vách.")
    add_bullet("Cấu trúc: Bản cánh siêu rộng 200mm đỡ kính cường lực khổ lớn đại sảnh, đóng mở piston thủy lực sàn êm ái.")
    add_bullet("Tải trọng cánh: Lên đến 800kg/cánh chuyên biệt cho cửa đại sảnh uy nghi.")

    add_p("4. Hệ cửa trượt lùa SD83 / SD115 (Chống bão)", bold=True, size=12, color=RGBColor(0x70,0x5D,0x30), space_before=6)
    add_bullet("Hệ trượt phẳng SD83: Ray inox SUS304 mượt mà, cánh mỏng tối giản kiểu Nhật Bản.")
    add_bullet("Hệ chịu bão SD115: Vách nhôm dày 2.0mm-2.5mm chịu được sức gió bão giật cấp 14 ngoài khơi biển.")
    add_bullet("Ứng dụng: Lối ra ban công hướng biển biệt thự, resort biển duyên hải.")

    # Bảng tóm tắt
    add_p("BẢNG TỔNG HỢP THÔNG SỐ KỸ THUẬT HỆ NHÔM MAXPRO.JP", bold=True, size=12, color=RGBColor(0x0D,0x22,0x40), space_before=12)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    
    hdr = tbl.rows[0].cells
    hdr[0].text = "STT"
    hdr[1].text = "Mã Hệ"
    hdr[2].text = 'Tên Hệ'
    hdr[3].text = 'Độ Dày (mm)'
    hdr[4].text = "Loại Sản Phẩm & Công Năng Tiêu Biểu"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    maxpro_tbl_data = [
        ("1", "R55", "Cửa đi & sổ mở quay hệ 55", "1.4mm", "Cửa đi, cửa sổ mở quay bản khung 55mm tối giản Nhật Bản."),
        ("2", "R65", "Cửa đi & sổ mở quay hệ 65", "1.6mm", "Cửa mở quay hất rãnh C Châu Âu đa dạng màu sắc mạ ED."),
        ("3", "R65 Plus", "Cửa mở quay R65 Plus cao cấp", "1.6mm", "Thêm gioăng trung tâm cách âm cách nước hoàn hảo."),
        ("4", "R70", "Cửa đi tân cổ điển R70", "2.0mm", "Khung vát cạnh tân cổ điển tạo hiệu ứng 3D sang trọng."),
        ("5", "R83", "Cửa đi uốn vòm nghệ thuật R83", "1.8mm", "Khung 83mm siêu dày thích hợp uốn vòm biệt thự Pháp."),
        ("6", "R200 Thủy lực", "Cửa đi thủy lực đại sảnh R200", "2.0mm", "Bản cánh rộng 200mm dùng piston thủy lực đóng mở tự động."),
        ("7", "SW55", "Cửa sổ trượt lùa hệ SW55", "1.4mm", "Cửa sổ lùa nhẹ nhàng tối ưu không gian nhà phố nhỏ."),
        ("8", "SD83", "Cửa đi trượt lùa phẳng hệ SD83", "1.6mm", "Cánh mỏng lướt nhẹ ray inox SUS304 phong cách Nhật."),
        ("9", "SD115", "Cửa lùa chịu bão biển SD115", "2.0-2.5mm", "Hệ lùa siêu chịu lực chịu bão cấp 14 ngoài khơi ven biển."),
        ("10", "SFD80", "Cửa đi xếp trượt hệ SFD80", "1.8mm", "Cửa đi xếp gọn cánh treo mở thoáng 100% ban công sân vườn.")
    ]
    
    for row in maxpro_tbl_data:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph()
    add_p("Tài liệu được bảo mật và tham chiếu catalogue Đại Phúc 2026 bởi Sao Vàng Group.", italic=True, size=10)
    
    path = os.path.join(BASE, "ThuVien_HeCuaNhom_Maxpro.docx")
    try:
        doc.save(path)
        print(f"  >> Docx Maxpro saved: {path}")
    except PermissionError:
        alt_path = path.replace(".docx", "_temp.docx")
        doc.save(alt_path)
        print(f"  [!] Locked. Saved Docx Maxpro to: {alt_path}")

if __name__ == "__main__":
    print("=" * 60)
    print("RECREATING standalone MAXPRO docx & xlsx in BASE directory...")
    print("=" * 60)
    create_maxpro_xlsx()
    create_maxpro_docx()
    print("=" * 60)
