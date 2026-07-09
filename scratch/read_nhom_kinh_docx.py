import os
import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document

filename = "Báo cáo nội dung chiến lược lĩnh vực nhôm kính.docx"
path = os.path.join(r"d:\Sao Vàng\Website-SaoVang", filename)
output_path = r"d:\Sao Vàng\Website-SaoVang\scratch\nhom_kinh_content.txt"

with open(output_path, "w", encoding="utf-8") as out:
    if os.path.exists(path):
        out.write(f"==================== {filename} ====================\n")
        doc = Document(path)
        for p in doc.paragraphs:
            if p.text.strip():
                out.write(p.text + "\n")
        
        # Also print tables if any
        for table in doc.tables:
            out.write("\n--- TABLE ---\n")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                out.write(" | ".join(row_text) + "\n")
        out.write("=============================================\n\n")
    else:
        out.write(f"File not found: {path}\n")

print(f"Content written to {output_path}")
