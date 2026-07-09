import os
import sys
import subprocess

# Auto-install python-docx if not present
try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_comprehensive_report():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)

    # Base Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO TOÀN DIỆN: Ý TƯỞNG THIẾT KẾ, ĐỊNH HƯỚNG NỘI DUNG & GIẢI PHÁP KỸ THUẬT WEBSITE SAO VÀNG")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(15)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Deep Navy
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(18)

    # Metadata Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    metadata = [
        ("Dự án:", "Xây dựng và Vận hành Website Thương hiệu Sao Vàng"),
        ("Đơn vị thực hiện:", "Ban kỹ thuật nội bộ (Lê Văn Hải phụ trách)"),
        ("Mục tiêu tài liệu:", "Trình bày chi tiết chiến lược nội dung, trải nghiệm người dùng (UX), cấu trúc thông tin và giải pháp công nghệ mở rộng hệ thống"),
        ("Trạng thái phê duyệt:", "Tài liệu đệ trình Ban Giám Đốc và Cố vấn chuyên môn đánh giá")
    ]
    
    for i, (label, val) in enumerate(metadata):
        row = table.rows[i]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        cell_lbl.width = docx.shared.Inches(2.2)
        cell_val.width = docx.shared.Inches(4.3)
        
        lbl_p = cell_lbl.paragraphs[0]
        lbl_run = lbl_p.add_run(label)
        lbl_run.bold = True
        lbl_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        val_p = cell_val.paragraphs[0]
        val_p.add_run(val)
        
        for cell in (cell_lbl, cell_val):
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>' % nsdecls('w'))
            tcPr.append(tcBorders)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Helper styling functions
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x5C, 0x76, 0x8D)
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        p.add_run(text)
        return p

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.bold = bold
        return p

    # ================= PHẦN I =================
    add_heading_1("PHẦN I: TẦM NHÌN CHIẾN LƯỢC, VAI TRÒ & MỤC TIÊU CỦA WEBSITE")
    
    add_heading_2("1. Vai trò của website đối với doanh nghiệp Sao Vàng")
    add_body("Đối với Công ty Sao Vàng – một đơn vị hoạt động đa lĩnh vực bao gồm gia công cơ khí kết cấu, cơ khí mỹ thuật và thi công nhôm kính cao cấp – website đóng vai trò là xương sống trong chiến lược chuyển đổi số và tiếp thị:")
    add_bullet("Văn phòng đại diện số 24/7: ", "Là bộ mặt thương hiệu chính thức của công ty trên Internet. Khách hàng, đối tác và nhà thầu có thể tìm hiểu toàn bộ thông tin pháp lý, cơ sở vật chất và hồ sơ năng lực của công ty bất kỳ lúc nào mà không cần gặp mặt trực tiếp.")
    add_bullet("Kênh phễu thu hút khách hàng chủ động: ", "Website đóng vai trò là đích đến của mọi chiến dịch quảng cáo. Đây là nơi tiếp nhận nhu cầu, phân loại thông tin khách hàng tiềm năng trước khi chuyển giao cho bộ phận kinh doanh xử lý.")
    add_bullet("Công cụ tối ưu hóa quy trình tư vấn: ", "Thay vì nhân viên phải gửi thủ công từng tệp ảnh sản phẩm hay tệp báo giá PDF nặng qua Zalo, website sẽ là thư viện mở chứa đầy đủ thông số kỹ thuật, bảng giá và hình ảnh thực tế được sắp xếp khoa học để gửi nhanh cho khách hàng.")

    add_heading_2("2. Mục tiêu chiến lược của website")
    add_bullet("Mục tiêu ngắn hạn (1-3 tháng): ", "Xây dựng hoàn thiện giao diện trực quan, thân thiện với di động; thiết lập các công cụ liên hệ nhanh và đo lường quảng cáo nhằm tối ưu hóa chi phí chạy Google Ads ngay khi website vận hành.")
    add_bullet("Mục tiêu trung hạn (3-6 tháng): ", "Đạt vị trí cao trên công cụ tìm kiếm Google (SEO) với các từ khóa mũi nhọn như: 'cầu thang xoắn Sao Vàng', 'gia công cơ khí nghệ thuật', 'cửa nhôm kính cao cấp'. Đưa website trở thành kênh tạo ra nguồn khách hàng tự nhiên ổn định.")
    add_bullet("Mục tiêu dài hạn (trên 6 tháng): ", "Tích hợp sâu hệ thống website vào các công cụ quản lý nội bộ của công ty (CRM/ERP), đồng bộ hóa dữ liệu khách hàng và thông số sản phẩm.")

    add_heading_2("3. Định hướng phát triển thương hiệu Sao Vàng trên môi trường số")
    add_body("Thương hiệu Sao Vàng trên website sẽ được định vị nhất quán qua hai giá trị cốt lõi:")
    add_bullet("Sức mạnh và Sự chính xác (Mảng Cơ khí): ", "Thể hiện qua hình ảnh nhà xưởng rộng lớn, máy móc CNC hiện đại, các kết cấu thép vững chãi. Tone giọng (Tone of voice) chuyên nghiệp, đáng tin cậy, nhấn mạnh vào chất lượng kỹ thuật.")
    add_bullet("Thẩm mỹ và Sự tinh tế (Mảng Nhôm kính & Mỹ thuật): ", "Thể hiện qua hình ảnh các công trình biệt thự sang trọng, các mẫu cửa nhôm kính cao cấp đón ánh sáng tự nhiên và các chi tiết uốn lượn nghệ thuật của cầu thang xoắn. Thiết kế giao diện thoáng đạt, hiện đại, mang tính định hướng giải pháp không gian sống.")

    # ================= PHẦN II =================
    add_heading_1("PHẦN II: ĐỐI TƯỢNG PHỤC VỤ & PHÂN NHÓM NGƯỜI DÙNG (Target Audience)")
    add_body("Website Sao Vàng được thiết kế để phục vụ ba nhóm người dùng mục tiêu với những hành vi và nhu cầu thông tin chuyên biệt:")

    add_heading_2("1. Nhóm đối tác B2B (Nhà thầu xây dựng, Chủ đầu tư dự án, Kiến trúc sư)")
    add_bullet("Đặc điểm hành vi: ", "Đòi hỏi thông tin kỹ thuật chính xác, tính pháp lý rõ ràng và năng lực sản xuất thực tế để đánh giá rủi ro trước khi hợp tác.")
    add_bullet("Thông tin cần tiếp cận: ", "Hồ sơ năng lực (Profile công ty dưới dạng PDF để tải về), chứng chỉ chất lượng sản phẩm (CO/CQ), danh mục máy móc thiết bị nhà xưởng, danh sách các dự án lớn đã hoàn thành kèm theo hình ảnh thực tế tại công trường.")

    add_heading_2("2. Nhóm khách hàng lẻ B2C (Chủ đầu tư tư nhân, Chủ nhà)")
    add_bullet("Đặc điểm hành vi: ", "Quan tâm nhiều đến tính thẩm mỹ, sự phù hợp với không gian kiến trúc và giá cả. Thường truy cập bằng thiết bị di động vào buổi tối và cần sự tư vấn nhanh chóng.")
    add_bullet("Thông tin cần tiếp cận: ", "Hình ảnh thực tế lắp đặt tại các công trình dân dụng (ảnh chụp đẹp, trực quan), bảng báo giá chi tiết rõ ràng, công cụ tính giá sơ bộ, các mẫu thiết kế xu hướng và cổng liên hệ tư vấn trực tiếp (Zalo, Hotline).")

    add_heading_2("3. Đại lý và Đối tác liên kết cung ứng")
    add_bullet("Đặc điểm hành vi: ", "Cần thông tin cập nhật về chính sách chiết khấu, thông tin phụ kiện và danh mục sản phẩm mới.")
    add_bullet("Thông tin cần tiếp cận: ", "Danh mục chi tiết sản phẩm và phụ kiện (nhôm, kính, khóa, bản lề...), chính sách bảo hành chính thức của Sao Vàng.")

    # ================= PHẦN III =================
    add_heading_1("PHẦN III: KIẾN TRÚC THÔNG TIN & SƠ ĐỒ WEBSITE (Sitemap & Content)")
    
    add_heading_2("1. Sơ đồ cấu trúc Website (Sitemap)")
    add_body("Website được tổ chức theo mô hình phân cấp tối ưu hóa trải nghiệm người dùng và SEO:")
    add_bullet("Trang Chủ (Index): ", "Tóm tắt năng lực, phân luồng dịch vụ, hiển thị dự án tiêu biểu và form liên hệ nhanh.")
    add_bullet("Trang Giới Thiệu & Năng Lực: ", "Giới thiệu lịch sử thành lập, tầm nhìn sứ mệnh, hồ sơ năng lực sản xuất (nhà xưởng, máy móc) và chứng nhận chất lượng.")
    add_bullet("Trang Lĩnh Vực Hoạt Động (Danh mục cấp 1): ", "Chia làm hai chuyên trang chuyên biệt:")
    add_bullet("   - Lĩnh vực Cơ khí (Danh mục cấp 2): ", "Gia công cơ khí kết cấu, kết cấu thép nhà xưởng và Cơ khí nghệ thuật (cầu thang xoắn, cửa cổng mỹ thuật CNC).")
    add_bullet("   - Lĩnh vực Nhôm kính (Danh mục cấp 2): ", "Cửa nhôm kính cao cấp (hệ Xingfa, PMI, Aluprof...), vách kính mặt dựng và lan can kính.")
    add_bullet("Trang Sản Phẩm & Phụ Kiện: ", "Trình bày chi tiết các mẫu cửa, mẫu cầu thang, lan can kèm theo danh mục phụ kiện đi kèm.")
    add_bullet("Trang Dự Án Thực Tế: ", "Thư viện lưu trữ các công trình đã thi công dưới dạng nhật ký hình ảnh thực tế.")
    add_bullet("Trang Bảng Báo Giá: ", "Bảng tra cứu đơn giá thi công sơ bộ các hạng mục nhôm kính và cơ khí.")
    add_bullet("Trang Liên Hệ: ", "Thông tin bản đồ nhà xưởng, văn phòng, số điện thoại hotline các bộ phận và form gửi yêu cầu.")

    add_heading_2("2. Chi tiết nội dung hiển thị tại các trang chính")
    add_bullet("Trang Chủ: ", "Banner lớn giới thiệu tổng quan nhà xưởng hoạt động; hai nút chuyển hướng lớn sang hai mảng Cơ khí và Nhôm kính; khối số liệu chứng minh năng lực; khối hiển thị 4 dự án gần nhất và form đăng ký tư vấn.")
    add_bullet("Trang Lĩnh Vực Cơ Khí: ", "Tập trung hình ảnh máy cắt CNC plasma, máy hàn laser đang hoạt động; mô tả chi tiết quy trình gia công kết cấu thép; thư viện hình ảnh các mẫu cầu thang xoắn độc bản kèm phân tích kỹ thuật về độ chịu lực và thẩm mỹ uốn.")
    add_bullet("Trang Lĩnh Vực Nhôm Kính: ", "Phân loại cửa theo cách mở (mở quay, mở lùa, xếp trượt); thông tin chi tiết về độ dày thanh nhôm, loại kính sử dụng (kính hộp cách âm, kính cường lực) và hệ phụ kiện đi kèm.")
    add_bullet("Trang Dự Án: ", "Trình bày dưới dạng bài viết ngắn gọn bao gồm: Tên công trình, địa chỉ thi công, hạng mục thực hiện, hình ảnh trước/sau khi thi công và đánh giá ngắn từ chủ nhà hoặc đại diện nhà thầu.")

    # ================= PHẦN IV =================
    add_heading_1("PHẦN IV: ĐỊNH HƯỚNG TRẢI NGHIỆM NGƯỜI DÙNG (UX) & CHIẾN LƯỢC NỘI DUNG")
    
    add_heading_2("1. Định hướng trải nghiệm người dùng (UX)")
    add_bullet("Thiết kế ưu tiên thiết bị di động (Mobile-First): ", "Hơn 80% khách hàng lẻ tìm kiếm thông tin bằng điện thoại di động. Giao diện được tối ưu hóa để các nút bấm to rõ ràng, hình ảnh sản phẩm vuốt mượt mà và form đăng ký tối giản nhất.")
    add_bullet("Tối ưu tốc độ tải trang cực nhanh: ", "Loại bỏ toàn bộ các hiệu ứng chuyển động phức tạp không cần thiết và nén tối đa tài nguyên hình ảnh. Đảm bảo thời gian phản hồi trang dưới 1.5 giây để giữ chân người dùng và tối ưu điểm chất lượng quảng cáo.")
    add_bullet("Hệ thống nút liên hệ 1-chạm thông minh: ", "Widget Hotline và Zalo luôn hiển thị ở góc dưới màn hình di động, cho phép khách hàng kết nối ngay lập tức với nhân viên tư vấn mà không cần tìm kiếm thông tin liên hệ.")
    add_bullet("Công cụ tính giá tương tác: ", "Cho phép khách hàng tự nhập kích thước và loại nhôm/kính để nhận ngay mức giá dự toán tham khảo, tạo trải nghiệm tương tác thú vị và thúc đẩy khách hàng để lại thông tin tư vấn chi tiết.")

    add_heading_2("2. Chiến lược nội dung dài hạn (Content Strategy)")
    add_bullet("Chuyên mục tư vấn kỹ thuật chuyên sâu (Blog): ", "Liên tục cập nhật các bài viết giải đáp thắc mắc của khách hàng như: Cách phân biệt nhôm thật/nhái; Nên chọn cửa lùa hay cửa mở quay cho phòng ngủ; Các lưu ý kỹ thuật khi lắp đặt cầu thang xoắn ngoài trời... Điều này giúp website tự động lên top Google không tốn phí.")
    add_bullet("Chuẩn hóa thông tin minh bạch: ", "Đăng tải rõ ràng chính sách bảo hành sản phẩm (5 năm cho kết cấu cơ khí, 2 năm cho phụ kiện nhôm kính) và cam kết tiến độ thi công nhằm củng cố lòng tin của khách hàng đối với thương hiệu Sao Vàng.")

    # ================= PHẦN V =================
    add_heading_1("PHẦN V: GIẢI PHÁP CÔNG NGHỆ & KẾ HOẠCH PHÁT TRIỂN (Tech Stack & Roadmap)")
    
    add_heading_2("1. Công nghệ áp dụng và phát triển")
    add_bullet("Frontend (Giao diện): ", "Sử dụng HTML5, CSS3 và Javascript thuần (Vanilla JS) được viết sạch, tối ưu hóa responsive hoàn toàn bằng CSS Grid/Flexbox. Việc tránh sử dụng các thư viện cồng kềnh giúp tối ưu hóa hiệu năng tối đa.")
    add_bullet("Backend & CMS (Trang quản trị): ", "Xây dựng trên nền tảng Node.js (Express.js). Đây là công nghệ hiện đại, xử lý yêu cầu bất đồng bộ cực nhanh và tiêu tốn rất ít tài nguyên máy chủ.")
    add_bullet("Khả năng tích hợp mở rộng: ", "Kiến trúc hệ thống được thiết kế dưới dạng module. Khi công ty mở rộng quy mô, hệ thống hoàn toàn sẵn sàng để mở thêm các API kết nối trực tiếp dữ liệu khách hàng về các phần mềm CRM nội bộ hoặc đồng bộ dữ liệu tồn kho từ ERP.")

    add_heading_2("2. Kế hoạch phát triển chi tiết (Roadmap)")
    add_bullet("Giai đoạn 1: Chuẩn hóa giao diện & Bố cục (Hiện tại): ", "Hoàn thiện toàn bộ khung giao diện tĩnh của các trang con, kiểm thử trải nghiệm hiển thị trên nhiều dòng điện thoại khác nhau.")
    add_bullet("Giai đoạn 2: Tích hợp dữ liệu & Tính năng tương tác (2 tuần tiếp theo): ", "Đổ dữ liệu sản phẩm, hình ảnh công trình thực tế của Sao Vàng lên hệ thống; hoàn thiện tính năng tính giá tự động và trang quản trị CMS nội bộ.")
    add_bullet("Giai đoạn 3: Kiểm thử & Vận hành chính thức (1 tuần cuối): ", "Tiến hành chạy thử nghiệm hệ thống, đo lường tốc độ tải trang, cài đặt mã theo dõi quảng cáo (Google Analytics, GTM), trỏ tên miền chính thức và bàn giao hệ thống.")

    # Save to the target path
    output_path = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ChiTiet_Website_SaoVang.docx"
    doc.save(output_path)
    print(f"Comprehensive report successfully saved to {output_path}")

if __name__ == "__main__":
    create_comprehensive_report()
