# -*- coding: utf-8 -*-
import os
import sys
import shutil
import subprocess

# Auto-install python-docx and openpyxl if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import openpyxl
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
    import openpyxl

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Formatting helper functions for Word
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=180, after=60):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(before / 20)
    heading.paragraph_format.space_after = Pt(after / 20)
    heading.paragraph_format.keep_with_next = True
    
    # Custom colors for headings
    run = heading.runs[0]
    run.font.name = "Arial"
    if level == 1:
        run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40) # Deep Navy
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold/Bronze
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63) # Gray
        run.font.size = Pt(11.5)
    return heading

def add_paragraph_with_spacing(doc, text="", before=0, after=90, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

# 1. SETUP DIRECTORIES
base_dir = r"d:\Sao Vàng\Website-SaoVang\BaoCao_ThuVien_CuaNhom"
master_dir = os.path.join(base_dir, "THƯ VIỆN HỆ NHÔM SAO VÀNG")

# 2. GENERATE MASTER DATABASE XLSX
def create_master_database():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Master Database"
    ws.views.sheetView[0].showGridLines = True
    
    font_title = Font(name="Arial", size=15, bold=True, color="0D2240")
    font_header = Font(name="Arial", size=9.5, bold=True, color="FFFFFF")
    font_body = Font(name="Arial", size=9)
    font_body_bold = Font(name="Arial", size=9, bold=True)
    
    fill_header = PatternFill(start_color="0D2240", end_color="0D2240", fill_type="solid")
    fill_premium = PatternFill(start_color="FFF9E6", end_color="FFF9E6", fill_type="solid") # Gold/Luxury
    fill_standard = PatternFill(start_color="E6F4EA", end_color="E6F4EA", fill_type="solid") # Green/Standard
    fill_economy = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Pink/Economy
    fill_zebra = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    
    border_thin = Side(border_style="thin", color="D1D5DB")
    border_thick = Side(border_style="medium", color="0D2240")
    border_double = Side(border_style="double", color="4B5563")
    
    box_border = Border(left=border_thin, right=border_thin, top=border_thin, bottom=border_thin)
    header_border = Border(left=border_thin, right=border_thin, top=border_thick, bottom=border_thick)
    bottom_heavy_border = Border(bottom=border_double, left=border_thin, right=border_thin)
    
    ws.merge_cells("A1:I2")
    ws["A1"] = "CƠ SỞ DỮ LIỆU GỐC HỆ THỐNG CỬA NHÔM KÍNH SAO VÀNG (MASTER DATABASE)"
    ws["A1"].font = font_title
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    
    headers = ["ID", "Nhóm", "Hệ Nhôm", "Hãng Sản Xuất", "Độ Dày (mm)", "Phụ Kiện Tương Thích", "Gioăng & Keo", "Bề Mặt & Đặc Tính", "Link Bản Vẽ Tham Khảo"]
    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col_idx)
        cell.value = h
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = header_border
        
    db_rows = [
        # Civro
        ("1", "Cầu cách nhiệt", "AW55IN/OU", "Civro (Đức)", "1.4mm - 1.8mm", "Sobinco, CMECH rãnh C", "EPDM 3 lớp, Dow Corning", "MED mịn mờ, cách âm 45dB", "03_Thư viện CAD/Cửa sổ/AW55.dwg"),
        ("2", "Cầu cách nhiệt", "AW65IN/OU", "Civro (Đức)", "1.4mm - 1.8mm", "Sobinco, CMECH rãnh C", "EPDM 3 lớp, Dow Corning", "Cầu cách nhiệt Technoform, MED", "03_Thư viện CAD/Cửa sổ/AW65.dwg"),
        ("3", "Cầu cách nhiệt", "AW75IN", "Civro (Đức)", "1.8mm", "Sobinco, tay nắm đặc chủng", "EPDM 3 lớp, Dow Corning", "Cầu cách nhiệt, lưới chống muỗi", "03_Thư viện CAD/Cửa sổ/AW75.dwg"),
        ("4", "Cầu cách nhiệt", "AD55/65/75", "Civro (Đức)", "2.0mm", "Sobinco, CMECH khóa an ninh", "EPDM 3 lớp, Dow Corning", "Cửa đi mở quay cách nhiệt, khóa RC2", "03_Thư viện CAD/Cửa đi/AD65.dwg"),
        ("5", "Slim Sliding", "CSD80", "Civro (Đức)", "1.8mm - 2.2mm", "CMECH tay khóa dẹt", "EPDM nén phẳng sàn, Dow", "Cửa trượt nâng Minimalist panorama", "03_Thư viện CAD/Cửa Slim/CSD80.dwg"),
        ("6", "Cửa đi lùa", "CSD130", "Civro (Đức)", "2.0mm - 3.0mm", "Sobinco nâng hạ, tải 400kg", "EPDM đa khoang nén chặt", "Cửa đi trượt nâng Lift & Slide khủng", "03_Thư viện CAD/Cửa lùa/CSD130.dwg"),
        # Maxpro JP
        ("7", "Classic/Luxury", "Maxpro Hệ 55", "Maxpro JP (Nhật)", "1.4mm - 2.0mm", "Maxpro đồng bộ, Draho", "EPDM kép, Dow Corning", "Sơn Anodise ED bảo hành 25 năm", "08_Catalogue theo Hãng/Maxpro.pdf"),
        ("8", "Premium", "Maxpro Hệ 65", "Maxpro JP (Nhật)", "1.6mm - 2.0mm", "Cmech, Roto rãnh C", "EPDM 3 tầng kín khít", "Hệ rãnh C chuẩn Âu cánh phẳng", "03_Thư viện CAD/Cửa đi/Maxpro65.dwg"),
        ("9", "Tân Cổ Điển", "Maxpro Hệ 83", "Maxpro JP (Nhật)", "1.6mm - 2.0mm", "Cmech, Roto bản lề 3D", "EPDM cao cấp, keo bọt PU", "Cánh phào chỉ bo tròn, kính hộp 32mm", "03_Thư viện CAD/Cửa đi/Maxpro83.dwg"),
        ("10", "Cửa đi lùa", "Maxpro Hệ 115", "Maxpro JP (Nhật)", "2.0mm - 3.0mm", "Cmech, Roto nâng hạ", "Gioăng nén EPDM chống bão", "Lift & Slide chống bão gió 2500 Pa", "03_Thư viện CAD/Cửa lùa/Maxpro115.dwg"),
        # Viralwindow
        ("11", "Classic/Luxury", "VRA55", "Viralwindow (VN)", "1.4mm - 2.0mm", "Viralwindow, Draho", "EPDM, silicone", "Cánh phẳng khung, kính lõm", "08_Catalogue theo Hãng/Viralwindow.pdf"),
        ("12", "Cửa đi lùa", "VRA94", "Viralwindow (VN)", "1.4mm", "Viralwindow đồng bộ", "Gioăng lông, silicone", "Cửa lùa lướt êm 2 ray", "03_Thư viện CAD/Cửa lùa/VRA94.dwg"),
        ("13", "Cửa đi lùa", "VRA120", "Viralwindow (VN)", "1.4mm", "Viralwindow đồng bộ", "Gioăng lông, silicone", "Cửa lùa lướt êm 3 ray", "03_Thư viện CAD/Cửa lùa/VRA120.dwg"),
        ("14", "Premium", "VRE65", "Viralwindow (VN)", "1.6mm - 2.0mm", "Cmech, Roto rãnh C", "EPDM 3 lớp kín khít", "Rãnh C chuẩn Châu Âu, cánh ôm khung", "03_Thư viện CAD/Cửa đi/VRE65.dwg"),
        ("15", "Cửa đi lùa", "VRE120", "Viralwindow (VN)", "2.0mm", "Viralwindow nâng hạ", "EPDM chống bão", "Trượt nâng Lift & Slide VIP", "03_Thư viện CAD/Cửa lùa/VRE120.dwg"),
        # Owin
        ("16", "Trượt Quay", "Owin Trượt Quay", "Owin (Đức/VN)", "1.6mm - 3.5mm", "Owin trượt quay đồng bộ", "EPDM kép chèn nén", "Không ray dưới sàn, mở rộng 100%", "03_Thư viện CAD/Cửa lùa/OwinTurn.dwg"),
        ("17", "Thủy Lực", "Owin Thủy Lực", "Owin (Đức/VN)", "2.0mm", "Adler/Hafele bản lề sàn", "Keo kết cấu, gioăng kép", "Cánh phào chỉ rộng 120-180mm", "03_Thư viện CAD/Cửa đi/OwinThuyLuc.dwg"),
        ("18", "Premium", "Owin Hệ 55", "Owin (Đức/VN)", "1.4mm - 2.0mm", "Cmech, Roto rãnh C", "EPDM kép Đức", "Rãnh C chuẩn Đức mở quay", "03_Thư viện CAD/Cửa sổ/Owin55.dwg"),
        # Yangli
        ("19", "Cạnh Vát", "Yangli Hệ 55", "Yangli (VN)", "1.2mm - 1.4mm", "Kinlong, PMA, Draho", "EPDM, silicone", "Vát cạnh nghiêng 35 độ thoát nước", "03_Thư viện CAD/Cửa sổ/Yangli55.dwg"),
        ("20", "Cửa đi lùa", "Yangli Hệ 93", "Yangli (VN)", "1.4mm - 1.8mm", "Kinlong, Draho", "Gioăng lông, silicone", "Cửa trượt lùa 2 ray phổ thông", "03_Thư viện CAD/Cửa lùa/Yangli93.dwg"),
        # Topal
        ("21", "Premium", "Topal Prima", "Topal - Austdoor", "1.3mm - 2.0mm", "Prima đồng bộ, Cmech", "EPDM 2 thành phần", "Cánh trơn phẳng 63mm, cách âm EN140", "03_Thư viện CAD/Cửa đi/TopalPrima.dwg"),
        ("22", "Standard", "Topal Slima", "Topal - Austdoor", "1.2mm - 1.5mm", "Slima đồng bộ, Sigico", "EPDM kép", "Cánh phẳng không gân tối giản", "03_Thư viện CAD/Cửa đi/TopalSlima.dwg"),
        ("23", "Standard", "Topal XFAD", "Topal - Austdoor", "1.4mm - 2.0mm", "Kinlong, Draho rãnh 22", "EPDM tiêu chuẩn", "Biên dạng Xingfa có gân tăng cứng", "03_Thư viện CAD/Cửa đi/TopalXFAD.dwg"),
        # EuroVN
        ("24", "Trượt Quay", "EuroVN Trượt Quay", "EuroVN (QueenViet)", "1.6mm - 2.8mm", "EuroVN đồng bộ, Janus", "EPDM kép dẻo dai", "Trượt quay 4 cánh không ray dưới", "03_Thư viện CAD/Cửa lùa/EuroVNTurn.dwg"),
        ("25", "Thủy Lực", "EuroVN Thủy Lực", "EuroVN (QueenViet)", "2.0mm", "Adler bản lề sàn đồng bộ", "Keo kết cấu, gioăng kép", "Cánh phào chỉ soi hèm giả gỗ 120-180", "03_Thư viện CAD/Cửa đi/EuroVNThuyLuc.dwg"),
        ("26", "An Ninh", "EuroVN Chấn Song", "EuroVN (QueenViet)", "1.5mm - 2.0mm", "Phụ kiện chấn song đồng bộ", "EPDM, silicone", "Tích hợp sẵn chấn song nhôm bảo vệ", "03_Thư viện CAD/Cửa sổ/EuroVNChanSong.dwg"),
        ("27", "Nội Thất", "EuroVN Omega", "EuroVN (QueenViet)", "1.0mm - 1.2mm", "Bản lề giảm chấn, tay nắm", "Silicon, gioăng kính", "Nhôm tủ bếp giả gỗ Omega Deco", "04_Thư viện Profile Nhôm/OmegaDeco.dwg"),
        # PAG
        ("28", "Cầu cách nhiệt", "PAG Hệ 60", "PAG (Đức/VN)", "1.8mm", "Hopo rãnh C, Sigico", "EPDM đa khoang cách âm 40dB", "Cầu cách nhiệt PA66, Hàn góc seamless", "03_Thư viện CAD/Cửa sổ/PAG60.dwg"),
        ("29", "Cầu cách nhiệt", "PAG Hệ 65", "PAG (Đức/VN)", "2.0mm", "Hopo, CMECH rãnh C", "EPDM đa khoang, keo PU", "Cầu cách nhiệt PA66, Hàn góc liền khối", "03_Thư viện CAD/Cửa đi/PAG65.dwg"),
        ("30", "Cửa đi lùa", "PAG Hệ 115/120", "PAG (Đức/VN)", "2.0mm - 3.0mm", "Hopo rãnh C nâng hạ", "Gioăng nén EPDM phẳng sàn", "Lift & Slide cách âm cách nhiệt tốt", "03_Thư viện CAD/Cửa lùa/PAG115.dwg"),
        # Slim Papo/Heckler
        ("31", "Slim Interior", "Slim Lùa Treo", "Slim (Đức/VN)", "1.6mm - 2.0mm", "OPK, Papo bánh xe treo", "Gioăng cao su kẹp kính", "Lùa treo không ray dưới sàn, có giảm chấn", "03_Thư viện CAD/Cửa Slim/SlimLuaTreo.dwg"),
        ("32", "Slim Interior", "Slim Ray Dẹt", "Slim (Đức/VN)", "1.6mm", "OPK bánh xe dưới, dẫn hướng", "Gioăng cao su, silicone", "Lùa liên động đa cánh ray phẳng dẹt", "03_Thư viện CAD/Cửa Slim/SlimRayDet.dwg"),
        ("33", "Slim Exterior", "Slim Ngoại Thất", "Slim (Đức/VN)", "2.0mm", "Khóa tay nắm dẹt đa điểm Slim", "Gioăng EPDM nhiều lớp chống bão", "Cửa đi lùa ngoại thất viền mảnh kính hộp", "03_Thư viện CAD/Cửa Slim/SlimNgoaiThat.dwg"),
        # Soco System
        ("34", "Anodized", "Soco 65 Quay", "Soco (VN/Yongxing)", "1.6mm - 2.0mm", "Cmech, Hopo rãnh C", "EPDM kép nhiều khoang", "Anodized Champagne Ý, rãnh C chuẩn Âu", "03_Thư viện CAD/Cửa đi/Soco65.dwg"),
        ("35", "Anodized", "Soco 94 Lùa", "Soco (VN/Yongxing)", "1.4mm - 1.6mm", "Hopo, Draho bánh xe chốt sập", "Gioăng lông, silicone", "Anodized Champagne lướt ray inox 2 ray", "03_Thư viện CAD/Cửa lùa/Soco94.dwg"),
        ("36", "Anodized", "Soco 180 Lùa", "Soco (VN/Yongxing)", "1.4mm - 2.0mm", "Cmech bánh xe chốt khóa tải", "EPDM đệm cao su kép dẻo dai", "Anodized Champagne trượt lùa 3 ray", "03_Thư viện CAD/Cửa lùa/Soco180.dwg")
    ]
    
    for row_idx, r_data in enumerate(db_rows, start=5):
        fill_to_apply = fill_zebra if row_idx % 2 == 1 else None
        seg = r_data[1]
        
        if "Cầu cách nhiệt" in seg or "Anodized" in seg or "Luxury" in seg:
            seg_fill = fill_premium
        elif "Slim" in seg or "Trượt Quay" in seg or "Thủy Lực" in seg:
            seg_fill = fill_standard
        else:
            seg_fill = fill_economy
            
        for col_idx, val in enumerate(r_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = val
            cell.font = font_body
            cell.border = box_border
            
            if col_idx == 2:
                cell.fill = seg_fill
                cell.font = font_body_bold
            elif fill_to_apply:
                cell.fill = fill_to_apply
                
            if col_idx in [1, 2, 3, 5]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left
                
    # Add heavy bottom border
    last_row_idx = len(db_rows) + 4
    for col in range(1, 10):
        ws.cell(row=last_row_idx, column=col).border = bottom_heavy_border
        
    # Auto-fit column widths
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.row in [1, 2]:
                continue
            val_str = str(cell.value or '')
            if len(val_str) > max_len:
                max_len = len(val_str)
        calculated_width = max(max_len * 1.05 + 4, 10)
        ws.column_dimensions[col_letter].width = min(calculated_width, 42)
        
    # Save Excel with Vietnamese name
    output_path = os.path.join(master_dir, "02_Cơ sở Dữ liệu Gốc.xlsx")
    wb.save(output_path)
    print("Master Database XLSX saved successfully in Vietnamese.")

# 3. GENERATE MASTER TECHNICAL HANDBOOK DOCX
def create_master_handbook():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page ---
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(120)
    
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("CÔNG TY CỔ PHẦN SẢN XUẤT CƠ KHÍ SAO VÀNG")
    run_org.font.name = "Arial"
    run_org.font.size = Pt(12)
    run_org.bold = True
    run_org.font.color.rgb = RGBColor(0x70, 0x5D, 0x30)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main_title.paragraph_format.space_before = Pt(40)
    p_main_title.paragraph_format.space_after = Pt(20)
    run_main_title = p_main_title.add_run("THƯ VIỆN HỆ NHÔM SAO VÀNG\n01_SỔ TAY KỸ THUẬT TỔNG HỢP")
    run_main_title.font.name = "Arial"
    run_main_title.font.size = Pt(22)
    run_main_title.bold = True
    run_main_title.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Cẩm nang tra cứu kỹ thuật hệ cửa nhôm kính toàn diện - Sách tra cứu nội bộ\nTổng hợp so sánh chéo 10 hãng nhôm và 80 hệ sản phẩm trên thị trường Việt Nam")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(150)
    run_date = p_date.add_run("Năm 2026 - Lưu hành nội bộ")
    run_date.font.name = "Arial"
    run_date.font.size = Pt(10.5)
    run_date.bold = True
    
    doc.add_page_break()
    
    # --- Table of Contents ---
    add_heading_with_spacing(doc, "MỤC LỤC TRA CỨU HỆ THỐNG", level=1)
    toc_lines = [
        "Volume 01: Giới thiệu chung & Quy chuẩn thuật ngữ kỹ thuật",
        "Volume 02: Hệ Cửa đi (Mở quay, Trượt quay, Thủy lực, Xếp trượt)",
        "Volume 03: Hệ Cửa sổ (Mở quay, Mở hất, Mở lật, Trượt lùa)",
        "Volume 04: Chuyên đề Hệ cửa nhôm Slim (Nội thất & Ngoại thất)",
        "Volume 05: Chuyên đề Nhôm cầu cách nhiệt (Thermal Break & Passive House)",
        "Volume 06: Chuyên đề Vách mặt dựng kính (Hệ vách lớn)",
        "Volume 07: Hệ thống Lan can kính không chôn chân đế",
        "Volume 08: Hệ thống Mái kính lấy sáng & Pergola tự động",
        "Volume 09: Thư viện Phụ kiện kim khí đồng bộ (Rãnh C vs Rãnh 22)",
        "Volume 10: Quy chuẩn mặt cắt thanh Profile nhôm định hình",
        "Volume 11: Bản vẽ kỹ thuật & Quy chuẩn Shopdrawing",
        "Volume 12: Hướng dẫn Lắp dựng & Nghiệm thu công trường",
        "Volume 13: Bảng so sánh ma trận hãng & Tra cứu nhanh giải pháp"
    ]
    for line in toc_lines:
        add_paragraph_with_spacing(doc, line, size=10.5, bold=True)
        
    doc.add_page_break()
    
    # --- Volume 01 ---
    add_heading_with_spacing(doc, "VOLUME 01: GIỚI THIỆU CHUNG & QUY CHUẨN THUẬT NGỮ KỸ THUẬT", level=1)
    add_paragraph_with_spacing(doc, 
        "Chào mừng bạn đến với THƯ VIỆN HỆ NHÔM SAO VÀNG. Đây là tài liệu cẩm nang kỹ thuật tích hợp "
        "nhằm chuẩn hóa toàn bộ các hệ cửa nhôm, vách mặt dựng, lan can kính và mái kính đang lưu hành. "
        "Quyển tài liệu này được biên soạn không chia theo hãng riêng lẻ, mà chia theo HỆ SẢN PHẨM làm trọng tâm, "
        "tổng hợp tất cả giải pháp thiết kế của các hãng trong cùng một chương để đội ngũ Kỹ sư, Kiến trúc sư, "
        "Gia công và Shopdrawing dễ tra cứu, so sánh trực quan.\n\n"
        "Thuật ngữ chuẩn hóa:\n"
        "• Thermal Break (Nhôm cầu cách nhiệt): Thanh profile có dải Polyamide cách nhiệt Technoform chia đôi khoang nhôm.\n"
        "• MED Coating (Sơn điện di mờ): Lớp phủ bảo vệ siêu cứng kháng muối biển mặn.\n"
        "• Slide & Turn (Trượt quay): Cửa kết hợp trượt và quay không ray dưới sàn.\n"
        "• Soft Closing (Giảm chấn thủy lực): Piston tự động hãm tốc khi cửa gần đóng khít."
    )
    
    # --- Volume 02 ---
    add_heading_with_spacing(doc, "VOLUME 02: HỆ CỬA ĐI (DOOR SYSTEMS)", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ thống cửa đi bao gồm các giải pháp từ phổ thông mở quay hệ 55, cửa trượt lùa ray inox hệ 93, "
        "đến các giải pháp đặc chủng cao cấp như cửa mở trượt quay Owin/EuroVN không ray dưới, và cửa đi thủy lực bản lớn phào nổi đại sảnh."
    )
    
    # Mẫu chuẩn hệ Cửa trượt quay
    add_heading_with_spacing(doc, "CHƯƠNG: CỬA TRƯỢT QUAY THÔNG MINH (SLIDE & TURN)", level=2)
    
    form_lines = [
        ("Tên hệ / Mã hệ", "Hệ cửa mở trượt quay Owin / EuroVN Trượt quay"),
        ("Loại", "Cửa đi đặc chủng thông minh"),
        ("Công năng", "Kết hợp đồng thời cơ cấu trượt lùa và mở quay trên cùng một bộ cửa, dồn xếp gọn cánh về góc tường mở rộng 100% diện tích thông thủy."),
        ("Ứng dụng", "Cửa ra vào mặt tiền nhà phố, lối thông phòng khách ra sân vườn biệt thự."),
        ("Ảnh hoàn thiện / Mặt cắt / Profile", "[Xem chi tiết file đính kèm thư mục: 10_Hình ảnh & Video/SlideTurn_Finished.png và 04_Thư viện Profile Nhôm/SlideTurn_Section.dwg]"),
        ("CAD / BIM / Shopdrawing", "[Đường dẫn: 03_Thư viện CAD/Cửa lùa/OwinTurn.dwg và 03_Thư viện CAD/Cửa lùa/EuroVNTurn.dwg]"),
        ("Phụ kiện tương thích", "Phụ kiện đồng bộ đặc chủng trượt quay Owin, Janus, Sigico chịu lực ray treo tốt."),
        ("Kính & Gioăng chèn", "Kính cường lực Temper 8mm - 12mm; Gioăng cao su đệm kép EPDM đàn hồi tốt."),
        ("Thông số kỹ thuật", "Độ dày nhôm từ 1.6mm - 2.0mm (riêng thanh ray động dày tới 2.8mm - 3.5mm); không có ray dưới sàn (treo trên hoàn toàn)."),
        ("Ưu điểm", "Mở rộng 100% không gian; không có ray dưới sàn tránh vấp ngã chân, quét dọn dễ dàng; lướt êm ái đầm tay."),
        ("Nhược điểm", "Yêu cầu kỹ thuật gia cố trần treo sắt hộp cực kỳ khắt khe để chống xệ cánh; chi phí lắp đặt cao."),
        ("Các hãng đang sản xuất", "Owin (tiêu chuẩn Đức), EuroVN (QueenViet Group), Viralwindow VRX75."),
        ("So sánh & Đánh giá", "Cửa trượt quay Owin có ray inox tròn đùn dày dặn hơn; trong khi EuroVN có giá thành kinh tế hơn và có màu vân gỗ phủ film bóng đẹp."),
        ("Nguồn tài liệu gốc", "[Tải catalogue hãng tại thư mục: 08_Catalogue theo Hãng/ThuVien_HeCuaNhom_Owin.docx]")
    ]
    
    table = doc.add_table(rows=len(form_lines), cols=2)
    table.alignment = 1
    for idx, (label, val) in enumerate(form_lines):
        row = table.rows[idx]
        bg = "F7F9FC" if idx % 2 == 1 else "FFFFFF"
        
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(1.8)
        set_cell_background(cell_lbl, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.name = "Arial"
        run_lbl.font.size = Pt(9.5)
        if idx == 0:
            run_lbl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        cell_val = row.cells[1]
        cell_val.width = Inches(4.7)
        set_cell_background(cell_val, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        run_val.font.name = "Arial"
        run_val.font.size = Pt(9)
        if idx == 0:
            run_val.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run_val.bold = True

    doc.add_page_break()
    
    # --- Volume 04 ---
    add_heading_with_spacing(doc, "VOLUME 04: CHUYÊN ĐỀ HỆ SLIM NỘI THẤT & NGOẠI THẤT", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ nhôm Slim đại diện cho xu thế thiết kế tối giản hiện đại (Minimalism) với khung viền siêu mảnh chỉ rộng từ 16mm - 25mm. "
        "Hệ lùa treo trong nhà tích hợp piston giảm chấn thủy lực tự động phanh hãm tốc độ đóng chậm lại chống kẹt tay va đập kính."
    )
    
    # Mẫu chuẩn hệ Slim
    add_heading_with_spacing(doc, "CHƯƠNG: CỬA LÙA TREO SLIM GIẢM CHẤN (SOFT CLOSING SLIM)", level=2)
    
    slim_form_lines = [
        ("Tên hệ / Mã hệ", "Hệ cửa lùa Slim giảm chấn thủy lực nội thất"),
        ("Loại", "Cửa đi trượt lùa siêu mỏng"),
        ("Công năng", "Lướt trượt nhẹ nhàng trên ray treo trên không ray dưới sàn, tích hợp piston giảm chấn đóng mở 2 chiều êm ái."),
        ("Ứng dụng", "Cửa đi phòng bếp, phòng thay đồ, vách ngăn phòng họp văn phòng hẹp."),
        ("CAD / BIM / Shopdrawing", "[Đường dẫn: 03_Thư viện CAD/Cửa Slim/SlimLuaTreo.dwg và 03_Thư viện CAD/Cửa Slim/SlimRayDet.dwg]"),
        ("Phụ kiện tương thích", "Phụ kiện trượt treo giảm chấn thủy lực OPK, Papo, Heckler nhập khẩu."),
        ("Kính & Gioăng chèn", "Kính cường lực an toàn Temper 8mm - 10mm; gioăng cao su kẹp kính chèn mép."),
        ("Thông số kỹ thuật", "Độ dày nhôm 1.6mm - 2.0mm đùn bằng hợp kim 6063-T5; viền nhôm đứng mảnh 16mm; tải trọng treo trần treo từ 80kg - 150kg."),
        ("Ưu điểm", "Thiết kế tối giản sang trọng, phẳng mặt sàn đi lại an toàn tuyệt đối; piston tự hãm chống va đập vỡ kính cực tốt."),
        ("Nhược điểm", "Bắt buộc phải gia cố dầm sắt hộp chịu tải trọng cánh kính treo trên trần thạch cao; không thích hợp làm cửa ngoài trời."),
        ("Các hãng sản xuất", "Heckler, Papo, Kogen, Yangli Slim, Soco Slim, Owin Slim."),
        ("So sánh & Đánh giá", "Hệ Slim Heckler và Papo dẫn đầu về độ tinh xảo của bánh xe và piston giảm chấn. Yangli Slim có giá thành kinh tế hơn, phù hợp cho nhà chung cư trung cấp."),
        ("Nguồn tài liệu gốc", "[Tải catalogue hãng tại thư mục: 08_Catalogue theo Hãng/ThuVien_HeCuaNhom_Slim.docx]")
    ]
    
    table_slim = doc.add_table(rows=len(slim_form_lines), cols=2)
    table_slim.alignment = 1
    for idx, (label, val) in enumerate(slim_form_lines):
        row = table_slim.rows[idx]
        bg = "F7F9FC" if idx % 2 == 1 else "FFFFFF"
        
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(1.8)
        set_cell_background(cell_lbl, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.name = "Arial"
        run_lbl.font.size = Pt(9.5)
        if idx == 0:
            run_lbl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        cell_val = row.cells[1]
        cell_val.width = Inches(4.7)
        set_cell_background(cell_val, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        run_val.font.name = "Arial"
        run_val.font.size = Pt(9)
        if idx == 0:
            run_val.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run_val.bold = True

    doc.add_page_break()
    
    # --- Volume 05 ---
    add_heading_with_spacing(doc, "VOLUME 05: CHUYÊN ĐỀ NHÔM CẦU CÁCH NHIỆT (THERMAL BREAK)", level=1)
    add_paragraph_with_spacing(doc, 
        "Hệ thống nhôm cầu cách nhiệt (Thermal Break) là giải pháp cao cấp nhất cản sự truyền nhiệt từ môi trường. "
        "Bằng việc sử dụng dải cầu Polyamide Technoform đặt giữa 2 khoang nhôm, kết hợp gioăng EPDM đa khoang kín khít đạt độ cách âm tới 45dB."
    )
    
    # Mẫu chuẩn hệ Cầu cách nhiệt
    add_heading_with_spacing(doc, "CHƯƠNG: CỬA SỔ MỞ QUAY CẦU CÁCH NHIỆT (THERMAL BREAK WINDOW)", level=2)
    
    tb_form_lines = [
        ("Tên hệ / Mã hệ", "Hệ cửa sổ mở quay/hất cầu cách nhiệt AW55 / AW65 / AW75 / PAG 60"),
        ("Loại", "Cửa sổ cách nhiệt cách âm siêu cao cấp"),
        ("Công năng", "Cản nhiệt truyền qua khung nhôm, cách âm chống ồn tuyệt đối đạt tới 45dB, tiết kiệm 30-40% điện năng điều hòa."),
        ("Ứng dụng", "Cửa sổ phòng ngủ biệt thự triệu đô, resort ven biển hạng sang, penthouses cao tầng."),
        ("CAD / BIM / Shopdrawing", "[Đường dẫn: 03_Thư viện CAD/Cửa sổ/AW65.dwg và 03_Thư viện CAD/Cửa sổ/PAG60.dwg]"),
        ("Phụ kiện tương thích", "Phụ kiện rãnh C Châu Âu chốt khóa đa điểm âm của Sobinco (Bỉ), CMECH (Mỹ), Hopo rãnh C."),
        ("Kính & Gioăng chèn", "Kính hộp Low-E cách âm cách nhiệt Double Glazing Low-E; gioăng cao su EPDM 3 lớp tiêu chuẩn hàng không."),
        ("Thông số kỹ thuật", "Độ dày nhôm 1.8mm - 2.0mm đùn bằng hợp kim 6060-T6/6063-T6; dải cầu cách nhiệt PA66 Technoform Đức; sơn phủ MED mờ chống xước muối mặn."),
        ("Ưu điểm", "Khả năng chống truyền nhiệt tuyệt hảo, chống đọng nước sương mép kính; cách âm vượt trội; hàn góc liền khối seamless mịn màng tuyệt đẹp (ở hệ PAG)."),
        ("Nhược điểm", "Giá thành thuộc phân khúc siêu đắt đỏ; gia công đòi hỏi máy móc chuyên dụng và thợ kỹ thuật tay nghề cao."),
        ("Các hãng sản xuất", "Civro (Đức), PAG (Đức/VN), Maxpro JP (Nhật), Topal Prima, EuroVN VIP (Gold)."),
        ("So sánh & Đánh giá", "Civro đứng đầu về độ thẩm mỹ lớp sơn MED mờ và tính đồng bộ cao cấp. PAG nổi bật với công nghệ hàn góc liền khối liền mạch không khe hở góc. Maxpro JP có độ bền màu điện di ED Anodize bảo hành 25 năm."),
        ("Nguồn tài liệu gốc", "[Tải catalogue hãng tại thư mục: 08_Catalogue theo Hãng/ThuVien_HeCuaNhom_Civro.docx]")
    ]
    
    table_tb = doc.add_table(rows=len(tb_form_lines), cols=2)
    table_tb.alignment = 1
    for idx, (label, val) in enumerate(tb_form_lines):
        row = table_tb.rows[idx]
        bg = "F7F9FC" if idx % 2 == 1 else "FFFFFF"
        
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(1.8)
        set_cell_background(cell_lbl, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.name = "Arial"
        run_lbl.font.size = Pt(9.5)
        if idx == 0:
            run_lbl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        cell_val = row.cells[1]
        cell_val.width = Inches(4.7)
        set_cell_background(cell_val, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        run_val.font.name = "Arial"
        run_val.font.size = Pt(9)
        if idx == 0:
            run_val.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run_val.bold = True

    doc.add_page_break()
    
    # --- Volume 13 ---
    add_heading_with_spacing(doc, "VOLUME 13: BẢNG SO SÁNH MA TRẬN HÃNG & TRA CỨU NHANH GIẢP PHÁP", level=1)
    
    add_heading_with_spacing(doc, "1. BẢNG SO SÁNH MA TRẬN GIẢI PHÁP HỆ NHÔM", level=2)
    add_paragraph_with_spacing(doc, "Dưới đây là ma trận chéo đánh giá khả năng cung ứng và đặc thù giải pháp giữa các hãng nhôm tiêu biểu:")
    
    # Mat trận so sanh chéo
    matrix_headers = ["Tiêu Chí Đánh Giá", "Civro", "Soco", "Maxpro JP", "Topal", "Viralwindow", "Owin", "Yangli", "EuroVN", "PAG", "Xingfa QD"]
    matrix_data = [
        ("Cầu cách nhiệt (Thermal)", "✓", "✗", "✓", "✓", "✗", "✗", "✗", "✓", "✓", "✓"),
        ("Xử lý bề mặt Anodized", "✓", "✓", "✓", "✗", "✗", "✓", "✗", "✗", "✗", "✗"),
        ("Thiết kế rãnh C chuẩn Âu", "✓", "✓", "✓", "✓", "✓", "✓", "✗", "✓", "✓", "✗"),
        ("Trượt quay không ray dưới", "✗", "✗", "✗", "✗", "✓", "✓", "✗", "✓", "✗", "✗"),
        ("Cửa thủy lực bản lớn phào", "✗", "✗", "✓", "✗", "✗", "✓", "✗", "✓", "✗", "✓"),
        ("Tích hợp sẵn chấn song nhôm", "✗", "✗", "✗", "✗", "✗", "✗", "✗", "✓", "✗", "✗"),
        ("Hàn góc liền khối Seamless", "✗", "✗", "✗", "✗", "✗", "✗", "✗", "✗", "✓", "✗"),
        ("Bảo hành sơn phủ bề mặt", "30 năm", "30 năm", "25 năm", "5-10 năm", "10-20 năm", "10 năm", "15 năm", "10 năm", "15 năm", "5 năm"),
        ("Phân khúc khách hàng", "Siêu cao cấp", "Cao cấp", "Cao cấp", "Trung-Cao", "Trung cấp", "Trung cấp", "Tiết kiệm", "Tiết kiệm", "Cao cấp", "Phổ thông")
    ]
    
    table_matrix = doc.add_table(rows=len(matrix_data)+1, cols=11)
    table_matrix.alignment = 1
    
    # Populate matrix header
    hdr_row = table_matrix.rows[0]
    for idx, name in enumerate(matrix_headers):
        cell = hdr_row.cells[idx]
        set_cell_background(cell, "0D2240")
        set_cell_margins(cell, top=100, bottom=100, left=60, right=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    for row_idx, r_data in enumerate(matrix_data, start=1):
        row = table_matrix.rows[row_idx]
        bg = "F7F9FC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, val in enumerate(r_data):
            cell = row.cells[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
            p = cell.paragraphs[0]
            if col_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(8)
            if val in ["✓", "Siêu cao cấp", "Cao cấp"]:
                run.bold = True
                run.font.color.rgb = RGBColor(0x70, 0x5D, 0x30) # Gold
                
    add_paragraph_with_spacing(doc, "\n")
    
    # 2. Hướng dẫn Tra cứu nhanh
    add_heading_with_spacing(doc, "2. HƯỚNG DẪN TRA CỨU NHANH THEO CÔNG TRÌNH", level=2)
    
    quick_search_lines = [
        ("Mục tiêu công trình / Ngân sách", "Hãng đề xuất & Mã hệ tiêu biểu"),
        ("Biệt thự / Resort biển siêu cao cấp", "Civro AW75/AD75 (Đức) hoặc Soco 180 Anodized (phong cách Ý), chịu mặn cực cao."),
        ("Căn hộ Penthouse / Tối giản Panorama", "Cửa Slim papo/heckler lùa treo không ray dưới sàn, Civro CSD80 Minimalist."),
        ("Nhà phố / Chung cư trung cấp", "Topal Slima cánh trơn phẳng tối giản, Owin trượt quay mở rộng 100% tiện ích."),
        ("Công trình kinh tế / Ngân sách thấp", "Yangli Hệ 55 vát cạnh mưa dốc nước nhanh, EuroVN Xingfa có gân nổi chịu lực."),
        ("Cửa đại sảnh / Lâu đài tân cổ điển", "Maxpro JP Hệ 83 phào bo tròn kính hộp 32mm, Owin Thủy lực phào chỉ nổi 180mm."),
        ("Chung cư tầng cao an toàn trẻ nhỏ", "EuroVN tích hợp chấn song bảo vệ nhôm đúc đồng bộ cứng vững thay chấn song sắt.")
    ]
    
    table_search = doc.add_table(rows=len(quick_search_lines), cols=2)
    table_search.alignment = 1
    for idx, (label, val) in enumerate(quick_search_lines):
        row = table_search.rows[idx]
        bg = "F7F9FC" if idx % 2 == 1 else "FFFFFF"
        
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.2)
        set_cell_background(cell_lbl, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.name = "Arial"
        run_lbl.font.size = Pt(9.5)
        if idx == 0:
            run_lbl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        cell_val = row.cells[1]
        cell_val.width = Inches(4.3)
        set_cell_background(cell_val, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        run_val.font.name = "Arial"
        run_val.font.size = Pt(9)
        if idx == 0:
            run_val.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run_val.bold = True
            
    add_paragraph_with_spacing(doc, "\n")
    
    # 3. Bảng phân loại
    add_heading_with_spacing(doc, "3. BẢNG PHÂN LOẠI NHÓM HỆ SẢN PHẨM", level=2)
    
    class_lines = [
        ("Nhóm hệ nhôm chính", "Hệ chi tiết / Cấu kiện"),
        ("Cửa đi (Cửa đi)", "Hệ 55, Hệ 60, Hệ 65, Hệ 70, Hệ 75, Hệ 80, Hệ 90, Trượt quay, Thủy lực, Xếp trượt"),
        ("Cửa sổ (Cửa sổ)", "Cửa sổ mở quay, mở hất, mở lật, trượt lùa, cố định, louvre thông gió"),
        ("Hệ Slim (Minimal)", "Slim Interior (lùa treo giảm chấn), Slim Exterior (kính hộp chịu lực), Slim Pocket"),
        ("Cầu cách nhiệt", "Thermal break door, Thermal break window, Thermal break sliding, Vách kính cách nhiệt"),
        ("Vách mặt dựng", "Vách mặt dựng Stick, Unitized, Semi-unitized, Spider glass"),
        ("Cấu kiện phụ trợ", "Lan can nhôm đúc, Lan can kính không chôn chân đế, Mái kính lấy sáng, Canopy, Pergola")
    ]
    
    table_class = doc.add_table(rows=len(class_lines), cols=2)
    table_class.alignment = 1
    for idx, (label, val) in enumerate(class_lines):
        row = table_class.rows[idx]
        bg = "F7F9FC" if idx % 2 == 1 else "FFFFFF"
        
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.2)
        set_cell_background(cell_lbl, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_lbl, top=80, bottom=80, left=100, right=100)
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.name = "Arial"
        run_lbl.font.size = Pt(9.5)
        if idx == 0:
            run_lbl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
        cell_val = row.cells[1]
        cell_val.width = Inches(4.3)
        set_cell_background(cell_val, "0D2240" if idx == 0 else bg)
        set_cell_margins(cell_val, top=80, bottom=80, left=100, right=100)
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        run_val.font.name = "Arial"
        run_val.font.size = Pt(9)
        if idx == 0:
            run_val.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run_val.bold = True

    add_paragraph_with_spacing(doc, "\n[KẾT THÚC TÀI LIỆU SỔ TAY KỸ THUẬT TỔNG HỢP - SAO VÀNG]", before=200, italic=True)
    
    # Save Word with Vietnamese name
    output_path = os.path.join(master_dir, "01_Sổ tay Kỹ thuật Tổng hợp.docx")
    doc.save(output_path)
    print("Master Technical Handbook DOCX saved successfully in Vietnamese.")

if __name__ == "__main__":
    create_master_database()
    create_master_handbook()
    print("Sao Vang Master Library generation completed successfully!")
